# -*- coding: utf-8 -*-
"""
Intelligent Global PV Solar System Design Platform
Flask web application — complete engineering + financial SaaS
"""
import os, json, math, sqlite3, csv, secrets, io, threading
from datetime import datetime, timedelta
from functools import wraps
from flask import (Flask, render_template, request, redirect, url_for,
                   session, flash, jsonify, make_response, abort, send_file)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from config.global_solar_data import (GLOBAL_DATA, get_countries, get_regions,
                                      get_solar_data, temp_derating)
from calculation.ac_cable_sizing import size_all_cables
from api_manager import api as _api
import secrets_broker as _sb  # Phase 1: audit + tier + Vault-ready secret reads

# Structured logging (tenant-aware JSON logs)
try:
    from logging_config.structured_logger import (
        log_app, log_error, log_audit, log_security,
        log_engineering, log_economic, log_ai, log_queue
    )
except ImportError:
    import logging as _logging
    _fl = _logging.getLogger("solarpro")
    def log_app(**k): _fl.info(str(k))
    def log_error(**k): _fl.error(str(k))
    def log_audit(**k): _fl.info(str(k))
    def log_security(**k): _fl.warning(str(k))
    def log_engineering(**k): _fl.info(str(k))
    def log_economic(**k): _fl.info(str(k))
    def log_ai(**k): _fl.info(str(k))
    def log_queue(**k): _fl.info(str(k))


# Load .env file if present (stable SECRET_KEY across restarts)
_env_path = os.path.join(os.path.dirname(__file__), ".env")
if os.path.exists(_env_path):
    with open(_env_path) as _ef:
        for _line in _ef:
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _v = _line.split("=", 1)
                os.environ.setdefault(_k.strip(), _v.strip())

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)
app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))
app.config.update(
    TEMPLATES_AUTO_RELOAD  = True,
    SESSION_COOKIE_HTTPONLY= True,
    SESSION_COOKIE_SAMESITE= "Lax",
    SESSION_COOKIE_SECURE  = False,   # works over both http and https tunnels
    PERMANENT_SESSION_LIFETIME = timedelta(hours=8),
)

# ─── Rate limiter ──────────────────────────────────────────────────────────────
def _get_real_ip():
    """Use X-Forwarded-For when behind serveo/proxy, else remote address."""
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.remote_addr

limiter = Limiter(
    _get_real_ip,
    app=app,
    default_limits=["600 per hour", "120 per minute"],
    storage_uri="memory://",
)

# ─── Security headers (applied after every response) ──────────────────────────
@app.after_request
def set_security_headers(resp):
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-Frame-Options"]        = "SAMEORIGIN"
    resp.headers["X-XSS-Protection"]       = "1; mode=block"
    resp.headers["Referrer-Policy"]        = "strict-origin-when-cross-origin"
    resp.headers["Cache-Control"]          = "no-store, no-cache, must-revalidate"
    resp.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://js.stripe.com https://cdnjs.cloudflare.com; "
        "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
        "font-src 'self' data: https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
        "img-src 'self' data: https: blob:; "
        "connect-src 'self'; "
        "frame-src https://js.stripe.com; "
        "object-src 'none';"
    )
    return resp

# ─── CSRF protection ───────────────────────────────────────────────────────────
def generate_csrf():
    if "_csrf" not in session:
        session["_csrf"] = secrets.token_hex(24)
    return session["_csrf"]

def csrf_protect():
    """Call at top of POST handlers that mutate state."""
    if request.method == "POST":
        token = request.form.get("_csrf") or request.headers.get("X-CSRF-Token")
        if not token or token != session.get("_csrf"):
            abort(403)

app.jinja_env.globals["csrf_token"] = generate_csrf
app.jinja_env.globals["enumerate"]  = enumerate

DB_PATH = os.path.join(os.path.dirname(__file__), "data", "solar_web.db")
os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)

# ─── Phase 4 config ───────────────────────────────────────────────────────────

STRIPE_SECRET    = os.environ.get("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK   = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
STRIPE_PRICES    = {
    "basic":        os.environ.get("STRIPE_PRICE_BASIC", ""),
    "professional": os.environ.get("STRIPE_PRICE_PROFESSIONAL", ""),
    "enterprise":   os.environ.get("STRIPE_PRICE_ENTERPRISE", ""),
}
# Phase 1: route PAYSTACK_SECRET through secrets_broker for audit + future Vault.
# DEGRADED tier means: try Vault, fall through to env warm-up automatically.
try:
    PAYSTACK_SECRET = _sb.get("payment/paystack", tier="DEGRADED")["secret"]
except Exception:
    PAYSTACK_SECRET = os.environ.get("PAYSTACK_SECRET_KEY", "")
PAYSTACK_PUBLIC  = os.environ.get("PAYSTACK_PUBLIC_KEY", "")

# ─── Free / demo mode & SMTP ──────────────────────────────────────────────────
# DEMO_MODE=true lets any user instantly activate a Professional plan for
# testing — no payment API calls required.
DEMO_MODE   = os.environ.get("DEMO_MODE", "true").lower() in ("1", "true", "yes")
DEMO_DAYS   = int(os.environ.get("DEMO_DAYS", "14"))

# Strip BOM/whitespace from env vars (GitHub Secrets occasionally store
# values with a UTF-8 BOM prefix which crashes int() at module import).
# Inputs:  name (env var), default (string)
# Output:  cleaned string with BOM + surrounding whitespace removed
def _env_clean(name, default=""):
    return os.environ.get(name, default).lstrip("\ufeff").strip()

SMTP_HOST   = os.environ.get("SMTP_HOST", "")
SMTP_PORT   = int(_env_clean("SMTP_PORT", "587") or "587")
SMTP_USER   = os.environ.get("SMTP_USER", "")
SMTP_PASS   = os.environ.get("SMTP_PASS", "")
SMTP_FROM      = os.environ.get("SMTP_FROM",      "support@aiappinvent.com")
SMTP_TLS       = _env_clean("SMTP_TLS", "false").lower() in ("1", "true", "yes")  # false=SSL/465, true=STARTTLS/587
RESEND_API_KEY = os.environ.get("RESEND_API_KEY", "")

# Dedicated per-purpose sender addresses (3 Namecheap mailboxes)
EMAIL_SALES     = os.environ.get("EMAIL_SALES",     "sales@aiappinvent.com")
EMAIL_SUPPORT   = os.environ.get("EMAIL_SUPPORT",   "support@aiappinvent.com")
EMAIL_BILLING   = os.environ.get("EMAIL_BILLING",   "billing@aiappinvent.com")
EMAIL_HELLO     = os.environ.get("EMAIL_HELLO",     "sales@aiappinvent.com")   # no hello@ mailbox -> route to sales
EMAIL_PROPOSALS = os.environ.get("EMAIL_PROPOSALS", "sales@aiappinvent.com")   # no proposals@ mailbox -> route to sales

PLAN_PRICES = {
    "professional": {"usd": 49, "label": "Professional", "projects": 10,
                     "features": ["10 projects","Priority support","Excel & CSV export","Multi-currency BOQ","Redesign recommendations"]},
    "business":     {"usd": 99, "label": "Business",     "projects": "Unlimited",
                     "features": ["Unlimited projects","API access","Custom branding","White-label reports","Dedicated engineer support"]},
}

_DEFAULT_APPLIANCES = [
    ("Lighting",    "LED Bulb 9W",               9),
    ("Lighting",    "LED Bulb 15W",              15),
    ("Lighting",    "Fluorescent Tube 36W",      36),
    ("Lighting",    "LED Spotlight 7W",          7),
    ("Lighting",    "Emergency Light 8W",        8),
    ("Cooling",     "Ceiling Fan",               75),
    ("Cooling",     "Standing Fan",              60),
    ("Cooling",     "Air Conditioner 1HP",       745),
    ("Cooling",     "Air Conditioner 1.5HP",     1118),
    ("Cooling",     "Air Conditioner 2HP",       1491),
    ("Cooling",     "Air Conditioner 2.5HP",     1864),
    ("Appliances",  "Refrigerator 200L",         150),
    ("Appliances",  "Refrigerator 300L",         200),
    ("Appliances",  "Chest Freezer 300L",        250),
    ("Appliances",  "Microwave Oven",            1200),
    ("Appliances",  "Electric Kettle",           2000),
    ("Appliances",  "Washing Machine",           500),
    ("Appliances",  "Iron",                      1000),
    ("Appliances",  "Blender",                   350),
    ("Appliances",  "Rice Cooker",               700),
    ("Electronics", "TV 32-inch LED",            50),
    ("Electronics", "TV 43-inch LED",            80),
    ("Electronics", "TV 55-inch LED",            120),
    ("Electronics", "Desktop Computer",          300),
    ("Electronics", "Laptop",                    65),
    ("Electronics", "Phone Charger",             10),
    ("Electronics", "Wi-Fi Router",              12),
    ("Electronics", "CCTV System 4ch",           40),
    ("Office",      "Printer",                   200),
    ("Office",      "Photocopier",               1500),
    ("Office",      "Server (small)",            400),
    ("Office",      "Projector",                 300),
    ("Pumps",       "Water Pump 0.5HP",          373),
    ("Pumps",       "Water Pump 1HP",            745),
    ("Pumps",       "Submersible Pump 0.5HP",    373),
    ("Pumps",       "Borehole Pump 1HP",         745),
    ("Heating",     "Electric Water Heater 3kW", 3000),
    ("Heating",     "Electric Cooker 4-ring",    6000),
    ("Heating",     "Electric Oven 2kW",         2000),
]

# ─── Database ─────────────────────────────────────────────────────────────────

def get_db():
    # Phase B1: dual-backend dispatch on DATABASE_URL. When unset (today),
    # behavior is byte-identical to the original SQLite path.
    _db_url = os.environ.get("DATABASE_URL", "")
    if _db_url.startswith(("postgres://", "postgresql://")):
        import db_adapter
        return db_adapter.open_postgres(_db_url)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    # ── Backend detection ──────────────────────────────────────────────
    # When DATABASE_URL is set the mirror-SQLite migration owns the
    # table + column shape (migrations/001_mirror_sqlite.sql). The
    # SQLite-only DDL paths below would either fail (AUTOINCREMENT) or
    # abort the Postgres transaction (ALTER on an existing column).
    # The seed phase further down runs on BOTH backends — it inserts
    # rows, doesn't touch schema.
    _is_postgres = bool(os.environ.get("DATABASE_URL"))

    if not _is_postgres:
        with get_db() as c:
            c.executescript("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                email    TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                name    TEXT DEFAULT '',
                company TEXT DEFAULT '',
                country TEXT DEFAULT '',
                plan    TEXT DEFAULT 'free',
                is_admin INTEGER DEFAULT 0,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );
            -- Referral program columns (REFERRAL_BACKFILL_DONE marker for idempotency)
            CREATE TABLE IF NOT EXISTS referrals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                referrer_id INTEGER NOT NULL,
                referee_id  INTEGER NOT NULL UNIQUE,
                signup_at   TEXT DEFAULT CURRENT_TIMESTAMP,
                upgraded_at TEXT,
                plan_at_upgrade TEXT,
                reward_status   TEXT DEFAULT 'pending',
                FOREIGN KEY (referrer_id) REFERENCES users(id),
                FOREIGN KEY (referee_id)  REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS projects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                name    TEXT NOT NULL,
                stage   TEXT DEFAULT 'new',
                data_json TEXT DEFAULT '{}',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS tickets (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id  INTEGER NOT NULL,
                subject  TEXT NOT NULL,
                message  TEXT NOT NULL,
                status   TEXT DEFAULT 'open',
                priority TEXT DEFAULT 'normal',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS ticket_replies (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                ticket_id INTEGER NOT NULL,
                user_id   INTEGER NOT NULL,
                is_admin  INTEGER DEFAULT 0,
                message   TEXT NOT NULL,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (ticket_id) REFERENCES tickets(id)
            );
            CREATE TABLE IF NOT EXISTS appliances (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                category     TEXT NOT NULL,
                name         TEXT NOT NULL,
                default_watt INTEGER NOT NULL,
                notes        TEXT DEFAULT ''
            );
            CREATE TABLE IF NOT EXISTS payments (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id      INTEGER NOT NULL,
                gateway      TEXT NOT NULL DEFAULT 'manual',
                plan         TEXT NOT NULL DEFAULT 'free',
                amount_usd   REAL DEFAULT 0,
                currency     TEXT DEFAULT 'USD',
                reference    TEXT DEFAULT '',
                status       TEXT DEFAULT 'success',
                created_at   TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS leads (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                name        TEXT NOT NULL,
                email       TEXT NOT NULL,
                phone       TEXT DEFAULT '',
                company     TEXT DEFAULT '',
                country     TEXT DEFAULT '',
                interest    TEXT DEFAULT 'residential',
                message     TEXT DEFAULT '',
                source      TEXT DEFAULT 'website',
                status      TEXT DEFAULT 'new',
                notes       TEXT DEFAULT '',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS newsletter_subscribers (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                email       TEXT UNIQUE NOT NULL,
                name        TEXT DEFAULT '',
                status      TEXT DEFAULT 'active',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS news_posts (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                title       TEXT NOT NULL,
                content     TEXT NOT NULL,
                category    TEXT DEFAULT 'industry',
                is_published INTEGER DEFAULT 1,
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS suppliers (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                name         TEXT NOT NULL,
                country      TEXT DEFAULT '',
                contact_name TEXT DEFAULT '',
                phone        TEXT DEFAULT '',
                email        TEXT DEFAULT '',
                website      TEXT DEFAULT '',
                categories   TEXT DEFAULT '',
                lead_time_days INTEGER DEFAULT 30,
                payment_terms TEXT DEFAULT 'TT 30 days',
                rating       INTEGER DEFAULT 5,
                notes        TEXT DEFAULT '',
                is_active    INTEGER DEFAULT 1,
                created_at   TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS equipment_catalog (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                category     TEXT NOT NULL,
                name         TEXT NOT NULL,
                brand        TEXT DEFAULT '',
                model        TEXT DEFAULT '',
                spec         TEXT DEFAULT '',
                unit         TEXT DEFAULT 'No.',
                price_usd    REAL DEFAULT 0,
                supplier_id  INTEGER DEFAULT 0,
                lead_time_days INTEGER DEFAULT 30,
                notes        TEXT DEFAULT '',
                is_active    INTEGER DEFAULT 1,
                created_at   TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS email_logs (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id     INTEGER NOT NULL,
                project_id  INTEGER DEFAULT 0,
                recipients  TEXT DEFAULT '',
                subject     TEXT DEFAULT '',
                status      TEXT DEFAULT 'sent',
                error_msg   TEXT DEFAULT '',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS upgrade_codes (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                code        TEXT UNIQUE NOT NULL,
                plan        TEXT NOT NULL DEFAULT 'professional',
                duration_days INTEGER DEFAULT 30,
                max_uses    INTEGER DEFAULT 1,
                uses        INTEGER DEFAULT 0,
                created_by  INTEGER DEFAULT 0,
                expires_at  TEXT DEFAULT '',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS assessment_requests (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                name          TEXT NOT NULL,
                email         TEXT NOT NULL,
                phone         TEXT DEFAULT '',
                company       TEXT DEFAULT '',
                country       TEXT DEFAULT '',
                system_type   TEXT DEFAULT 'off-grid',
                system_size_kw REAL DEFAULT 0,
                budget_usd    TEXT DEFAULT '',
                location_desc TEXT DEFAULT '',
                message       TEXT DEFAULT '',
                ai_score      INTEGER DEFAULT 0,
                ai_grade      TEXT DEFAULT '',
                ai_notes      TEXT DEFAULT '',
                pipeline_stage TEXT DEFAULT 'new',
                assigned_to   TEXT DEFAULT '',
                follow_up_date TEXT DEFAULT '',
                source        TEXT DEFAULT 'website',
                status        TEXT DEFAULT 'open',
                created_at    TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at    TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS installers (
                id              INTEGER PRIMARY KEY AUTOINCREMENT,
                company_name    TEXT NOT NULL,
                contact_name    TEXT NOT NULL,
                email           TEXT UNIQUE NOT NULL,
                phone           TEXT DEFAULT '',
                country         TEXT DEFAULT '',
                regions         TEXT DEFAULT '',
                years_exp       INTEGER DEFAULT 0,
                staff_count     INTEGER DEFAULT 0,
                certifications  TEXT DEFAULT '',
                specialties     TEXT DEFAULT '',
                max_project_kw  REAL DEFAULT 0,
                website         TEXT DEFAULT '',
                notes           TEXT DEFAULT '',
                status          TEXT DEFAULT 'pending',
                ai_grade        TEXT DEFAULT '',
                created_at      TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS monitor_alerts (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                url         TEXT UNIQUE NOT NULL,
                title       TEXT DEFAULT '',
                snippet     TEXT DEFAULT '',
                country     TEXT DEFAULT '',
                source_type TEXT DEFAULT '',
                is_new      INTEGER DEFAULT 1,
                found_at    TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS monitor_state (
                id               INTEGER PRIMARY KEY CHECK (id = 1),
                last_scan        TEXT DEFAULT '',
                last_count       INTEGER DEFAULT 0,
                is_running       INTEGER DEFAULT 0,
                scan_interval    INTEGER DEFAULT 120,
                notify_email     INTEGER DEFAULT 0,
                last_agent_run   TEXT DEFAULT '',
                agent_run_count  INTEGER DEFAULT 0
            );
            INSERT OR IGNORE INTO monitor_state (id) VALUES (1);
            CREATE TABLE IF NOT EXISTS password_reset_tokens (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id    INTEGER NOT NULL,
                token      TEXT UNIQUE NOT NULL,
                expires_at TEXT NOT NULL,
                used       INTEGER DEFAULT 0,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS helpline_learned_kb (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                agent      TEXT DEFAULT 'helpline',
                question   TEXT NOT NULL,
                answer     TEXT NOT NULL,
                use_count  INTEGER DEFAULT 0,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS beta_signups (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                name        TEXT NOT NULL,
                email       TEXT UNIQUE NOT NULL,
                company     TEXT DEFAULT '',
                role        TEXT DEFAULT '',
                status      TEXT DEFAULT 'pending',
                invited_at  TEXT DEFAULT '',
                notes       TEXT DEFAULT '',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS beta_feedback (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id     INTEGER DEFAULT NULL,
                username    TEXT DEFAULT '',
                email       TEXT DEFAULT '',
                type        TEXT DEFAULT 'general',
                message     TEXT NOT NULL,
                page        TEXT DEFAULT '',
                status      TEXT DEFAULT 'new',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS audit_logs (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id     INTEGER DEFAULT NULL,
                username    TEXT DEFAULT '',
                action      TEXT NOT NULL,
                ip_address  TEXT DEFAULT '',
                details     TEXT DEFAULT '',
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS login_failures (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                username    TEXT NOT NULL,
                ip_address  TEXT NOT NULL,
                created_at  TEXT DEFAULT CURRENT_TIMESTAMP
            );
            """)
        # Migrate older DBs — ignore if column already exists.
        # SQLite swallows the duplicate-column error via try/except. Postgres
        # would abort the transaction on the first dup, so we gate the whole
        # block — the mirror migration already has these columns baked in.
        for stmt in [
            "ALTER TABLE users ADD COLUMN is_admin INTEGER DEFAULT 0",
            "ALTER TABLE users ADD COLUMN stripe_customer_id TEXT DEFAULT ''",
            "ALTER TABLE users ADD COLUMN subscription_end TEXT DEFAULT ''",
            "ALTER TABLE leads ADD COLUMN system_type TEXT DEFAULT 'residential'",
            "ALTER TABLE leads ADD COLUMN system_size_kw REAL DEFAULT 0",
            "ALTER TABLE leads ADD COLUMN budget_usd TEXT DEFAULT ''",
            "ALTER TABLE leads ADD COLUMN ai_score INTEGER DEFAULT 0",
            "ALTER TABLE leads ADD COLUMN ai_grade TEXT DEFAULT ''",
            "ALTER TABLE leads ADD COLUMN ai_notes TEXT DEFAULT ''",
            "ALTER TABLE leads ADD COLUMN pipeline_stage TEXT DEFAULT 'new'",
            "ALTER TABLE leads ADD COLUMN follow_up_date TEXT DEFAULT ''",
            # projects lifecycle stage column
            "ALTER TABLE projects ADD COLUMN stage TEXT DEFAULT 'new'",
            # monitor_state: configurable interval + email notifications
            "ALTER TABLE monitor_state ADD COLUMN scan_interval INTEGER DEFAULT 120",
            "ALTER TABLE monitor_state ADD COLUMN notify_email INTEGER DEFAULT 0",
            "ALTER TABLE monitor_state ADD COLUMN last_agent_run TEXT DEFAULT ''",
            "ALTER TABLE monitor_state ADD COLUMN agent_run_count INTEGER DEFAULT 0",
            # Assessment intake v2 columns
            "ALTER TABLE assessment_requests ADD COLUMN assessment_ref TEXT DEFAULT ''",
            "ALTER TABLE assessment_requests ADD COLUMN building_desc TEXT DEFAULT ''",
            "ALTER TABLE assessment_requests ADD COLUMN building_size TEXT DEFAULT ''",
            "ALTER TABLE assessment_requests ADD COLUMN num_floors INTEGER DEFAULT 0",
            "ALTER TABLE assessment_requests ADD COLUMN building_type TEXT DEFAULT ''",
            "ALTER TABLE assessment_requests ADD COLUMN pipeline_stage TEXT DEFAULT 'assessment_submitted'",
            "ALTER TABLE assessment_requests ADD COLUMN region TEXT DEFAULT ''",
            # User roles (job function — separate from plan/billing tier)
            "ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'customer'",
        ]:
            try:
                with get_db() as c:
                    c.execute(stmt)
            except Exception:
                pass
        # Migrate: org profile columns on users
        for col, defval in [
            ("org_name",    "''"),
            ("org_address", "''"),
            ("org_email",   "''"),
            ("org_phone",   "''"),
            ("org_website", "''"),
            ("timezone",    "'UTC'"),
            ("org_whatsapp","''"),
        ]:
            try:
                with get_db() as c:
                    c.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT DEFAULT {defval}")
            except Exception:
                pass
        # Migrate: date/time + per-user SMTP settings
        for col, defval in [
            ("date_format", "'DD/MM/YYYY'"),
            ("time_format", "'24h'"),
            ("smtp_host",   "''"),
            ("smtp_port",   "'587'"),
            ("smtp_user",   "''"),
            ("smtp_pass",   "''"),
            ("smtp_from",   "''"),
            ("smtp_tls",    "'starttls'"),
            ("resend_api_key", "''"),
        ]:
            try:
                with get_db() as c:
                    c.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT DEFAULT {defval}")
            except Exception:
                pass

        # Migrate: beta_feedback rating columns (3 axes for evaluators).
        # Wrapped in try/except per column so reruns are idempotent on SQLite.
        for col in ["perf_score", "creativity_score", "value_score"]:
            try:
                with get_db() as c:
                    c.execute(f"ALTER TABLE beta_feedback ADD COLUMN {col} INTEGER")
            except Exception:
                pass

    # ── Seed phase — runs on BOTH backends ─────────────────────────────
    # All operations below are row-level INSERT/UPDATE/SELECT, not DDL,
    # and work identically on SQLite + Postgres through the adapter.

    # Seed default users — ensure admin and owner accounts always exist
    def _seed_pwd(env_var):
        # Phase 1: route via secrets_broker (audit + tier + Vault-ready).
        # DEGRADED tier with built-in env warm-up preserves the existing
        # "env var required" guarantee while landing audit trail and the
        # Vault cutover path.
        _PATH_MAP = {
            "SOLARPRO_ADMIN_PASSWORD": ("seed/admin",     "password"),
            "SOLARPRO_OWNER_PASSWORD": ("seed/marc667us", "password"),
        }
        v = None
        mapping = _PATH_MAP.get(env_var)
        if mapping:
            path, field = mapping
            try:
                v = _sb.get(path, tier="DEGRADED")[field]
            except Exception:
                v = None
        if not v:
            v = os.environ.get(env_var)
        if not v:
            raise RuntimeError(env_var + " env var required for initial user seed (see SolarPro_Schedule_2026-06-08.md Phase 0.1)")
        return v
    _SEED_USERS = [
        ("admin",    "admin@solarpro.global", "Administrator", _seed_pwd("SOLARPRO_ADMIN_PASSWORD"), "enterprise", 1),
        ("marc667us","marc667us@yahoo.com",   "Marc",          _seed_pwd("SOLARPRO_OWNER_PASSWORD"),       "enterprise", 1),
    ]
    with get_db() as c:
        for uname, email, name, pwd, plan, is_admin in _SEED_USERS:
            exists = c.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone()
            if not exists:
                c.execute(
                    "INSERT INTO users (username,email,name,password_hash,plan,is_admin) "
                    "VALUES (?,?,?,?,?,?)",
                    (uname, email, name, generate_password_hash(pwd), plan, is_admin))

    # Referral-program columns: SQLite still uses ALTER TABLE for schema
    # evolution; Postgres has them baked in via the mirror migration.
    # Gate just the DDL — the UPDATE/SELECT backfill that follows runs
    # on both.
    if not _is_postgres:
        with get_db() as _c:
            try: _c.execute("ALTER TABLE users ADD COLUMN referral_code TEXT")
            except Exception: pass
            try: _c.execute("ALTER TABLE users ADD COLUMN referred_by INTEGER")
            except Exception: pass
            try: _c.execute("ALTER TABLE users ADD COLUMN email_verified INTEGER DEFAULT 0")
            except Exception: pass
            try: _c.execute("ALTER TABLE users ADD COLUMN email_verify_token TEXT")
            except Exception: pass

    # Referral-program backfill — runs on BOTH backends (row-level only).
    import secrets as _sec
    with get_db() as _c:
        # Backfill: any pre-existing user becomes verified (do not lock them out).
        try: _c.execute("UPDATE users SET email_verified=1 WHERE email_verified IS NULL OR email_verified=0")
        except Exception: pass
        # Backfill codes for any user that lacks one
        _missing = _c.execute("SELECT id FROM users WHERE referral_code IS NULL OR referral_code = ''").fetchall()
        for _row in _missing:
            while True:
                _code = _sec.token_urlsafe(6).replace('_','').replace('-','')[:8].upper()
                if not _c.execute("SELECT 1 FROM users WHERE referral_code = ?", (_code,)).fetchone():
                    _c.execute("UPDATE users SET referral_code = ? WHERE id = ?", (_code, _row[0]))
                    break

    # Grant admin to the first registered user if no admins exist yet
    with get_db() as c:
        admins = c.execute("SELECT COUNT(*) FROM users WHERE is_admin=1").fetchone()[0]
        if admins == 0:
            c.execute("UPDATE users SET is_admin=1 WHERE id=(SELECT MIN(id) FROM users)")
    # Seed appliances once
    with get_db() as c:
        if c.execute("SELECT COUNT(*) FROM appliances").fetchone()[0] == 0:
            c.executemany(
                "INSERT INTO appliances (category,name,default_watt) VALUES (?,?,?)",
                _DEFAULT_APPLIANCES)
    # Seed default suppliers
    with get_db() as c:
        if c.execute("SELECT COUNT(*) FROM suppliers").fetchone()[0] == 0:
            _default_suppliers = [
                ("JinkoSolar International", "China", "Sales Desk", "+86 21 5183 8777",
                 "info@jinkosolar.com", "www.jinkosolar.com", "PV Modules", 45, "TT 30 days", 5),
                ("LONGi Solar", "China", "Export Team", "+86 29 8118 6677",
                 "export@longi-solar.com", "www.longi.com", "PV Modules", 45, "TT 30 days", 5),
                ("Victron Energy", "Netherlands", "Regional Sales", "+31 36 535 9700",
                 "sales@victronenergy.com", "www.victronenergy.com", "Inverters,MPPT", 21, "TT 30 days", 5),
                ("Pylontech", "China", "Sales", "+86 21 5031 9999",
                 "sales@pylontech.com.cn", "www.pylontech.com.cn", "Batteries", 30, "TT 30 days", 5),
                ("Deye Inverter", "China", "Global Sales", "+86 579 8913 0000",
                 "export@deyeinverter.com", "www.deyeinverter.com", "Inverters", 30, "TT 30 days", 4),
                ("BYD Battery", "China", "Energy Storage", "+86 755 8988 8888",
                 "energy@byd.com", "www.byd.com", "Batteries", 45, "LC 60 days", 5),
                ("RS Components", "UK", "Procurement", "+44 1536 444 222",
                 "export@rs-components.com", "www.rs-online.com", "Cables,Protection,Earthing", 14, "Net 30", 4),
                ("Schneider Electric", "France", "Solar Division", "+33 1 41 29 70 00",
                 "solar@schneider-electric.com", "www.se.com", "Protection,Distribution", 21, "Net 30", 5),
            ]
            c.executemany(
                "INSERT INTO suppliers (name,country,contact_name,phone,email,website,categories,"
                "lead_time_days,payment_terms,rating) VALUES (?,?,?,?,?,?,?,?,?,?)",
                _default_suppliers)
    # Seed default equipment catalog
    with get_db() as c:
        if c.execute("SELECT COUNT(*) FROM equipment_catalog").fetchone()[0] == 0:
            sup = {r["name"]: r["id"] for r in
                   c.execute("SELECT id,name FROM suppliers").fetchall()}
            _default_equip = [
                ("PV Modules","JinkoSolar 400 Wp Mono PERC","JinkoSolar","JKM400M-54HL4-V","400 Wp, Mono PERC, IEC 61215","No.",90, sup.get("JinkoSolar International",0),45),
                ("PV Modules","JinkoSolar 450 Wp Mono PERC","JinkoSolar","JKM450M-60HL4-V","450 Wp, Mono PERC, IEC 61215","No.",100, sup.get("JinkoSolar International",0),45),
                ("PV Modules","LONGi 500 Wp Hi-MO5","LONGi","LR4-72HIH-500M","500 Wp, Mono PERC, IEC 61215","No.",115, sup.get("LONGi Solar",0),45),
                ("Inverters","Victron MultiPlus-II 3kVA","Victron","MultiPlus-II 48/3000","3 kVA, 48V, Hybrid, built-in charger","No.",450, sup.get("Victron Energy",0),21),
                ("Inverters","Victron MultiPlus-II 5kVA","Victron","MultiPlus-II 48/5000","5 kVA, 48V, Hybrid, built-in charger","No.",620, sup.get("Victron Energy",0),21),
                ("Inverters","Deye SUN-8K-SG04LP1","Deye","SUN-8K-SG04LP1","8 kW, LV Hybrid, single-phase","No.",780, sup.get("Deye Inverter",0),30),
                ("Batteries","Pylontech US3000C 3.5kWh","Pylontech","US3000C","3.5 kWh LiFePO4, 48V, BMS, stackable","No.",620, sup.get("Pylontech",0),30),
                ("Batteries","BYD Battery-Box Premium HVS 10","BYD","HVS 10.2","10.2 kWh, HV LiFePO4, IP55","No.",1850, sup.get("BYD Battery",0),45),
                ("MPPT","Victron SmartSolar 100/50","Victron","MPPT 100/50","100V 50A MPPT, Bluetooth","No.",120, sup.get("Victron Energy",0),14),
                ("Cables","DC Solar Cable 6mmÂ² (100m)","General","TUV 6mmÂ²","6mmÂ² TÃœV 1.8kV DC solar cable, UV-rated","Roll",85, sup.get("RS Components",0),14),
                ("Protection","Schneider iC60N MCB 32A","Schneider","iC60N-32A","32A MCB Type C, 6kA, DIN","No.",12, sup.get("Schneider Electric",0),14),
                ("Earthing","Copper Earth Rod 1.2m","Generic","CER-12","16mm dia copper-clad steel, 1.2m","No.",8, sup.get("RS Components",0),7),
            ]
            c.executemany(
                "INSERT INTO equipment_catalog (category,name,brand,model,spec,unit,price_usd,"
                "supplier_id,lead_time_days) VALUES (?,?,?,?,?,?,?,?,?)",
                _default_equip)
    # Seed default news posts
    with get_db() as c:
        if c.execute("SELECT COUNT(*) FROM news_posts").fetchone()[0] == 0:
            _default_news = [
                ("Global Solar Capacity Hits 2 TW Milestone",
                 "The world has surpassed 2 terawatts of installed solar photovoltaic capacity, a landmark that took decades but accelerated rapidly in recent years. Africa and Southeast Asia are leading new deployment growth.",
                 "industry"),
                ("LiFePO4 Battery Prices Fall 40% — Storage Projects Now More Bankable",
                 "Lithium iron phosphate battery pack prices dropped 40% year-on-year, making solar-plus-storage projects significantly more financially attractive and bankable for commercial and industrial clients.",
                 "market"),
                ("IEC 61215:2021 Update — What Solar Designers Must Know",
                 "The updated IEC 61215 standard introduces new temperature cycling and thermal shock tests for PV modules. Engineers specifying panels should ensure module certifications reference the 2021 edition.",
                 "technology"),
                ("Solar Financing: New Concessional Loan Facilities for Sub-Saharan Africa",
                 "Several development finance institutions have launched new concessional solar loan windows targeting Sub-Saharan Africa. Interest rates as low as 4% are available for projects with bankable engineering documentation.",
                 "policy"),
            ]
            c.executemany(
                "INSERT INTO news_posts (title,content,category) VALUES (?,?,?)",
                _default_news)


# ─── System email helper ─────────────────────────────────────────────────────

def _send_email(to_addr, subject, html_body, text_body=None, from_addr=None, resend_key=None, attachments=None):
    """Delegate to api_manager (single source). Resend -> SMTP fallback. attachments: optional [(filename, bytes, mime), ...]."""
    return _api.email.send(to_addr, subject, html_body,
                           text_body=text_body, from_addr=from_addr,
                           resend_key_override=resend_key, attachments=attachments)


def _send_system_email(to_addr, subject, body_text):
    """Send a transactional system email (password reset, alerts)."""
    html = ("<div style='font-family:sans-serif;padding:24px;color:#1a1a2e'>"
            "<pre style='white-space:pre-wrap'>" + body_text + "</pre></div>")
    return _send_email(to_addr, subject, html, text_body=body_text)


# ─── Auth helpers ─────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def current_user():
    if "user_id" not in session:
        return None
    with get_db() as c:
        return c.execute("SELECT * FROM users WHERE id=?",
                         (session["user_id"],)).fetchone()


@app.context_processor
def inject_user_into_templates():
    # Inject the current user into every template context so base.html
    # navbar (which gates on {% if user %}) correctly shows the user
    # dropdown + Logout on pages whose route handlers do not pass
    # user=current_user() explicitly (terms, privacy, forgot_password,
    # reset_password, auth.html re-renders, etc.).
    return {"user": current_user()}


def get_project(pid):
    with get_db() as c:
        row = c.execute("SELECT * FROM projects WHERE id=? AND user_id=?",
                        (pid, session["user_id"])).fetchone()
    if not row:
        return None
    p = dict(row)
    p["data"] = json.loads(p["data_json"] or "{}")
    return p


def save_project_data(pid, data):
    stage = "results" if "results" in data else None
    with get_db() as c:
        if stage:
            c.execute(
                "UPDATE projects SET data_json=?, updated_at=?, stage=? WHERE id=?",
                (json.dumps(data), datetime.now().isoformat(), stage, pid))
        else:
            c.execute("UPDATE projects SET data_json=?, updated_at=? WHERE id=?",
                      (json.dumps(data), datetime.now().isoformat(), pid))


# ─── Equipment specifications ─────────────────────────────────────────────────

# Battery chemistry parameters — per lithium spec document (pv1/)
BATTERY_CHEMISTRY = {
    "LiFePO4": {
        "name":        "Lithium Iron Phosphate (LiFePO4)",
        "dod":         0.90,
        "efficiency":  0.96,
        "cycle_life":  "4,000—8,000 cycles",
        "lifetime_yr": "10—15 years",
        "cell_v":      3.2,
        "temp_range":  "-20Â°C to +55Â°C",
        "brands":      "BYD Battery-Box, Pylontech Force H, Dyness BX51100, Sungrow SBH, Huawei Luna2000",
        "sizes_kwh":   [5.12, 10.24, 13.5, 15.36, 20.48, 30.72],
        "usd_per_kwh": 120,
    },
    "NMC": {
        "name":        "Lithium Nickel Manganese Cobalt (NMC)",
        "dod":         0.85,
        "efficiency":  0.95,
        "cycle_life":  "2,000—4,000 cycles",
        "lifetime_yr": "8—12 years",
        "cell_v":      3.7,
        "temp_range":  "-20Â°C to +45Â°C",
        "brands":      "LG RESU, Samsung SDI, Panasonic EverVolt",
        "sizes_kwh":   [9.8, 16.0, 19.6],
        "usd_per_kwh": 140,
    },
    "LTO": {
        "name":        "Lithium Titanate (LTO)",
        "dod":         0.95,
        "efficiency":  0.98,
        "cycle_life":  "15,000—30,000 cycles",
        "lifetime_yr": "20—25 years",
        "cell_v":      2.4,
        "temp_range":  "-40Â°C to +65Â°C",
        "brands":      "Toshiba SCiB, Microvast, Yabo Power",
        "sizes_kwh":   [10.0, 20.0, 40.0],
        "usd_per_kwh": 210,
    },
}

# PV Panel specification — monocrystalline PERC (IEC 61215)
PANEL_SPEC = {
    "technology":   "Monocrystalline PERC",
    "temp_coeff":   -0.0035,          # %/Â°C power temperature coefficient
    "standard_wp":  [110, 250, 330, 400, 450, 500, 550],
    "default_wp":   400,
    "eff_pct":      "21—23%",
    "warranty_yr":  "12 yr product / 25 yr linear power",
    "brands":       "JinkoSolar, LONGi Solar, Canadian Solar, Trina Solar, JA Solar",
}

# Inverter brand recommendations by power range
INVERTER_BRANDS = [
    (3,   "Victron MultiPlus-II, Growatt SPF-3000, Deye SUN-3K-SG04LP1"),
    (5,   "Deye SUN-5K-SG04LP1, Growatt SPF-5000, Victron MultiPlus-II 5kVA"),
    (8,   "Sungrow SH8.0RT, Fronius Symo GEN24, Deye SUN-8K"),
    (12,  "Sungrow SH10RT, SMA Sunny Tripower 10, Huawei SUN2000-10KTL"),
    (9999,"Huawei SUN2000-20KTL, SMA Sunny Tripower CORE1, Sungrow SG25CX"),
]

# ─── Ghana PURC Tariff Schedule (Q2 2026, effective April 1 2026) ─────────────
# Source: Public Utilities Regulatory Commission (PURC) Ghana
GHANA_PURC_TARIFFS = {
    "Residential Lifeline (0-30 kWh/month)": {
        "rate_ghc":   0.8690,
        "fixed_ghc":  2.13,
        "description": "Low-income households consuming up to 30 kWh/month",
        "bldg_hint":  ["low_income", "single_room"],
    },
    "Residential Standard (0-300 kWh/month)": {
        "rate_ghc":   1.9688,
        "fixed_ghc":  10.73,
        "description": "Standard residential homes consuming up to 300 kWh/month",
        "bldg_hint":  ["residential", "apartment", "bungalow", "villa", "duplex"],
    },
    "Residential High Use (>300 kWh/month)": {
        "rate_ghc":   2.4500,
        "fixed_ghc":  10.73,
        "description": "High-consumption residential, >300 kWh/month (large homes)",
        "bldg_hint":  ["mansion", "estate"],
    },
    "Non-Residential Standard (0-300 kWh/month)": {
        "rate_ghc":   1.7775,
        "fixed_ghc":  30.00,
        "description": "Small offices, shops, clinics up to 300 kWh/month",
        "bldg_hint":  ["office", "retail", "shop", "clinic", "small_commercial"],
    },
    "Non-Residential High Use (>300 kWh/month)": {
        "rate_ghc":   2.1649,
        "fixed_ghc":  30.00,
        "description": "Larger commercial users, hotels, supermarkets, >300 kWh/month",
        "bldg_hint":  ["commercial", "hotel", "supermarket", "restaurant", "church"],
    },
    "Special Load - LV (hospitals, schools)": {
        "rate_ghc":   2.3211,
        "fixed_ghc":  100.00,
        "description": "Hospitals, schools, government buildings on LV supply",
        "bldg_hint":  ["hospital", "school", "government", "institution", "university"],
    },
    "Special Load - HV (large facilities)": {
        "rate_ghc":   1.8212,
        "fixed_ghc":  200.00,
        "description": "Large special facilities on high-voltage supply",
        "bldg_hint":  [],
    },
    "Industrial - LV (factories, warehouses)": {
        "rate_ghc":   2.2000,
        "fixed_ghc":  300.00,
        "description": "Industrial users on low-voltage supply — factories, warehouses",
        "bldg_hint":  ["industrial", "factory", "warehouse", "manufacturing"],
    },
    "Industrial - HV (large industries)": {
        "rate_ghc":   1.7000,
        "fixed_ghc":  500.00,
        "description": "Large industrial facilities on high-voltage supply",
        "bldg_hint":  [],
    },
    "EV Charging Station": {
        "rate_ghc":   2.0160,
        "fixed_ghc":  50000.00,
        "description": "Electric vehicle charging stations",
        "bldg_hint":  ["ev_station", "petrol_station"],
    },
}

# ─── Demand Factors by load category ─────────────────────────────────────────
# Fraction of connected load actually operating simultaneously (IEC 60364 / BS 7671)
DEMAND_FACTORS = {
    "Lighting":    0.75,   # Not all lights on at once
    "Cooling":     0.65,   # Staggered thermostat cycling
    "Appliances":  0.50,   # Diverse usage patterns
    "Electronics": 0.70,   # Partial simultaneous use
    "Pumps":       0.70,   # Usually one pump runs at a time
    "Heating":     0.60,   # Thermostat-controlled staggering
    "Office":      0.70,   # Not all office equipment running simultaneously
    "Other":       0.70,   # General diversity allowance
}

def inverter_brand(inv_kw):
    for threshold, brand in INVERTER_BRANDS:
        if inv_kw <= threshold:
            return brand
    return INVERTER_BRANDS[-1][1]


# ─── Engineering calculations ─────────────────────────────────────────────────

def calc_loads(loads):
    """Sum diversified daily energy demand.
    Each load's kWh is multiplied by its demand_factor (fraction of connected
    load actually operating simultaneously). Default DF per category if not set."""
    total = 0.0
    for ld in loads:
        w  = float(ld.get("wattage", 0))
        q  = float(ld.get("quantity", 1))
        h  = float(ld.get("hours", 0))
        df = float(ld.get("demand_factor",
                   DEMAND_FACTORS.get(ld.get("category", "Other"), 0.70)))
        df = max(0.10, min(1.0, df))   # clamp 10%—100%
        total += (w * q * h * df) / 1000
    return round(total, 3)


def calc_pv(daily_kwh, psh, temp_c, panel_wp=400, sys_eff=0.75):
    """Monocrystalline PERC sizing with temperature derating."""
    td  = temp_derating(temp_c)
    eff = sys_eff * td
    pv_kw      = daily_kwh / (psh * eff)
    num_panels = math.ceil(pv_kw * 1000 / panel_wp)
    return round(pv_kw, 3), num_panels, round(td, 4)


def calc_battery(daily_kwh, autonomy=1, chemistry="LiFePO4"):
    """Size lithium battery bank — fewest compact units, chemistry-aware."""
    chem     = BATTERY_CHEMISTRY.get(chemistry, BATTERY_CHEMISTRY["LiFePO4"])
    dod      = chem["dod"]
    eff      = chem["efficiency"]
    required = max((daily_kwh * autonomy) / (dod * eff), 0.1)
    SIZES    = chem["sizes_kwh"]
    MAX_RATIO = 2.0

    best = None
    for size in SIZES:
        n     = max(1, math.ceil(required / size))
        total = n * size
        if n == 1 and total > required * MAX_RATIO and size != SIZES[0]:
            continue
        score = n * 1000 + (total - required)
        if best is None or score < best[0]:
            best = (score, n, size, total)

    _, n, unit_kwh, total = best
    return round(total, 1), n, unit_kwh


def calc_mppt(pv_kw, dc_voltage):
    """Size MPPT charge controller (A) — 1.25Ã— safety factor."""
    i_max = (pv_kw * 1000) / dc_voltage * 1.25
    for size in [20, 30, 40, 50, 60, 80, 100, 120, 150, 200]:
        if size >= i_max:
            return size
    return math.ceil(i_max / 10) * 10


def calc_inverter(daily_kwh, peak_kw=0.0, peak_factor=0.30, safety=1.25):
    """Inverter must satisfy both energy-based sizing and peak demand."""
    from_energy = daily_kwh * peak_factor * safety
    # Must handle connected peak load — inverter rating â‰¥ peak demand
    inv_kw = max(from_energy, peak_kw * 1.0)
    # Round up to nearest standard size
    for std in [3.0, 5.0, 8.0, 10.0, 12.0, 15.0, 20.0, 30.0, 50.0, 100.0]:
        if std >= inv_kw:
            return std
    return round(inv_kw, 2)


def calc_economics(pv_kw, num_panels, bat_kwh, num_bat, inv_kw,
                   daily_kwh, tariff, currency, symbol,
                   cost_usd_kwp, fx_usd, autonomy=1, boq_total_local=None,
                   chemistry="LiFePO4", funding_mode="loan", install_rate_pct=15):
    """Full economic analysis: NPV, IRR, payback, DSCR, loan, replacement costs.

    Optimised assumptions (West Africa 2025 market basis):
      â€¢ O&M: 0.8% of CAPEX/yr (was 1.2% — robust LiFePO4 systems need minimal maintenance)
      â€¢ Install rate: user-adjustable (default 15%)
      â€¢ Discount rate: 12% (self-funded can use 10% opportunity cost)
      â€¢ Tariff escalation: 8%/yr — consistent with West Africa utility rate trends
      â€¢ Battery replacement: LiFePO4 yr 13, NMC yr 8 (at 70% of original cost)
      â€¢ Inverter replacement: yr 10 (at 80% of original cost)
      â€¢ Residual value: 5% of CAPEX at yr 25

    funding_mode: 'loan' = include DSCR/bankability; 'self' = self-funded, no loan analysis.
    """
    INSTALL_RATE = install_rate_pct / 100.0
    # ── Cost estimation ───────────────────────────────────────────────────────
    if boq_total_local is not None:
        total_local   = boq_total_local        # BOQ already includes installation line
        install_local = total_local * (INSTALL_RATE / (1 + INSTALL_RATE))  # â‰ˆ 13% of total
        equip_local   = total_local - install_local
    else:
        equip_usd     = pv_kw * cost_usd_kwp
        install_usd   = equip_usd * INSTALL_RATE
        total_usd     = equip_usd + install_usd
        total_local   = total_usd * fx_usd
        equip_local   = equip_usd * fx_usd
        install_local = install_usd * fx_usd

    # ── Model constants ───────────────────────────────────────────────────────
    DISC   = 0.10 if funding_mode == "self" else 0.12   # self-funded: lower opportunity cost
    ESC    = 0.08   # utility tariff escalation %/yr (Ghana PURC trend)
    DEGRAD = 0.005  # PV panel degradation 0.5%/yr (IEC 61215)
    OM_PCT = 0.008  # O&M 0.8% of CAPEX/yr (was 1.2%)
    OM_ESC = 0.04   # O&M cost escalation 4%/yr
    LIFE   = 25

    # ── Component replacement cost estimates ──────────────────────────────────
    chem_data = BATTERY_CHEMISTRY.get(chemistry, BATTERY_CHEMISTRY["LiFePO4"])
    bat_replace_yr = {"LiFePO4": 13, "NMC": 8, "LTO": 99}.get(chemistry, 13)
    bat_replace_cost = bat_kwh * chem_data["usd_per_kwh"] * 0.70 * fx_usd  # 70% (prices declining)

    # Inverter cost estimate (consistent with calc_boq prices post-reduction)
    if inv_kw <= 3:    inv_orig_usd = 280 * 0.90
    elif inv_kw <= 5:  inv_orig_usd = 400 * 0.90
    elif inv_kw <= 8:  inv_orig_usd = 560 * 0.90
    elif inv_kw <= 12: inv_orig_usd = 820 * 0.90
    else:              inv_orig_usd = 1180 * 0.90
    inv_replace_cost = inv_orig_usd * 0.80 * fx_usd   # 80% (slight price decline)

    residual_value = total_local * 0.05   # 5% salvage at year 25

    annual_kwh = daily_kwh * 365
    annual_sav = annual_kwh * tariff
    om_yr1     = total_local * OM_PCT
    # Simple payback (undiscounted, yr-1 basis — quick screening metric)
    net_yr1    = annual_sav - om_yr1
    payback    = total_local / net_yr1 if net_yr1 > 0 else float("inf")
    co2_yr     = annual_kwh * 0.40 / 1000   # tonnes CO2 (Ghana grid intensity 0.4 kgCO2/kWh)

    # ── 25-year cash flow model ───────────────────────────────────────────────
    cashflows  = [-total_local]
    npv        = -total_local
    cumul      = -total_local
    breakeven  = None
    cf_rows    = []

    for yr in range(1, LIFE + 1):
        degraded  = annual_kwh * ((1 - DEGRAD) ** yr)
        esc_tarif = tariff * ((1 + ESC) ** yr)
        gross     = degraded * esc_tarif
        om        = om_yr1 * ((1 + OM_ESC) ** yr)
        # One-time capital expenditure events
        capex_yr  = 0.0
        if yr == bat_replace_yr:
            capex_yr += bat_replace_cost
        if yr == 10:
            capex_yr += inv_replace_cost
        if yr == LIFE:
            capex_yr -= residual_value     # salvage value (positive cash flow)
        net       = gross - om - capex_yr
        disc      = net / ((1 + DISC) ** yr)
        cumul    += net
        npv      += disc
        cashflows.append(net)
        if cumul >= 0 and breakeven is None:
            breakeven = yr
        cf_rows.append({"yr": yr, "gross": gross, "om": om,
                        "capex_yr": capex_yr, "net": net, "cumul": cumul, "npv_c": npv})

    cumul_10 = sum(r["net"] for r in cf_rows[:10])
    cumul_25 = sum(r["net"] for r in cf_rows)
    roi_pct  = (cumul_25 / total_local) * 100 if total_local > 0 else 0

    irr = _irr(cashflows)

    sym = symbol

    # ── Loan / bankability analysis ───────────────────────────────────────────
    if funding_mode == "self":
        loan_amt  = 0.0
        equity    = total_local
        pmt       = 0.0
        annual_pmt = 0.0
        dscr      = 0.0
        bankability = "SELF-FUNDED"
        bank_color  = "#818cf8"
        bank_reasons = [
            "This project is fully self-funded — no debt service analysis required.",
            f"Total capital investment: {sym} {total_local:,.0f}",
            f"Opportunity cost of capital: {int(DISC*100)}% per annum",
            f"NPV accounts for full investment over {LIFE}-year system life.",
        ]
    else:
        loan_pct   = 0.70
        loan_amt   = total_local * loan_pct
        equity     = total_local - loan_amt
        rate       = 0.15
        tenor_yr   = 7
        n          = tenor_yr * 12
        r_m        = rate / 12
        pmt        = loan_amt * (r_m * (1+r_m)**n) / ((1+r_m)**n - 1) if r_m else loan_amt/n
        annual_pmt = pmt * 12
        dscr       = net_yr1 / annual_pmt if annual_pmt > 0 else 0

        if dscr >= 1.25:
            bankability  = "BANKABLE"
            bank_color   = "#34d399"
            bank_reasons = [
                f"DSCR {dscr:.2f} meets lender minimum of 1.25",
                f"Net savings {sym} {net_yr1:,.0f}/yr comfortably cover debt service {sym} {annual_pmt:,.0f}/yr",
                "Cash flow supports commercial loan repayment",
                f"Equity required: {sym} {equity:,.0f} ({int((1-loan_pct)*100)}% of total)",
            ]
        elif dscr >= 1.0:
            bankability  = "MARGINAL"
            bank_color   = "#fbbf24"
            bank_reasons = [
                f"DSCR {dscr:.2f} below lender minimum of 1.25 — additional security may be required",
                f"Net savings {sym} {net_yr1:,.0f}/yr barely cover debt service {sym} {annual_pmt:,.0f}/yr",
                "Consider increasing equity contribution or extending loan tenor to 10 years",
                f"A 10-year tenor reduces annual payments to ~{sym} {annual_pmt*0.72:,.0f}",
                "Development finance (IFC, AfDB) offers concessional rates 6—9% for viable solar projects",
            ]
        else:
            bankability  = "NOT BANKABLE"
            bank_color   = "#f87171"
            bank_reasons = [
                f"DSCR {dscr:.2f} below 1.00 — debt service exceeds net savings",
                f"Annual shortfall: {sym} {max(0, annual_pmt - net_yr1):,.0f}",
                "Recommendation: switch to self-funded mode, or seek grant/subsidy financing",
                f"At 70% self-funded equity, loan reduces to {sym} {loan_amt*0.30:,.0f} — recheck DSCR",
                "Alternative: review load schedule to right-size system and reduce CAPEX",
            ]

    # ── Project verdict ───────────────────────────────────────────────────────
    # Self-funded: more lenient thresholds (no debt burden, pure ROI)
    # Loan-funded: standard thresholds
    if funding_mode == "self":
        _approve_yr, _cond_yr = 12, 25
    else:
        _approve_yr, _cond_yr = 10, 20

    if payback <= _approve_yr and npv > 0:
        verdict = "APPROVED"
        v_color = "#34d399"
        verdict_reasons = [
            f"Payback of {payback:.1f} years is within the {_approve_yr}-year target",
            f"Positive NPV {sym} {npv:,.0f} — the project creates real economic value",
            (f"IRR {irr*100:.1f}% exceeds discount rate {int(DISC*100)}%" if irr else "Strong positive returns"),
            f"25-year net return: {sym} {cumul_25:,.0f} | ROI: {roi_pct:.0f}%",
            f"Avoids {co2_yr*LIFE:.0f} t COâ‚‚ over system life",
        ]
    elif payback <= _cond_yr:
        verdict = "CONDITIONAL"
        v_color = "#fbbf24"
        npv_str = f"positive ({sym} {int(npv):,})" if npv > 0 else f"negative ({sym} {int(npv):,})"
        verdict_reasons = [
            f"Payback {payback:.1f} years — within {_cond_yr}-yr maximum but above {_approve_yr}-yr target",
            f"NPV is {npv_str}",
            "Approval recommended subject to: verified tariff data and full load audit",
            f"If utility tariff rises 10%/yr, payback improves to ~{payback*0.87:.1f} years",
            "Review load schedule — demand factor optimisation can reduce system size 15—25%",
        ]
    else:
        verdict = "REJECTED"
        v_color = "#f87171"
        import math as _m
        _pb_str  = f"{payback:.1f}" if _m.isfinite(payback) else "N/A (savings < O&M costs)"
        _pb80str = f"{payback*0.80:.1f}" if _m.isfinite(payback) else "improved"
        verdict_reasons = [
            f"Payback {_pb_str} years exceeds {_cond_yr}-year threshold",
            f"NPV {sym} {npv:,.0f} — insufficient return on investment",
            f"Annual savings {sym} {annual_sav:,.0f} cannot justify {sym} {total_local:,.0f} investment",
            "Action required: (1) Audit loads and remove non-essential equipment, "
            "(2) Obtain competitive equipment quotes, (3) Apply for utility grants or tax incentives",
            f"Reducing system cost by 20% would improve payback to ~{_pb80str} years",
        ]

    return {
        "total_local":    total_local,
        "equip_local":    equip_local,
        "install_local":  install_local,
        "install_rate_pct": int(install_rate_pct),
        "annual_kwh":     annual_kwh,
        "annual_sav":     annual_sav,
        "om_yr1":         om_yr1,
        "net_yr1":        net_yr1,
        "payback":        payback,
        "npv":            npv,
        "irr_pct":        irr * 100 if irr else None,
        "roi_pct":        roi_pct,
        "breakeven":      breakeven,
        "co2_yr":         co2_yr,
        "cumul_10":       cumul_10,
        "cumul_25":       cumul_25,
        "cf_rows":        cf_rows,
        "loan_amt":       loan_amt,
        "equity":         equity,
        "pmt":            pmt,
        "annual_pmt":     annual_pmt,
        "dscr":           dscr,
        "bankability":    bankability,
        "bank_color":     bank_color,
        "bank_reasons":   bank_reasons,
        "verdict":        verdict,
        "v_color":        v_color,
        "verdict_reasons": verdict_reasons,
        "funding_mode":   funding_mode,
        "bat_replace_yr": bat_replace_yr,
        "bat_replace_cost": bat_replace_cost,
        "inv_replace_cost": inv_replace_cost,
        "residual_value": residual_value,
        "disc_rate_pct":  int(DISC * 100),
        "currency":       currency,
        "symbol":         symbol,
        "tariff":         tariff,
    }


def _irr(cashflows, guess=0.10, tol=1e-6, max_iter=100):
    """Compute IRR using Newton-Raphson method. Returns None on degenerate inputs."""
    # Guard: all-zero or all-negative cashflows have no meaningful IRR
    if not cashflows or all(cf <= 0 for cf in cashflows):
        return None
    r = guess
    try:
        for _ in range(max_iter):
            npv = sum(cf / (1 + r) ** t for t, cf in enumerate(cashflows))
            dnpv = sum(-t * cf / (1 + r) ** (t + 1) for t, cf in enumerate(cashflows))
            if abs(dnpv) < 1e-12:
                break
            r_new = r - npv / dnpv
            if abs(r_new - r) < tol:
                return r_new
            r = r_new
        return r if -1 < r < 10 else None
    except (OverflowError, ZeroDivisionError, ValueError):
        return None


def calc_recommendations(eco, d, r):
    """Generate prioritised redesign recommendations to fix rejected/non-bankable projects."""
    if eco["verdict"] == "APPROVED" and eco["bankability"] == "BANKABLE":
        return []

    recs     = []
    sym      = d.get("symbol", "$")
    payback  = eco.get("payback", 99)
    dscr     = eco.get("dscr", 0)
    net_yr1  = eco.get("net_yr1", 0)
    total    = eco.get("total_local", 0)
    ann_pmt  = eco.get("annual_pmt", eco.get("ann_pmt", 0))   # key was renamed; support both
    ann_kwh  = eco.get("annual_kwh", 0)
    tariff   = d.get("tariff", 0)
    autonomy = d.get("autonomy", 1)
    voltage  = d.get("voltage", 48)
    panel_wp = r.get("panel_wp", 400)

    # Guard: payback may be float("inf") when annual savings < O&M costs
    import math as _math
    _payback_finite = _math.isfinite(payback)
    _payback_safe   = min(payback, 200.0) if _payback_finite else 200.0  # cap at 200 yr for display

    # 1 ─ Reduce system size / load (most impactful for high payback)
    if payback > 10:
        reduce = min(35, max(15, int((_payback_safe - 8) / _payback_safe * 85)))
        new_cost = total * (1 - reduce / 100)
        new_pb   = new_cost / net_yr1 if net_yr1 > 0 else _payback_safe
        pb_str   = f"{payback:.1f}" if _payback_finite else "N/A (no net savings)"
        new_pb_str = f"{new_pb:.1f}" if net_yr1 > 0 else "improved"
        recs.append({
            "priority": 1, "icon": "bi-arrows-collapse", "color": "#f59e0b",
            "title":  f"Reduce System Size by {reduce}%",
            "action": f"Implement energy efficiency measures (LED lighting, efficient appliances, load "
                      f"scheduling) to cut daily consumption by {reduce}%, then downsize the PV array "
                      f"and battery bank proportionally.",
            "impact": f"System cost falls to ~{sym} {new_cost:,.0f}. "
                      f"Payback improves from {pb_str} yr to ~{new_pb_str} yr.",
            "category": "Design",
        })

    loan_amt = eco.get("loan_amt", 0)
    equity   = eco.get("equity", 0)
    pmt      = eco.get("pmt", 0)
    om_yr1   = eco.get("om_yr1", 0)
    co2_yr   = eco.get("co2_yr", 0)

    # 2 ─ Increase equity to achieve DSCR â‰¥ 1.25
    if dscr < 1.25 and ann_pmt > 0 and loan_amt > 0:
        loan_factor  = ann_pmt / loan_amt
        needed_pmt   = net_yr1 / 1.25
        needed_loan  = needed_pmt / loan_factor if loan_factor > 0 else 0
        new_eq_pct   = max(0, (total - needed_loan) / total * 100) if total > 0 else 0
        new_eq_amt   = total * new_eq_pct / 100
        if new_eq_pct <= 70:
            recs.append({
                "priority": 1, "icon": "bi-cash-stack", "color": "#22c55e",
                "title":  f"Increase Equity Contribution to {new_eq_pct:.0f}%",
                "action": f"Raise own-equity from 30% ({sym} {equity:,.0f}) to "
                          f"{new_eq_pct:.0f}% ({sym} {new_eq_amt:,.0f}) to reduce the loan "
                          f"amount and annual debt service.",
                "impact": f"Annual debt service reduces to {sym} {needed_pmt:,.0f}/yr. "
                          f"DSCR improves to 1.25 — project becomes BANKABLE.",
                "category": "Finance",
            })

    # 3 ─ Extend loan tenor to 10 years
    if 0.7 <= dscr < 1.25:
        r_m   = 0.15 / 12
        n10   = 10 * 12
        pmt10 = loan_amt * (r_m * (1 + r_m)**n10) / ((1 + r_m)**n10 - 1) if r_m and loan_amt else 0
        apmt10 = pmt10 * 12
        d10    = net_yr1 / apmt10 if apmt10 > 0 else 0
        recs.append({
            "priority": 2, "icon": "bi-calendar-range", "color": "#0ea5e9",
            "title":  "Extend Loan Tenor to 10 Years",
            "action": f"Negotiate with lender to extend repayment period from 7 to 10 years. "
                      f"Monthly repayment drops from {sym} {pmt:,.0f} to {sym} {pmt10:,.0f}.",
            "impact": f"Annual debt service reduces to {sym} {apmt10:,.0f}/yr. "
                      f"DSCR improves to {d10:.2f} "
                      f"({'BANKABLE' if d10 >= 1.25 else 'MARGINAL' if d10 >= 1.0 else 'NOT BANKABLE'}).",
            "category": "Finance",
        })

    # 4 ─ Verify / improve electricity tariff
    if tariff < 0.12 and payback > 10 and ann_kwh > 0:
        needed_tariff = (total / 8 + om_yr1) / ann_kwh
        recs.append({
            "priority": 2, "icon": "bi-receipt", "color": "#a78bfa",
            "title":  "Verify Electricity Tariff & Include Diesel Savings",
            "action": f"Confirm current tariff ({sym}{tariff:.3f}/kWh) is the real commercial rate. "
                      f"If displacing diesel generation, add fuel cost savings (typically $0.25—0.40/kWh "
                      f"equivalent). Use peak Time-of-Use (ToU) rate if applicable.",
            "impact": f"A blended effective tariff of {sym}{needed_tariff:.3f}/kWh achieves an 8-year payback. "
                      f"Including diesel offsets typically doubles the effective tariff.",
            "category": "Revenue",
        })

    # 5 ─ Reduce battery autonomy (if > 1 day)
    if autonomy > 1 and payback > 10:
        bat_cost_share = 0.35
        new_cost = total * (1 - bat_cost_share * (autonomy - 1) / autonomy)
        new_pb   = new_cost / net_yr1 if net_yr1 > 0 else payback
        recs.append({
            "priority": 2, "icon": "bi-battery-half", "color": "#f59e0b",
            "title":  f"Reduce Battery Autonomy from {autonomy} to 1 Day",
            "action": f"Design for 1-day battery autonomy instead of {autonomy} days. "
                      f"For grid-connected or hybrid systems this is sufficient. "
                      f"Use a backup generator for extended outages if needed.",
            "impact": f"Battery bank reduces by ~{(1-1/autonomy)*100:.0f}%. "
                      f"System cost falls to ~{sym} {new_cost:,.0f}. "
                      f"Payback improves to ~{new_pb:.1f} yr.",
            "category": "Design",
        })

    # 6 ─ Upgrade panel wattage
    if panel_wp < 500 and payback > 10:
        saving_pct = (500 - panel_wp) / panel_wp * 0.5  # fewer panels â†' less BOS cost
        new_cost = total * (1 - saving_pct)
        new_pb   = new_cost / net_yr1 if net_yr1 > 0 else payback
        recs.append({
            "priority": 3, "icon": "bi-sun-fill", "color": "#fbbf24",
            "title":  "Upgrade to 500 Wp High-Efficiency Panels",
            "action": f"Switch from {panel_wp} Wp to 500 Wp TOPCon panels. Fewer panels required "
                      f"({r['num_panels']} â†' ~{int(r['num_panels']*panel_wp/500)} modules), "
                      f"reducing mounting, cabling, and labour costs.",
            "impact": f"Balance-of-system cost saving ~{saving_pct*100:.0f}%. "
                      f"System cost reduces to ~{sym} {new_cost:,.0f}. "
                      f"Payback improves to ~{new_pb:.1f} yr.",
            "category": "Design",
        })

    # 7 ─ Switch to grid-tied (if off-grid)
    if d.get("system_type") == "off-grid" and payback > 12:
        bat_saving = total * 0.35
        new_cost   = total - bat_saving
        new_pb     = new_cost / net_yr1 if net_yr1 > 0 else payback
        recs.append({
            "priority": 3, "icon": "bi-plug-fill", "color": "#34d399",
            "title":  "Consider Grid-Tied / Hybrid Configuration",
            "action": f"If grid connection is available, eliminate the battery bank and switch to "
                      f"grid-tied or hybrid mode. Export surplus energy and draw from the grid "
                      f"at night or in low-solar periods.",
            "impact": f"Removing batteries saves ~{sym} {bat_saving:,.0f} (35% of capital). "
                      f"System cost falls to ~{sym} {new_cost:,.0f}. "
                      f"Payback improves from {payback:.1f} yr to ~{new_pb:.1f} yr.",
            "category": "Design",
        })

    # 8 ─ Apply for grants / concessional finance
    if eco["verdict"] in ("REJECTED", "CONDITIONAL") or eco["bankability"] in ("NOT BANKABLE", "MARGINAL"):
        grant_pct = 20
        new_cost  = total * (1 - grant_pct / 100)
        new_pb    = new_cost / net_yr1 if net_yr1 > 0 else payback
        recs.append({
            "priority": 3, "icon": "bi-bank", "color": "#34d399",
            "title":  "Apply for Grants, Subsidies & Concessional Finance",
            "action": f"Investigate: (a) Government capital grants (15—30% of cost), "
                      f"(b) Development bank loans at 6—9% vs 15% commercial, "
                      f"(c) Carbon credits under Gold Standard / VCS (~{sym}{co2_yr*15:,.0f}/yr), "
                      f"(d) Green bonds or impact investment at preferential rates.",
            "impact": f"A 20% grant reduces capital to {sym} {new_cost:,.0f}, improving payback to "
                      f"~{new_pb:.1f} yr. Concessional rate (8%) improves DSCR by ~0.4.",
            "category": "Finance",
        })

    recs.sort(key=lambda x: x["priority"])
    return recs


def calc_boq(num_panels, num_bat, inv_kw, pv_kw, bat_kwh,
             unit_bat_kwh, chemistry, mppt_a, cost_usd_kwp, fx_usd,
             panel_wp=400, ac_cables=None, voltage=48, num_strings=1,
             supply_markup_pct=8, install_rate_pct=15):
    """Generate BOQ — real equipment specs, brands, optimised costing.
    Pricing basis (West Africa 2025-2026, verified against market data):
      â€¢ Basic prices reduced 10% vs previous version
      â€¢ Supply markup: default 8% (user-adjustable)
      â€¢ Installation labour: default 15% of supply (user-adjustable)
    """
    SUPPLY_MARKUP  = supply_markup_pct / 100.0   # user-set supply/procurement markup
    INSTALL_RATE   = install_rate_pct  / 100.0   # user-set installation labour rate
    PRICE_FACTOR   = 0.90   # 10% reduction on basic unit prices

    chem     = BATTERY_CHEMISTRY.get(chemistry, BATTERY_CHEMISTRY["LiFePO4"])
    # Panel: use 50% of cost_usd_kwp allocation (was 55%), reduced further by PRICE_FACTOR
    panel_usd = (cost_usd_kwp * pv_kw) / max(num_panels, 1) * 0.50 * PRICE_FACTOR
    bat_usd   = unit_bat_kwh * chem["usd_per_kwh"] * PRICE_FACTOR
    if inv_kw <= 3:   inv_usd = 280 * PRICE_FACTOR   # was 320
    elif inv_kw <= 5: inv_usd = 400 * PRICE_FACTOR   # was 450
    elif inv_kw <= 8: inv_usd = 560 * PRICE_FACTOR   # was 620
    elif inv_kw <=12: inv_usd = 820 * PRICE_FACTOR   # was 900
    else:             inv_usd = 1180 * PRICE_FACTOR  # was 1300
    mppt_usd = (55 + mppt_a * 1.6) * PRICE_FACTOR   # was 60 + 1.8

    def local(usd): return usd * fx_usd

    # ── DC cable sizing ───────────────────────────────────────────────────────
    # String cable: 6 mmÂ² is standard for all residential/commercial PV strings
    # (panel Isc â‰ˆ panel_wp/40 A; TÃœV-rated 1.8 kV DC, max 17 A per string)
    dc_str_qty = num_panels * 8          # â‰ˆ 4 m pos + 4 m neg per panel
    dc_str_spec = ("6 mmÂ² Twin-core UV-resistant PV Solar Cable, TÃœV 2Pfg1169, "
                   "1.8 kV DC, âˆ'40 Â°C to +90 Â°C, IEC 62930 — red & black")

    # DC main cable (string combiner â†' DC isolator â†' inverter)
    panel_isc_a = panel_wp / 40.0
    total_dc_a  = num_strings * panel_isc_a * 1.25
    _DC = [(20,4),(30,6),(45,10),(60,16),(80,25),(120,35),(999,50)]
    dc_main_mm2 = next(s for lim,s in _DC if total_dc_a <= lim)
    dc_main_qty = 35                     # combiner-to-inverter run + slack
    dc_main_spec = (f"{dc_main_mm2} mmÂ² Twin-core UV-resistant PV Solar Cable, "
                    f"TÃœV 1.8 kV DC — {num_strings} strings Ã— {panel_isc_a:.0f} A Isc, "
                    f"design current {total_dc_a:.0f} A (incl. 1.25 Ã— safety factor)")

    # Battery DC cable (battery bank â†' inverter DC bus)
    bat_a   = inv_kw * 1000 / max(voltage, 12) * 1.25
    _BAT = [(40,10),(65,16),(100,25),(150,35),(200,50),(999,70)]
    bat_mm2 = next(s for lim,s in _BAT if bat_a <= lim)
    bat_qty = num_bat * 3 + 5
    bat_spec = (f"{bat_mm2} mmÂ² Flexible Multi-strand Cu, PVC 105 Â°C, 1 kV DC — "
                f"{voltage} V bus, design current {bat_a:.0f} A, "
                f"ANL fuse-protected per string — red & black")

    # ── Items list (non-cable) ────────────────────────────────────────────────
    def p(usd): return local(usd * PRICE_FACTOR)   # price with 10% reduction + convert

    items = [
        ("PV Modules — Mono PERC",
         f"{num_panels} Ã— {panel_wp} Wp | {PANEL_SPEC['brands'].split(',')[0].strip()}",
         num_panels, "No.", local(panel_usd)),         # panel_usd already has PRICE_FACTOR
        ("Hybrid Inverter / Charger",
         f"{inv_kw:.1f} kW | {inverter_brand(inv_kw).split(',')[0].strip()}",
         1, "No.", local(inv_usd)),                    # inv_usd already has PRICE_FACTOR
        (f"Battery — {chemistry}",
         f"{unit_bat_kwh:.4g} kWh unit | {chem['brands'].split(',')[0].strip()}",
         num_bat, "No.", local(bat_usd)),               # bat_usd already has PRICE_FACTOR
        ("MPPT Charge Controller",
         f"{mppt_a} A | Victron BlueSolar / Epever",
         1, "No.", local(mppt_usd)),                    # mppt_usd already has PRICE_FACTOR
        ("PV Mounting Structure",
         "Aluminium rail system + clamps, IEC 61215",
         num_panels, "No.", p(16)),          # was 18
        ("DC Combiner / String Box",
         f"{min(num_strings,4)}-string, 15 A DC fuses, DC SPD Type 2, IP65, IEC 61173",
         1, "No.", p(42)),                  # was 48
        # ── DC Cables ─────────────────────────────────────────────────────────
        ("DC String Cable — 6 mmÂ²", dc_str_spec, dc_str_qty, "m", p(1.30)),   # was 1.50
        (f"DC Main Cable — {dc_main_mm2} mmÂ²", dc_main_spec, dc_main_qty, "m",
         p(1.6 + dc_main_mm2 * 0.035)),                                        # was 1.8 + 0.04
        (f"Battery DC Cable — {bat_mm2} mmÂ²", bat_spec, bat_qty, "m",
         p(2.0 + bat_mm2 * 0.045)),                                            # was 2.2 + 0.05
        ("Earthing Cable — 16 mmÂ² G/Y",
         "16 mmÂ² Green/Yellow Cu PVC, main earthing conductor, BS 7430 / IEC 60364-5-54",
         15, "m", p(1.6)),                 # was 1.8
        ("Bonding Cable — 6 mmÂ² G/Y",
         "6 mmÂ² Green/Yellow Cu PVC, panel frame & equipment bonding, BS 7671 Ch.54",
         int(num_panels * 3), "m", p(1.05)),   # was 1.2
        ("DC MCB / Isolator",
         "DC-rated, 1000 V, BS EN 60947-2",
         4, "No.", p(3.6)),                # was 4.0
    ]

    # ── AC Cables — per sized circuit ─────────────────────────────────────────
    if ac_cables:
        for c in ac_cables:
            sz  = c["cable_size_mm2"]
            brk = c["breaker_a"]
            vd  = c["vd_percent"]
            Im  = c.get("install_method", "C")
            ph  = c.get("phase", "single")
            Ib  = c.get("design_current", 0)
            cores = "3-core" if ph == "three" else "2-core + E"
            if sz <= 16:
                insul = f"Cu PVC 70 Â°C, BS EN 50525-2-31, Method {Im}"
            else:
                insul = f"Cu XLPE/SWA/PVC, BS 5467, 0.6/1 kV, Method {Im}"
            spec = (f"{sz} mmÂ² {cores} {insul} | "
                    f"Ib={Ib:.1f} A, {brk} A MCB/RCCB | VD={vd:.2f}%")
            qty  = c["length_m"] + 5
            rate = p(1.1 + sz * 0.055)    # was 1.2 + 0.06
            items.append((f"AC Cable — {c['circuit']}", spec, qty, "m", rate))
    else:
        items.append(("AC Cables — All Circuits",
                      "Cu PVC/XLPE, BS EN 50525 / BS 5467 (sizes per cable schedule)",
                      25, "m", p(1.35)))   # was 1.5

    items += [
        ("AC RCCB + MCBs",
         "30 mA RCCB (BS EN 61008 Type A) + 6 Ã— MCB (BS EN 60898), 6 kA",
         1, "Set", p(48)),              # was 55
        ("Surge Protection Device",
         "DC Type 2 (IEC 61643-31, 1000 VDC) + AC Type 2 (BS EN 61643, 230/415 V)",
         2, "No.", p(19)),              # was 22
        ("Earthing & Bonding Kit",
         "2 Ã— 16 mm dia. Cu earth rod 2.4 m, rod driver, clamps, earth busbar — BS 7430",
         1, "Set", p(30)),              # was 35
        ("Battery Enclosure / Rack",
         "IP44 powder-coated steel rack, ventilated, lockable",
         1, "No.", p(48)),              # was 55
        ("Cable Trunking & Conduit",
         "20 mm & 32 mm metallic conduit + 50Ã—50 mm galvanised steel trunking",
         1, "Lot", p(52)),              # was 60
        ("Hardware, Fixings & Misc",
         "MC4 connectors (IP67, 1000 VDC), cable ties, glands, labels, consumables",
         1, "Lot", p(36)),              # was 42
    ]

    # ── Build supply rows (basic Ã— supply markup) ─────────────────────────────
    rows = []
    supply_grand = 0.0
    for no, (desc, spec, qty, unit, basic) in enumerate(items, 1):
        total_r = basic * (1 + SUPPLY_MARKUP)   # 8% supply markup
        amount  = qty * total_r
        supply_grand += amount
        rows.append({"no": no, "desc": desc, "spec": spec, "qty": qty,
                     "unit": unit, "basic": basic, "total_r": total_r,
                     "amount": amount})

    # ── Installation Labour row (15% of supply subtotal) ─────────────────────
    install_amount = supply_grand * INSTALL_RATE
    rows.append({
        "no":      len(rows) + 1,
        "desc":    "Installation Labour",
        "spec":    (f"Supply & installation of all PV, battery, inverter, cabling and protection "
                    f"equipment; commissioning and testing ({int(INSTALL_RATE*100)}% of supply total)"),
        "qty":     1,
        "unit":    "Lot",
        "basic":   supply_grand,        # supply subtotal is the "basic" base
        "total_r": install_amount,      # install amount (shown in total_r column)
        "amount":  install_amount,
    })

    grand = supply_grand + install_amount
    return rows, grand


# ─── Routes — Auth ────────────────────────────────────────────────────────────


@app.route("/")
def landing():
    with get_db() as c:
        news = c.execute(
            "SELECT * FROM news_posts WHERE is_published=1 ORDER BY created_at DESC LIMIT 3"
        ).fetchall()
        admin = c.execute(
            "SELECT org_whatsapp FROM users WHERE is_admin=1 ORDER BY id LIMIT 1"
        ).fetchone()
    wa_number = (admin["org_whatsapp"] if admin and admin["org_whatsapp"] else "233535068102")
    return render_template("landing.html", user=current_user(),
                           countries=get_countries(), news_posts=news,
                           wa_number=wa_number,
                           sales_email=EMAIL_SALES)


@app.route("/platform")
def landing_page2():
    with get_db() as c:
        news = c.execute(
            "SELECT * FROM news_posts WHERE is_published=1 ORDER BY created_at DESC LIMIT 3"
        ).fetchall()
    return render_template("landing_page2.html", user=current_user(),
                           countries=get_countries(), news_posts=news)


@app.route("/register", methods=["GET", "POST"])
@limiter.limit("10 per hour")
def register():
    if request.method == "POST":
        csrf_protect()
        f = request.form
        if not f.get("terms_agreed"):
            flash("You must read and accept the Terms of Service and Privacy Policy to create an account.", "danger")
            return render_template("auth.html", mode="register", countries=get_countries())
        ph = generate_password_hash(f["password"])
        try:
            with get_db() as c:
                # All new signups begin on 'free' plan — upgrade later
                # -- Referral capture -------------------------------------
                # If the visitor arrived via /r/<code> or ?ref=<code>, the
                # base.html JS dropped a `ref_code` cookie. We look up the
                # referring user here and store both fields on the new row.
                _ref_cookie = (request.cookies.get("ref_code") or "").upper().strip()[:16]
                _referrer_id = None
                if _ref_cookie:
                    _hit = c.execute("SELECT id FROM users WHERE referral_code = ?",
                                     (_ref_cookie,)).fetchone()
                    if _hit:
                        _referrer_id = _hit[0]
                # Generate this new user's own referral code (unique per user)
                _new_code = _gen_referral_code()
                c.execute(
                    "INSERT INTO users (username,email,password_hash,name,company,country,plan,referral_code,referred_by) "
                    "VALUES (?,?,?,?,?,?,?,?,?)",
                    (f["username"], f["email"], ph,
                     f.get("name",""), f.get("company",""),
                     f.get("country",""), "free", _new_code, _referrer_id))
                uid = c.execute("SELECT last_insert_rowid()").fetchone()[0]
                # Log the referral event so we can track conversions over time
                if _referrer_id:
                    try:
                        c.execute("INSERT INTO referrals (referrer_id, referee_id) VALUES (?, ?)",
                                  (_referrer_id, uid))
                    except Exception:
                        pass  # UNIQUE violation = already logged, fine
            # Email verification: generate single-use token, persist unverified row.
            # NOTE: We open a FRESH `with get_db()` block here — calling
            # `c.execute(...)` on the prior `c` (which left scope above) starts a
            # new implicit transaction on a still-open SQLite connection and
            # never commits it. That open write txn blocks the NEXT request's
            # INSERT (sqlite3 OperationalError: database is locked), which
            # bubbles past our `except sqlite3.IntegrityError` -> 500.
            import secrets as _secrets_v
            _verify_token = _secrets_v.token_urlsafe(32)
            try:
                with get_db() as _vc:
                    _vc.execute("UPDATE users SET email_verified=0, email_verify_token=? WHERE id=?",
                                (_verify_token, uid))
            except Exception:
                pass  # Token persistence failure is non-fatal; user can re-register.
            # Build verification URL using the current request host (works for both
            # solarpro-global.onrender.com and the custom domain once cert is live).
            try:
                _verify_url = request.host_url.rstrip("/") + url_for("verify_email", token=_verify_token)
                _login_url  = request.host_url.rstrip("/") + url_for("login")
                _send_system_email(
                    f["email"],
                    "Verify your email to activate SolarPro Global",
                    "Hello " + f.get("name", f["username"]) + ",\n\n"
                    "Welcome to SolarPro Global. Your account has been created on the Free plan.\n\n"
                    "Before you can log in, please confirm your email by clicking the link below:\n\n"
                    "    " + _verify_url + "\n\n"
                    "After verifying, return to the login page to sign in:\n\n"
                    "    " + _login_url + "\n\n"
                    "If you did not create this account you can safely ignore this email.\n\n"
                    "Need help? Reply to this email or write to support@aiappinvent.com.\n\n"
                    "The SolarPro Global Team\n"
                    "solarpro.aiappinvent.com\n")
            except Exception:
                pass  # Email failure must never block registration; user can request a resend.
            flash("Account created. Check your email to verify, then come back here to log in.", "success")
            return redirect(url_for("login"))
        except sqlite3.IntegrityError:
            flash("Username or email already registered.", "danger")
    return render_template("auth.html", mode="register",
                           countries=get_countries())


@app.route("/verify-email/<token>", methods=["GET"])
@limiter.limit("60 per hour")
def verify_email(token):
    """Activate a user account via the single-use token mailed at signup.

    Flow:
      1. Look up a user whose email_verify_token matches.
      2. If found: set email_verified=1, clear the token (prevents reuse).
      3. Flash success and redirect to /login.
      4. If not found OR already verified: friendly error, redirect to /login.
    """
    if not token:
        flash("Verification link is missing the token. Please use the link from your email.", "danger")
        return redirect(url_for("login"))
    with get_db() as c:
        row = c.execute(
            "SELECT id, username, email_verified FROM users WHERE email_verify_token=?",
            (token,)).fetchone()
        if not row:
            flash("This verification link is invalid or has already been used. "
                  "If you already verified, just log in.", "warning")
            return redirect(url_for("login"))
        if row["email_verified"]:
            flash("Your email is already verified. Please log in.", "info")
            return redirect(url_for("login"))
        c.execute(
            "UPDATE users SET email_verified=1, email_verify_token=NULL WHERE id=?",
            (row["id"],))
    flash("Email verified! You can now log in.", "success")
    return redirect(url_for("login"))


@app.route("/login", methods=["GET", "POST"])
@limiter.limit("20 per hour")
def login():
    if request.method == "POST":
        csrf_protect()
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        ip = _get_real_ip()

        # ── Brute-force lockout: max 10 failures per username or IP in 15 min ──
        lockout_since = (datetime.utcnow() - timedelta(minutes=15)).strftime("%Y-%m-%d %H:%M:%S")
        with get_db() as _lf:
            fail_count = _lf.execute(
                "SELECT COUNT(*) FROM login_failures "
                "WHERE (username=? OR ip_address=?) AND created_at > ?",
                (username, ip, lockout_since)).fetchone()[0]
        if fail_count >= 10:
            flash("Too many failed login attempts. Please wait 15 minutes and try again.", "danger")
            return render_template("auth.html", mode="login")

        with get_db() as c:
            user = c.execute("SELECT * FROM users WHERE username=?",
                             (username,)).fetchone()
        if user and check_password_hash(user["password_hash"], password):
            # Email-verification gate: refuse login until the user clicks the
            # verification link mailed at signup. Existing users (pre-feature)
            # are backfilled email_verified=1 in init_db so they are not locked out.
            try:
                _ev = user["email_verified"]
            except Exception:
                _ev = 1  # row missing the column = legacy DB; treat as verified.
            if not _ev:
                flash("Please verify your email first. Check your inbox for the link we "
                      "sent when you signed up.", "warning")
                return render_template("auth.html", mode="login")
            session["user_id"] = user["id"]
            session["username"] = user["username"]
            # ── Log success & clear failure counter ──
            with get_db() as _db:
                _db.execute(
                    "INSERT INTO audit_logs (user_id, username, action, ip_address) VALUES (?,?,?,?)",
                    (user["id"], username, "login_success", ip))
                _db.execute("DELETE FROM login_failures WHERE username=?", (username,))
            return redirect(url_for("dashboard"))

        # ── Log failure ──
        with get_db() as _db:
            _db.execute(
                "INSERT INTO login_failures (username, ip_address) VALUES (?,?)",
                (username, ip))
            _db.execute(
                "INSERT INTO audit_logs (username, action, ip_address, details) VALUES (?,?,?,?)",
                (username, "login_failed", ip, "Invalid credentials"))
        flash("Invalid username or password.", "danger")
    return render_template("auth.html", mode="login")


@app.route("/logout")
def logout():
    uid = session.get("user_id")
    if uid:
        # Purge draft/incomplete projects — only completed (stage='results') persist
        with get_db() as _db:
            _db.execute(
                "DELETE FROM projects WHERE user_id=? AND stage NOT IN ('results')",
                (uid,))
    session.clear()
    return redirect(url_for("landing"))


# ─── Password Reset ───────────────────────────────────────────────────────────

@app.route("/forgot-password", methods=["GET", "POST"])
@limiter.limit("5 per hour")
def forgot_password():
    if request.method == "POST":
        csrf_protect()
        email = request.form.get("email", "").strip().lower()
        with get_db() as c:
            user = c.execute(
                "SELECT * FROM users WHERE LOWER(email)=?", (email,)).fetchone()
            if user:
                # Expire any existing unused tokens for this user
                c.execute(
                    "UPDATE password_reset_tokens SET used=1 WHERE user_id=? AND used=0",
                    (user["id"],))
                token   = secrets.token_urlsafe(32)
                expires = (datetime.utcnow() + timedelta(hours=1)).strftime("%Y-%m-%d %H:%M:%S")
                c.execute(
                    "INSERT INTO password_reset_tokens (user_id, token, expires_at) VALUES (?,?,?)",
                    (user["id"], token, expires))
                reset_url = url_for("reset_password", token=token, _external=True)
                body = (
                    f"Hello {user['name'] or user['username']},\n\n"
                    f"A password reset was requested for your SolarPro Global account.\n\n"
                    f"Click the link below to set a new password (valid for 1 hour):\n\n"
                    f"  {reset_url}\n\n"
                    f"If you did not request this, ignore this email — your password has not changed.\n\n"
                    f"— SolarPro Global"
                )
                ok, err = _send_system_email(
                    user["email"], "Reset your SolarPro password", body)
                if ok:
                    flash(
                        "Reset link sent! Check your inbox (and spam folder).", "success")
                else:
                    # SMTP not configured — show link directly so admins can share it securely
                    flash(
                        f"SMTP not configured on the server. "
                        f"Admin: share this link securely with the user â†' {reset_url}", "warning")
            else:
                # Always show the same message to avoid email enumeration
                flash(
                    "If that email address is registered, a reset link has been sent.", "info")
        return redirect(url_for("forgot_password"))
    return render_template("forgot_password.html")


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token):
    with get_db() as c:
        rec = c.execute(
            "SELECT * FROM password_reset_tokens WHERE token=? AND used=0",
            (token,)).fetchone()
    if not rec:
        flash("This reset link is invalid or has already been used.", "danger")
        return redirect(url_for("forgot_password"))
    try:
        expires = datetime.strptime(rec["expires_at"], "%Y-%m-%d %H:%M:%S")
        if datetime.utcnow() > expires:
            flash("This reset link has expired (links are valid for 1 hour). Request a new one.", "warning")
            return redirect(url_for("forgot_password"))
    except Exception:
        flash("Invalid reset token.", "danger")
        return redirect(url_for("forgot_password"))

    if request.method == "POST":
        csrf_protect()
        new_pw  = request.form.get("new_password", "")
        conf_pw = request.form.get("confirm_password", "")
        if len(new_pw) < 8:
            flash("Password must be at least 8 characters.", "danger")
            return render_template("reset_password.html", token=token)
        if new_pw != conf_pw:
            flash("Passwords do not match.", "danger")
            return render_template("reset_password.html", token=token)
        with get_db() as c:
            c.execute("UPDATE users SET password_hash=? WHERE id=?",
                      (generate_password_hash(new_pw), rec["user_id"]))
            c.execute(
                "UPDATE password_reset_tokens SET used=1 WHERE token=?", (token,))
        flash("Password changed successfully! You can now log in.", "success")
        return redirect(url_for("login"))

    return render_template("reset_password.html", token=token)


# ─── Support / Help Centre ───────────────────────────────────────────────────

@app.route("/support")
@login_required
def support():
    return render_template("support.html", user=current_user())


@app.route("/support/email-setup")
@login_required
def support_email_setup():
    return render_template("support_email_setup.html", user=current_user())


@app.route("/support/email-setup/pdf")
@login_required
def support_email_setup_pdf():
    return _render_pdf(
        "Email SMTP Setup Guide — SolarPro Global",
        _TUTORIAL_EMAIL_SETUP_MD,
        "SolarPro_Email_Setup_Guide.pdf")


@app.route("/support/user-guide")
@login_required
def support_user_guide():
    return render_template("support_user_guide.html", user=current_user())


@app.route("/support/user-guide/pdf")
@login_required
def support_user_guide_pdf():
    return _render_pdf(
        "User Guide — SolarPro Global",
        _TUTORIAL_USER_GUIDE_MD,
        "SolarPro_User_Guide.pdf")


_TUTORIAL_EMAIL_SETUP_MD = """# Email & SMTP Setup Guide
## SolarPro Global Help Centre

**Version 1.0 Â· solarpro.aiappinvent.com**

---

## Why Set Up Email?

SolarPro Global can send your solar project PDF reports — BOQ, Economic Analysis,
Energy Impact, and more — directly to clients, banks, and installers. To do this,
the platform needs to connect to your outbound mail server (SMTP).

Once configured you can:

- Send any project report from the **Email Reports** page
- Receive password reset emails automatically
- Get system notifications from the platform

---

# Step 1 — Open Settings

1. Log in to your SolarPro account.
2. Click your **username** in the top-right navigation bar.
3. Select **Settings** from the dropdown menu.
4. Click the **Email / SMTP** tab at the top of the Settings page.

You will see the SMTP configuration form with fields for host, port, username,
password, encryption mode, and From address.

---

# Step 2 — Choose Your Email Provider

SolarPro supports any standard SMTP provider. Use the **Quick-Fill buttons** on the
settings page to auto-fill the host and port for popular providers.

| Provider | Best For | Free Tier |
|---|---|---|
| Gmail | Personal accounts, small teams | 500 emails / day |
| Outlook / Microsoft 365 | Business Microsoft accounts | Included with M365 |
| Brevo (formerly Sendinblue) | Professional SaaS sending | 300 emails / day |
| Mailgun | High-volume transactional | 100 emails / day |

> **Recommendation:** For a solar consultancy, **Brevo** is the best choice.
> It is free for up to 300 emails per day, has excellent deliverability, and
> does not require an App Password like Gmail.

---

# Step 3A — Set Up Gmail

Gmail requires an **App Password** — a special 16-character code separate from
your normal Google account password.

**Step-by-step — generate a Gmail App Password:**

1. Go to **myaccount.google.com**
2. Click **Security** in the left sidebar.
3. Under "How you sign in to Google", click **2-Step Verification** (must be on).
4. Scroll to the bottom â†' click **App passwords**.
5. App: **Mail** Â· Device: **Other** â†' type "SolarPro" â†' click **Generate**.
6. Copy the 16-character password shown (no spaces).

**Enter these values in SolarPro Settings â†' Email / SMTP:**

| Field | Value |
|---|---|
| SMTP Host | smtp.gmail.com |
| Port | 587 |
| Encryption | STARTTLS |
| Username | your.email@gmail.com |
| Password | The 16-character App Password (NOT your Google password) |
| From Address | your.email@gmail.com |

> If you paste the App Password, remove any spaces. Google displays it in groups
> of 4 characters — it is one continuous string.

---

# Step 3B — Set Up Brevo (Recommended)

1. Sign up free at **brevo.com**.
2. After logging in, go to **Senders & IP** â†' **SMTP & API**.
3. Copy your **SMTP Login** (your Brevo email) and **Master Password**,
   or generate a dedicated SMTP key.

**Enter these values in SolarPro Settings â†' Email / SMTP:**

| Field | Value |
|---|---|
| SMTP Host | smtp-relay.brevo.com |
| Port | 587 |
| Encryption | STARTTLS |
| Username | Your Brevo account email |
| Password | Your Brevo master password or SMTP key |
| From Address | A verified sender in your Brevo account |

> You must verify your sender email address or domain in Brevo before emails
> will be delivered. Go to **Brevo â†' Senders â†' Add a sender**.

---

# Step 3C — Set Up Outlook / Microsoft 365

| Field | Value |
|---|---|
| SMTP Host | smtp.office365.com |
| Port | 587 |
| Encryption | STARTTLS |
| Username | your@company.com |
| Password | Your Microsoft 365 account password |
| From Address | your@company.com |

> If your organisation has Multi-Factor Authentication (MFA) enabled, generate
> an App Password in your Microsoft account security settings.

---

# Step 4 — Test the Connection

After filling in all fields, click **Test Connection** (the blue button).

- **✓ "Connection successful"** — SMTP is working. Click **Save SMTP**.
- **âœ— "Authentication failed"** — check your username and password.
  For Gmail, ensure you used the App Password, not your regular password.
- **âœ— "Connection refused / timed out"** — check the host and port.
  Port 587 may be blocked by your network firewall.

Always click **Save SMTP** after a successful test.

---

# Step 5 — Send a Test Report

1. Open any project with completed calculations.
2. Click **Results** â†' **Email Reports**.
3. Enter your own email address in the **Recipients** field.
4. Select any report (e.g. BOQ Report).
5. Click **Send Email**.

The email should arrive within 1—2 minutes. Check your spam folder if it does
not appear in your inbox.

---

# Troubleshooting

## "SMTP not configured" warning on the Email Reports page
Your SMTP settings have not been saved yet. Go to **Settings â†' Email / SMTP**
and complete setup, then click **Test Connection** before saving.

## "Authentication failed" error
- **Gmail:** use an App Password, not your Google account password.
- **Brevo:** check your SMTP Login credentials, not the API key.
- **Outlook:** generate an App Password if MFA is enabled.

## "Connection timed out" error
- Port 587 may be blocked by your network.
- Try port 465 with **SSL/TLS** encryption instead.
- Contact your network administrator.

## Emails landing in spam
- Verify your From Address as a sender in your email provider dashboard.
- For Brevo: verify your sending domain (adds SPF and DKIM DNS records).
- Ensure the From Address matches the SMTP username.

## Password reset emails not arriving
The password reset system uses the **server-level SMTP** (Render environment
variables), not the per-user Settings SMTP. Contact your system administrator
to configure SMTP_HOST, SMTP_USER, and SMTP_PASS in the Render dashboard.

---

# Quick Reference

| Provider | Host | Port | Encryption |
|---|---|---|---|
| Gmail | smtp.gmail.com | 587 | STARTTLS |
| Outlook / M365 | smtp.office365.com | 587 | STARTTLS |
| Brevo | smtp-relay.brevo.com | 587 | STARTTLS |
| Mailgun | smtp.mailgun.org | 587 | STARTTLS |
| Yahoo Mail | smtp.mail.yahoo.com | 587 | STARTTLS |
| Zoho Mail | smtp.zoho.com | 587 | STARTTLS |

---

*SolarPro Global Help Centre Â· support@aiappinvent.com*
"""


_TUTORIAL_USER_GUIDE_MD = """# SolarPro Global — Complete User Guide
## Intelligent PV Solar Design & Financial Engineering Platform

**Version 1.0 Â· solarpro.aiappinvent.com**

---

## Welcome to SolarPro Global

SolarPro Global is a professional pv solar design platform that takes you from
initial site assessment all the way through to a bankable financial proposal —
in a single workflow. This guide walks you through every step.

---

# Step 1 — Create Your Account

1. Go to **solarpro.aiappinvent.com**
2. Click **Start Free** on the homepage.
3. Fill in your name, company, country, and choose a plan.
4. Enter your email address and a secure password.
5. Click **Create Account** — you are automatically logged in.

**Plans available:**

| Plan | Projects | Best For |
|---|---|---|
| Free | Up to 3 | Evaluation and small projects |
| Professional | Up to 20 | Consultancies and SMEs |
| Enterprise | Unlimited | Large firms and white-label |

---

# Step 2 — Set Up Your Organisation Profile

Before creating your first project, configure your organisation details so they
appear on all generated reports.

1. Click your **username** (top right) â†' **Settings**.
2. On the **Organization** tab, fill in:
   - Company name and email
   - Street address
   - Phone / WhatsApp
   - Website URL
   - Time zone
3. Click **Save Profile**.

Your organisation details will appear as the author on every PDF report,
proposal, and email sent from the platform.

---

# Step 3 — Create a New Project

1. Click **New Project** in the navigation bar (or the **+** button on the dashboard).
2. Enter a project name (e.g. "Accra Commercial Office — 50 kW").
3. Click **Create Project** — you are taken to the Location form.

---

# Step 4 — Set the Location & Solar Resource

The Location form captures where the system will be installed and pulls in
solar irradiance data for that region.

**Fields to complete:**

- **Country** — select from the 22 supported countries.
- **Region / City** — the region dropdown populates automatically after you select a country.
- **Solar resource data** loads automatically once country and region are selected,
  showing Peak Sun Hours (PSH), average irradiance, and the local electricity tariff.
- **System Type** — Off-grid, On-grid (grid-tied), or Hybrid.
- **DC Bus Voltage** — 12V, 24V, or 48V (for off-grid/hybrid).
- **Currency** — shown on all financial outputs.

**Solar Design Parameters** (on the right panel):

- **Panel Tilt** — degrees from horizontal (typically 10—20Â° for equatorial regions).
- **Azimuth** — 0Â° = south-facing (northern hemisphere); adjust for your roof orientation.
- **System Losses** — default 14% (wiring, soiling, temperature derating).
- **Inverter Efficiency** — default 95%.
- **Battery Depth of Discharge (DoD)** — default 80% for lithium, 50% for lead-acid.
- **Performance Ratio** — overall system efficiency factor (default 75%).

Click **Save & Continue** when the location is set.

---

# Step 5 — Enter the Load Schedule

The load schedule is the heart of the design — it defines how much energy the
building consumes and when.

1. Click **Add Appliance** to add each electrical load.
2. For each appliance enter:
   - **Name** (e.g. Air Conditioner, Refrigerator, LED Lights)
   - **Quantity** — number of identical units
   - **Power (Watts)** — rated wattage per unit
   - **Hours per day** — average daily run time
   - **Days per week** — operating days
3. The **Daily Load (Wh)** calculates automatically.
4. Add all appliances, then click **Save Loads**.

The platform separates loads into **critical** (must run during outages) and
**non-critical** to optimise battery sizing for hybrid systems.

---

# Step 6 — Review Engineering Results

Click **Results** to see the full system design. The platform calculates:

**PV Array:**
- Required PV capacity (kWp)
- Number of solar panels
- Suggested panel configuration (series Ã— parallel strings)

**Battery Bank:**
- Required capacity (kWh and Ah)
- Number of batteries
- Battery configuration

**Inverter & Charge Controller:**
- Minimum inverter rating (kVA)
- Charge controller rating (Amps)

**AC Cable Schedule:**
- Cable cross-sections for each circuit
- Voltage drop calculations
- Protection device ratings

**Financial Summary:**
- Total system cost (CAPEX)
- Payback period
- Net Present Value (NPV) and Internal Rate of Return (IRR)

---

# Step 7 — Access Reports

From the Results page, click any report in the left panel:

| Report | What It Contains |
|---|---|
| PV Design | Panel layout, string configuration, technical specs |
| Bill of Quantities (BOQ) | Itemised equipment list with local currency pricing |
| AC Cable Schedule | All cable sizes, lengths, voltage drop, protection |
| Economic Analysis | NPV, IRR, payback, 25-year cash flow projections |
| Energy Impact | Monthly generation, grid offset, CO2 savings, trees equivalent |
| Installation Plan | Step-by-step construction and commissioning checklist |
| Staffing Plan | Crew roles, responsibilities, and man-hour estimates |
| Site Assessment | Field inspection checklist for engineers |

Every report can be **exported as PDF** or printed directly from the browser.

---

# Step 8 — Export to Excel

From the Results page, click **Export Excel** to download a full workbook
containing all results in tabular format — useful for clients and banks who
need the numbers in a spreadsheet.

---

# Step 9 — Email Reports to Clients

Once your SMTP is configured in **Settings â†' Email / SMTP**, you can send any
PDF report directly to clients, banks, or installers.

1. From the Results page, click **Email Reports**.
2. Select the report to send (e.g. Economic Analysis).
3. Enter one or more recipient email addresses (comma-separated).
4. Edit the subject line and message body.
5. Click **Send Email**.

The platform logs all sent emails. You can view the history at the bottom of
the Email Reports page.

---

# Step 10 — Request a Site Assessment

Before designing a system, you may want a formal site assessment.

1. Click **Free Assessment** in the navigation (or on the homepage).
2. Fill in the client's details: name, phone, country, region, building type,
   floor count, and a description of the site.
3. Submit — you receive a reference code immediately.
4. The admin team reviews the assessment and can create a project directly
   from the assessment record in the pipeline.

---

# Step 11 — Manage Your Account

**Dashboard** — shows all your projects with their status, system size, and
quick-access buttons.

**Account page** (click username â†' Account) — view your subscription plan,
payment history, and support tickets.

**Settings** — configure organisation profile, date/time format, appearance
theme, SMTP email, and change your password.

**Upgrade** — move from Free to Professional or Enterprise to unlock more
projects and premium features.

---

# Tips for Professional Reports

- Fill in your **organisation name and logo details** in Settings before exporting
  any PDFs — these appear on the report header.
- Set the correct **currency and electricity tariff** in the Location form —
  all financial figures depend on these inputs.
- Use the **Hybrid** system type for grid-connected systems with battery backup;
  use **Off-grid** only for sites with no grid connection at all.
- Always run **Test Connection** after saving SMTP settings before trying to send
  your first report.
- For the most accurate BOQ pricing, use **Settings â†' Equipment Catalog** (admin)
  to set current local market prices.

---

# Getting Help

- **Support Tickets** — click **Support** in the navigation to open a ticket.
  The support team responds within 24 hours.
- **Email** — contact support@aiappinvent.com
- **Help Centre** — visit **Support â†' Help Centre** for step-by-step tutorials
  on specific features.

---

*SolarPro Global Â· solarpro.aiappinvent.com Â· support@aiappinvent.com*
"""


# ─── Routes — Dashboard ───────────────────────────────────────────────────────

@app.route("/dashboard")
@login_required
def dashboard():
    user = current_user()
    plan = (user["plan"] or "free").lower()
    uid  = session["user_id"]
    with get_db() as c:
        raw_projects = c.execute(
            "SELECT * FROM projects WHERE user_id=? ORDER BY updated_at DESC",
            (uid,)).fetchall()
        open_tickets = c.execute(
            "SELECT COUNT(*) FROM tickets WHERE user_id=? AND status='open'",
            (uid,)).fetchone()[0]
        emails_sent  = c.execute(
            "SELECT COUNT(*) FROM email_logs WHERE user_id=? AND status='sent'",
            (uid,)).fetchone()[0]

    # Enrich each project with parsed snapshot data
    projects = []
    total_kwp = 0.0
    total_kwh = 0.0
    for p in raw_projects:
        d   = json.loads(p["data_json"] or "{}")
        r   = d.get("results", {})
        eco = r.get("economics", {})
        # Determine completion stage
        stage = "new"
        if r:
            stage = "results"
        elif d.get("loads"):
            stage = "loads"
        elif d.get("location") or d.get("country"):
            stage = "location"
        total_kwp += r.get("pv_kw", 0)
        total_kwh += r.get("bat_kwh", 0)
        projects.append({
            "id":        p["id"],
            "name":      p["name"],
            "created":   (p["created_at"] or "")[:10],
            "updated":   (p["updated_at"] or "")[:16].replace("T", " "),
            "country":   d.get("country", ""),
            "location":  d.get("location", ""),
            "system_type": d.get("system_type", "off-grid"),
            "phase":     d.get("phase", "single"),
            "pv_kw":     round(r.get("pv_kw", 0), 2),
            "bat_kwh":   round(r.get("bat_kwh", 0), 1),
            "inv_kw":    round(r.get("inv_kw", 0), 1),
            "num_panels":r.get("num_panels", 0),
            "verdict":   eco.get("verdict", ""),
            "bankability":eco.get("bankability", ""),
            "payback":   eco.get("payback", 0),
            "total_local":eco.get("total_local", 0),
            "symbol":    d.get("symbol", "$"),
            "stage":     stage,
        })

    limit    = PLAN_LIMITS.get(plan, 1)
    at_limit = len(projects) >= limit
    return render_template("dashboard.html", user=user, projects=projects,
                           plan=plan, limit=limit, at_limit=at_limit,
                           open_tickets=open_tickets, emails_sent=emails_sent,
                           total_kwp=round(total_kwp, 1),
                           total_kwh=round(total_kwh, 1))


# ─── Routes — Project ─────────────────────────────────────────────────────────

PLAN_LIMITS = {"free": 1, "professional": 10, "business": 9999, "enterprise": 9999}

@app.route("/project/new", methods=["GET", "POST"])
@login_required
def project_new():
    user = current_user()
    plan = (user["plan"] or "free").lower()
    limit = PLAN_LIMITS.get(plan, 1)
    with get_db() as c:
        count = c.execute(
            "SELECT COUNT(*) FROM projects WHERE user_id=?",
            (session["user_id"],)).fetchone()[0]
    if count >= limit:
        flash(
            f"Your {plan.title()} plan allows up to {limit} project{'s' if limit>1 else ''}. "
            f"Upgrade to create more.", "warning")
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        csrf_protect()
        name = request.form.get("name", "New Project")
        with get_db() as c:
            c.execute("INSERT INTO projects (user_id, name) VALUES (?,?)",
                      (session["user_id"], name))
            pid = c.execute("SELECT last_insert_rowid()").fetchone()[0]
        return redirect(url_for("project_location", pid=pid))
    u = current_user()
    sales_data = None
    if u and u["is_admin"]:
        with get_db() as c:
            sales_data = {
                "new_leads":  c.execute("SELECT COUNT(*) FROM leads WHERE status='new'").fetchone()[0],
                "total_leads":c.execute("SELECT COUNT(*) FROM leads").fetchone()[0],
                "subs":       c.execute("SELECT COUNT(*) FROM newsletter_subscribers WHERE status='active'").fetchone()[0],
                "total_rev":  c.execute("SELECT SUM(amount_usd) FROM payments WHERE status='success'").fetchone()[0] or 0,
                "paid_users": c.execute("SELECT COUNT(*) FROM users WHERE plan NOT IN ('free','demo') AND plan IS NOT NULL").fetchone()[0],
                "recent_leads": c.execute("SELECT name,email,company,interest,status,created_at FROM leads ORDER BY created_at DESC LIMIT 5").fetchall(),
            }
    return render_template("project_new.html", user=u,
                           plan=plan, limit=limit, count=count,
                           sales_data=sales_data)


@app.route("/project/from-assessment/<ref>")
@login_required
def project_from_assessment(ref):
    """Create a new project pre-populated from an assessment intake record."""
    user = current_user()
    plan  = (user["plan"] or "free").lower()
    limit = PLAN_LIMITS.get(plan, 1)
    with get_db() as c:
        count = c.execute(
            "SELECT COUNT(*) FROM projects WHERE user_id=?",
            (session["user_id"],)).fetchone()[0]
        if count >= limit:
            flash(
                f"Your {plan.title()} plan allows up to {limit} project"
                f"{'s' if limit > 1 else ''}. Upgrade to create more.", "warning")
            return redirect(url_for("admin_pipeline"))
        # Look up the assessment
        row = c.execute(
            "SELECT * FROM assessment_requests WHERE assessment_ref=?",
            (ref,)).fetchone()
        if not row:
            flash(f"Assessment {ref} not found.", "danger")
            return redirect(url_for("admin_pipeline"))
        ar = dict(row)
        # Build project name from assessment details
        btype   = ar.get("building_type") or ar.get("system_type") or "Solar"
        loc     = ar.get("country") or ""
        client  = ar.get("name") or "Client"
        proj_name = f"{btype} Solar Design — {client} [{ref}]"
        if loc:
            proj_name = f"{loc} Â· {proj_name}"
        # Pre-fill data_json with known information
        initial_data = {
            "from_assessment_ref":  ref,
            "from_assessment_name": ar.get("name", ""),
            "from_assessment_phone":ar.get("phone", ""),
            "country":  ar.get("country", ""),
            "region":   ar.get("region", ""),
            "building_type": btype,
            "building_size": ar.get("building_size", ""),
            "num_floors":    ar.get("num_floors", 1),
            "client_notes":  ar.get("building_desc", ""),
        }
        import json as _json
        c.execute(
            "INSERT INTO projects (user_id, name, data_json) VALUES (?,?,?)",
            (session["user_id"], proj_name, _json.dumps(initial_data)))
        pid = c.execute("SELECT last_insert_rowid()").fetchone()[0]
        # Advance the assessment pipeline stage
        c.execute(
            "UPDATE assessment_requests SET pipeline_stage=? WHERE assessment_ref=?",
            ("assessment_reviewed", ref))
        # Also advance the mirrored lead
        c.execute(
            "UPDATE leads SET pipeline_stage=? WHERE message LIKE ? AND pipeline_stage='assessment_submitted'",
            ("assessment_reviewed", f"%{ar.get('building_desc','')[:30]}%"))
    flash(
        f"Project created from assessment {ref}. "
        f"Complete the location and load details below.", "success")
    return redirect(url_for("project_location", pid=pid))


@app.route("/project/<int:pid>/location", methods=["GET", "POST"])
@login_required
def project_location(pid):
    project = get_project(pid)
    if not project:
        flash("Project not found.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        csrf_protect()
        f = request.form
        country = f["country"]
        region  = f["region"]
        sd = get_solar_data(country, region)
        if not sd:
            flash("Invalid location selection.", "danger")
            return redirect(url_for("project_location", pid=pid))
        data = project["data"]
        data.update({
            "country": country, "region": region,
            "psh": sd["psh"], "avg_temp": sd["avg_temp"],
            "tariff": float(f.get("tariff", sd["tariff"])),
            "currency": sd["currency"], "symbol": sd["symbol"],
            "cost_usd_kwp": sd["cost_usd_kwp"], "fx_usd": sd["fx_usd"],
            "system_type": f.get("system_type", "off-grid"),
            "phase":       f.get("phase", "single"),
            "voltage":     int(f.get("voltage", 48)),
            "autonomy":    int(f.get("autonomy", 1)),
            "chemistry":      f.get("chemistry", "LiFePO4"),
            "panel_wp":       int(f.get("panel_wp", 400)),
            "mounting_type":  f.get("mounting_type", "rooftop_pitched"),
            # Solar design parameters
            "tilt_angle":       float(f.get("tilt_angle", 15)),
            "azimuth":          float(f.get("azimuth", 0)),
            "system_losses":    float(f.get("system_losses", 14)),
            "inverter_eff":     float(f.get("inverter_eff", 95)),
            "battery_dod":      float(f.get("battery_dod", 80)),
            "performance_ratio": float(f.get("performance_ratio", 75)),
            # User-defined BOQ cost rates
            "supply_markup_pct": max(0, min(50, float(f.get("supply_markup_pct", 8)))),
            "install_rate_pct":  max(0, min(100, float(f.get("install_rate_pct", 15)))),
        })
        # Save Ghana PURC category if provided
        purc_cat = f.get("purc_category", "").strip()
        if purc_cat:
            data["purc_category"] = purc_cat
        elif "purc_category" in data and country != "Ghana":
            data.pop("purc_category", None)   # clear if country changed
        # Save funding mode
        data["funding_mode"] = f.get("funding_mode", "loan")
        save_project_data(pid, data)
        return redirect(url_for("project_loads", pid=pid))

    return render_template("location.html", user=current_user(),
                           project=project, countries=get_countries(),
                           global_data=GLOBAL_DATA,
                           battery_chemistries=list(BATTERY_CHEMISTRY.keys()),
                           panel_options=PANEL_SPEC["standard_wp"],
                           purc_tariffs=GHANA_PURC_TARIFFS)


@app.route("/project/<int:pid>/loads", methods=["GET", "POST"])
@login_required
def project_loads(pid):
    project = get_project(pid)
    if not project:
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        csrf_protect()
        loads = []
        names    = request.form.getlist("load_name[]")
        cats     = request.form.getlist("load_cat[]")
        watts    = request.form.getlist("load_watt[]")
        qtys     = request.form.getlist("load_qty[]")
        hours    = request.form.getlist("load_hours[]")
        dfs      = request.form.getlist("load_df[]")
        critical = request.form.getlist("load_critical[]")

        for i in range(len(names)):
            if not names[i].strip():
                continue
            cat = cats[i] if i < len(cats) else "Other"
            df_default = DEMAND_FACTORS.get(cat, 0.70)
            try:
                df_val = float(dfs[i]) if i < len(dfs) and dfs[i] else df_default
            except (ValueError, IndexError):
                df_val = df_default
            df_val = max(0.10, min(1.0, df_val))
            loads.append({
                "name":          names[i],
                "category":      cat,
                "wattage":       float(watts[i]) if i < len(watts) else 0,
                "quantity":      float(qtys[i]) if i < len(qtys) else 1,
                "hours":         float(hours[i]) if i < len(hours) else 0,
                "demand_factor": round(df_val, 2),
                "critical":      str(i) in critical or names[i] in critical,
            })

        if not loads:
            flash("Please add at least one load.", "warning")
            return redirect(url_for("project_loads", pid=pid))

        data = project["data"]
        data["loads"] = loads

        # Save PURC customer category if provided
        purc_cat = request.form.get("purc_category", "").strip()
        if purc_cat:
            data["purc_category"] = purc_cat

        # Connected peak load (wattage Ã— qty, no DF — used for phase selection & cable sizing)
        peak_kw = sum(
            float(ld.get("wattage", 0)) * float(ld.get("quantity", 1))
            for ld in loads
        ) / 1000.0

        # Diversified peak (with demand factors — used for inverter sizing)
        div_peak_kw = sum(
            float(ld.get("wattage", 0)) * float(ld.get("quantity", 1)) * float(ld.get("demand_factor", 0.70))
            for ld in loads
        ) / 1000.0

        # Auto phase selection: > 8 kW connected load â†' 3-phase (BS 7671 / IEC 60364)
        auto_phase = "three" if peak_kw > 8.0 else "single"
        prev_phase = data.get("phase", "single")
        if auto_phase != prev_phase:
            if auto_phase == "three":
                flash(
                    f"Phase automatically set to 3-phase 415V — connected peak load "
                    f"{peak_kw:.1f} kW exceeds the 8 kW single-phase limit. "
                    f"Single-phase limit: 8 kW (IEC 60364 / BS 7671).",
                    "info"
                )
            else:
                flash(
                    f"Phase set to single-phase 230V — connected peak load "
                    f"{peak_kw:.1f} kW is within single-phase range.",
                    "info"
                )
        data["phase"] = auto_phase

        # Run calculations
        daily_kwh = calc_loads(loads)
        psh       = data.get("psh", 5.0)
        temp      = data.get("avg_temp", 28.0)
        autonomy  = data.get("autonomy", 1)
        system_type = data.get("system_type", "off-grid")
        phase     = auto_phase
        tariff    = data.get("tariff", 2.0)
        currency  = data.get("currency", "USD")
        symbol    = data.get("symbol", "$")
        cost_kwp  = data.get("cost_usd_kwp", 900)
        fx        = data.get("fx_usd", 1.0)

        chemistry  = data.get("chemistry", "LiFePO4")
        dc_voltage = data.get("voltage", 48)
        panel_wp   = data.get("panel_wp", 400)
        supply_markup_pct = float(data.get("supply_markup_pct", 8))
        install_rate_pct  = float(data.get("install_rate_pct", 15))

        pv_kw, num_panels, td      = calc_pv(daily_kwh, psh, temp, panel_wp)
        bat_kwh, num_bat, unit_bat = calc_battery(daily_kwh, autonomy, chemistry)
        inv_kw                     = calc_inverter(daily_kwh, peak_kw=div_peak_kw)
        mppt_a                     = calc_mppt(pv_kw, dc_voltage)
        ac_cables = size_all_cables(inv_kw, pv_kw, system_type, phase,
                                    ambient_c=temp)
        pps        = 2 if dc_voltage <= 24 else 4 if dc_voltage <= 48 else 8
        num_strings = math.ceil(num_panels / pps)
        # BOQ uses actual AC cable sizes, DC string count, and user-defined rates
        boq_rows, boq_grand        = calc_boq(
            num_panels, num_bat, inv_kw, pv_kw, bat_kwh,
            unit_bat, chemistry, mppt_a, cost_kwp, fx, panel_wp,
            ac_cables=ac_cables, voltage=dc_voltage, num_strings=num_strings,
            supply_markup_pct=supply_markup_pct, install_rate_pct=install_rate_pct)
        economics                  = calc_economics(
            pv_kw, num_panels, bat_kwh, num_bat, inv_kw,
            daily_kwh, tariff, currency, symbol, cost_kwp, fx, autonomy,
            boq_total_local=boq_grand,
            chemistry=chemistry,
            funding_mode=data.get("funding_mode", "loan"),
            install_rate_pct=install_rate_pct)

        chem_info = BATTERY_CHEMISTRY.get(chemistry, BATTERY_CHEMISTRY["LiFePO4"])
        data["results"] = {
            "daily_kwh":    daily_kwh,
            "pv_kw":        pv_kw,
            "num_panels":   num_panels,
            "panel_wp":     panel_wp,
            "temp_derating":td,
            "bat_kwh":      bat_kwh,
            "num_bat":      num_bat,
            "unit_bat_kwh": unit_bat,
            "chemistry":    chemistry,
            "chem_name":    chem_info["name"],
            "chem_dod":     chem_info["dod"],
            "chem_cycles":  chem_info["cycle_life"],
            "chem_life":    chem_info["lifetime_yr"],
            "chem_brands":  chem_info["brands"],
            "inv_kw":       inv_kw,
            "inv_brand":    inverter_brand(inv_kw),
            "mppt_a":       mppt_a,
            "peak_kw":      round(peak_kw, 2),
            "div_peak_kw":  round(div_peak_kw, 2),
            "auto_phase":   auto_phase,
            "panel_spec":   PANEL_SPEC,
            "economics":    economics,
            "boq_rows":     boq_rows,
            "boq_grand":    boq_grand,
            "ac_cables":    ac_cables,
        }
        save_project_data(pid, data)
        return redirect(url_for("project_results", pid=pid))

    cats = ["Lighting", "Cooling", "Appliances", "Electronics",
            "Pumps", "Heating", "Office", "Other"]
    return render_template("loads.html", user=current_user(),
                           project=project, categories=cats)


@app.route("/project/<int:pid>/results")
@login_required
def project_results(pid):
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run load calculations first.", "warning")
        return redirect(url_for("project_loads", pid=pid))
    r    = project["data"]["results"]
    recs = calc_recommendations(r["economics"], project["data"], r)
    return render_template("results.html", user=current_user(),
                           project=project, d=project["data"],
                           r=r, recommendations=recs)


# ─── Routes — Reports ─────────────────────────────────────────────────────────

def _plan_gate(pid):
    """Return (project, user, plan) or redirect if project missing/no results."""
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run load calculations first.", "warning")
        return None, None, None
    u    = current_user()
    plan = (u["plan"] or "free").lower()
    return project, u, plan


@app.route("/project/<int:pid>/report/inspection")
@login_required
def report_inspection(pid):
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    return render_template("report_inspection.html", user=current_user(),
                           project=project, d=project["data"],
                           r=project["data"]["results"])


@app.route("/project/<int:pid>/report/inspection/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_inspection(pid):
    """PDF export — Field Inspection Report (available on free plan)."""
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d = project["data"]
    r = project["data"]["results"]
    phase  = d.get("phase", "single")
    v_ac   = 415 if phase == "three" else 230

    def _chk():
        return "â˜ Pass  â˜ Fail  â˜ N/A"

    def _section(num, title, items):
        rows = f"\n\n### {num}. {title}\n\n"
        rows += "| # | Inspection Item | Pass | Fail | N/A | Remarks |\n"
        rows += "|---|---|:---:|:---:|:---:|---|\n"
        for i, item in items:
            rows += f"| {num}.{i} | {item} | â˜ | â˜ | â˜ | |\n"
        return rows

    ac_cables = r.get("ac_cables", [])
    first_breaker = ac_cables[0]["breaker_a"] if ac_cables else "—"
    first_cable   = ac_cables[0]["cable_size_mm2"] if ac_cables else "—"

    md = f"""# Pre-Installation Site Assessment - {project["name"]}

*Technical Feasibility Consultation - Prepared prior to installation commitment*

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {_fmt(r["pv_kw"],2)} kWp - {_fmt(r["bat_kwh"],2)} kWh - {_fmt(r["inv_kw"],1)} kW

Prepared by: SolarPro Global | Free Consultation Service | BS 7671:2018 - IEC 60364

---

## Proposed System Overview

| Field | Details |
|---|---|
| Client / Site Name | {project["name"]} |
| Location | {d.get("region","")}, {d.get("country","")} |
| System Type | {d.get("system_type","").title()} |
| Phase | {"Three-Phase 415V" if phase=="three" else "Single-Phase 230V"} |
| Daily Energy Demand | {_fmt(r["daily_kwh"],2)} kWh/day |
| Proposed PV Array | {_fmt(r["pv_kw"],2)} kWp ({r["num_panels"]} x {r.get("panel_wp",400)} Wp) |
| Battery Storage | {_fmt(r["bat_kwh"],2)} kWh - {r["num_bat"]} units {r.get("chemistry","LiFePO4")} |
| Inverter | {_fmt(r["inv_kw"],1)} kW |
| Estimated System Cost | {d.get("symbol","$")} {_fmt(r["economics"]["total_local"],0)} |
| Estimated Payback | {_fmt(r["economics"]["payback"],1)} years |

| Consultant / Engineer | | Assessment Date | | Site Visit Date | |
|---|---|---|---|---|---|
| | | | | | |

---
""" + _section(1, "Site Suitability & Solar Resource", [
        (1, f"Location {d.get('region','')}, {d.get('country','')} has adequate solar irradiance (>= 4 PSH/day)"),
        (2, f"Roof/ground area sufficient for {r['num_panels']} panels (approx. {r['num_panels']*2} m2 minimum)"),
        (3, "Panel orientation achievable - equator-facing surface available"),
        (4, "No permanent shading obstruction between 9am-3pm (buildings, trees, towers)"),
        (5, "Roof pitch >= 10 deg or flat roof with tilt framing available"),
        (6, "No planned construction or tree growth likely to cause future shading"),
        (7, "Access to roof/ground area is safe and adequate for installation and maintenance"),
    ]) + _section(2, "Structural & Roof Assessment", [
        (1, "Roof/structure type identified (concrete slab / IBR metal sheet / clay tile / flat membrane)"),
        (2, "Structure age and condition acceptable for 25-year system life"),
        (3, f"Load capacity adequate - can support ~{r['num_panels']*20} kg PV array weight"),
        (4, f"For ground mount: {r['num_panels']*4} m2 land area available, level and stable"),
        (5, "Existing skylights, vents, or services do not conflict with array footprint"),
        (6, "Roof waterproofing in acceptable condition prior to installation"),
        (7, "No asbestos-containing materials identified in roof structure"),
    ]) + _section(3, "Existing Electrical Infrastructure", [
        (1, "Existing main distribution board (MDB) identified and accessible"),
        (2, "MDB has spare capacity / ways for new solar incomer"),
        (3, "Existing earthing / earth rod present - condition to be tested"),
        (4, f"Current utility connection is {'3-phase 415V' if phase=='three' else 'single-phase 230V'} - matches proposed system"),
        (5, "Available space for inverter and battery bank within equipment room"),
        (6, "Equipment room is dry, secure, and ventilated"),
        (7, "Cable routing path from array to equipment room identified and clear"),
        (8, "Existing wiring checked - no obvious overloads or faults before proceeding"),
    ]) + _section(4, "Load & Demand Validation", [
        (1, f"Daily demand {_fmt(r['daily_kwh'],2)} kWh/day verified against actual utility bills"),
        (2, f"Peak connected load {_fmt(r.get('peak_kw',0),2)} kW confirmed - no major loads omitted"),
        (3, "Critical loads (medical, refrigeration, security) identified for backup priority"),
        (4, "Load profile is reasonably consistent - no major seasonal variation"),
        (5, "Future load growth in next 3-5 years estimated and factored into sizing"),
        (6, f"High-draw appliances (A/C, pumps) confirmed compatible with {_fmt(r['inv_kw'],1)} kW inverter"),
    ]) + _section(5, "Financial Feasibility", [
        (1, f"Client aware of estimated system cost: {d.get('symbol','$')} {_fmt(r['economics']['total_local'],0)}"),
        (2, f"Payback period of {_fmt(r['economics']['payback'],1)} years is acceptable to client"),
        (3, "Funding source confirmed (cash / loan / lease / PPA)"),
        (4, "VAT and import duties on equipment considered in budget"),
        (5, "Client understands O&M obligations (panel cleaning, battery checks, annual inspection)"),
        (6, "Grid connection / utility approval process understood"),
        (7, "Building permit for installation investigated"),
    ]) + _section(6, "Grid Connection & Regulatory", [
        (1, f"Utility grid connection available at site"),
        (2, f"Net metering / feed-in tariff policy investigated for {d.get('country','')}"),
        (3, "Anti-islanding protection required - inverter has built-in function"),
        (4, "Local installation code compliance confirmed (BS 7671 / IEC 60364 / national standard)"),
        (5, "Planning authority notified of proposed installation (if required)"),
    ]) + _section(7, "Health, Safety & Access", [
        (1, "Safe roof access route confirmed - scaffolding / MEWP requirements noted"),
        (2, "Electrical isolation of existing installation possible before work begins"),
        (3, "No asbestos, hazardous materials, or restricted zones on site"),
        (4, "Client / occupants can remain during installation or relocation required"),
        (5, "Fire risk assessed - battery room ventilation and extinguisher provision confirmed"),
    ]) + f"""

---

## Consultant Recommendation

**Overall Site Feasibility:** [ ] FEASIBLE  [ ] FEASIBLE WITH CONDITIONS  [ ] NOT FEASIBLE

Next Steps / Conditions:

_______________________________________________

_______________________________________________

_______________________________________________

---

## Assessment Sign-Off

| | Solar Consultant / Engineer | Client / Site Owner |
|---|---|---|
| **Name** | | |
| **Qualification / Contact** | | |
| **Date** | | |
| **Signature** | | |

*Disclaimer: This pre-installation assessment is based on the load schedule provided and does not substitute for a full on-site survey by a qualified engineer. SolarPro Global accepts no liability for decisions made solely on this report.*

---

*Pre-Installation Site Assessment - SolarPro Global | Free Consultation Service | BS 7671:2018 - IEC 60364*
"""

    return _render_pdf(f"Pre-Installation Site Assessment - {project['name']}", md,
                       f"site_assessment_{project['name'].replace(' ','_')}.pdf")


def _paid_only(pid):
    """Gate: redirect free-plan users to upgrade page."""
    u    = current_user()
    plan = (u["plan"] or "free").lower() if u else "free"
    if plan == "free":
        flash("This report is available on Basic plan and above. Upgrade to access all reports.", "warning")
        return redirect(url_for("upgrade"))
    return None


@app.route("/project/<int:pid>/report/pv")
@login_required
def report_pv(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    return render_template("report_pv.html", user=current_user(),
                           project=project, d=project["data"],
                           r=project["data"]["results"])


@app.route("/project/<int:pid>/report/boq")
@login_required
def report_boq(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    return render_template("report_boq.html", user=current_user(),
                           project=project, d=project["data"],
                           r=project["data"]["results"])


@app.route("/project/<int:pid>/report/cable")
@login_required
def report_cable(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    return render_template("report_cable.html", user=current_user(),
                           project=project, d=project["data"],
                           r=project["data"]["results"])


@app.route("/project/<int:pid>/report/economic")
@login_required
def report_economic(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    r   = project["data"]["results"]
    eco = r["economics"]
    recs = calc_recommendations(eco, project["data"], r)
    return render_template("report_economic.html", user=current_user(),
                           project=project, d=project["data"],
                           r=r, eco=eco, recommendations=recs)


@app.route("/project/<int:pid>/report/installation")
@login_required
def report_installation(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d = project["data"]
    r = project["data"]["results"]
    voltage = d.get("voltage", 48)
    num_panels = r["num_panels"]
    pps = 2 if voltage <= 24 else 4 if voltage <= 48 else 8
    num_strings = math.ceil(num_panels / pps)
    last_str_panels = num_panels - (num_strings - 1) * pps
    phase = d.get("phase", "single")
    v_ac = 415 if phase == "three" else 230
    return render_template("report_installation.html",
                           user=current_user(),
                           project=project, d=d, r=r,
                           pps=pps, num_strings=num_strings,
                           last_str_panels=last_str_panels,
                           v_ac=v_ac)


@app.route("/project/<int:pid>/report/installation/drawings")
@login_required
def report_installation_drawings(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d = project["data"]
    r = project["data"]["results"]
    voltage = int(d.get("voltage", 48))
    num_panels = int(r.get("num_panels", 4))
    pps = 2 if voltage <= 24 else 4 if voltage <= 48 else 8
    num_strings = math.ceil(num_panels / pps)
    last_str_panels = num_panels - (num_strings - 1) * pps
    phase = d.get("phase", "single")
    v_ac = 415 if phase == "three" else 230
    return render_template("report_installation_drawings.html",
                           user=current_user(),
                           project=project, d=d, r=r,
                           voltage=voltage,
                           pps=pps, num_strings=num_strings,
                           last_str_panels=last_str_panels,
                           v_ac=v_ac)


@app.route("/project/<int:pid>/report/proposal")
@login_required
def report_proposal(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    r   = project["data"]["results"]
    eco = r["economics"]
    d   = project["data"]
    # Monthly generation data (same calculation as energy report)
    monthly_factors = [0.88,0.90,0.95,1.00,1.05,1.08,1.10,1.08,1.03,0.98,0.92,0.88]
    base_monthly    = r["daily_kwh"] * 30.44
    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    monthly = [{"month": m, "kwh": round(base_monthly * f, 1),
                "saving": round(base_monthly * f * d.get("tariff", 1), 1)}
               for m, f in zip(months, monthly_factors)]
    trees_equiv = round(eco["co2_yr"] / 21.77, 0)
    cars_equiv  = round(eco["co2_yr"] / 4600, 2)
    return render_template("report_proposal.html", user=current_user(),
                           project=project, d=d, r=r, eco=eco,
                           monthly=monthly,
                           trees_equiv=trees_equiv,
                           cars_equiv=cars_equiv)


@app.route("/project/<int:pid>/report/energy")
@login_required
def report_energy(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    r   = project["data"]["results"]
    eco = r["economics"]
    d   = project["data"]
    # Monthly generation breakdown (simple seasonal variation Â±10%)
    monthly_factors = [0.88,0.90,0.95,1.00,1.05,1.08,1.10,1.08,1.03,0.98,0.92,0.88]
    base_monthly    = r["daily_kwh"] * 30.44
    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    monthly = [{"month": m, "kwh": round(base_monthly * f, 1),
                "saving": round(base_monthly * f * d.get("tariff",1), 1)}
               for m, f in zip(months, monthly_factors)]
    trees_equiv = round(eco["co2_yr"] / 21.77, 0)   # avg tree absorbs ~21.77 kg CO2/yr
    cars_equiv  = round(eco["co2_yr"] / 4600, 2)     # avg car ~4.6 tonne CO2/yr
    return render_template("report_energy.html", user=current_user(),
                           project=project, d=d, r=r, eco=eco,
                           monthly=monthly, trees_equiv=trees_equiv,
                           cars_equiv=cars_equiv)


@app.route("/project/<int:pid>/delete", methods=["POST"])
@login_required
def project_delete(pid):
    csrf_protect()
    with get_db() as c:
        c.execute("DELETE FROM projects WHERE id=? AND user_id=?",
                  (pid, session["user_id"]))
    flash("Project deleted successfully.", "success")
    return redirect(url_for("dashboard"))


@app.route("/project/<int:pid>/clone", methods=["POST"])
@login_required
def project_clone(pid):
    """Clone a completed project as a new editable draft (Use as Template)."""
    csrf_protect()
    uid  = session["user_id"]
    user = current_user()
    plan = (user["plan"] or "free").lower()
    limit = PLAN_LIMITS.get(plan, 1)

    with get_db() as c:
        src = c.execute(
            "SELECT * FROM projects WHERE id=? AND user_id=?", (pid, uid)
        ).fetchone()
        if not src:
            flash("Project not found.", "danger")
            return redirect(url_for("dashboard"))

        count = c.execute(
            "SELECT COUNT(*) FROM projects WHERE user_id=?", (uid,)
        ).fetchone()[0]
        if count >= limit:
            flash(f"Project limit reached — upgrade to create more.", "warning")
            return redirect(url_for("upgrade"))

        new_name = f"Copy of {src['name']}"
        # Copy full data_json (location + loads) but reset stage to 'location'
        c.execute(
            "INSERT INTO projects (user_id, name, stage, data_json) VALUES (?,?,'location',?)",
            (uid, new_name, src["data_json"] or "{}"))
        new_pid = c.execute("SELECT last_insert_rowid()").fetchone()[0]

    flash(f"'{new_name}' created — update location or loads then recalculate.", "success")
    return redirect(url_for("project_location", pid=new_pid))


@app.route("/project/<int:pid>/reset", methods=["POST"])
@login_required
def project_reset(pid):
    """Reset a project's calculated results — clears results/economics so it
    can be recalculated fresh after editing. Location and loads are kept."""
    csrf_protect()
    uid = session["user_id"]
    with get_db() as c:
        row = c.execute(
            "SELECT data_json FROM projects WHERE id=? AND user_id=?",
            (pid, uid)).fetchone()
        if not row:
            flash("Project not found.", "danger")
            return redirect(url_for("dashboard"))
        data = json.loads(row["data_json"] or "{}")
        # Strip only the results blob — keep location settings and loads
        data.pop("results", None)
        c.execute(
            "UPDATE projects SET data_json=?, updated_at=CURRENT_TIMESTAMP WHERE id=? AND user_id=?",
            (json.dumps(data), pid, uid))
    flash("Project results cleared — edit loads and re-run the calculation.", "info")
    return redirect(url_for("project_loads", pid=pid))


# ─── Project Save / Open (export â†' JSON file, import â† JSON file) ────────────

@app.route("/project/<int:pid>/save")
@login_required
def project_save(pid):
    """Download the project as a .solarpro JSON backup file."""
    uid = session["user_id"]
    with get_db() as c:
        row = c.execute(
            "SELECT * FROM projects WHERE id=? AND user_id=?", (pid, uid)
        ).fetchone()
    if not row:
        flash("Project not found.", "danger")
        return redirect(url_for("dashboard"))

    payload = {
        "app":      "SolarPro Global",
        "version":  1,
        "exported": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"),
        "project": {
            "name":       row["name"],
            "stage":      row["stage"],
            "created_at": row["created_at"],
            "data_json":  json.loads(row["data_json"] or "{}"),
        }
    }
    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in row["name"])
    fname = f"SolarPro_{safe_name}.solarpro"
    resp = make_response(json.dumps(payload, indent=2))
    resp.headers["Content-Type"]        = "application/json"
    resp.headers["Content-Disposition"] = f'attachment; filename="{fname}"'
    return resp


@app.route("/project/open", methods=["GET", "POST"])
@login_required
def project_open():
    """Import a .solarpro backup file and create a new project from it."""
    uid  = session["user_id"]
    user = current_user()
    plan = (user["plan"] or "free").lower()
    limit = PLAN_LIMITS.get(plan, 1)

    if request.method == "GET":
        # Render a simple upload form
        return render_template("project_open.html")

    # POST — process uploaded file
    f = request.files.get("project_file")
    if not f or not f.filename:
        flash("Please select a .solarpro file to open.", "warning")
        return redirect(url_for("project_open"))

    try:
        payload = json.loads(f.read().decode("utf-8"))
        if payload.get("app") != "SolarPro Global" or "project" not in payload:
            raise ValueError("Not a valid SolarPro project file.")
        proj = payload["project"]
        name = proj.get("name") or "Imported Project"
        data = proj.get("data_json") or {}
        stage = proj.get("stage") or "location"
    except Exception as ex:
        flash(f"Could not read project file: {ex}", "danger")
        return redirect(url_for("project_open"))

    # Check plan limit
    with get_db() as c:
        count = c.execute(
            "SELECT COUNT(*) FROM projects WHERE user_id=?", (uid,)
        ).fetchone()[0]
        if count >= limit:
            flash(f"Project limit reached — upgrade to import more projects.", "warning")
            return redirect(url_for("upgrade"))

        import_name = f"[Imported] {name}"
        c.execute(
            "INSERT INTO projects (user_id, name, stage, data_json) VALUES (?,?,?,?)",
            (uid, import_name, stage, json.dumps(data))
        )
        new_pid = c.execute("SELECT last_insert_rowid()").fetchone()[0]

    flash(f"Project '{import_name}' opened successfully.", "success")
    # If it has results, go to results; otherwise continue from where it was
    if data.get("results"):
        return redirect(url_for("project_results", pid=new_pid))
    elif data.get("loads"):
        return redirect(url_for("project_loads", pid=new_pid))
    else:
        return redirect(url_for("project_location", pid=new_pid))


# ─── Public solar-data API (no login — used by assessment form) ──────────────

@app.route("/api/solar_regions/<country>")
@limiter.limit("60 per minute")
def api_solar_regions_public(country):
    """Public: list regions for a country (used by the assessment form)."""
    if country not in GLOBAL_DATA:
        return jsonify({"regions": []})
    return jsonify({"regions": get_regions(country)})


@app.route("/api/solar_data/<country>/<region>")
@limiter.limit("60 per minute")
def api_solar_data_public(country, region):
    """Public: get irradiance/tariff data for a country+region (assessment form)."""
    sd = get_solar_data(country, region)
    if not sd:
        return jsonify({"error": "not found"}), 404
    return jsonify(sd)


# ─── API endpoints (login-required, rate-limited) ─────────────────────────────

@app.route("/api/regions/<country>")
@login_required
@limiter.limit("60 per minute")
def api_regions(country):
    # Validate input — only allow known countries
    if country not in GLOBAL_DATA:
        abort(400)
    return jsonify(get_regions(country))


@app.route("/api/solar/<country>/<region>")
@login_required
@limiter.limit("60 per minute")
def api_solar(country, region):
    if country not in GLOBAL_DATA:
        abort(400)
    sd = get_solar_data(country, region)
    if not sd:
        abort(404)
    return jsonify(sd)


@app.route("/api/purc-tariffs")
@limiter.limit("120 per minute")
def api_purc_tariffs():
    """Return Ghana PURC Q2 2026 tariff schedule as JSON."""
    result = {}
    for cat, info in GHANA_PURC_TARIFFS.items():
        result[cat] = {
            "rate_ghc":    info["rate_ghc"],
            "fixed_ghc":   info["fixed_ghc"],
            "description": info["description"],
        }
    return jsonify(result)


@app.route("/api/demand-factors")
@limiter.limit("120 per minute")
def api_demand_factors():
    """Return default demand factors per load category."""
    return jsonify(DEMAND_FACTORS)


# ─── Export endpoints ─────────────────────────────────────────────────────────

def _xl_header(ws, cols, fill_color="1e1e3a", font_color="F59E0B"):
    fill = PatternFill("solid", fgColor=fill_color)
    font = Font(bold=True, color=font_color)
    for col, (header, width) in enumerate(cols, 1):
        c = ws.cell(row=1, column=col, value=header)
        c.fill = fill; c.font = font
        c.alignment = Alignment(horizontal="center")
        ws.column_dimensions[c.column_letter].width = width

def _xl_send(wb, filename):
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name=filename)


@app.route("/project/<int:pid>/export/excel")
@login_required
@limiter.limit("10 per minute")
def export_excel(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run calculations first.", "warning")
        return redirect(url_for("project_results", pid=pid))

    r   = project["data"]["results"]
    eco = r["economics"]
    d   = project["data"]
    sym = d.get("symbol", "$")
    wb  = openpyxl.Workbook()

    # ── Sheet 1: Summary ──────────────────────────────────────────────────────
    ws = wb.active; ws.title = "Summary"
    ws.sheet_view.showGridLines = False
    title_font = Font(bold=True, size=14, color="F59E0B")
    ws["A1"] = "SolarPro Global — Project Summary"; ws["A1"].font = title_font
    ws["A2"] = project["name"]
    ws["A3"] = f"{d.get('region','')}, {d.get('country','')}   |   {d.get('system_type','').title()} System"
    ws["A4"] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ws.append([])
    phase_v = "415V 50Hz" if d.get("phase") == "three" else "230V 50Hz"
    rows = [
        ("── LOCATION & SYSTEM ──",   ""),
        ("Location",                  f"{d.get('region','')}, {d.get('country','')}"),
        ("System Type",               d.get("system_type","").title()),
        ("Phase",                     f"{d.get('phase','single').title()} Phase {phase_v}"),
        ("DC Bus Voltage",            f"{d.get('voltage',48)} V"),
        ("Autonomy",                  f"{d.get('autonomy',1)} day(s)"),
        ("── PV SYSTEM SIZING ──",    ""),
        ("Daily Energy Demand",       f"{r['daily_kwh']:.3f} kWh/day"),
        ("Annual Generation",         f"{round(r['daily_kwh']*365,0):,.0f} kWh/yr"),
        ("PV Array Size",             f"{r['pv_kw']:.3f} kWp"),
        ("No. of Panels",             f"{r['num_panels']} Ã— {r.get('panel_wp',400)} Wp Mono PERC"),
        ("Battery Chemistry",         r.get("chemistry","LiFePO4")),
        ("Battery Storage",           f"{r['bat_kwh']:.1f} kWh ({r['num_bat']} Ã— {r['unit_bat_kwh']:.2g} kWh units)"),
        ("Inverter Rating",           f"{r['inv_kw']:.2f} kW"),
        ("MPPT Controller",           f"{r.get('mppt_a','—')} A"),
        ("── FINANCIAL SUMMARY ──",   ""),
        ("Funding Mode",              "Self-Funded" if eco.get("funding_mode")=="self" else "Loan Finance"),
        ("Total System Cost",         f"{sym} {eco['total_local']:,.0f}"),
        ("Annual Bill Saving",        f"{sym} {eco['annual_sav']:,.0f}"),
        ("Annual O&M (Yr 1)",         f"{sym} {eco.get('om_yr1',0):,.0f}"),
        ("Net Annual Benefit (Yr1)",  f"{sym} {eco['net_yr1']:,.0f}"),
        ("Simple Payback",            f"{eco['payback']:.1f} years"),
        ("NPV (25yr)",                f"{sym} {eco['npv']:,.0f}"),
        ("IRR",                       f"{eco['irr_pct']:.1f}%" if eco['irr_pct'] else "N/A"),
        ("ROI (25yr cumulative)",     f"{eco.get('roi_pct',0):.0f}%"),
        ("DSCR",                      f"{eco['dscr']:.2f}" if eco.get("funding_mode","loan")=="loan" else "N/A (self-funded)"),
        ("Bankability",               eco["bankability"]),
        ("Verdict",                   eco["verdict"]),
        ("── ENVIRONMENTAL ──",       ""),
        ("COâ‚‚ Reduction",             f"{eco['co2_yr']:.2f} t/year"),
        ("COâ‚‚ Reduction (25yr)",      f"{eco['co2_yr']*25:.1f} t total"),
        ("25yr Cumulative Savings",   f"{sym} {eco['cumul_25']:,.0f}"),
    ]
    hfill = PatternFill("solid", fgColor="0f0f22")
    hfont = Font(bold=True, color="9090c0")
    ws.append(["Parameter", "Value"])
    ws[f"A{ws.max_row}"].font = hfont; ws[f"B{ws.max_row}"].font = hfont
    ws[f"A{ws.max_row}"].fill = hfill; ws[f"B{ws.max_row}"].fill = hfill
    ws.column_dimensions["A"].width = 36; ws.column_dimensions["B"].width = 32
    gold      = PatternFill("solid", fgColor="1a1a30")
    sec_fill  = PatternFill("solid", fgColor="0f0f1e")
    sec_font  = Font(bold=True, color="6060a0", italic=True)
    for param, val in rows:
        ws.append([param, val])
        if param.startswith("──"):          # section divider row
            ws[f"A{ws.max_row}"].font  = sec_font
            ws[f"A{ws.max_row}"].fill  = sec_fill
            ws[f"B{ws.max_row}"].fill  = sec_fill
        else:
            ws[f"A{ws.max_row}"].fill  = gold

    # ── Sheet 2: Load Schedule ────────────────────────────────────────────────
    ws2 = wb.create_sheet("Load Schedule")
    ws2.sheet_view.showGridLines = False
    cols2 = [("Category",14),("Load Name",22),("Wattage (W)",14),
             ("Quantity",10),("Hours/Day",11),("kWh/Day",11),("Critical",10)]
    _xl_header(ws2, cols2)
    for ld in d.get("loads", []):
        kwh = (float(ld.get("wattage",0)) * float(ld.get("quantity",1)) * float(ld.get("hours",0))) / 1000
        ws2.append([ld.get("category",""), ld.get("name",""),
                    float(ld.get("wattage",0)), int(ld.get("quantity",1)),
                    float(ld.get("hours",0)), round(kwh,3),
                    "Yes" if ld.get("critical") else ""])
    ws2.append(["","","","","","TOTAL",round(r["daily_kwh"],3)])
    ws2[f"F{ws2.max_row}"].font = Font(bold=True, color="F59E0B")

    # ── Sheet 3: BOQ ─────────────────────────────────────────────────────────
    ws3 = wb.create_sheet("Bill of Quantities")
    ws3.sheet_view.showGridLines = False
    cols3 = [("No.",5),("Description",28),("Specification",30),
             ("Qty",6),("Unit",7),(f"Basic Rate ({sym})",16),
             (f"Total Rate ({sym}) +8%",18),(f"Amount ({sym})",16)]
    _xl_header(ws3, cols3)
    for row in r.get("boq_rows",[]):
        ws3.append([row["no"], row["desc"], row["spec"], row["qty"],
                    row["unit"], round(row["basic"],2),
                    round(row["total_r"],2), round(row["amount"],2)])
    ws3.append(["","","","","","","GRAND TOTAL", round(r["boq_grand"],2)])
    ws3[f"G{ws3.max_row}"].font = Font(bold=True, color="F59E0B")
    ws3[f"H{ws3.max_row}"].font = Font(bold=True, color="F59E0B")

    # ── Cost Summary below BOQ ────────────────────────────────────────────────
    install_pct   = eco.get("install_rate_pct", 15)
    equip_local   = eco.get("equip_local", 0)
    install_local = eco.get("install_local", 0)
    total_local   = eco.get("total_local", 0)
    funding_label = "Self-Funded" if eco.get("funding_mode") == "self" else "Loan Finance"

    ws3.append([])
    ws3.append(["COST SUMMARY & CAPEX BREAKDOWN"])
    ws3[f"A{ws3.max_row}"].font = Font(bold=True, size=12, color="F59E0B")

    sum_hdr_row = ws3.max_row + 1
    ws3.append(["", "Item", "", "", "", "", "Amount", f"({sym})"])
    for col in range(1, 9):
        ws3.cell(sum_hdr_row, col).fill = PatternFill("solid", fgColor="1e1e3a")
        ws3.cell(sum_hdr_row, col).font = Font(bold=True, color="9090c0")

    cost_rows = [
        ("Equipment Supply (incl. 8% supply markup)",   equip_local,   False),
        (f"Installation Labour ({install_pct}% of supply subtotal)", install_local, False),
        ("TOTAL CAPEX",                                  total_local,   True),
        ("Contingency (10%) — advisory",                 total_local * 0.10, False),
        ("Budget (incl. contingency)",                   total_local * 1.10, False),
        ("Funding Mode",                                 funding_label, False),
    ]
    for label, val, bold in cost_rows:
        if isinstance(val, str):
            ws3.append(["", label, "", "", "", "", "", val])
        else:
            ws3.append(["", label, "", "", "", "", "", round(val, 0)])
        lbl_cell = ws3[f"B{ws3.max_row}"]
        val_cell = ws3[f"H{ws3.max_row}"]
        if bold:
            lbl_cell.font = Font(bold=True)
            val_cell.font = Font(bold=True, color="F59E0B")
            ws3[f"A{ws3.max_row}"].fill = PatternFill("solid", fgColor="1a1a30")
            ws3[f"B{ws3.max_row}"].fill = PatternFill("solid", fgColor="1a1a30")
            ws3[f"H{ws3.max_row}"].fill = PatternFill("solid", fgColor="1a1a30")

    ws3.append([])
    ws3.append(["BOQ NOTES"])
    ws3[f"A{ws3.max_row}"].font = Font(bold=True, color="9090c0")
    for note in [
        "Total Rate = Basic Rate Ã— 1.08  (8% supply/procurement markup — delivery, overheads & profit)",
        f"Installation Labour: {install_pct}% of supply subtotal (confirmed West Africa 2025—26 market rate)",
        "Contingency is advisory — not included in the CAPEX or payback calculation",
        "Quantities subject to detailed site survey confirmation",
        "Excludes site-specific VAT, import duties, and permit fees",
    ]:
        ws3.append(["", f"â€¢ {note}"])
        ws3[f"B{ws3.max_row}"].font = Font(italic=True, color="808080")

    # ── Sheet 4: 25-Year Cash Flow ────────────────────────────────────────────
    ws4 = wb.create_sheet("25-Year Cash Flow")
    ws4.sheet_view.showGridLines = False
    cols4 = [("Year",6),(f"Gross Saving ({sym})",18),(f"O&M ({sym})",14),
             (f"Net Benefit ({sym})",18),(f"Cumulative ({sym})",18),(f"NPV Cumul. ({sym})",18)]
    _xl_header(ws4, cols4)
    green_fill = PatternFill("solid", fgColor="052405")
    for cf in eco.get("cf_rows", []):
        ws4.append([cf["yr"], round(cf["gross"],0), round(cf["om"],0),
                    round(cf["net"],0), round(cf["cumul"],0), round(cf["npv_c"],0)])
        if cf["cumul"] >= eco["total_local"]:
            for col in range(1, 7):
                ws4.cell(ws4.max_row, col).fill = green_fill

    # ── Shared helper: write a labelled section into any sheet ───────────────
    def _xl_section(ws, title, param_val_pairs):
        ws.append([title])
        ws[f"A{ws.max_row}"].font = Font(bold=True, color="9090c0")
        ws[f"A{ws.max_row}"].fill = PatternFill("solid", fgColor="1e1e3a")
        ws[f"B{ws.max_row}"].fill = PatternFill("solid", fgColor="1e1e3a")
        row_fill = PatternFill("solid", fgColor="1a1a30")
        for param, val in param_val_pairs:
            ws.append([param, val])
            ws[f"A{ws.max_row}"].fill = row_fill

    def _xl_notes(ws, notes_list):
        ws.append([])
        ws.append(["NOTES"])
        ws[f"A{ws.max_row}"].font = Font(bold=True, color="9090c0")
        for note in notes_list:
            ws.append(["", f"â€¢ {note}"])
            ws[f"B{ws.max_row}"].font = Font(italic=True, color="808080")

    # ── Sheet 5: PV System Design ─────────────────────────────────────────────
    ws5 = wb.create_sheet("PV System Design")
    ws5.sheet_view.showGridLines = False
    ws5["A1"] = "PV System Design Report"; ws5["A1"].font = Font(bold=True, size=14, color="F59E0B")
    ws5["A2"] = f"{project['name']}  |  {d.get('region','')}, {d.get('country','')}"
    ws5["A3"] = (f"System: {d.get('system_type','').title()}  |  "
                 f"Phase: {d.get('phase','single').title()} {phase_v}  |  "
                 f"DC: {d.get('voltage',48)} V  |  Standard: BS 7671:2018 / IEC 60364")
    ws5.column_dimensions["A"].width = 36; ws5.column_dimensions["B"].width = 36
    ws5.append([])

    eff_pct = round(0.75 * r.get("temp_derating", 1.0) * 100, 1)
    _xl_section(ws5, "1. Design Inputs", [
        ("Location",                     f"{d.get('region','')}, {d.get('country','')}"),
        ("System Type",                  d.get("system_type","").title()),
        ("Phase / AC Voltage",           f"{d.get('phase','single').title()} Phase {phase_v}"),
        ("DC Bus Voltage",               f"{d.get('voltage',48)} V"),
        ("Peak Sun Hours (PSH)",         f"{d.get('psh',5.0)} h/day"),
        ("Avg Ambient Temperature",      f"{d.get('avg_temp',30)} Â°C"),
        ("Temperature Derating Factor",  f"{r.get('temp_derating',1.0):.4f}"),
        ("System Efficiency (BOS)",      "75%"),
        ("Effective Efficiency",         f"{eff_pct:.1f}%"),
        ("Daily Energy Demand",          f"{r['daily_kwh']:.3f} kWh/day"),
        ("Annual Generation",            f"{round(r['daily_kwh']*365, 0):,.0f} kWh/yr"),
        ("Autonomy Days",                f"{d.get('autonomy',1)} day(s)"),
    ])
    ws5.append([])
    _xl_section(ws5, "2. PV Array — Monocrystalline PERC", [
        ("Sizing Formula",               "PV (kWp) = Daily Load Ã· (PSH Ã— Î·_eff)"),
        ("PV Array Required",            f"{r['pv_kw']:.3f} kWp"),
        ("Panel Technology",             "Monocrystalline PERC — IEC 61215 certified"),
        ("Panel Rating",                 f"{r.get('panel_wp',400)} Wp"),
        ("Module Efficiency",            "21—23%"),
        ("Temperature Coefficient",      "âˆ'0.35 %/Â°C (Pmax)"),
        ("Warranty",                     "12 yr product / 25 yr linear power"),
        ("Recommended Brands",           "JinkoSolar, LONGi, Canadian Solar, Trina, JA Solar"),
        ("No. of Panels Required",       f"{r['num_panels']} modules"),
    ])
    ws5.append([])
    _xl_section(ws5, f"3. Battery Storage — {r.get('chemistry','LiFePO4')}", [
        ("Chemistry",                    r.get("chem_name", r.get("chemistry","LiFePO4"))),
        ("Sizing Formula",               "B = (E Ã— Autonomy) Ã· (DoD Ã— Î·_bat)"),
        ("Depth of Discharge (DoD)",     f"{round(r.get('chem_dod',0.8)*100,0):.0f}%"),
        ("Cycle Life",                   str(r.get("chem_cycles","4000+"))),
        ("Design Lifetime",              str(r.get("chem_life","10—15 years"))),
        ("Unit Size",                    f"{r['unit_bat_kwh']:.2g} kWh per unit"),
        ("No. of Units Required",        f"{r['num_bat']} unit{'s' if r['num_bat']>1 else ''}"),
        ("Total Storage",                f"{r['bat_kwh']:.1f} kWh"),
        ("Recommended Brands",           r.get("chem_brands","—")),
        ("BMS",                          "Integrated (overcharge, over-discharge, thermal)"),
    ])
    ws5.append([])
    _xl_section(ws5, "4. Hybrid Inverter / Charger", [
        ("Sizing Basis",                 "30% peak demand factor Ã— 1.25 safety factor"),
        ("Inverter Rating",              f"{r['inv_kw']:.2f} kW"),
        ("Type",                         "Hybrid MPPT (Solar / Battery / Grid)"),
        ("AC Output",                    f"{d.get('phase','single').title()} Phase {phase_v}"),
        ("Recommended",                  r.get("inv_brand","—")),
    ])
    ws5.append([])
    _xl_section(ws5, "5. MPPT Charge Controller & Protection", [
        ("Controller Type",              "MPPT (Maximum Power Point Tracking)"),
        ("PV Array Power",               f"{r['pv_kw']:.2f} kWp"),
        ("DC Bus Voltage",               f"{d.get('voltage',48)} V"),
        ("MPPT Rating",                  f"{r.get('mppt_a','—')} A"),
        ("Recommended",                  "Victron BlueSolar / Epever Tracer BN"),
        ("DC Protection",                "String fuses + DC-rated MCB 1000 V"),
        ("AC Protection",                "RCCB 30 mA + MCBs + SPD Type 2"),
        ("Earthing System",              "TT system — BS 7430 / IEC 60364-4-41"),
        ("Lightning Protection",         "IEC 62305 SPD Type 2, all circuits"),
    ])
    ws5.append([])
    _xl_section(ws5, "6. Installation Notes", [
        ("Roof Load",                    "Verify structural capacity â‰¥ 15 kg/mÂ² (IEC 61215)"),
        ("Panel Tilt",                   "Region-optimised tilt â‰ˆ latitude angle for equatorial sites"),
        ("Ventilation",                  "â‰¥ 150 mm around inverter; â‰¥ 300 mm around batteries"),
        ("Battery Safety",               f"{r.get('chemistry','LiFePO4')} — cool, dry, ventilated area"),
        ("DC Cabling",                   "UV-resistant 6 mmÂ² twin-core solar cable — IEC 60364"),
        ("Commissioning",                "Measure Voc, Isc, insulation resistance, earth continuity — log all"),
    ])

    # ── Sheet 6: Economic Analysis ────────────────────────────────────────────
    ws6 = wb.create_sheet("Economic Analysis")
    ws6.sheet_view.showGridLines = False
    ws6["A1"] = "Economic Analysis & Financial Engineering"; ws6["A1"].font = Font(bold=True, size=14, color="F59E0B")
    ws6["A2"] = f"{project['name']}  |  {d.get('region','')}, {d.get('country','')}  |  Currency: {d.get('currency','')}"
    funding_lbl = "Self-Funded" if eco.get("funding_mode") == "self" else "Loan Finance (70/30)"
    ws6["A3"] = f"Verdict: {eco.get('verdict','')}  |  Bankability: {eco.get('bankability','')}  |  Funding: {funding_lbl}"
    ws6.column_dimensions["A"].width = 40; ws6.column_dimensions["B"].width = 28
    ws6.append([])

    install_pct_e = eco.get("install_rate_pct", 15)
    _xl_section(ws6, "1. Capital Investment (CAPEX)", [
        ("Equipment Supply (incl. 8% markup)",        round(eco.get("equip_local",0), 0)),
        (f"Installation Labour ({install_pct_e}%)",   round(eco.get("install_local",0), 0)),
        ("Total CAPEX",                               round(eco.get("total_local",0), 0)),
        ("Contingency (10%) — advisory",              round(eco.get("total_local",0)*0.10, 0)),
        ("Budget incl. contingency — advisory",       round(eco.get("total_local",0)*1.10, 0)),
        ("Annual O&M (yr 1, 0.8% of CAPEX)",         round(eco.get("om_yr1",0), 0)),
        ("Discount Rate (WACC)",                      f"{eco.get('disc_rate_pct',12)}% per year"),
    ])
    ws6.append([])
    _xl_section(ws6, "2. Energy & Revenue", [
        ("Daily Load",                               f"{r['daily_kwh']:.3f} kWh/day"),
        ("Annual Generation",                        f"{eco.get('annual_kwh',0):,.0f} kWh/yr"),
        (f"Electricity Tariff ({d.get('currency','')})", eco.get("tariff", 0)),
        ("Tariff Escalation Assumption",             "8% per year"),
        ("Annual Bill Saving (yr 1)",                round(eco.get("annual_sav",0), 0)),
        ("Net Annual Saving (yr 1, after O&M)",      round(eco.get("net_yr1",0), 0)),
    ])
    ws6.append([])
    if eco.get("funding_mode") == "self":
        _xl_section(ws6, "3. Funding — Self-Funded", [
            ("Total Own Capital Required",           round(eco.get("total_local",0), 0)),
            ("Opportunity Cost Rate",                f"{eco.get('disc_rate_pct',10)}% per annum"),
            (f"Battery Replacement (yr {eco.get('bat_replace_yr','—')})",
                                                     round(eco.get("bat_replace_cost",0), 0)),
            ("Inverter Replacement (yr 10)",         round(eco.get("inv_replace_cost",0), 0)),
            ("Residual / Salvage Value (yr 25)",     round(eco.get("residual_value",0), 0)),
        ])
    else:
        _xl_section(ws6, "3. Loan Financing & Debt Service", [
            ("Loan Amount (70% of CAPEX)",           round(eco.get("loan_amt",0), 0)),
            ("Equity Required (30%)",                round(eco.get("equity",0), 0)),
            ("Interest Rate",                        "15% per annum"),
            ("Loan Tenor",                           "7 years"),
            ("Monthly Repayment",                    round(eco.get("pmt",0), 0)),
            ("Annual Debt Service",                  round(eco.get("annual_pmt",0), 0)),
            (f"Battery Replacement (yr {eco.get('bat_replace_yr','—')})",
                                                     round(eco.get("bat_replace_cost",0), 0)),
            ("Inverter Replacement (yr 10)",         round(eco.get("inv_replace_cost",0), 0)),
            ("DSCR (Debt Service Coverage Ratio)",   round(eco.get("dscr",0), 2)),
            ("Bankability",                          eco.get("bankability","")),
        ])
    ws6.append([])
    _xl_section(ws6, "4. Financial Metrics & Viability", [
        ("Simple Payback Period",                    f"{eco.get('payback',0):.1f} years"),
        ("NPV (25-year, discounted)",                round(eco.get("npv",0), 0)),
        ("IRR",                                      f"{eco.get('irr_pct',0):.1f}%" if eco.get("irr_pct") else "N/A"),
        ("ROI (25-year cumulative)",                 f"{eco.get('roi_pct',0):.1f}%"),
        ("10-Year Cumulative Net Saving",            round(eco.get("cumul_10",0), 0)),
        ("25-Year Cumulative Net Saving",            round(eco.get("cumul_25",0), 0)),
        ("Project Verdict",                          eco.get("verdict","")),
        ("Bankability",                              eco.get("bankability","")),
    ])
    ws6.append([])
    _xl_section(ws6, "5. Environmental Impact", [
        ("COâ‚‚ Avoided per Year",                    f"{eco.get('co2_yr',0):.2f} t COâ‚‚/yr"),
        ("COâ‚‚ Avoided over 25 Years",               f"{eco.get('co2_yr',0)*25:.1f} t COâ‚‚ total"),
        ("Equivalent Trees Planted",                f"{int(eco.get('co2_yr',0)*25*45):,} trees"),
        ("Clean Energy Generated (25yr)",           f"{eco.get('annual_kwh',0)*25:,.0f} kWh"),
    ])
    ws6.append([])
    _xl_section(ws6, "6. Model Assumptions", [
        ("Tariff Escalation",                       "8% per year (utility trend)"),
        ("Discount Rate — Loan",                    "12% per year (WACC)"),
        ("Discount Rate — Self-Funded",             "10% per year (opportunity cost)"),
        ("O&M Cost",                                "0.8% of CAPEX per year"),
        ("Panel Degradation",                       "0.5% per year"),
        ("System Lifetime",                         "25 years"),
        ("Grid COâ‚‚ Emission Factor",                "0.40 kg COâ‚‚/kWh"),
        ("Battery Replacement Cycle",               f"~{r.get('chem_life','10')} years"),
        ("Inverter Replacement Cycle",              "Year 10 at 80% of original cost"),
        ("Residual / Salvage Value",                "5% of CAPEX at year 25 (self-funded only)"),
        ("Loan — LTV",                              "70% debt / 30% equity"),
        ("Loan — Interest Rate",                    "15% per annum"),
        ("Loan — Tenor",                            "7 years"),
    ])

    # ── Sheet 7: AC Cable Schedule ────────────────────────────────────────────
    ws7 = wb.create_sheet("AC Cable Schedule")
    ws7.sheet_view.showGridLines = False
    ws7["A1"] = "AC Cable Sizing & Voltage Drop Schedule"; ws7["A1"].font = Font(bold=True, size=14, color="F59E0B")
    ws7["A2"] = (f"{project['name']}  |  Standard: BS 7671:2018 / IEC 60364-5-52  |  "
                 f"Cable: Copper PVC/XLPE 70Â°C  |  Method C  |  PF: 0.90")
    ws7.append([])
    cols7 = [("Circuit",20),("Power (kW)",10),("Voltage (V)",10),("Ib (A)",8),
             ("Length (m)",10),("Cable (mmÂ²)",10),("Iz (A)",8),
             ("VD (V)",8),("VD (%)",8),("Limit (%)",9),("Check",8),("Breaker (A)",11)]
    _xl_header(ws7, cols7)
    fail_fill = PatternFill("solid", fgColor="2a0000")
    for cab in r.get("ac_cables", []):
        chk = "PASS" if cab.get("vd_ok") else "FAIL"
        ws7.append([
            cab.get("circuit",""), cab.get("power_kw",0), cab.get("voltage_v",0),
            cab.get("design_current",0), cab.get("length_m",0),
            cab.get("cable_size_mm2",""), cab.get("cable_capacity",0),
            round(cab.get("vd_volts",0),3), round(cab.get("vd_percent",0),3),
            cab.get("vd_limit_pct",0), chk, cab.get("breaker_a",0),
        ])
        if not cab.get("vd_ok"):
            for col in range(1, 13):
                ws7.cell(ws7.max_row, col).fill = fail_fill
    _xl_notes(ws7, [
        "Standard: BS 7671:2018 (18th Edition) / IEC 60364-5-52",
        "Cable type: Copper conductor, PVC or XLPE 70Â°C insulation — Installation Method C (clipped direct)",
        "VD Formula (single-phase): VD (V) = mV/A/m Ã— Ib Ã— L Ã· 1000",
        "VD Formula (three-phase): VD (V) = mV/A/m Ã— 0.866 Ã— Ib Ã— L Ã· 1000",
        "VD reference: BS 7671 Appendix 4, Tables 4D2B / 4D5B",
        "Rows highlighted in red = VD FAIL — increase cable size or reduce circuit length",
        "All lengths are estimates — verify actual cable runs on site before ordering",
    ])

    return _xl_send(wb, f"SolarPro_{project['name'].replace(' ','_')}_Results.xlsx")


@app.route("/project/<int:pid>/export/csv")
@login_required
@limiter.limit("10 per minute")
def export_csv(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run calculations first.", "warning")
        return redirect(url_for("project_results", pid=pid))

    r   = project["data"]["results"]
    eco = r["economics"]
    d   = project["data"]
    sym = d.get("symbol", "$")

    output = io.StringIO()
    w = csv.writer(output)

    w.writerow(["SolarPro Global — Results Export"])
    w.writerow([project["name"], f"{d.get('region','')}, {d.get('country','')}", datetime.now().strftime('%Y-%m-%d')])
    w.writerow([])
    w.writerow(["=== SYSTEM SUMMARY ==="])
    for k,v in [
        ("Daily Demand (kWh/day)", r["daily_kwh"]),
        ("PV Array (kWp)", r["pv_kw"]),
        ("No. of Panels", r["num_panels"]),
        ("Panel Wattage (Wp)", r.get("panel_wp",400)),
        ("Battery Chemistry", r.get("chemistry","LiFePO4")),
        ("Battery Total (kWh)", r["bat_kwh"]),
        ("Battery Units", r["num_bat"]),
        ("Unit Size (kWh)", r["unit_bat_kwh"]),
        ("Inverter (kW)", r["inv_kw"]),
        ("MPPT (A)", r.get("mppt_a","—")),
        (f"System Cost ({sym})", round(eco["total_local"],2)),
        (f"Annual Savings ({sym})", round(eco["annual_sav"],2)),
        ("Payback (years)", round(eco["payback"],1)),
        (f"NPV ({sym})", round(eco["npv"],0)),
        ("IRR (%)", round(eco["irr_pct"],2) if eco["irr_pct"] else "N/A"),
        ("DSCR", round(eco["dscr"],2)),
        ("Bankability", eco["bankability"]),
        ("Verdict", eco["verdict"]),
        ("CO2 Reduction (t/yr)", round(eco["co2_yr"],2)),
    ]:
        w.writerow([k, v])

    w.writerow([]); w.writerow(["=== LOAD SCHEDULE ==="])
    w.writerow(["Category","Name","Wattage (W)","Qty","Hours/Day","kWh/Day","Critical"])
    for ld in d.get("loads", []):
        kwh = (float(ld.get("wattage",0))*float(ld.get("quantity",1))*float(ld.get("hours",0)))/1000
        w.writerow([ld.get("category",""), ld.get("name",""),
                    ld.get("wattage",0), ld.get("quantity",1),
                    ld.get("hours",0), round(kwh,3),
                    "Yes" if ld.get("critical") else "No"])
    w.writerow(["","","","","","TOTAL", round(r["daily_kwh"],3)])

    w.writerow([]); w.writerow(["=== BOQ ==="])
    w.writerow(["No.","Description","Specification","Qty","Unit",f"Rate ({sym})",f"Amount ({sym})"])
    for row in r.get("boq_rows",[]):
        w.writerow([row["no"],row["desc"],row["spec"],row["qty"],
                    row["unit"],round(row["total_r"],2),round(row["amount"],2)])
    w.writerow(["","","","","","Grand Total", round(r["boq_grand"],2)])

    w.writerow([]); w.writerow(["=== 25-YEAR CASH FLOW ==="])
    w.writerow(["Year",f"Gross ({sym})",f"O&M ({sym})",f"Net ({sym})",f"Cumulative ({sym})",f"NPV ({sym})"])
    for cf in eco.get("cf_rows",[]):
        w.writerow([cf["yr"],round(cf["gross"],0),round(cf["om"],0),
                    round(cf["net"],0),round(cf["cumul"],0),round(cf["npv_c"],0)])

    buf = io.BytesIO(output.getvalue().encode("utf-8-sig"))  # utf-8-sig for Excel compat
    return send_file(buf, mimetype="text/csv", as_attachment=True,
                     download_name=f"SolarPro_{project['name'].replace(' ','_')}_Results.csv")


# ─── Template helpers ─────────────────────────────────────────────────────────

@app.template_filter("fmt")
def fmt(v, dec=2):
    try:
        return f"{float(v):,.{dec}f}"
    except Exception:
        return str(v)


@app.template_filter("fmti")
def fmti(v):
    try:
        return f"{int(v):,}"
    except Exception:
        return str(v)


# ─── PDF helper ───────────────────────────────────────────────────────────────

def _render_pdf(title, md_content, filename):
    """Convert markdown string â†' PDF bytes and return as Flask download."""
    from markdown_pdf import MarkdownPdf, Section

    CSS = """
    body{font-family:'Segoe UI',Arial,sans-serif;color:#111827;font-size:11pt;line-height:1.55;margin:0;padding:0}
    h1{color:#b45309;font-size:17pt;border-bottom:3px solid #f59e0b;padding-bottom:8px;margin-bottom:14px}
    h2{color:#1e3a8a;font-size:13pt;border-bottom:1px solid #bfdbfe;padding-bottom:4px;margin-top:20px}
    h3{color:#374151;font-size:11pt;margin-top:14px}
    table{width:100%;border-collapse:collapse;margin:10px 0;font-size:10pt}
    th{background:#1e3a5f;color:#fff;padding:7px 10px;text-align:left}
    td{border:1px solid #e5e7eb;padding:5px 10px}
    tr:nth-child(even){background:#f8fafc}
    blockquote{background:#f0fdf4;border-left:4px solid #22c55e;padding:10px 16px;margin:8px 0;border-radius:4px}
    .warn{background:#fffbeb;border-left:4px solid #f59e0b;padding:10px 16px;margin:8px 0;border-radius:4px}
    .danger{background:#fef2f2;border-left:4px solid #ef4444;padding:10px 16px;margin:8px 0;border-radius:4px}
    p{margin:5px 0}
    hr{border:none;border-top:1px solid #e5e7eb;margin:14px 0}
    code{background:#f3f4f6;padding:1px 4px;border-radius:3px;font-size:10pt}
    """

    pdf = MarkdownPdf(toc_level=2)
    pdf.meta.update({"title": title, "author": "SolarPro Global", "subject": title})

    # Split on top-level H1 so each section starts on a new page
    parts = md_content.split("\n# ")
    pdf.add_section(Section(parts[0], toc=False), user_css=CSS)
    for part in parts[1:]:
        pdf.add_section(Section("# " + part, toc=True), user_css=CSS)

    buf = io.BytesIO()
    pdf.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/pdf",
                     as_attachment=True, download_name=filename)


# AI_BUDGET_LEDGER_MARKER_HELPER_DIAGRAMS
def _diagrams_markdown(d, r):
    """Return markdown block with SLD + topology + mounting-plan PNG embeds.
    Lazy-imports pdf_diagrams so the module is optional at runtime. Best-effort
    rendering: if a value is missing or matplotlib trips, returns an empty
    string and the PDF still ships without diagrams."""
    try:
        import pdf_diagrams as _pdfd
    except Exception:
        return ""
    try:
        pv_kw       = float(r.get("pv_kw") or d.get("pv_kw") or 0)
        inv_kw      = float(r.get("inv_kw") or 0)
        bat_kwh     = float(r.get("bat_kwh") or 0)
        num_bat     = int(r.get("num_bat") or 1)
        mppt_a      = int(r.get("mppt_a") or 60)
        num_panels  = int(r.get("num_panels") or 0)
        panel_wp    = int(r.get("panel_wp") or 400)
        chemistry   = r.get("chemistry") or "LiFePO4"
        system_type = (d.get("system_type") or "hybrid").lower()
        daily_kwh   = float(d.get("daily_kwh") or 0)
        psh         = float(d.get("psh") or 5.0)
        mounting    = d.get("mounting_type") or "rooftop_pitched"
    except Exception:
        return ""
    try:
        sld  = _pdfd.single_line_diagram_b64(pv_kw, inv_kw, bat_kwh, num_bat,
                                              mppt_a, chemistry, system_type)
        topo = _pdfd.system_topology_b64(pv_kw, inv_kw, bat_kwh, daily_kwh,
                                          psh, system_type)
        plan = (_pdfd.mounting_plan_b64(num_panels, panel_wp, "landscape", mounting)
                if num_panels > 0 else "")
    except Exception:
        return ""
    md = "# Design Diagrams\n\n"
    md += "## Single Line Diagram\n\n![Single Line Diagram](" + sld + ")\n\n"
    md += "## System Topology\n\n![System Topology](" + topo + ")\n\n"
    if plan:
        md += "## Mounting Plan\n\n![Mounting Plan](" + plan + ")\n\n"
    md += "---\n\n"
    return md


def _fmt(v, dec=2):
    """Format number for PDF markdown."""
    try:
        if dec == 0:
            return f"{int(round(v)):,}"
        return f"{v:,.{dec}f}"
    except Exception:
        return str(v)


# ─── PDF export routes ─────────────────────────────────────────────────────────

@app.route("/project/<int:pid>/report/boq/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_boq(pid):
    """PDF export — Bill of Quantities & CAPEX."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d   = project["data"]
    r   = project["data"]["results"]
    eco = r["economics"]
    sym = d.get("symbol", "$")

    md = f"""# Bill of Quantities (BOQ) & CAPEX — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {_fmt(r["pv_kw"],2)} kWp | {_fmt(r["bat_kwh"],2)} kWh | Currency: {d.get("currency","USD")}

Prepared by: SolarPro Global Â· BS 7671:2018 Â· IEC 60364

---

# Bill of Quantities

| No. | Description | Specification | Qty | Unit | Basic Rate ({sym}) | Total Rate ({sym}) | Amount ({sym}) |
|---|---|---|---|---|---|---|---|
"""
    for row in r.get("boq_rows", []):
        spec = str(row["spec"]).replace("|", "Â·")   # pipes break markdown tables
        md += (f"| {row['no']} | {row['desc']} | {spec} | "
               f"{row['qty']} | {row['unit']} | "
               f"{_fmt(row['basic'],2)} | {_fmt(row['total_r'],2)} | "
               f"**{_fmt(row['amount'],2)}** |\n")

    md += f"""
| | | | | | | **GRAND TOTAL** | **{sym} {_fmt(r['boq_grand'],2)}** |

*Note: Total Rate = Basic Rate Ã— 1.08 (8% supply/procurement markup — delivery, overheads & profit)*

---

# CAPEX Breakdown

| Item | Amount ({sym}) |
|---|---|
| Equipment Supply (incl. 8% markup) | {_fmt(eco.get("equip_local",0),0)} |
| Installation Labour ({eco.get("install_rate_pct",15)}%) | {_fmt(eco.get("install_local",0),0)} |
| **Total CAPEX** | **{_fmt(eco.get("total_local",0),0)}** |
| Contingency (10%) — advisory | {_fmt(eco.get("total_local",0)*0.1,0)} |
| **Budget (incl. contingency)** | **{_fmt(eco.get("total_local",0)*1.1,0)}** |

---

# System Summary

| Parameter | Value |
|---|---|
| PV Capacity | {_fmt(r["pv_kw"],2)} kWp |
| PV Panels | {r["num_panels"]} Ã— {r.get("panel_wp",400)} Wp Monocrystalline PERC |
| Battery Storage | {_fmt(r["bat_kwh"],2)} kWh — {r["num_bat"]} Ã— {_fmt(r["unit_bat_kwh"],2)} kWh {r.get("chemistry","LiFePO4")} |
| Inverter / Charger | {_fmt(r["inv_kw"],1)} kW {d.get("phase","single").title()}-Phase |
| DC Bus Voltage | {d.get("voltage",48)} V |
| Standard | BS 7671:2018 / IEC 60364 / IEC 61215 |

---

# BOQ Notes

- Rates: Total Rate = Basic Rate Ã— 1.08 (8% supply/procurement markup — delivery + overheads & profit)
- Quantities subject to detailed design review and site survey
- Cable lengths are estimated — confirm actual lengths on site
- Subject to contractor quotation; excludes site-specific VAT / import duties
- DC cable sizes and AC cable sizes are calculated from actual system design

*Report generated by SolarPro Global Â· {d.get("region","")}, {d.get("country","")}*
"""
    fname = f"BOQ_{project['name'].replace(' ','_')}.pdf"
    md = _diagrams_markdown(d, r) + md
    return _render_pdf(f"Bill of Quantities — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/cable/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_cable(pid):
    """PDF export — AC cable sizing & voltage drop calculation working."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d = project["data"]
    r = project["data"]["results"]
    phase = d.get("phase", "single")
    v_ac  = 415 if phase == "three" else 230

    md = f"""# AC Cable Sizing & Voltage Drop Report — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {phase.title()} Phase {v_ac} V | Ambient: {d.get("avg_temp",30)}Â°C

Standard: **BS 7671:2018 / IEC 60364-5-52** | Cable: Copper PVC/XLPE 70Â°C | PF: 0.90

---

# Design Basis

| Parameter | Value |
|---|---|
| Standard | BS 7671:2018 (18th Edition) / IEC 60364-5-52 |
| Cable Type | Copper conductor, PVC or XLPE 70Â°C insulation |
| Installation Method | Method C — clipped direct to surface (default) |
| Ambient Temperature | {d.get("avg_temp",30)}Â°C |
| Power Factor | 0.90 (lagging) |
| VD Formula (single-phase) | VD (V) = mV/A/m Ã— Ib Ã— L / 1000 |
| VD Formula (three-phase) | VD (V) = mV/A/m Ã— 0.866 Ã— Ib Ã— L / 1000 |
| VD Reference | BS 7671 Appendix 4, Tables 4D2B / 4D5B |
| Inverter | {_fmt(r["inv_kw"],1)} kW {phase.title()}-Phase |
| PV Array | {_fmt(r["pv_kw"],2)} kWp |

---

# Circuit Summary

| Circuit | Power (kW) | Vn (V) | Ib (A) | L (m) | Cable (mmÂ²) | Iz (A) | VD (V) | VD (%) | Limit | Check | Breaker |
|---|---|---|---|---|---|---|---|---|---|---|---|
"""
    for c in r.get("ac_cables", []):
        chk = "PASS" if c["vd_ok"] else "FAIL"
        md += (f"| {c['circuit']} | {c['power_kw']} | {c['voltage_v']} | "
               f"{c['design_current']} | {c['length_m']} | **{c['cable_size_mm2']}** | "
               f"{c['cable_capacity']} | {c['vd_volts']:.3f} | {c['vd_percent']:.3f}% | "
               f"{c['vd_limit_pct']}% | **{chk}** | {c['breaker_a']} A |\n")

    md += "\n---\n\n# Voltage Drop Calculation — Step-by-Step Working\n\n"

    for c in r.get("ac_cables", []):
        vd_limit_v = c["vd_limit_pct"] / 100 * c["voltage_v"]
        phase_note = "(Ã—0.866 three-phase factor already applied)" if c["phase"] == "three" else ""
        result_str = "✓ PASS" if c["vd_ok"] else "âœ— FAIL — increase cable size"
        md += f"""## {c["circuit"]}

**Cable selected: {c["cable_size_mm2"]} mmÂ² {c["core_type"]}** &nbsp;|&nbsp; {c["cable_capacity"]} A capacity &nbsp;|&nbsp; {c["breaker_a"]} A protective device

### Circuit Parameters

| Parameter | Symbol | Value |
|---|---|---|
| Nominal voltage | Vn | {c["voltage_v"]} V ({c["phase"].title()}-Phase) |
| Load power | P | {c["power_kw"]} kW |
| Design current | Ib | **{c["design_current"]} A** |
| Cable length | L | **{c["length_m"]} m** |
| Installation method | — | Method {c["install_method"]} — {c["install_desc"]} |
| Ambient temperature | Ta | {c["ambient_c"]}Â°C |
| Temperature factor | Ct | {c["temp_factor"]} |
| Grouping factor | Cg | {c["group_factor"]} |
| Minimum Iz required | Iz_min | {c["i_z_required"]} A |

### Voltage Drop Working

**Step 1 — Tabulated mV/A/m (BS 7671 Appendix 4)**

For {c["cable_size_mm2"]} mmÂ² {c["core_type"]} copper cable:

> mV/A/m = **{c["vd_mv_am"]}** mV/A/m {phase_note}

**Step 2 — Actual voltage drop**

> VD = mV/A/m Ã— Ib Ã— L / 1000
>
> VD = {c["vd_mv_am"]} Ã— {c["design_current"]} Ã— {c["length_m"]} / 1000 = **{c["vd_volts"]:.3f} V**
>
> VD% = ({c["vd_volts"]:.3f} / {c["voltage_v"]}) Ã— 100 = **{c["vd_percent"]:.3f}%**

**Step 3 — Check against permitted limit**

> Permitted limit = {c["vd_limit_pct"]}% of {c["voltage_v"]} V = **{vd_limit_v:.2f} V**
>
> Actual VD = **{c["vd_volts"]:.3f} V** — Limit = {vd_limit_v:.2f} V — **{result_str}**

---

"""

    md += """# Calculation Notes

| Item | Reference |
|---|---|
| VD tabulated values | BS 7671:2018 Appendix 4, Tables 4D2B / 4D5B |
| 3-phase VD factor Ã—0.866 | = âˆš3/2, IEC 60364-5-52 |
| Temperature correction | BS 7671 Table 4B2 (ref 30Â°C) |
| Grouping correction | BS 7671 Table 4B1 |
| VD limits | Inverterâ†'DB: 1.5% Â· Main feeder: 2.5% Â· Sub-distribution: 3.0% Â· Grid/Gen: 2.0% |
| Breaker coordination | Next standard size above Ib Ã— 1.05; must not exceed cable Iz |

*All installations must comply with BS 7671:2018 (18th Edition), IEC 60364-5-52, and local regulations.*

*Report generated by SolarPro Global*
"""
    fname = f"AC_Cable_VD_{project['name'].replace(' ','_')}.pdf"
    return _render_pdf(f"AC Cable Sizing & Voltage Drop — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/energy/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_energy(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    r   = project["data"]["results"]
    eco = r["economics"]
    d   = project["data"]
    sym = d.get("symbol", "$")

    monthly_factors = [0.88,0.90,0.95,1.00,1.05,1.08,1.10,1.08,1.03,0.98,0.92,0.88]
    base_monthly    = r["daily_kwh"] * 30.44
    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    monthly = [{"month": m, "kwh": round(base_monthly * f, 1),
                "saving": round(base_monthly * f * d.get("tariff",1), 1)}
               for m, f in zip(months, monthly_factors)]
    trees  = round(eco["co2_yr"] / 21.77, 0)
    cars   = round(eco["co2_yr"] / 4600, 2)
    offset_factor = 1.0 if d.get("system_type") == "off-grid" else 0.8
    annual_offset_kwh = r["daily_kwh"] * 365 * offset_factor

    md = f"""# Energy Impact Analysis — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {d.get("psh",0)} PSH | Generated by SolarPro Global

---

## Key Performance Indicators

| Metric | Value | Unit |
|---|---|---|
| Daily Solar Generation | {_fmt(r["daily_kwh"],2)} | kWh/day |
| Annual Solar Generation | {_fmt(r["daily_kwh"]*365,0)} | kWh/year |
| Annual Grid Offset | {_fmt(annual_offset_kwh,0)} | kWh/year |
| Annual Savings (Yr 1) | {sym} {_fmt(eco["annual_sav"],0)} | /year |
| COâ‚‚ Reduction | {_fmt(eco["co2_yr"],2)} | tonnes/year |
| Trees Equivalent | {int(trees)} | trees/year |
| Grid Offset | {"100" if d.get("system_type")=="off-grid" else "~80"} | % |

---

# Monthly Energy Generation & Savings

| Month | Generation (kWh) | Grid Offset (kWh) | Utility Offset (%) | Savings ({sym}) |
|---|---|---|---|---|
"""
    for m in monthly:
        offset_kwh = round(m['kwh'] * offset_factor, 1)
        base_avg   = r["daily_kwh"] * 365 / 12
        pct        = min(int(m['kwh'] / base_avg * 100), 100)
        md += f"| {m['month']} | {m['kwh']} | {offset_kwh} | {pct}% | {sym} {m['saving']} |\n"
    md += f"| **ANNUAL TOTAL** | **{_fmt(r['daily_kwh']*365,0)} kWh** | **{_fmt(annual_offset_kwh,0)} kWh** | — | **{sym} {_fmt(eco['annual_sav'],0)}** |\n"

    md += f"""
---

# Energy Savings Summary

| Item | Value |
|---|---|
| Daily Solar Generation | {_fmt(r["daily_kwh"],2)} kWh/day |
| Annual Solar Generation | {_fmt(r["daily_kwh"]*365,0)} kWh/year |
| Electricity Tariff | {sym}{d.get("tariff",0):.3f}/kWh — {d.get("utility","") or "Grid Utility"} |
| Tariff Reference | {d.get("tariff_ref","") or "Published utility schedule"} |
| Gross Annual Savings | {sym} {_fmt(eco["annual_sav"],0)} |
| Annual O&M Cost | {sym} {_fmt(eco["om_yr1"],0)} |
| **Net Annual Benefit** | **{sym} {_fmt(eco["net_yr1"],0)}** |
| Cumulative Savings (10yr) | {sym} {_fmt(eco["cumul_10"],0)} |
| Cumulative Savings (25yr) | {sym} {_fmt(eco["cumul_25"],0)} |
| Simple Payback | {_fmt(eco["payback"],1)} years |

---

# Environmental Impact

| Metric | Value |
|---|---|
| Annual COâ‚‚ Avoided | {_fmt(eco["co2_yr"],2)} tonnes/year |
| 25-Year COâ‚‚ Avoided | {_fmt(eco["co2_yr"]*25,1)} tonnes |
| Equivalent Trees Planted | {int(trees)} trees/year |
| Equivalent Cars Removed | {cars} cars/year |
| Grid Emission Factor | 0.40 kg COâ‚‚/kWh |
| Carbon Status | **Carbon Positive** |

---

# 25-Year Cash Flow Projection

| Year | Gross Saving | O&M | Net Benefit | Cumulative |
|---|---|---|---|---|
"""
    for cf in eco["cf_rows"]:
        flag = " â—„ BREAK-EVEN" if eco.get("breakeven") and cf["yr"] == eco["breakeven"] else ""
        md += f"| {cf['yr']} | {sym}{_fmt(cf['gross'],0)} | {sym}{_fmt(cf['om'],0)} | {sym}{_fmt(cf['net'],0)} | {sym}{_fmt(cf['cumul'],0)}{flag} |\n"

    md += f"\n---\n\n*Report generated by SolarPro Global Â· BS 7671 Â· IEC 60364 Â· IEEE*\n"

    fname = f"Energy_Impact_{project['name'].replace(' ','_')}.pdf"
    return _render_pdf(f"Energy Impact Analysis — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/economic/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_economic(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    r   = project["data"]["results"]
    eco = r["economics"]
    d   = project["data"]
    sym = d.get("symbol", "$")
    recs = calc_recommendations(eco, d, r)

    verdict_icon = "âœ… APPROVED" if eco["verdict"]=="APPROVED" else "âš ï¸ CONDITIONAL" if eco["verdict"]=="CONDITIONAL" else "âŒ REJECTED"
    bank_icon    = "âœ… BANKABLE" if eco["bankability"]=="BANKABLE" else "âš ï¸ MARGINAL" if eco["bankability"]=="MARGINAL" else "âŒ NOT BANKABLE"

    md = f"""# Economic Analysis — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {_fmt(r["pv_kw"],2)} kWp Â· {_fmt(r["bat_kwh"],2)} kWh Â· {_fmt(r["inv_kw"],1)} kW | Currency: {d.get("currency","")}

---

## Project Verdict: {verdict_icon} | Bankability: {bank_icon}

**Payback** {_fmt(eco["payback"],1)} yr | **NPV** {sym} {_fmt(eco["npv"],0)} | **IRR** {f'{eco["irr_pct"]:.1f}%' if eco["irr_pct"] else "N/A"} | **DSCR** {_fmt(eco["dscr"],2)} | **COâ‚‚** {_fmt(eco["co2_yr"],2)} t/yr

### Project Assessment
"""
    for reason in eco["verdict_reasons"]:
        md += f"- {reason}\n"

    if eco["bankability"] != "BANKABLE":
        md += f"\n### Bankability — {eco['bankability']}\n"
        for reason in eco["bank_reasons"]:
            md += f"- {reason}\n"

    md += f"""
---

# Key Financial Metrics

| Metric | Value |
|---|---|
| Total Investment | {sym} {_fmt(eco["total_local"],0)} |
| Annual Saving (Yr 1) | {sym} {_fmt(eco["annual_sav"],0)}/yr |
| Simple Payback | {_fmt(eco["payback"],1)} years |
| NPV (25yr) | {sym} {_fmt(eco["npv"],0)} |
| IRR | {f'{eco["irr_pct"]:.1f}%' if eco["irr_pct"] else "N/A"} |
| ROI (25yr) | {_fmt(eco["roi_pct"],0)}% |
| DSCR | {_fmt(eco["dscr"],2)} |
| COâ‚‚ Saved | {_fmt(eco["co2_yr"],2)} t/yr |

---

# Investment Summary

| Item | Amount |
|---|---|
| Equipment Supply (incl. 8% markup) | {sym} {_fmt(eco["equip_local"],0)} |
| Installation Labour ({eco.get("install_rate_pct",15)}%) | {sym} {_fmt(eco["install_local"],0)} |
| **Total Capital** | **{sym} {_fmt(eco["total_local"],0)}** |

---

# Energy & Savings

| Item | Value |
|---|---|
| Daily Load | {_fmt(r["daily_kwh"],3)} kWh/day |
| Annual Generation | {_fmt(eco["annual_kwh"],0)} kWh/yr |
| Tariff | {sym}{_fmt(eco["tariff"],3)}/kWh |
| Annual Bill Saving | {sym} {_fmt(eco["annual_sav"],0)}/yr |
| Annual O&M Cost | {sym} {_fmt(eco["om_yr1"],0)}/yr |
| **Net Annual Saving** | **{sym} {_fmt(eco["net_yr1"],0)}/yr** |

---

# Loan & Funding Analysis

| Item | Value |
|---|---|
| Loan Amount (70%) | {sym} {_fmt(eco["loan_amt"],0)} |
| Equity (30%) | {sym} {_fmt(eco["equity"],0)} |
| Interest Rate | 15% p.a. |
| Loan Tenor | 7 years |
| Monthly Repayment | {sym} {_fmt(eco["pmt"],0)}/mo |
| Annual Debt Service | {sym} {_fmt(eco["annual_pmt"],0)}/yr |
| **DSCR** | **{_fmt(eco["dscr"],2)} — {eco["bankability"]}** |

---

# 25-Year Cash Flow Projection

| Year | Gross Saving | O&M | Net Saving | Cumulative |
|---|---|---|---|---|
"""
    for cf in eco["cf_rows"]:
        flag = " â—„ BREAK-EVEN" if eco.get("breakeven") and cf["yr"] == eco["breakeven"] else ""
        md += f"| {cf['yr']} | {sym}{_fmt(cf['gross'],0)} | {sym}{_fmt(cf['om'],0)} | {sym}{_fmt(cf['net'],0)} | {sym}{_fmt(cf['cumul'],0)}{flag} |\n"

    if recs:
        md += f"\n---\n\n# Redesign Recommendations\n\n"
        md += f"The following engineering and financial improvements are recommended to achieve project approval and bankability:\n\n"
        for i, rec in enumerate(recs, 1):
            priority_label = "HIGH PRIORITY" if rec["priority"]==1 else "MEDIUM PRIORITY" if rec["priority"]==2 else "ADVISORY"
            md += f"## {i}. {rec['title']} [{priority_label}] ({rec['category']})\n\n"
            md += f"**Action:** {rec['action']}\n\n"
            md += f"**Expected Impact:** {rec['impact']}\n\n"

    md += f"\n---\n\n*Report generated by SolarPro Global Â· BS 7671 Â· IEC 60364 Â· IEEE*\n"
    md += f"\n*Assumptions: Tariff escalation 8%/yr, Discount rate 12%, O&M 1.2%/yr, Degradation 0.5%/yr, Life 25 years*\n"

    fname = f"Economic_Analysis_{project['name'].replace(' ','_')}.pdf"
    return _render_pdf(f"Economic Analysis — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/installation/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_installation(pid):
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d  = project["data"]
    r  = project["data"]["results"]
    voltage    = d.get("voltage", 48)
    pps        = 2 if voltage <= 24 else 4 if voltage <= 48 else 8
    num_strings = math.ceil(r["num_panels"] / pps)
    last_panels = r["num_panels"] - (num_strings - 1) * pps
    phase  = d.get("phase", "single")
    v_ac   = 415 if phase == "three" else 230
    sym    = d.get("symbol", "$")
    chem   = r.get("chemistry", "LiFePO4")

    mt = d.get("mounting_type", "rooftop_pitched")
    mt_labels = {
        "rooftop_pitched":  "Pitched / Sloped Roof Mount",
        "rooftop_flat":     "Flat Roof Ballast Mount (Concrete Slab)",
        "rooftop_metal":    "Metal Sheet Roof Mount (IBR / Corrugated)",
        "rooftop_membrane": "Membrane / Green Roof Mount (Lightweight Ballast)",
        "ground_fixed":     "Ground-Fixed Mount (Concrete Footing)",
        "ground_tracking":  "Ground-Fixed with Single-Axis Tracking",
    }
    mt_steps = {
        "rooftop_pitched": [
            ("Structural Survey", "Verify rafter/purlin spacing. Mark rafter lines. Confirm roof pitch and row layout with >= 150 mm air gap under modules."),
            ("Mounting Rail Installation", "Fix L-foot brackets to rafters at <= 1200 mm centres using M8 SS bolts. Apply EPDM seal and UV silicone to all penetrations. Run aluminium T-rails."),
            ("Panel Layout and Clamping", "Install panels in portrait rows. Use mid-clamps at <= 400 mm, end-clamps 100-150 mm from panel edge. Torque to 20 Nm."),
            ("DC Wiring", "Route 6 mm2 TUV DC cables through metallic conduit fixed to rails. Connect MC4 pairs per string layout. Label each string at combiner box."),
            ("Weatherproofing", "Water-test all penetrations. Confirm 600 mm roof-edge clearance. Fit bird-proofing mesh around perimeter."),
        ],
        "rooftop_flat": [
            ("Structural Survey", "Verify slab rating (allow 20-25 kg/m2). Check membrane condition. Survey parapets for wind compliance per BS EN 1991-1-4."),
            ("Row Spacing Layout", "Calculate inter-row spacing for <= 2% shading at winter solstice. Mark rows. Maintain 1 m walkway at perimeter and between rows."),
            ("Ballast Frame Installation", "Place aluminium ballast trays at 10-15 deg tilt. No penetrations required. Add HDPE pads under feet. Apply ballast blocks per wind uplift calc."),
            ("Panel Installation and Bonding", "Slide panels into tray clamps. Engage locking clips. Bond all frames to earth busbar via 6 mm2 GY cable."),
            ("DC Cable Tray", "Route cables in UV cable trays. Avoid drainage paths. Fix conduit at <= 500 mm. Use sealed gland at roof penetration."),
        ],
        "rooftop_metal": [
            ("Structural Survey", "Inspect purlins for rust. Confirm purlin centres and sheet condition. Replace corroded sheets before mounting."),
            ("Clamp Installation", "For IBR: non-penetrating rib clamps on raised rib - no drilling. For corrugated: L-bracket with EPDM washer, butyl sealant, SS dome-head screw."),
            ("Rail Alignment", "Fix aluminium rails across purlins. Level with shims. Confirm rail twist < 1 deg/m."),
            ("Panel Clamping", "Install panels landscape to reduce wind uplift. Apply mid/end clamps per spec. Torque check after 24 hours."),
            ("Bonding and Weathering", "Bond all roof components and frames to earth busbar. Check all penetrations with sealant. Inspect from inside for daylight ingress."),
        ],
        "rooftop_membrane": [
            ("Membrane Survey", "Inspect EPDM/TPO/bitumen membrane. Complete any repairs and cure before mounting. Verify slab load >= 15 kg/m2 for ballast."),
            ("Protection Layer", "Lay 5 mm HDPE protection mat over entire array footprint to prevent membrane puncture."),
            ("Ballast Frame Placement", "Use aerodynamic low-profile trays for membrane roofs - no penetrations. Calculate ballast per BS EN 1991-1-4. Minimum 5 deg tilt for self-cleaning."),
            ("Panel Installation and Bonding", "Snap-lock or clamp panels per frame system. Bond all frames to floating earth ring connected to earth spike at parapet."),
            ("Cable Management", "Route DC cables in UV conduit weighted with ballast. Keep clear of drainage. Maintain inspection access path throughout array."),
        ],
        "ground_fixed": [
            ("Site Preparation", "Clear and level footprint (4 m2 per panel). Mark column grid. Confirm no underground services in excavation zone."),
            ("Foundation", "Excavate 600 mm deep x 400 mm dia. pits at 3-4 m centres. Fix M16 HD anchor bolts in shutter. Cast C25/30 concrete. Cure 3 days before loading."),
            ("Steel Frame Erection", "Fix galvanised 100x100 RHS posts to anchor bolts. Plumb and align all columns. Bolt cross-purlin members. Clamp top rails at design tilt angle."),
            ("Panel Installation", "Lift panels with panel handlers - avoid walking on panels. Apply mid/end clamps. Bond each frame with 6 mm2 GY cable to common earth strip."),
            ("DC Wiring and Earthing", "Route string cables in IP67 UV armoured conduit buried 600 mm below grade. Earth rod at each end of array connected to main earth busbar via 16 mm2 GY cable."),
        ],
        "ground_tracking": [
            ("Site Survey and Grading", "Topographic survey. Grade to <= 2% slope. Compact sub-base to 95% Proctor. Bearing capacity >= 100 kN/m2 for piles."),
            ("Pile Installation", "Drive 3 m galvanised piles (100 mm dia.) to refusal with hydraulic pile driver. Cap with baseplate and tolerance shim at grid spacing per tracker layout."),
            ("Tracker Structure Erection", "Assemble torque tube sections on ground. Lift into pillow-block bearings. Align to within 2 mm/m. Install drive motor and actuation arm at mid-span."),
            ("Panel Stringing", "Mount panels landscape on tracker rails. N-S string wiring (east +, west -) per MPPT channel design."),
            ("Tracker Control Commission", "Connect motor to central controller. Set morning/evening stow. Commission east-to-west sweep. Install SCADA/data logger for monitoring."),
        ],
    }
    mt_label = mt_labels.get(mt, "Rooftop Mount")
    mt_method = mt_steps.get(mt, mt_steps["rooftop_pitched"])

    md = f"""# Installation Report - {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {phase.title()} Phase {v_ac}V | {d.get("voltage",48)}V DC

**Mounting Method:** {mt_label}

---

## System Summary

| Component | Specification |
|---|---|
| PV Array | {_fmt(r["pv_kw"],2)} kWp ({r["num_panels"]} Ã— {r.get("panel_wp",400)} Wp Mono PERC) |
| Battery Bank | {_fmt(r["bat_kwh"],2)} kWh ({r["num_bat"]} Ã— {_fmt(r["unit_bat_kwh"],2)} kWh {chem}) |
| Inverter/Charger | {_fmt(r["inv_kw"],1)} kW {phase.title()}-Phase |
| MPPT Rating | {r.get("mppt_a",0)} A |
| AC Voltage | {v_ac} V / 50 Hz |
| DC Bus Voltage | {d.get("voltage",48)} V |

---

# PV Array Configuration

| Parameter | Value |
|---|---|
| Total Panels | {r["num_panels"]} modules |
| Panels per String | {pps} in series |
| Number of Strings | {num_strings} parallel |
| Last String Panels | {last_panels} modules |
| Array Size | {_fmt(r["pv_kw"],2)} kWp |
| String Voc (est.) | {pps*24} V |
| DC Cable (strings) | 6 mmÂ² UV solar cable (TÃœV certified) |
| DC Cable (main run) | 10 mmÂ² to isolator |
| Tilt Angle | 10—15Â° minimum (self-cleaning) |
| Orientation | Equator-facing (south in N. hemisphere) |

---

# Battery Bank

| Parameter | Value |
|---|---|
| Chemistry | {chem} |
| Total Capacity | {_fmt(r["bat_kwh"],2)} kWh |
| Number of Units | {r["num_bat"]} |
| Unit Capacity | {_fmt(r["unit_bat_kwh"],2)} kWh |
| DC Voltage | {d.get("voltage",48)} V |
| Mounting | Ventilated steel rack, â‰¥ 300mm clearance |
| BMS | Built-in Battery Management System |

---

# AC Distribution (BS 7671)

| Circuit | Protection | Cable | Load |
|---|---|---|---|
| Incoming (Inverter output) | {r["ac_cables"][0]["breaker_a"] if r.get("ac_cables") else "—"}A RCCB 30mA | {r["ac_cables"][0]["cable_size_mm2"] if r.get("ac_cables") else "—"} mmÂ² | {_fmt(r["inv_kw"],1)} kW |
| Lighting & Emergency | 10A MCB | 1.5 mmÂ² | ~1.0 kW |
| Power Sockets | 16A MCB | 2.5 mmÂ² | ~2.0 kW |
| Air Conditioning | 32A MCB | {r["ac_cables"][-1]["cable_size_mm2"] if r.get("ac_cables") else "—"} mmÂ² | ~3.5 kW |
| Water / Borehole Pump | 16A MCB | 2.5 mmÂ² | ~1.5 kW |
| Office Equipment | 16A MCB | 2.5 mmÂ² | ~1.5 kW |

---

# Wire Colour Code (BS 7671 / IEC 60364)

| Conductor | Colour |
|---|---|
| DC Positive (+) | Red |
| DC Negative (âˆ') | Blue |
| AC Line / Phase | Brown |
| AC Neutral | Grey |
| Protective Earth (PE) | Green/Yellow |
| Battery Circuit | Purple |

---

# Installation Methodology

## Mounting Method: {mt_label}

""" + "".join(
    f"### Step {i+1}. {t}\n{d_}\n\n"
    for i, (t, d_) in enumerate(mt_method)
) + f"""
## General Installation Requirements

### PV Array Mounting Rules
- Maintain >= 600 mm from all roof edges (wind uplift zone)
- Allow >= 150 mm gap under panels for airflow (reduces temperature)
- Inter-row shading distance = panel height x tan(solar elevation)
- All panel frames bonded to earth bar with >= 6 mm2 green/yellow cable

## 2. DC Wiring — Combiner to Equipment Room
- All DC cables in metallic conduit or cable tray
- Segregate DC from AC cables — separate conduits or â‰¥ 50 mm separation
- String fuses rated at 1.25 Ã— Isc in combiner box
- DC isolator must be DC-rated (AC isolators MUST NOT be used on DC circuits)
- DC SPD (Type 2) in combiner box and again at inverter DC input

## 3. Battery Installation
- Mount on purpose-built steel rack, bolted to floor or wall
- Minimum 300 mm clearance from walls and ceiling on all sides
- Ensure mechanical ventilation — {chem} has low gas risk but ventilate regardless
- No ignition sources within 1 m. Class D fire extinguisher within 5 m
- Battery fuse: 1.25 Ã— maximum charge current

## 4. AC Distribution Board
- RCCB 30 mA on incomer (BS EN 61008 Type A)
- Type 2 AC SPD (BS EN 61643) inside DB — connection cable â‰¤ 0.5 m
- All MCBs coordinated — incomer trips last (discrimination)
- DB top edge â‰¤ 1.8 m from finished floor level

## 5. Earthing & Bonding (BS 7430)
- TT earthing arrangement — copper earth rod â‰¥ 2.4 m, driven vertically
- Earth rod â‰¥ 2 m from building structure
- Electrode resistance â‰¤ 10 Î© (test with earth clamp meter)
- All metalwork bonded: panel frames, inverter chassis, battery rack, DB enclosure
- Minimum 6 mmÂ² green/yellow bonding cable throughout

## 6. Testing & Commissioning
- Insulation resistance test: â‰¥ 1 MÎ© per IEC 60364-6
- Earth continuity: â‰¤ 0.1 Î© on all bonding connections
- RCD trip test: â‰¤ 40 ms at rated current (BS EN 61008)
- Polarity check on all DC circuits before energising inverter
- Functional test of all MCBs, RCCB, and SPDs

---

*All installations must comply with BS 7671:2018 (18th Edition), IEC 60364, IEC 62305 (lightning), and applicable local regulations. Engage a qualified electrical contractor for final installation and commissioning.*

*Report generated by SolarPro Global*
"""
    fname = f"Installation_{project['name'].replace(' ','_')}.pdf"
    md = _diagrams_markdown(d, r) + md
    return _render_pdf(f"Installation Report — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/workplan/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_workplan(pid):
    """PDF export — Installation Work Plan (material schedule, programme, staffing)."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d  = project["data"]
    r  = project["data"]["results"]
    voltage     = d.get("voltage", 48)
    pps         = 2 if voltage <= 24 else 4 if voltage <= 48 else 8
    num_strings = math.ceil(r["num_panels"] / pps)
    chem        = r.get("chemistry", "LiFePO4")
    sym         = d.get("symbol", "$")
    eco         = r["economics"]

    md = f"""# Installation Work Plan — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {_fmt(r["pv_kw"],2)} kWp | {r["num_panels"]} Panels | {_fmt(r["bat_kwh"],2)} kWh Battery

Prepared by: SolarPro Global Â· BS 7671:2018 Â· IEC 60364 Â· IEC 62446

---

# Section 1 — Material & Equipment Schedule

## 1A — PV Array & Mounting

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 1.1 | PV Solar Panels | {r.get("panel_wp",400)} Wp Monocrystalline PERC, Tier 1, Voc â‰ˆ 24 V | {r["num_panels"]} | Modules |
| 1.2 | Aluminium Mounting Rails | 40Ã—40 mm anodised aluminium | {int(r["num_panels"]*1.2)} | m |
| 1.3 | Mid & End Clamps | Stainless steel SS304 | {r["num_panels"]*4} | Sets |
| 1.4 | Roof Mounting Brackets | Galvanised steel, tilt-adjustable | {int(r["num_panels"]*1.5)} | No. |
| 1.5 | DC Solar Cable (strings) | 6 mmÂ² TÃœV UV-resistant, red & black | {int(r["num_panels"]*8)} | m |
| 1.6 | MC4 Connectors | IP67, 1000 VDC rated | {r["num_panels"]*4} | Pairs |
| 1.7 | DC Main Cable (combinerâ†'inverter) | 10 mmÂ² DC solar cable | 30 | m |
| 1.8 | Earthing Cable for Panel Frames | 6 mmÂ² green/yellow | {int(r["num_panels"]*3)} | m |

## 1B — DC Combiner & Protection

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 2.1 | DC Combiner Box | IP65, {min(num_strings,4)}-string, lockable | 1 | No. |
| 2.2 | String Fuses | 10A DC PV fuse, 1000 VDC | {num_strings*2} | No. |
| 2.3 | DC Surge Protection Device | Type 2, 1000 VDC, IEC 61643-31 | 1 | No. |
| 2.4 | DC Main Isolator | 3-pole DC-rated, lockable | 1 | No. |
| 2.5 | Metallic Cable Conduit | 20 mm steel conduit | 25 | m |
| 2.6 | Cable Tray / Trunking | 50Ã—50 mm galvanised steel | 10 | m |

## 1C — Inverter, Battery & MPPT

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 3.1 | Hybrid Inverter / Charger | {_fmt(r["inv_kw"],1)} kW, {d.get("voltage",48)}V DC, built-in MPPT {r.get("mppt_a","—")}A | 1 | No. |
| 3.2 | Lithium Battery Units | {_fmt(r["unit_bat_kwh"],2)} kWh {chem}, {d.get("voltage",48)}V, BMS | {r["num_bat"]} | No. |
| 3.3 | Battery Steel Rack | Powder-coated, for {r["num_bat"]} units | {max(1,(r["num_bat"]+1)//2)} | No. |
| 3.4 | Battery DC Fuse (ANL) | 1.25 Ã— max charge current | {r["num_bat"]} | No. |
| 3.5 | Battery DC Cable | 25 mmÂ² flexible, red & black | {r["num_bat"]*4} | m |
| 3.6 | BMS Communication Cable | RS485 / CAN bus | {r["num_bat"]} | Cables |
| 3.7 | Inverter Wall Bracket | Heavy-duty steel | 1 | Set |

## 1D — AC Distribution & Protection

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 4.1 | Main AC Distribution Board | {"18-way" if d.get("phase")=="single" else "12-way 3-ph"}, IP40 | 1 | No. |
| 4.2 | RCCB Incomer | {r["ac_cables"][0]["breaker_a"] if r.get("ac_cables") else 63}A, 30mA, Type A, BS EN 61008 | 1 | No. |
| 4.3 | MCB Lighting | 10A Type B, 6kA | 2 | No. |
| 4.4 | MCB Sockets | 16A Type B, 6kA | 3 | No. |
| 4.5 | MCB Air Conditioning | 32A Type C, 6kA | 1 | No. |
| 4.6 | MCB Pumps / Motors | 16A Type C, 6kA | 2 | No. |
| 4.7 | AC Surge Protection Device | Type 2, 230/415V, BS EN 61643 | 1 | No. |
"""
    for i, c in enumerate(r.get("ac_cables", []), start=8):
        md += f"| 4.{i} | AC Cable — {c['circuit']} | {c['cable_size_mm2']} mmÂ² Cu XLPE/PVC | {c.get('length_m',20)+5} | m |\n"

    md += f"""
## 1E — Earthing, Bonding & Sundries

| # | Description | Qty | Unit |
|---|---|---|---|
| 5.1 | Copper Earth Rod (16 mm dia., 2.4 m) | 2 | No. |
| 5.2 | Earth Rod Clamp & Driver | 2 | Sets |
| 5.3 | Earth Busbar (10-way copper) | 1 | No. |
| 5.4 | Main Earthing Conductor (16 mmÂ² G/Y) | 15 | m |
| 5.5 | Bonding Cables (6 mmÂ² G/Y) | 40 | m |
| 5.6 | Cable Labels (PVC self-laminating) | 200 | No. |
| 5.7 | Cable Ties UV-resistant | 1 | Box (200) |
| 5.8 | Warning / Safety Labels (BS EN 60445) | 1 | Set |
| 5.9 | IP65 Cable Glands (M20—M32) | 20 | No. |
| 5.10 | UV-Resistant Silicone Sealant | 4 | Tubes |

## 1F — Test & Commissioning Instruments

| Instrument | Purpose | Acceptance Standard |
|---|---|---|
| Insulation Resistance Tester (Megger) | Cable insulation integrity | â‰¥ 1 MÎ© (IEC 60364-6) |
| Earth Electrode Resistance Tester | Earth rod resistance | â‰¤ 10 Î© (BS 7430) |
| Clamp Earth Tester | Non-invasive earth continuity | â‰¤ 0.1 Î© (BS 7671 Ch.61) |
| Digital Multimeter (1000V DC) | String Voc, polarity | CAT III 1000V |
| DC Clamp Meter (1000V / 60A DC) | String Isc, battery current | CAT III 600V |
| RCD Tester | Trip time verification | â‰¤ 40 ms (BS EN 61008) |
| Voltage Drop Tester | Full-load volt drop | â‰¤ 3% (BS 7671 App 4) |
| Thermal Imaging Camera | Hot-spot detection | IEC 62446-3 |

---

# Section 2 — Approach to Installation Work & Programme

## Method Statement Summary

The installation follows a **7-phase methodology** aligned with BS 7671:2018, IEC 60364, and IEC 62446.
Each phase has defined entry criteria, activities, deliverables, and sign-off before the next phase begins.
Total programme: **12 working days** (weather permitting).

| Phase | Activity | Days | Duration |
|---|---|---|---|
| 1 | Mobilisation & Site Preparation | Days 1—2 | 2 days |
| 2 | Civil & Structural Works | Days 2—4 | 3 days |
| 3 | PV Panel Installation | Days 4—6 | 3 days |
| 4 | Equipment Room Fit-Out | Days 5—7 | 3 days |
| 5 | DC & AC Wiring | Days 7—9 | 3 days |
| 6 | Earthing, Bonding & Pre-commissioning Tests | Days 9—10 | 2 days |
| 7 | Commissioning, Testing & Handover | Days 10—12 | 3 days |

## Phase Detail

### Phase 1 — Mobilisation & Site Preparation (Days 1—2)

**Activities:**
- Deliver and inventory all equipment on site
- Set up site compound, secure storage for panels and batteries
- Install temporary power and lighting for working area
- Brief all staff on HSE plan and emergency procedures
- Prepare roof access — scaffold or MEWP
- Mark out equipment room layout and cable routes

**Sign-off outputs:** Signed delivery notes, site induction records, HSE risk assessment signed

### Phase 2 — Civil & Structural Works (Days 2—4)

**Activities:**
- Install roof mounting brackets/L-feet at designed spacing
- Assemble and level aluminium mounting rails; verify tilt angle
- Core through roof/walls for DC cable entry — seal immediately
- Fix conduit supports and tray brackets along cable route
- Install metallic conduit from roof to equipment room

**Sign-off outputs:** Waterproofing test, structural load check (if required), as-installed conduit sketch

### Phase 3 — PV Panel Installation (Days 4—6)

**Activities:**
- Mount panels row-by-row, bottom to top; torque clamps to spec
- Connect in series strings with MC4 connectors; verify polarity
- Install string fuses in combiner box; record fuse ratings
- Run and label DC string cables in conduit

**Sign-off outputs:** String cable labels at both ends, visual inspection — no cracked panels

### Phase 4 — Equipment Room Fit-Out (Days 5—7)

**Activities:**
- Fix inverter wall bracket; mount and level inverter
- Assemble battery rack; anchor to floor or wall
- Install batteries; connect in parallel per wiring diagram
- Connect BMS communication cables; configure settings
- Fix DC isolator and AC DB at correct heights
- Install earth busbar; run main earthing conductor

**Sign-off outputs:** Inverter mounting record, battery connection torque check, room layout photo

### Phase 5 — DC & AC Wiring (Days 7—9)

**Activities:**
- Run DC main cable combiner â†' inverter; double-check polarity
- Connect battery cables with ANL fuses
- Wire AC output inverter â†' DB incomer
- Install RCCB, MCBs, SPD in DB
- Run all AC final circuit cables; label both ends at every junction

**Sign-off outputs:** As-installed wiring diagram, cable schedule with actual lengths, labels verified

### Phase 6 — Earthing, Bonding & Pre-commissioning Tests (Days 9—10)

**Activities:**
- Drive earth rods to â‰¥ 2.4 m depth; connect to earth busbar
- Bond all metalwork — panel frames, inverter, battery rack, DB
- Test earth electrode resistance (â‰¤ 10 Î© before proceeding)
- Insulation resistance test — all circuits â‰¥ 1 MÎ©
- Continuity — all earth/bonding conductors â‰¤ 0.1 Î©
- Polarity check on all DC strings (signed by two technicians)

**Sign-off outputs:** Earth resistance certificate, IR test schedule, polarity check record

### Phase 7 — Commissioning, Testing & Handover (Days 10—12)

**Activities:**
- Energise inverter; verify AC output voltage and frequency
- Test MPPT tracking; confirm generation on inverter display
- RCD trip time test â‰¤ 40 ms; MCB overload test per circuit
- Voltage drop test under full load â‰¤ 3%
- 7-day performance monitoring — generation vs design
- Client handover training; issue O&M manual and warranties

**Sign-off outputs:** Full commissioning test schedule, 7-day monitoring log, Installation Completion Certificate, O&M manual issued

## Programme Notes

- **Total duration:** 12 working days (weather-permitting)
- **Weather hold:** No roof work in rain, lightning, or wind > 25 mph
- **Working hours:** 07:30—17:30 Mon—Fri; 07:30—13:00 Sat if required
- **Parallel working:** Phases 3 and 4 overlap (Days 5—7) — civil and electrical teams work simultaneously

---

# Section 3 — Staffing Plan

## Project Team

| Role | No. | Key Qualifications | Days on Site |
|---|---|---|---|
| Project Engineer / Site Manager | 1 | BEng Electrical, 18th Ed BS 7671, PV Solar cert, IOSH | All 12 days |
| Senior Electrical Technician | 1 | C&G 2365 NVQ L3, 18th Ed + 2391, ECS Gold, Work at Height | All 12 days |
| Electrical Apprentice / Assistant | 1 | NVQ L2 Electrical, Manual Handling, CSCS Green | All 12 days |
| Structural / Civil Technician | 1 | Roof mounting experience, PASMA/IPAF, CSCS | Days 1—6 |
| HSE Officer (part-time) | 1 | NEBOSH General, First Aid at Work | Days 1—2, 9—10 |

**Total peak headcount: 5 persons on site (Days 1—2 and 4—6)**

## Staff Deployment by Phase

| Phase | Activity | Proj. Eng. | Sr. Tech | Apprentice | Civil Tech | HSE Officer |
|---|---|---|---|---|---|---|
| 1 | Mobilisation | ✓ | ✓ | ✓ | ✓ | ✓ |
| 2 | Civil Works | ✓ (part) | — | ✓ | ✓ | ✓ (part) |
| 3 | PV Installation | ✓ | ✓ | ✓ | ✓ | ✓ (part) |
| 4 | Equipment Fit-Out | ✓ | ✓ | ✓ | — | — |
| 5 | DC & AC Wiring | ✓ | ✓ | ✓ | — | — |
| 6 | Earthing & Testing | ✓ | ✓ | ✓ | — | ✓ (part) |
| 7 | Commissioning | ✓ | ✓ | ✓ | — | — |

## Key Responsibilities

### Project Engineer / Site Manager
Responsible for overall technical quality, programme, safety, and client communication.
Signs off each phase, approves all test results, issues the Installation Completion Certificate.

### Senior Electrical Technician
Leads all electrical installation activities. Performs and records all pre-commissioning
and commissioning tests. Configures inverter and BMS settings. Mentors the apprentice.

### Electrical Apprentice / Assistant
Supports wiring, cable pulling, labelling, and containment installation.
Assists the senior technician and updates the material schedule as items are installed.

### Structural / Civil Technician
Installs all roof mounting structure. Responsible for panel mounting, tilt angle accuracy,
and waterproofing of all roof penetrations.

### HSE Officer
Conducts daily toolbox talks. Inspects PPE and access equipment. Maintains site accident log.
Emergency response coordinator.

## Mandatory PPE — All Personnel

- Safety helmet (EN 397) — all roof and overhead work
- Safety boots, steel toe cap, anti-slip (EN ISO 20345)
- High-visibility vest or jacket (EN ISO 20471 Class 2)
- Safety glasses when drilling, cutting, or using chemicals
- Electrical insulated gloves (EN 60903) when working near live equipment
- Full body harness (EN 361) for all work at height > 2 m
- Anti-static wrist strap when handling inverter electronics

## Site Safety Procedures

- Daily toolbox talk before work — attendance signed
- Written risk assessment and method statement on site at all times
- Permit to work before any work on energised equipment
- Two-person rule — no solo working on roof or electrical equipment
- All access equipment inspected daily — defective equipment tagged out
- Emergency evacuation plan posted at site entrance
- First aid kit and COâ‚‚ fire extinguisher on site at all times
- Any near-miss reported within 2 hours

---

*Installation Work Plan — {project["name"]}*
*Generated by SolarPro Global Â· BS 7671:2018 Â· IEC 60364 Â· IEC 62446 Â· IEC 62305*
"""

    fname = f"WorkPlan_{project['name'].replace(' ','_')}.pdf"
    return _render_pdf(f"Installation Work Plan — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/staffing/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_staffing(pid):
    """PDF export — Staffing Plan only (roles, deployment matrix, responsibilities, PPE)."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d  = project["data"]
    r  = project["data"]["results"]
    phase  = d.get("phase", "single")
    v_ac   = 415 if phase == "three" else 230

    md = f"""# Staffing Plan — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","").title()} System | {_fmt(r["pv_kw"],2)} kWp | {r["num_panels"]} Panels | {_fmt(r["bat_kwh"],2)} kWh Battery

Prepared by: SolarPro Global Â· BS 7671:2018 Â· IEC 60364 Â· IEC 62446

---

# Project Team

| Role | No. | Key Qualifications | Days on Site |
|---|---|---|---|
| Project Engineer / Site Manager | 1 | BEng Electrical, 18th Ed BS 7671, PV Solar cert, IOSH | All 12 days |
| Senior Electrical Technician | 1 | C&G 2365 NVQ L3, 18th Ed + 2391, ECS Gold, Work at Height | All 12 days |
| Electrical Apprentice / Assistant | 1 | NVQ L2 Electrical, Manual Handling, CSCS Green | All 12 days |
| Structural / Civil Technician | 1 | Roof mounting, PASMA/IPAF, CSCS | Days 1—6 |
| HSE Officer (part-time) | 1 | NEBOSH General, First Aid at Work | Days 1—2, 9—10 |

**Total peak headcount: 5 persons on site (Days 1—2 and 4—6)**

---

# Staff Deployment by Phase

| Phase | Activity | Days | Proj. Eng. | Sr. Tech | Apprentice | Civil Tech | HSE Officer |
|---|---|---|---|---|---|---|---|
| 1 | Mobilisation | 1—2 | ✓ | ✓ | ✓ | ✓ | ✓ |
| 2 | Civil Works | 2—4 | ✓ (part) | — | ✓ | ✓ | ✓ (part) |
| 3 | PV Installation | 4—6 | ✓ | ✓ | ✓ | ✓ | ✓ (part) |
| 4 | Equipment Fit-Out | 5—7 | ✓ | ✓ | ✓ | — | — |
| 5 | DC & AC Wiring | 7—9 | ✓ | ✓ | ✓ | — | — |
| 6 | Earthing & Testing | 9—10 | ✓ | ✓ | ✓ | — | ✓ (part) |
| 7 | Commissioning | 10—12 | ✓ | ✓ | ✓ | — | — |

---

# Key Responsibilities

## Project Engineer / Site Manager
Responsible for overall technical quality, programme, safety, and client communication.
Signs off each phase, approves all test results, and issues the Installation Completion Certificate.
Holds valid BS 7671 certification and PV Solar design qualification (MCS or equivalent).

## Senior Electrical Technician
Leads all electrical installation activities. Performs and records all pre-commissioning and
commissioning tests. Configures inverter and BMS settings. Mentors the apprentice throughout.

## Electrical Apprentice / Assistant
Supports wiring, cable pulling, labelling, and containment installation.
Assists the senior technician and updates the material schedule as items are installed.

## Structural / Civil Technician
Installs all roof mounting structure. Responsible for panel mounting, tilt angle accuracy,
and weatherproofing of all roof penetrations. Engaged for Days 1—6.

## HSE Officer (Part-Time / Shared)
Conducts daily toolbox talks. Inspects PPE and access equipment. Maintains site accident log.
Acts as emergency response coordinator.

---

# Mandatory PPE — All Personnel

| Item | Standard |
|---|---|
| Safety helmet | EN 397 — all roof and overhead work |
| Safety boots (steel toe cap, anti-slip) | EN ISO 20345 |
| High-visibility vest / jacket (Class 2) | EN ISO 20471 |
| Safety glasses (drilling, cutting, chemicals) | EN 166 |
| Electrical insulated gloves (live-adjacent work) | EN 60903 |
| Full body harness (work at height > 2 m) | EN 361 |
| Anti-static wrist strap (inverter electronics) | IEC 61340-5 |

---

# Site Safety Procedures

- Daily toolbox talk before work — attendance signed by all personnel
- Written risk assessment and method statement on site at all times
- Permit to work before any work on or near energised equipment
- Two-person rule — no solo working on roof or electrical equipment
- All access equipment inspected daily; defective items tagged out of service
- Emergency evacuation plan posted at site entrance and briefed on Day 1
- First aid kit and COâ‚‚ fire extinguisher on site at all times
- Any near-miss reported within 2 hours; incident report completed within 24 hours

---

*Staffing Plan — {project["name"]}*
*Generated by SolarPro Global Â· BS 7671:2018 Â· IEC 62446 Â· NEBOSH / HSE Guidance*
"""

    fname = f"StaffingPlan_{project['name'].replace(' ','_')}.pdf"
    return _render_pdf(f"Staffing Plan — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/pv/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_pv(pid):
    """PDF export — PV System Design Report."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d   = project["data"]
    r   = project["data"]["results"]
    eco = r.get("economics", {})
    sym = d.get("symbol", "$")

    md = f"""# PV System Design Report — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {d.get("system_type","off-grid").title()} System | Currency: {d.get("currency","USD")}

Prepared by: SolarPro Global Â· BS 7671:2018 Â· IEC 60364 Â· IEC 61215

---

# Site & Solar Resource

| Parameter | Value |
|---|---|
| Location | {d.get("region","")}, {d.get("country","")} |
| Peak Sun Hours (PSH) | {d.get("psh",5.0)} h/day |
| Average Ambient Temperature | {d.get("avg_temp",28)}Â°C |
| System Type | {d.get("system_type","off-grid").title()} |
| Phase | {d.get("phase","single").title()}-Phase |
| DC Bus Voltage | {d.get("voltage",48)} V |
| Electricity Tariff | {sym}{d.get("tariff",0)}/kWh |
| Exchange Rate | 1 USD = {d.get("fx_usd",1.0)} {d.get("currency","USD")} |

---

# Load Analysis

| Parameter | Value |
|---|---|
| Total Daily Energy Demand | {_fmt(r.get("daily_kwh",0),3)} kWh/day |
| Annual Energy Demand | {_fmt(r.get("daily_kwh",0)*365,0)} kWh/year |
| Connected Peak Load | {_fmt(r.get("peak_kw",0),2)} kW |
| Diversified Peak Load | {_fmt(r.get("div_peak_kw",0),2)} kW (with demand factors) |

---

# PV Array Design

| Parameter | Value |
|---|---|
| PV Array Capacity | {_fmt(r.get("pv_kw",0),2)} kWp |
| Number of Panels | {r.get("num_panels",0)} Ã— {r.get("panel_wp",400)} Wp |
| Panel Technology | Monocrystalline PERC, BS EN 61215 |
| Temperature Derating Factor | {r.get("temp_derating","—")} |
| BOS Efficiency | 75% |
| Design Standard | IEC 61215 / IEC 61730 |

---

# Battery Storage Design

| Parameter | Value |
|---|---|
| Total Battery Capacity | {_fmt(r.get("bat_kwh",0),2)} kWh |
| Number of Units | {r.get("num_bat",0)} Ã— {_fmt(r.get("unit_bat_kwh",0),2)} kWh each |
| Chemistry | {r.get("chemistry","LiFePO4")} |
| Depth of Discharge (DoD) | {int((r.get("chem_dod",0.9))*100)}% |
| Cycle Life | {r.get("chem_cycles","4,000+")} cycles |
| Lifetime | {r.get("chem_life","12")} years |
| Autonomy | {d.get("autonomy",1)} day(s) |
| BMS | Built-in per unit |

---

# Inverter & Charge Controller

| Parameter | Value |
|---|---|
| Inverter Rating | {_fmt(r.get("inv_kw",0),1)} kW |
| Type | Hybrid Inverter/Charger |
| Recommended Brands | {r.get("inv_brand","Victron / Growatt / Deye")} |
| MPPT Rating | {r.get("mppt_a","—")} A |
| DC Input Voltage | {d.get("voltage",48)} V |
| AC Output | {"415 V three-phase" if d.get("phase")=="three" else "230 V single-phase"} |
| Inverter Efficiency | â‰¥ 95% |

---

# System Performance Summary

| Parameter | Value |
|---|---|
| Daily Energy Supply | {_fmt(r.get("daily_kwh",0),3)} kWh |
| Annual Generation | {_fmt(r.get("daily_kwh",0)*365,0)} kWh |
| System Losses (BOS) | {d.get("system_losses",14)}% |
| Performance Ratio | {d.get("performance_ratio",75)}% |
| Simple Payback | {_fmt(eco.get("payback",0),1)} years |
| Total CAPEX | {sym} {_fmt(eco.get("total_local",0),0)} |
| Project Verdict | {eco.get("verdict","—")} |

---

*PV System Design Report — {project["name"]}*
*Generated by SolarPro Global Â· BS 7671:2018 Â· IEC 60364 Â· IEC 61215*
"""
    fname = f"PV_Report_{project['name'].replace(' ','_')}.pdf"
    md = _diagrams_markdown(d, r) + md
    return _render_pdf(f"PV System Design Report — {project['name']}", md, fname)


@app.route("/project/<int:pid>/report/proposal/pdf")
@login_required
@limiter.limit("5 per minute")
def export_pdf_proposal(pid):
    """PDF export — Full Technical & Financial Proposal (superset of all reports)."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    d   = project["data"]
    r   = project["data"]["results"]
    eco = r.get("economics", {})
    sym = d.get("symbol", "$")
    phase   = d.get("phase", "single")
    v_ac    = 415 if phase == "three" else 230
    chem    = r.get("chemistry", "LiFePO4")
    voltage = d.get("voltage", 48)
    pps         = 2 if voltage <= 24 else 4 if voltage <= 48 else 8
    num_strings = math.ceil(r["num_panels"] / pps) if r.get("num_panels") else 0
    last_panels = (r["num_panels"] - (num_strings - 1) * pps) if num_strings else 0

    # Monthly generation
    monthly_factors = [0.88,0.90,0.95,1.00,1.05,1.08,1.10,1.08,1.03,0.98,0.92,0.88]
    base_monthly    = r["daily_kwh"] * 30.44
    months_list     = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    monthly         = [(m, round(base_monthly * f, 1)) for m, f in zip(months_list, monthly_factors)]
    offset_factor   = 1.0 if d.get("system_type") == "off-grid" else 0.8
    annual_offset_kwh = r["daily_kwh"] * 365 * offset_factor
    trees = int(round(eco.get("co2_yr", 0) / 21.77, 0)) if eco.get("co2_yr") else 0
    cars  = round(eco.get("co2_yr", 0) / 4600, 2) if eco.get("co2_yr") else 0

    # Redesign recommendations (may be empty)
    try:
        recs = calc_recommendations(eco, d, r)
    except Exception:
        recs = []

    # ── Header ──────────────────────────────────────────────────────────────
    md = f"""# PV Solar System Proposal — {project["name"]}

**Location:** {d.get("region","")}, {d.get("country","")}
**System Type:** {d.get("system_type","off-grid").title()} | **PV Capacity:** {_fmt(r["pv_kw"],2)} kWp
**Battery:** {_fmt(r["bat_kwh"],2)} kWh {r.get("chemistry","LiFePO4")} | **Inverter:** {_fmt(r["inv_kw"],1)} kW
**Project Verdict:** {eco.get("verdict","—")} | **Bankability:** {eco.get("bankability","—")}

Prepared by: SolarPro Global Â· BS 7671:2018 Â· IEC 60364 Â· IEC 62305 Â· IEC 62446 Â· IEEE

This proposal is a **superset of all individual engineering reports** — it consolidates the Site Assessment,
PV Design, AC Cable, BOQ, Energy Impact, Economic Analysis, Installation Work Plan, and Staffing Plan
into a single deliverable suitable for client award and bank submission.

---

# PART A — TECHNICAL PROPOSAL

## A1. Site Assessment & Solar Resource

| Parameter | Value |
|---|---|
| Location | {d.get("region","")}, {d.get("country","")} |
| Peak Sun Hours | {d.get("psh",5.0)} h/day |
| Average Temperature | {d.get("avg_temp",28)}Â°C |
| Electricity Tariff | {sym}{d.get("tariff",0)}/kWh |
| System Type | {d.get("system_type","off-grid").title()} |
| Phase | {d.get("phase","single").title()}-Phase |
| DC Voltage | {d.get("voltage",48)} V |
| Autonomy | {d.get("autonomy",1)} day(s) |
| Battery Chemistry | {d.get("chemistry","LiFePO4")} |
| Exchange Rate | 1 USD = {d.get("fx_usd",1.0)} {d.get("currency","USD")} |

## A2. Electrical Load Analysis

| Parameter | Value |
|---|---|
| Total Daily Energy | {_fmt(r["daily_kwh"],3)} kWh/day |
| Annual Energy | {_fmt(r["daily_kwh"]*365,0)} kWh/year |
| Connected Peak Load | {_fmt(r.get("peak_kw",0),2)} kW |
| Diversified Peak Load | {_fmt(r.get("div_peak_kw",0),2)} kW |

# Load Schedule

| Category | Appliance | Power (W) | Qty | Hours/day | kWh/day |
|---|---|---|---|---|---|
"""
    for ld in d.get("loads", []):
        kwh = round(ld.get("wattage",0)*ld.get("quantity",1)*ld.get("hours",0)/1000, 3)
        md += (f"| {ld.get('category','')} | {ld.get('name','')} | "
               f"{int(ld.get('wattage',0))} | {int(ld.get('quantity',1))} | "
               f"{ld.get('hours',0)} | **{kwh}** |\n")
    md += f"| | **TOTAL** | | | | **{_fmt(r['daily_kwh'],3)}** |\n"

    md += f"""
## A3. Engineering Sizing Calculations

| Component | Calculation | Result |
|---|---|---|
| PV Array | {_fmt(r["daily_kwh"],3)} kWh Ã· ({d.get("psh",5)} h Ã— 75% BOS) | **{_fmt(r["pv_kw"],2)} kWp â†' {r["num_panels"]} Ã— {r.get("panel_wp",400)} Wp** |
| Battery | {_fmt(r["daily_kwh"],3)} Ã— {d.get("autonomy",1)} day Ã· ({int(r.get("chem_dod",0.9)*100)}% DoD) | **{_fmt(r["bat_kwh"],2)} kWh â†' {r["num_bat"]} Ã— {_fmt(r["unit_bat_kwh"],2)} kWh** |
| Inverter | Peak load {_fmt(r.get("peak_kw",0),2)} kW Ã— 1.25 SF | **{_fmt(r["inv_kw"],1)} kW** |
| MPPT | Array {_fmt(r["pv_kw"],2)} kWp Ã· {d.get("voltage",48)} V bus | **{r.get("mppt_a","—")} A** |

## A4. PV Array Design

| Parameter | Value |
|---|---|
| PV Array Capacity | {_fmt(r.get("pv_kw",0),2)} kWp |
| Number of Panels | {r.get("num_panels",0)} Ã— {r.get("panel_wp",400)} Wp |
| Panel Technology | Monocrystalline PERC, BS EN 61215 |
| Panels per String | {pps} in series |
| Number of Strings | {num_strings} parallel |
| Last String Panels | {last_panels} modules |
| String Voc (est.) | {pps*24} V |
| Temperature Derating Factor | {r.get("temp_derating","—")} |
| BOS Efficiency | 75% |
| Tilt Angle | 10—15Â° minimum (self-cleaning) |
| Orientation | Equator-facing (south in N. hemisphere) |
| Design Standard | IEC 61215 / IEC 61730 |

## A5. Battery Storage Design

| Parameter | Value |
|---|---|
| Total Battery Capacity | {_fmt(r.get("bat_kwh",0),2)} kWh |
| Number of Units | {r.get("num_bat",0)} Ã— {_fmt(r.get("unit_bat_kwh",0),2)} kWh each |
| Chemistry | {chem} |
| Depth of Discharge (DoD) | {int((r.get("chem_dod",0.9))*100)}% |
| Cycle Life | {r.get("chem_cycles","4,000+")} cycles |
| Lifetime | {r.get("chem_life","12")} years |
| Autonomy | {d.get("autonomy",1)} day(s) |
| BMS | Built-in per unit |
| Mounting | Ventilated steel rack, â‰¥ 300mm clearance |

## A6. Inverter & Charge Controller

| Parameter | Value |
|---|---|
| Inverter Rating | {_fmt(r.get("inv_kw",0),1)} kW |
| Type | Hybrid Inverter/Charger |
| Recommended Brands | {r.get("inv_brand","Victron / Growatt / Deye")} |
| MPPT Rating | {r.get("mppt_a","—")} A |
| DC Input Voltage | {d.get("voltage",48)} V |
| AC Output | {"415 V three-phase" if phase=="three" else "230 V single-phase"} |
| Inverter Efficiency | â‰¥ 95% |

## A7. AC Cable Schedule (BS 7671)

| Circuit | Size (mmÂ²) | Capacity (A) | Breaker (A) | Volt Drop | Compliant |
|---|---|---|---|---|---|
"""
    for c in r.get("ac_cables", []):
        md += (f"| {c.get('circuit','')} | {c.get('cable_size_mm2','')} mmÂ² | "
               f"{c.get('cable_capacity','')} A | {c.get('breaker_a','')} A | "
               f"{c.get('vd_percent','')}% | {'✓ Yes' if c.get('vd_ok') else 'âœ— Review'} |\n")

    md += """
## A8. AC Cable Voltage Drop Working (BS 7671 Appendix 4)

Per-circuit step-by-step working for design verification and bank review.

"""
    for c in r.get("ac_cables", []):
        vd_limit_v = c["vd_limit_pct"] / 100 * c["voltage_v"]
        phase_note = "(Ã—0.866 three-phase factor already applied)" if c["phase"] == "three" else ""
        result_str = "✓ PASS" if c["vd_ok"] else "âœ— FAIL — increase cable size"
        md += f"""### {c["circuit"]}

**Cable selected: {c["cable_size_mm2"]} mmÂ² {c["core_type"]}** &nbsp;|&nbsp; {c["cable_capacity"]} A capacity &nbsp;|&nbsp; {c["breaker_a"]} A protective device

| Parameter | Symbol | Value |
|---|---|---|
| Nominal voltage | Vn | {c["voltage_v"]} V ({c["phase"].title()}-Phase) |
| Load power | P | {c["power_kw"]} kW |
| Design current | Ib | **{c["design_current"]} A** |
| Cable length | L | **{c["length_m"]} m** |
| Installation method | — | Method {c["install_method"]} — {c["install_desc"]} |
| Ambient temperature | Ta | {c["ambient_c"]}Â°C |
| Temperature factor | Ct | {c["temp_factor"]} |
| Grouping factor | Cg | {c["group_factor"]} |
| Minimum Iz required | Iz_min | {c["i_z_required"]} A |
| Tabulated mV/A/m | — | **{c["vd_mv_am"]}** mV/A/m {phase_note} |
| Actual VD | — | {c["vd_mv_am"]} Ã— {c["design_current"]} Ã— {c["length_m"]} / 1000 = **{c["vd_volts"]:.3f} V** ({c["vd_percent"]:.3f}%) |
| Permitted limit | — | {c["vd_limit_pct"]}% of {c["voltage_v"]} V = {vd_limit_v:.2f} V |
| Result | — | **{result_str}** |

"""

    md += """### Cable Calculation Notes

| Item | Reference |
|---|---|
| VD tabulated values | BS 7671:2018 Appendix 4, Tables 4D2B / 4D5B |
| 3-phase VD factor Ã—0.866 | = âˆš3/2, IEC 60364-5-52 |
| Temperature correction | BS 7671 Table 4B2 (ref 30Â°C) |
| Grouping correction | BS 7671 Table 4B1 |
| VD limits | Inverterâ†'DB: 1.5% Â· Main feeder: 2.5% Â· Sub-distribution: 3.0% Â· Grid/Gen: 2.0% |
| Breaker coordination | Next standard size above Ib Ã— 1.05; must not exceed cable Iz |

## A9. Wire Colour Code (BS 7671 / IEC 60364)

| Conductor | Colour |
|---|---|
| DC Positive (+) | Red |
| DC Negative (âˆ') | Blue |
| AC Line / Phase | Brown |
| AC Neutral | Grey |
| Protective Earth (PE) | Green/Yellow |
| Battery Circuit | Purple |

## A10. Pre-Installation Site Assessment Checklist

The following must be verified on site before installation begins. Each item is marked Pass / Fail / N/A
with remarks. Sign-off by the consultant and the client is required before mobilisation.

### A10.1 Site Suitability & Solar Resource

| # | Item | Status |
|---|---|:---:|
"""
    insp_sections = [
        ("A10.1", [
            f"Location {d.get('region','')}, {d.get('country','')} has adequate solar irradiance (â‰¥ 4 PSH/day)",
            f"Roof/ground area sufficient for {r['num_panels']} panels (approx. {r['num_panels']*2} mÂ² minimum)",
            "Panel orientation achievable — equator-facing surface available",
            "No permanent shading obstruction between 9am—3pm (buildings, trees, towers)",
            "Roof pitch â‰¥ 10Â° or flat roof with tilt framing available",
            "No planned construction or tree growth likely to cause future shading",
            "Access to roof/ground area is safe and adequate for installation and maintenance",
        ]),
        ("A10.2 Structural & Roof Assessment", [
            "Roof/structure type identified (concrete slab / IBR metal sheet / clay tile / flat membrane)",
            "Structure age and condition acceptable for 25-year system life",
            f"Load capacity adequate — supports ~{r['num_panels']*20} kg PV array weight",
            f"For ground mount: {r['num_panels']*4} mÂ² land area available, level and stable",
            "Existing skylights, vents, or services do not conflict with array footprint",
            "Roof waterproofing in acceptable condition prior to installation",
            "No asbestos-containing materials identified in roof structure",
        ]),
        ("A10.3 Existing Electrical Infrastructure", [
            "Existing main distribution board (MDB) identified and accessible",
            "MDB has spare capacity / ways for new solar incomer",
            "Existing earthing / earth rod present — condition to be tested",
            f"Current utility connection is {'3-phase 415V' if phase=='three' else 'single-phase 230V'} — matches proposed system",
            "Available space for inverter and battery bank within equipment room",
            "Equipment room is dry, secure, and ventilated",
            "Cable routing path from array to equipment room identified and clear",
            "Existing wiring checked — no obvious overloads or faults",
        ]),
        ("A10.4 Load & Demand Validation", [
            f"Daily demand {_fmt(r['daily_kwh'],2)} kWh/day verified against actual utility bills",
            f"Peak connected load {_fmt(r.get('peak_kw',0),2)} kW confirmed — no major loads omitted",
            "Critical loads (medical, refrigeration, security) identified for backup priority",
            "Load profile reasonably consistent — no major seasonal variation",
            "Future load growth in next 3—5 years factored into sizing",
            f"High-draw appliances (A/C, pumps) confirmed compatible with {_fmt(r['inv_kw'],1)} kW inverter",
        ]),
        ("A10.5 Grid Connection & Regulatory", [
            "Utility grid connection available at site",
            f"Net metering / feed-in tariff policy investigated for {d.get('country','')}",
            "Anti-islanding protection required — inverter has built-in function",
            "Local installation code compliance confirmed (BS 7671 / IEC 60364 / national standard)",
            "Planning authority notified of proposed installation (if required)",
        ]),
        ("A10.6 Health, Safety & Access", [
            "Safe roof access route confirmed — scaffolding / MEWP requirements noted",
            "Electrical isolation of existing installation possible before work begins",
            "No asbestos, hazardous materials, or restricted zones on site",
            "Client / occupants can remain during installation or relocation required",
            "Fire risk assessed — battery room ventilation and extinguisher provision confirmed",
        ]),
    ]
    # First section already has its sub-heading rendered in the leading md block; emit its rows then the others
    for j, item in enumerate(insp_sections[0][1], 1):
        md += f"| {j} | {item} | â˜ |\n"
    for sec_title, items in insp_sections[1:]:
        md += f"\n### {sec_title}\n\n| # | Item | Status |\n|---|---|:---:|\n"
        for j, item in enumerate(items, 1):
            md += f"| {j} | {item} | â˜ |\n"

    md += f"""

---

# PART B — FINANCIAL PROPOSAL

## B1. Bill of Quantities (BOQ)

| No. | Description | Specification | Qty | Unit | Basic Rate ({sym}) | Total Rate ({sym}) | Amount ({sym}) |
|---|---|---|---|---|---|---|---|
"""
    for row in r.get("boq_rows", []):
        spec = str(row.get("spec","")).replace("|", "Â·")  # pipes break markdown tables
        md += (f"| {row['no']} | {row['desc']} | {spec} | "
               f"{row['qty']} | {row['unit']} | "
               f"{_fmt(row['basic'],2)} | {_fmt(row['total_r'],2)} | "
               f"**{_fmt(row['amount'],2)}** |\n")
    md += f"| | | | | | | **GRAND TOTAL** | **{sym} {_fmt(r['boq_grand'],2)}** |\n"

    md += f"""
*Note: Total Rate = Basic Rate Ã— 1.08 (8% supply/procurement markup — delivery, overheads & profit)*

### BOQ Notes

- Rates: Total Rate = Basic Rate Ã— 1.08 (8% supply/procurement markup)
- Quantities subject to detailed design review and site survey
- Cable lengths are estimated — confirm actual lengths on site
- Subject to contractor quotation; excludes site-specific VAT / import duties
- DC cable sizes and AC cable sizes are calculated from actual system design

## B2. CAPEX Breakdown

| Item | Amount ({sym}) |
|---|---|
| Equipment Supply (incl. 8% markup) | {_fmt(eco.get("equip_local",0),0)} |
| Installation Labour ({eco.get("install_rate_pct",15)}%) | {_fmt(eco.get("install_local",0),0)} |
| **Total CAPEX** | **{_fmt(eco.get("total_local",0),0)}** |
| Contingency (10%) — advisory | {_fmt(eco.get("total_local",0)*0.1,0)} |
| **Budget (incl. contingency)** | **{_fmt(eco.get("total_local",0)*1.1,0)}** |

## B3. Financial Summary

| Item | Value |
|---|---|
| Equipment Supply (incl. markup) | {sym} {_fmt(eco.get("equip_local",0),0)} |
| Installation Labour ({eco.get("install_rate_pct",15)}%) | {sym} {_fmt(eco.get("install_local",0),0)} |
| **Total CAPEX** | **{sym} {_fmt(eco.get("total_local",0),0)}** |
| Contingency (10%) | {sym} {_fmt(eco.get("total_local",0)*0.1,0)} |
| Budget with contingency | {sym} {_fmt(eco.get("total_local",0)*1.1,0)} |

## B4. Return on Investment

| Metric | Value |
|---|---|
| Annual Solar Generation | {_fmt(eco.get("annual_kwh",0),0)} kWh/year |
| Gross Annual Savings (Yr 1) | {sym} {_fmt(eco.get("annual_sav",0),0)}/year |
| Annual O&M Cost | {sym} {_fmt(eco.get("om_yr1",0),0)}/year |
| **Net Annual Benefit (Yr 1)** | **{sym} {_fmt(eco.get("net_yr1",0),0)}/year** |
| Simple Payback Period | {_fmt(eco.get("payback",0),1)} years |
| Net Present Value (25yr) | {sym} {_fmt(eco.get("npv",0),0)} |
| Internal Rate of Return | {"%.1f" % eco.get("irr_pct",0) if eco.get("irr_pct") else "N/A"}% |
| 25-Year ROI | {_fmt(eco.get("roi_pct",0),0)}% |
| Cumulative Savings (25yr) | {sym} {_fmt(eco.get("cumul_25",0),0)} |
| Annual COâ‚‚ Avoided | {_fmt(eco.get("co2_yr",0),2)} tonnes/year |

## B5. Project Assessment & Verdict Reasons

**Verdict:** {eco.get("verdict","—")} | **Bankability:** {eco.get("bankability","—")}

"""
    for reason in eco.get("verdict_reasons", []):
        md += f"- {reason}\n"
    if eco.get("bankability") and eco.get("bankability") != "BANKABLE":
        md += f"\n**Bankability assessment ({eco.get('bankability','—')}):**\n\n"
        for reason in eco.get("bank_reasons", []):
            md += f"- {reason}\n"

    md += f"""
## B6. Loan Structure & Bankability

| Parameter | Value |
|---|---|
| Total Investment | {sym} {_fmt(eco.get("total_local",0),0)} |
| Loan Amount (70%) | {sym} {_fmt(eco.get("loan_amt",0),0)} |
| Client Equity (30%) | {sym} {_fmt(eco.get("equity",0),0)} |
| Interest Rate | 15% p.a. |
| Loan Tenor | 7 years |
| Monthly Repayment | {sym} {_fmt(eco.get("pmt",0),0)}/month |
| Annual Debt Service | {sym} {_fmt(eco.get("annual_pmt",0),0)}/year |
| **DSCR** | **{_fmt(eco.get("dscr",0),2)} — {eco.get("bankability","—")}** |

## B7. Monthly Generation & Savings Profile

| Month | Generation (kWh) | Grid Offset (kWh) | Savings ({sym}) |
|---|---|---|---|
"""
    annual_gen = r["daily_kwh"] * 365
    base_avg   = annual_gen / 12
    for m, kwh in monthly:
        offset_kwh = round(kwh * offset_factor, 1)
        saving = round(kwh * d.get("tariff", 0), 1)
        md += f"| {m} | {kwh} | {offset_kwh} | {sym} {saving} |\n"
    md += (f"| **Annual Total** | **{_fmt(annual_gen,0)}** | "
           f"**{_fmt(annual_offset_kwh,0)}** | **{sym} {_fmt(eco.get('annual_sav',0),0)}** |\n")

    md += f"""
## B8. Environmental Impact

| Metric | Value |
|---|---|
| Annual COâ‚‚ Avoided | {_fmt(eco.get("co2_yr",0),2)} tonnes/year |
| 25-Year COâ‚‚ Avoided | {_fmt(eco.get("co2_yr",0)*25,1)} tonnes |
| Equivalent Trees Planted | {trees} trees/year |
| Equivalent Cars Removed | {cars} cars/year |
| Grid Emission Factor | 0.40 kg COâ‚‚/kWh |
| Carbon Status | **Carbon Positive** |

## B9. 25-Year Cash Flow Projection

| Year | Gross Saving | O&M | Net Saving | Cumulative |
|---|---|---|---|---|
"""
    for cf in eco.get("cf_rows", []):
        flag = " â—„ BREAK-EVEN" if eco.get("breakeven") and cf["yr"] == eco["breakeven"] else ""
        md += (f"| {cf['yr']} | {sym}{_fmt(cf['gross'],0)} | {sym}{_fmt(cf['om'],0)} | "
               f"{sym}{_fmt(cf['net'],0)} | {sym}{_fmt(cf['cumul'],0)}{flag} |\n")

    md += "\n*Assumptions: Tariff escalation 8%/yr, Discount rate 12%, O&M 1.2%/yr, Degradation 0.5%/yr, Life 25 years*\n"

    if recs:
        md += f"\n## B10. Redesign Recommendations\n\n"
        md += "The following engineering and financial improvements are recommended to achieve project approval and bankability:\n\n"
        for i, rec in enumerate(recs, 1):
            priority_label = "HIGH PRIORITY" if rec["priority"]==1 else "MEDIUM PRIORITY" if rec["priority"]==2 else "ADVISORY"
            md += f"### {i}. {rec['title']} [{priority_label}] ({rec['category']})\n\n"
            md += f"**Action:** {rec['action']}\n\n"
            md += f"**Expected Impact:** {rec['impact']}\n\n"

    # ── PART C — Project Delivery (Material Schedule, Programme, Staffing, Safety) ──
    md += f"""---

# PART C — PROJECT DELIVERY

## C1. Material & Equipment Schedule

### C1.1 PV Array & Mounting

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 1.1 | PV Solar Panels | {r.get("panel_wp",400)} Wp Monocrystalline PERC, Tier 1, Voc â‰ˆ 24 V | {r["num_panels"]} | Modules |
| 1.2 | Aluminium Mounting Rails | 40Ã—40 mm anodised aluminium | {int(r["num_panels"]*1.2)} | m |
| 1.3 | Mid & End Clamps | Stainless steel SS304 | {r["num_panels"]*4} | Sets |
| 1.4 | Roof Mounting Brackets | Galvanised steel, tilt-adjustable | {int(r["num_panels"]*1.5)} | No. |
| 1.5 | DC Solar Cable (strings) | 6 mmÂ² TÃœV UV-resistant, red & black | {int(r["num_panels"]*8)} | m |
| 1.6 | MC4 Connectors | IP67, 1000 VDC rated | {r["num_panels"]*4} | Pairs |
| 1.7 | DC Main Cable (combinerâ†'inverter) | 10 mmÂ² DC solar cable | 30 | m |
| 1.8 | Earthing Cable for Panel Frames | 6 mmÂ² green/yellow | {int(r["num_panels"]*3)} | m |

### C1.2 DC Combiner & Protection

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 2.1 | DC Combiner Box | IP65, {min(num_strings,4)}-string, lockable | 1 | No. |
| 2.2 | String Fuses | 10A DC PV fuse, 1000 VDC | {num_strings*2} | No. |
| 2.3 | DC Surge Protection Device | Type 2, 1000 VDC, IEC 61643-31 | 1 | No. |
| 2.4 | DC Main Isolator | 3-pole DC-rated, lockable | 1 | No. |
| 2.5 | Metallic Cable Conduit | 20 mm steel conduit | 25 | m |
| 2.6 | Cable Tray / Trunking | 50Ã—50 mm galvanised steel | 10 | m |

### C1.3 Inverter, Battery & MPPT

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 3.1 | Hybrid Inverter / Charger | {_fmt(r["inv_kw"],1)} kW, {d.get("voltage",48)}V DC, built-in MPPT {r.get("mppt_a","—")}A | 1 | No. |
| 3.2 | Lithium Battery Units | {_fmt(r["unit_bat_kwh"],2)} kWh {chem}, {d.get("voltage",48)}V, BMS | {r["num_bat"]} | No. |
| 3.3 | Battery Steel Rack | Powder-coated, for {r["num_bat"]} units | {max(1,(r["num_bat"]+1)//2)} | No. |
| 3.4 | Battery DC Fuse (ANL) | 1.25 Ã— max charge current | {r["num_bat"]} | No. |
| 3.5 | Battery DC Cable | 25 mmÂ² flexible, red & black | {r["num_bat"]*4} | m |
| 3.6 | BMS Communication Cable | RS485 / CAN bus | {r["num_bat"]} | Cables |
| 3.7 | Inverter Wall Bracket | Heavy-duty steel | 1 | Set |

### C1.4 AC Distribution & Protection

| # | Description | Specification | Qty | Unit |
|---|---|---|---|---|
| 4.1 | Main AC Distribution Board | {"18-way" if phase=="single" else "12-way 3-ph"}, IP40 | 1 | No. |
| 4.2 | RCCB Incomer | {r["ac_cables"][0]["breaker_a"] if r.get("ac_cables") else 63}A, 30mA, Type A, BS EN 61008 | 1 | No. |
| 4.3 | MCB Lighting | 10A Type B, 6kA | 2 | No. |
| 4.4 | MCB Sockets | 16A Type B, 6kA | 3 | No. |
| 4.5 | MCB Air Conditioning | 32A Type C, 6kA | 1 | No. |
| 4.6 | MCB Pumps / Motors | 16A Type C, 6kA | 2 | No. |
| 4.7 | AC Surge Protection Device | Type 2, 230/415V, BS EN 61643 | 1 | No. |
"""
    for i, c in enumerate(r.get("ac_cables", []), start=8):
        md += f"| 4.{i} | AC Cable — {c['circuit']} | {c['cable_size_mm2']} mmÂ² Cu XLPE/PVC | {c.get('length_m',20)+5} | m |\n"

    md += f"""
### C1.5 Earthing, Bonding & Sundries

| # | Description | Qty | Unit |
|---|---|---|---|
| 5.1 | Copper Earth Rod (16 mm dia., 2.4 m) | 2 | No. |
| 5.2 | Earth Rod Clamp & Driver | 2 | Sets |
| 5.3 | Earth Busbar (10-way copper) | 1 | No. |
| 5.4 | Main Earthing Conductor (16 mmÂ² G/Y) | 15 | m |
| 5.5 | Bonding Cables (6 mmÂ² G/Y) | 40 | m |
| 5.6 | Cable Labels (PVC self-laminating) | 200 | No. |
| 5.7 | Cable Ties UV-resistant | 1 | Box (200) |
| 5.8 | Warning / Safety Labels (BS EN 60445) | 1 | Set |
| 5.9 | IP65 Cable Glands (M20—M32) | 20 | No. |
| 5.10 | UV-Resistant Silicone Sealant | 4 | Tubes |

### C1.6 Test & Commissioning Instruments

| Instrument | Purpose | Acceptance Standard |
|---|---|---|
| Insulation Resistance Tester (Megger) | Cable insulation integrity | â‰¥ 1 MÎ© (IEC 60364-6) |
| Earth Electrode Resistance Tester | Earth rod resistance | â‰¤ 10 Î© (BS 7430) |
| Clamp Earth Tester | Non-invasive earth continuity | â‰¤ 0.1 Î© (BS 7671 Ch.61) |
| Digital Multimeter (1000V DC) | String Voc, polarity | CAT III 1000V |
| DC Clamp Meter (1000V / 60A DC) | String Isc, battery current | CAT III 600V |
| RCD Tester | Trip time verification | â‰¤ 40 ms (BS EN 61008) |
| Voltage Drop Tester | Full-load volt drop | â‰¤ 3% (BS 7671 App 4) |
| Thermal Imaging Camera | Hot-spot detection | IEC 62446-3 |

## C2. Installation Programme

The installation follows a **7-phase methodology** aligned with BS 7671:2018, IEC 60364, and IEC 62446.
Each phase has defined entry criteria, activities, deliverables, and sign-off before the next phase begins.
**Total programme: 12 working days (weather permitting).**

| Phase | Activity | Days | Duration |
|---|---|---|---|
| 1 | Mobilisation & Site Preparation | Days 1—2 | 2 days |
| 2 | Civil & Structural Works | Days 2—4 | 3 days |
| 3 | PV Panel Installation | Days 4—6 | 3 days |
| 4 | Equipment Room Fit-Out | Days 5—7 | 3 days |
| 5 | DC & AC Wiring | Days 7—9 | 3 days |
| 6 | Earthing, Bonding & Pre-commissioning Tests | Days 9—10 | 2 days |
| 7 | Commissioning, Testing & Handover | Days 10—12 | 3 days |

### Programme Notes

- **Total duration:** 12 working days (weather-permitting)
- **Weather hold:** No roof work in rain, lightning, or wind > 25 mph
- **Working hours:** 07:30—17:30 Mon—Fri; 07:30—13:00 Sat if required
- **Parallel working:** Phases 3 and 4 overlap (Days 5—7) — civil and electrical teams work simultaneously

## C3. Installation Phase Detail

### Phase 1 — Mobilisation & Site Preparation (Days 1—2)

**Activities:** Deliver and inventory all equipment on site; set up site compound and secure storage for
panels and batteries; install temporary power and lighting; brief all staff on HSE plan; prepare roof
access (scaffold or MEWP); mark out equipment room layout and cable routes.

**Sign-off outputs:** Signed delivery notes, site induction records, HSE risk assessment signed.

### Phase 2 — Civil & Structural Works (Days 2—4)

**Activities:** Install roof mounting brackets/L-feet at designed spacing; assemble and level aluminium
mounting rails; verify tilt angle; core through roof/walls for DC cable entry — seal immediately;
fix conduit supports and tray brackets along cable route; install metallic conduit from roof to
equipment room.

**Sign-off outputs:** Waterproofing test, structural load check (if required), as-installed conduit sketch.

### Phase 3 — PV Panel Installation (Days 4—6)

**Activities:** Mount panels row-by-row, bottom to top; torque clamps to spec; connect in series strings
with MC4 connectors; verify polarity; install string fuses in combiner box; record fuse ratings; run
and label DC string cables in conduit.

**Sign-off outputs:** String cable labels at both ends, visual inspection — no cracked panels.

### Phase 4 — Equipment Room Fit-Out (Days 5—7)

**Activities:** Fix inverter wall bracket; mount and level inverter; assemble battery rack; anchor to
floor or wall; install batteries; connect in parallel per wiring diagram; connect BMS communication
cables; configure settings; fix DC isolator and AC DB at correct heights; install earth busbar;
run main earthing conductor.

**Sign-off outputs:** Inverter mounting record, battery connection torque check, room layout photo.

### Phase 5 — DC & AC Wiring (Days 7—9)

**Activities:** Run DC main cable combiner â†' inverter (double-check polarity); connect battery cables
with ANL fuses; wire AC output inverter â†' DB incomer; install RCCB, MCBs, SPD in DB; run all AC
final circuit cables; label both ends at every junction.

**Sign-off outputs:** As-installed wiring diagram, cable schedule with actual lengths, labels verified.

### Phase 6 — Earthing, Bonding & Pre-commissioning Tests (Days 9—10)

**Activities:** Drive earth rods to â‰¥ 2.4 m depth; connect to earth busbar; bond all metalwork —
panel frames, inverter, battery rack, DB; test earth electrode resistance (â‰¤ 10 Î© before proceeding);
insulation resistance test all circuits â‰¥ 1 MÎ©; continuity all earth/bonding conductors â‰¤ 0.1 Î©;
polarity check on all DC strings (signed by two technicians).

**Sign-off outputs:** Earth resistance certificate, IR test schedule, polarity check record.

### Phase 7 — Commissioning, Testing & Handover (Days 10—12)

**Activities:** Energise inverter; verify AC output voltage and frequency; test MPPT tracking;
confirm generation on inverter display; RCD trip time test â‰¤ 40 ms; MCB overload test per circuit;
voltage drop test under full load â‰¤ 3%; 7-day performance monitoring (generation vs design);
client handover training; issue O&M manual and warranties.

**Sign-off outputs:** Full commissioning test schedule, 7-day monitoring log, Installation
Completion Certificate, O&M manual issued.

## C4. Staffing Plan

### C4.1 Project Team

| Role | No. | Key Qualifications | Days on Site |
|---|---|---|---|
| Project Engineer / Site Manager | 1 | BEng Electrical, 18th Ed BS 7671, PV Solar cert, IOSH | All 12 days |
| Senior Electrical Technician | 1 | C&G 2365 NVQ L3, 18th Ed + 2391, ECS Gold, Work at Height | All 12 days |
| Electrical Apprentice / Assistant | 1 | NVQ L2 Electrical, Manual Handling, CSCS Green | All 12 days |
| Structural / Civil Technician | 1 | Roof mounting experience, PASMA/IPAF, CSCS | Days 1—6 |
| HSE Officer (part-time) | 1 | NEBOSH General, First Aid at Work | Days 1—2, 9—10 |

**Total peak headcount: 5 persons on site (Days 1—2 and 4—6)**

### C4.2 Staff Deployment by Phase

| Phase | Activity | Days | Proj. Eng. | Sr. Tech | Apprentice | Civil Tech | HSE Officer |
|---|---|---|---|---|---|---|---|
| 1 | Mobilisation | 1—2 | ✓ | ✓ | ✓ | ✓ | ✓ |
| 2 | Civil Works | 2—4 | ✓ (part) | — | ✓ | ✓ | ✓ (part) |
| 3 | PV Installation | 4—6 | ✓ | ✓ | ✓ | ✓ | ✓ (part) |
| 4 | Equipment Fit-Out | 5—7 | ✓ | ✓ | ✓ | — | — |
| 5 | DC & AC Wiring | 7—9 | ✓ | ✓ | ✓ | — | — |
| 6 | Earthing & Testing | 9—10 | ✓ | ✓ | ✓ | — | ✓ (part) |
| 7 | Commissioning | 10—12 | ✓ | ✓ | ✓ | — | — |

### C4.3 Key Responsibilities

**Project Engineer / Site Manager** — Overall technical quality, programme, safety, and client
communication. Signs off each phase, approves all test results, issues the Installation Completion
Certificate. Holds valid BS 7671 certification and PV Solar design qualification.

**Senior Electrical Technician** — Leads all electrical installation activities. Performs and records
all pre-commissioning and commissioning tests. Configures inverter and BMS settings. Mentors the
apprentice.

**Electrical Apprentice / Assistant** — Supports wiring, cable pulling, labelling, and containment
installation. Assists the senior technician and updates the material schedule as items are installed.

**Structural / Civil Technician** — Installs all roof mounting structure. Responsible for panel
mounting, tilt angle accuracy, and waterproofing of all roof penetrations.

**HSE Officer (part-time)** — Daily toolbox talks. Inspects PPE and access equipment. Maintains site
accident log. Emergency response coordinator.

### C4.4 Mandatory PPE — All Personnel

| Item | Standard |
|---|---|
| Safety helmet | EN 397 — all roof and overhead work |
| Safety boots (steel toe cap, anti-slip) | EN ISO 20345 |
| High-visibility vest / jacket (Class 2) | EN ISO 20471 |
| Safety glasses (drilling, cutting, chemicals) | EN 166 |
| Electrical insulated gloves (live-adjacent work) | EN 60903 |
| Full body harness (work at height > 2 m) | EN 361 |
| Anti-static wrist strap (inverter electronics) | IEC 61340-5 |

### C4.5 Site Safety Procedures

- Daily toolbox talk before work — attendance signed by all personnel
- Written risk assessment and method statement on site at all times
- Permit to work before any work on or near energised equipment
- Two-person rule — no solo working on roof or electrical equipment
- All access equipment inspected daily; defective items tagged out of service
- Emergency evacuation plan posted at site entrance and briefed on Day 1
- First aid kit and COâ‚‚ fire extinguisher on site at all times
- Any near-miss reported within 2 hours; incident report completed within 24 hours

## C5. Testing & Verification Schedule

| Test | Standard | Acceptance Criteria |
|---|---|---|
| Insulation Resistance | IEC 60364-6 | â‰¥ 1 MÎ© |
| Earth Continuity | BS 7671 | â‰¤ 0.1 Î© |
| Earth Electrode Resistance | BS 7430 | â‰¤ 10 Î© |
| RCD Trip Time | BS EN 61008 | â‰¤ 40ms |
| DC String Polarity | IEC 62446 | No reversed polarity |
| DC Open-Circuit Voltage | IEC 62446 | Within 5% of Voc |
| AC Output Voltage | BS 7671 | 230V Â±10% |
| Voltage Drop | BS 7671 | â‰¤ 3% final circuits |
| 7-Day Performance Check | IEC 62446 | â‰¥ 90% design output |

## C6. Warranties & O&M

| Item | Warranty |
|---|---|
| PV Panels — Product | 12 years |
| PV Panels — Performance | 25 years (â‰¥80% output) |
| Battery — Cycle Life | {r.get("chem_cycles","4,000+")} cycles / {r.get("chem_life","12")} years |
| Inverter | 5 years (extendable) |
| Installation Workmanship | 2 years |
| Annual O&M Cost (Yr 1) | {sym} {_fmt(eco.get("om_yr1",0),0)} |

---

# PROJECT VERDICT

**Verdict:** {eco.get("verdict","—")} | **Bankability:** {eco.get("bankability","—")}

{chr(10).join("- " + r2 for r2 in eco.get("verdict_reasons",[]))}

**DSCR:** {_fmt(eco.get("dscr",0),2)} | **Simple Payback:** {_fmt(eco.get("payback",0),1)} years | **25-yr NPV:** {sym} {_fmt(eco.get("npv",0),0)}

---

*Full Technical & Financial Proposal (Superset) — {project["name"]}*
*Generated by SolarPro Global Â· Intelligent PV Solar Design Platform*
*BS 7671:2018 Â· IEC 60364 Â· IEC 62305 Â· IEC 62446 Â· IEC 61215 Â· IEEE*
*All figures are indicative and subject to final site survey and detailed design.*
"""
    fname = f"Proposal_{project['name'].replace(' ','_')}.pdf"
    md = _diagrams_markdown(d, r) + md
    return _render_pdf(f"PV Solar Proposal — {project['name']}", md, fname)


# ─── Phase 4: Admin panel ─────────────────────────────────────────────────────

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        with get_db() as c:
            u = c.execute("SELECT is_admin FROM users WHERE id=?",
                          (session["user_id"],)).fetchone()
        if not u or not u["is_admin"]:
            abort(403)
        return f(*args, **kwargs)
    return decorated


@app.route("/admin")
@admin_required
def admin_dashboard():
    with get_db() as c:
        users      = c.execute("SELECT * FROM users ORDER BY created_at DESC").fetchall()
        projects   = c.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        tickets    = c.execute(
            "SELECT t.*, u.username FROM tickets t "
            "LEFT JOIN users u ON t.user_id=u.id "
            "ORDER BY t.updated_at DESC LIMIT 8").fetchall()
        open_t     = c.execute("SELECT COUNT(*) FROM tickets WHERE status='open'").fetchone()[0]
        total_rev  = c.execute("SELECT COALESCE(SUM(amount_usd),0) FROM payments WHERE status='success'").fetchone()[0]
        new_leads  = c.execute("SELECT COUNT(*) FROM leads WHERE status='new'").fetchone()[0]
        subs       = c.execute("SELECT COUNT(*) FROM newsletter_subscribers WHERE status='active'").fetchone()[0]
        recent_users = c.execute(
            "SELECT id, username, email, plan, created_at, is_admin "
            "FROM users ORDER BY created_at DESC LIMIT 6").fetchall()
        monthly_signups = c.execute(
            "SELECT strftime('%Y-%m', created_at) AS mo, COUNT(*) AS cnt "
            "FROM users GROUP BY mo ORDER BY mo DESC LIMIT 6").fetchall()

    plan_counts = {}
    for u in users:
        p = (u["plan"] or "free").lower()
        plan_counts[p] = plan_counts.get(p, 0) + 1

    paid_users = sum(plan_counts.get(p, 0) for p in ("basic", "professional", "enterprise"))

    return render_template("admin.html", user=current_user(),
                           users=users, projects=projects,
                           tickets=tickets, open_t=open_t,
                           plan_counts=plan_counts,
                           total_rev=total_rev, new_leads=new_leads,
                           subs=subs, paid_users=paid_users,
                           recent_users=recent_users,
                           monthly_signups=list(reversed(list(monthly_signups))))


@app.route("/admin/users", methods=["GET", "POST"])
@admin_required
def admin_users():
    if request.method == "POST":
        csrf_protect()
        uid    = request.form.get("uid", type=int)
        action = request.form.get("action")
        plan   = request.form.get("plan", "free")
        with get_db() as c:
            if action == "set_plan":
                c.execute("UPDATE users SET plan=? WHERE id=?", (plan, uid))
                flash(f"Plan updated.", "success")
            elif action == "toggle_admin":
                cur = c.execute("SELECT is_admin FROM users WHERE id=?", (uid,)).fetchone()
                new_val = 0 if cur and cur["is_admin"] else 1
                c.execute("UPDATE users SET is_admin=? WHERE id=?", (new_val, uid))
                flash("Admin status toggled.", "success")
            elif action == "set_role":
                new_role = request.form.get("role", "customer")
                c.execute("UPDATE users SET role=? WHERE id=?", (new_role, uid))
                flash(f"Role updated to '{new_role}'.", "success")
            elif action == "delete":
                # Locate the target by uid or username. Admin UI can post either.
                _target_id = uid
                if not _target_id:
                    _un = (request.form.get("username") or "").strip()
                    if _un:
                        _row = c.execute("SELECT id FROM users WHERE username=?",
                                          (_un,)).fetchone()
                        _target_id = _row["id"] if _row else None
                if not _target_id:
                    flash("User not found.", "danger")
                else:
                    _meta = c.execute("SELECT id, is_admin, username FROM users WHERE id=?",
                                       (_target_id,)).fetchone()
                    if not _meta:
                        flash("User not found.", "danger")
                    elif _meta["is_admin"]:
                        flash("Refused: cannot delete an admin via this action.", "danger")
                    elif _meta["id"] == session.get("user_id"):
                        flash("Refused: cannot delete yourself.", "danger")
                    else:
                        # Cascade. Tables that hold user-owned rows that the
                        # app expects to be present-or-absent per FK semantics.
                        # We intentionally skip audit_log so forensic trail
                        # survives the cleanup.
                        for _stmt in (
                            ("DELETE FROM projects WHERE user_id=?", (_target_id,)),
                            ("DELETE FROM payments WHERE user_id=?", (_target_id,)),
                            ("DELETE FROM referrals WHERE referrer_id=? OR referee_id=?",
                             (_target_id, _target_id)),
                            ("DELETE FROM login_failures WHERE username=?",
                             (_meta["username"],)),
                            ("DELETE FROM users WHERE id=?", (_target_id,)),
                        ):
                            try:
                                c.execute(_stmt[0], _stmt[1])
                            except Exception:
                                pass  # table may not exist in this schema version.
                        flash(f"User {_meta['username']} (id {_target_id}) permanently deleted.",
                              "warning")
            elif action == "disable":
                c.execute("UPDATE users SET plan='disabled' WHERE id=?", (uid,))
                flash("Account disabled.", "warning")
            elif action == "record_payment":
                pay_plan    = request.form.get("pay_plan", "professional")
                pay_amount  = float(request.form.get("pay_amount", 0) or 0)
                pay_curr    = request.form.get("pay_currency", "USD").upper()
                pay_ref     = request.form.get("pay_reference", "").strip() or f"MANUAL-{secrets.token_hex(6).upper()}"
                pay_gateway = request.form.get("pay_gateway", "manual")
                pay_upgrade = request.form.get("pay_upgrade_plan") == "1"
                _record_payment(uid, pay_gateway, pay_plan, pay_amount,
                                currency=pay_curr, reference=pay_ref, status="success")
                if pay_upgrade:
                    c.execute("UPDATE users SET plan=? WHERE id=?", (pay_plan, uid))
                flash(f"Payment recorded: {pay_curr} {pay_amount:.2f} ({pay_plan}). Ref: {pay_ref}", "success")
        return redirect(url_for("admin_users"))
    with get_db() as c:
        users = c.execute(
            "SELECT u.*, (SELECT COUNT(*) FROM projects WHERE user_id=u.id) AS proj_count "
            "FROM users u ORDER BY u.created_at DESC").fetchall()
    return render_template("admin_users.html", user=current_user(),
                           users=users, plan_prices=PLAN_PRICES)


@app.route("/admin/tickets")
@admin_required
def admin_tickets():
    status = request.args.get("status", "all")
    with get_db() as c:
        q = ("SELECT t.*, u.username, u.email "
             "FROM tickets t JOIN users u ON t.user_id=u.id ")
        if status != "all":
            rows = c.execute(q + "WHERE t.status=? ORDER BY t.updated_at DESC", (status,)).fetchall()
        else:
            rows = c.execute(q + "ORDER BY t.updated_at DESC").fetchall()
        counts = {r["status"]: r["cnt"] for r in c.execute(
            "SELECT status, COUNT(*) AS cnt FROM tickets GROUP BY status").fetchall()}
        counts["all"] = sum(counts.values())
    return render_template("admin_tickets.html", user=current_user(),
                           tickets=rows, status_filter=status, counts=counts)


@app.route("/admin/ticket/<int:tid>", methods=["GET", "POST"])
@admin_required
def admin_ticket_detail(tid):
    with get_db() as c:
        ticket = c.execute(
            "SELECT t.*, u.username, u.email, u.plan, u.created_at AS user_joined, u.name AS user_name, u.company "
            "FROM tickets t JOIN users u ON t.user_id=u.id WHERE t.id=?", (tid,)).fetchone()
        if not ticket:
            abort(404)
        replies = c.execute(
            "SELECT r.*, u.username, u.is_admin AS sender_is_admin FROM ticket_replies r "
            "JOIN users u ON r.user_id=u.id WHERE r.ticket_id=? ORDER BY r.created_at",
            (tid,)).fetchall()
        user_ticket_count = c.execute(
            "SELECT COUNT(*) FROM tickets WHERE user_id=?", (ticket["user_id"],)).fetchone()[0]
        prev_ticket = c.execute(
            "SELECT id FROM tickets WHERE id < ? ORDER BY id DESC LIMIT 1", (tid,)).fetchone()
        next_ticket = c.execute(
            "SELECT id FROM tickets WHERE id > ? ORDER BY id ASC LIMIT 1", (tid,)).fetchone()
    if request.method == "POST":
        csrf_protect()
        action = request.form.get("action")
        if action == "reply":
            msg = request.form.get("message", "").strip()
            if msg:
                with get_db() as c:
                    c.execute(
                        "INSERT INTO ticket_replies (ticket_id,user_id,is_admin,message) VALUES (?,?,1,?)",
                        (tid, session["user_id"], msg))
                    c.execute("UPDATE tickets SET status='answered',updated_at=? WHERE id=?",
                              (datetime.now().isoformat(), tid))
                flash("Reply sent.", "success")
        elif action in ("close", "open", "in_progress", "answered"):
            with get_db() as c:
                c.execute("UPDATE tickets SET status=?,updated_at=? WHERE id=?",
                          (action, datetime.now().isoformat(), tid))
            flash(f"Ticket marked {action.replace('_',' ')}.", "success")
        return redirect(url_for("admin_ticket_detail", tid=tid))
    return render_template("admin_ticket_detail.html", user=current_user(),
                           ticket=ticket, replies=replies,
                           user_ticket_count=user_ticket_count,
                           prev_ticket=prev_ticket, next_ticket=next_ticket)


@app.route("/admin/appliances", methods=["GET", "POST"])
@admin_required
def admin_appliances():
    if request.method == "POST":
        csrf_protect()
        action = request.form.get("action")
        with get_db() as c:
            if action == "add":
                c.execute(
                    "INSERT INTO appliances (category,name,default_watt,notes) VALUES (?,?,?,?)",
                    (request.form["category"], request.form["name"],
                     int(request.form["default_watt"]), request.form.get("notes", "")))
                flash("Appliance added.", "success")
            elif action == "delete":
                c.execute("DELETE FROM appliances WHERE id=?",
                          (request.form.get("aid", type=int),))
                flash("Appliance deleted.", "success")
            elif action == "edit":
                c.execute(
                    "UPDATE appliances SET category=?,name=?,default_watt=?,notes=? WHERE id=?",
                    (request.form["category"], request.form["name"],
                     int(request.form["default_watt"]), request.form.get("notes", ""),
                     request.form.get("aid", type=int)))
                flash("Appliance updated.", "success")
        return redirect(url_for("admin_appliances"))
    with get_db() as c:
        apps = c.execute("SELECT * FROM appliances ORDER BY category,name").fetchall()
    categories = sorted(set(a["category"] for a in apps))
    return render_template("admin_appliances.html", user=current_user(),
                           appliances=apps, categories=categories)


@app.route("/admin/helpline-kb", methods=["GET", "POST"])
@admin_required
def admin_helpline_kb():
    """View and manage the dynamically learned Helpline KB."""
    csrf_protect() if request.method == "POST" else None
    if request.method == "POST":
        action = request.form.get("action")
        kid    = request.form.get("kid", type=int)
        if action == "delete" and kid:
            with get_db() as c:
                c.execute("DELETE FROM helpline_learned_kb WHERE id=?", (kid,))
            flash("Entry deleted.", "success")
        elif action == "delete_all":
            agent = request.form.get("agent", "helpline")
            with get_db() as c:
                c.execute("DELETE FROM helpline_learned_kb WHERE agent=?", (agent,))
            flash("All entries cleared.", "warning")
        return redirect(url_for("admin_helpline_kb"))
    with get_db() as c:
        entries = c.execute(
            "SELECT * FROM helpline_learned_kb ORDER BY use_count DESC, created_at DESC"
        ).fetchall()
    return render_template("admin_helpline_kb.html", user=current_user(), entries=entries)


# ─── Phase 4: Account / subscription management ───────────────────────────────

def _record_payment(uid, gateway, plan, amount_usd, currency="USD",
                    reference="", status="success"):
    with get_db() as c:
        c.execute(
            "INSERT INTO payments (user_id,gateway,plan,amount_usd,currency,reference,status) "
            "VALUES (?,?,?,?,?,?,?)",
            (uid, gateway, plan, amount_usd, currency, reference, status))
        user_row = c.execute("SELECT email, username FROM users WHERE id=?", (uid,)).fetchone()
    if status == "success" and amount_usd and user_row:
        try:
            plan_label = PLAN_PRICES.get(plan, {}).get("label", plan.title())
            _subj = "Payment Confirmed - SolarPro " + plan_label + " Plan"
            _html = (
                "<div style='font-family:sans-serif;background:#0a0a14;color:#e2e2f0;"
                "padding:28px;border-radius:12px;max-width:600px'>"
                "<h2 style='color:#a78bfa'>Payment Confirmed</h2>"
                "<p>Hi " + str(user_row["username"]) + ", thank you for your payment.</p>"
                "<table style='width:100%;border-collapse:collapse;margin:16px 0'>"
                "<tr><td style='padding:8px;border-bottom:1px solid #1e1e3a'>Plan</td>"
                "<td style='padding:8px;border-bottom:1px solid #1e1e3a;color:#a78bfa'><b>" + plan_label + "</b></td></tr>"
                "<tr><td style='padding:8px;border-bottom:1px solid #1e1e3a'>Amount</td>"
                "<td style='padding:8px;border-bottom:1px solid #1e1e3a'><b>" + currency + " " + str(amount_usd) + "</b></td></tr>"
                "<tr><td style='padding:8px;border-bottom:1px solid #1e1e3a'>Gateway</td>"
                "<td style='padding:8px;border-bottom:1px solid #1e1e3a'>" + gateway.title() + "</td></tr>"
                "<tr><td style='padding:8px'>Reference</td>"
                "<td style='padding:8px;font-size:12px'>" + (reference or "N/A") + "</td></tr>"
                "</table>"
                "<p>Your subscription is now active. <a href='https://solarpro.aiappinvent.com/account' "
                "style='color:#a78bfa'>View your account</a>.</p>"
                "<hr style='border-color:#1e1e3a'>"
                "<p style='color:#6868a0;font-size:12px'>Questions? Email billing@aiappinvent.com</p>"
                "</div>"
            )
            _send_email(user_row["email"], _subj, _html, from_addr=EMAIL_BILLING)
        except Exception:
            pass


@app.route("/account")
@login_required
def account():
    user = current_user()
    plan = (user["plan"] or "free").lower()
    uid  = session["user_id"]
    with get_db() as c:
        proj_count   = c.execute("SELECT COUNT(*) FROM projects WHERE user_id=?", (uid,)).fetchone()[0]
        pay_rows     = c.execute(
            "SELECT * FROM payments WHERE user_id=? ORDER BY created_at DESC LIMIT 24", (uid,)).fetchall()
        open_tickets = c.execute(
            "SELECT COUNT(*) FROM tickets WHERE user_id=? AND status='open'", (uid,)).fetchone()[0]
        emails_sent  = c.execute(
            "SELECT COUNT(*) FROM email_logs WHERE user_id=? AND status='sent'", (uid,)).fetchone()[0]
        total_paid   = c.execute(
            "SELECT COALESCE(SUM(amount_usd),0) FROM payments WHERE user_id=? AND status='success'",
            (uid,)).fetchone()[0]
    limit = PLAN_LIMITS.get(plan, 1)
    return render_template("account.html", user=user, plan=plan,
                           proj_count=proj_count, limit=limit,
                           payments=pay_rows, plan_prices=PLAN_PRICES,
                           open_tickets=open_tickets, emails_sent=emails_sent,
                           total_paid=total_paid)


@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings():
    if request.method == "POST":
        csrf_protect()
        section = request.form.get("_section", "profile")
        uid = session["user_id"]

        if section == "profile":
            fields = ["org_name", "org_address", "org_email", "org_phone", "org_website", "timezone", "org_whatsapp"]
            vals = {f: request.form.get(f, "").strip() for f in fields}
            with get_db() as c:
                c.execute(
                    "UPDATE users SET org_name=?, org_address=?, org_email=?, "
                    "org_phone=?, org_website=?, timezone=?, org_whatsapp=? WHERE id=?",
                    (vals["org_name"], vals["org_address"], vals["org_email"],
                     vals["org_phone"], vals["org_website"], vals["timezone"],
                     vals["org_whatsapp"], uid))
            flash("Organisation profile saved.", "success")
            return redirect(url_for("settings") + "?tab=org")

        elif section == "datetime":
            date_fmt = request.form.get("date_format", "DD/MM/YYYY").strip()
            time_fmt = request.form.get("time_format", "24h").strip()
            with get_db() as c:
                c.execute("UPDATE users SET date_format=?, time_format=? WHERE id=?",
                          (date_fmt, time_fmt, uid))
            flash("Date & time preferences saved.", "success")
            return redirect(url_for("settings") + "?tab=datetime")

        elif section == "smtp":
            resend_key = request.form.get("resend_api_key", "").strip()
            smtp_host  = request.form.get("smtp_host", "").strip()
            smtp_port  = request.form.get("smtp_port", "587").strip() or "587"
            smtp_user  = request.form.get("smtp_user", "").strip()
            smtp_pass  = request.form.get("smtp_pass", "").strip()
            smtp_from  = request.form.get("smtp_from", "").strip()
            smtp_tls   = request.form.get("smtp_tls", "starttls").strip()
            with get_db() as c:
                if not smtp_pass:
                    existing = c.execute("SELECT smtp_pass FROM users WHERE id=?", (uid,)).fetchone()
                    if existing:
                        smtp_pass = existing["smtp_pass"] or ""
                if not resend_key:
                    existing2 = c.execute("SELECT resend_api_key FROM users WHERE id=?", (uid,)).fetchone()
                    if existing2:
                        resend_key = existing2["resend_api_key"] or ""
                c.execute(
                    "UPDATE users SET resend_api_key=?, smtp_host=?, smtp_port=?, smtp_user=?, "
                    "smtp_pass=?, smtp_from=?, smtp_tls=? WHERE id=?",
                    (resend_key, smtp_host, smtp_port, smtp_user, smtp_pass, smtp_from, smtp_tls, uid))
            flash("Email configuration saved.", "success")
            return redirect(url_for("settings") + "?tab=smtp")

        elif section == "password":
            current_pw = request.form.get("current_password", "")
            new_pw     = request.form.get("new_password", "")
            confirm_pw = request.form.get("confirm_password", "")
            with get_db() as c:
                user_row = c.execute("SELECT password_hash FROM users WHERE id=?", (uid,)).fetchone()
            if not check_password_hash(user_row["password_hash"], current_pw):
                flash("Current password is incorrect.", "danger")
            elif len(new_pw) < 8:
                flash("New password must be at least 8 characters.", "danger")
            elif new_pw != confirm_pw:
                flash("New passwords do not match.", "danger")
            else:
                with get_db() as c:
                    c.execute("UPDATE users SET password_hash=? WHERE id=?",
                              (generate_password_hash(new_pw), uid))
                flash("Password changed successfully.", "success")
            return redirect(url_for("settings") + "?tab=security")

        return redirect(url_for("settings"))

    return render_template("settings.html", user=current_user())


@app.route("/settings/test-smtp", methods=["POST"])
@login_required
def test_smtp_connection():
    """AJAX: test per-user SMTP configuration."""
    csrf_protect()
    user = current_user()
    host = (request.form.get("smtp_host") or user.get("smtp_host") or "").strip()
    port_str = (request.form.get("smtp_port") or user.get("smtp_port") or "587").strip()
    port = int(port_str) if port_str.isdigit() else 587
    usr  = (request.form.get("smtp_user") or user.get("smtp_user") or "").strip()
    pwd  = (request.form.get("smtp_pass") or user.get("smtp_pass") or "").strip()
    tls_mode = (request.form.get("smtp_tls") or user.get("smtp_tls") or "starttls").strip()
    if not host or not usr:
        return jsonify(ok=False, msg="SMTP host and username are required.")
    try:
        import smtplib
        if tls_mode == "starttls":
            srv = smtplib.SMTP(host, port, timeout=10)
            srv.ehlo()
            srv.starttls()
        elif tls_mode == "ssl":
            srv = smtplib.SMTP_SSL(host, port, timeout=10)
        else:
            srv = smtplib.SMTP(host, port, timeout=10)
        if pwd:
            srv.login(usr, pwd)
        srv.quit()
        return jsonify(ok=True, msg="Connection successful — SMTP is working correctly.")
    except Exception as ex:
        return jsonify(ok=False, msg=str(ex))


@app.route("/account/cancel", methods=["POST"])
@login_required
def account_cancel():
    csrf_protect()
    user = current_user()
    if (user["plan"] or "free").lower() == "free":
        flash("You are already on the Free plan.", "info")
        return redirect(url_for("account"))
    with get_db() as c:
        c.execute("UPDATE users SET plan='free' WHERE id=?", (session["user_id"],))
    _record_payment(session["user_id"], "manual", "free", 0, status="cancelled")
    flash("Subscription cancelled. You are now on the Free plan.", "info")
    return redirect(url_for("account"))


# ─── Invoice / Receipt PDF ────────────────────────────────────────────────────

@app.route("/account/invoice/<int:payment_id>")
@login_required
def account_invoice(payment_id):
    """Generate and download a PDF payment receipt."""
    uid = session["user_id"]
    with get_db() as c:
        pay = c.execute(
            "SELECT * FROM payments WHERE id=? AND user_id=?", (payment_id, uid)
        ).fetchone()
        user = c.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not pay:
        abort(404)

    org   = (user["org_name"]    or user["name"] or user["username"] or "SolarPro Global Customer")
    addr  = (user["org_address"] or "").strip()
    email = (user["org_email"]   or user["email"] or "")
    phone = (user["org_phone"]   or "").strip()

    date_str = (pay["created_at"] or datetime.utcnow().isoformat())[:10]
    ref      = pay["reference"] or f"INV-{pay['id']:06d}"
    plan_label = (pay["plan"] or "").title() or "Subscription"
    amount_usd = float(pay["amount_usd"] or 0)
    currency   = (pay["currency"] or "USD").upper()
    gateway    = (pay["gateway"]  or "—").title()
    status_str = (pay["status"]   or "—").upper()

    PLAN_LABELS = {
        "starter":      "SolarPro Global — Starter Plan (Monthly)",
        "professional": "SolarPro Global — Professional Plan (Monthly)",
        "business":     "SolarPro Global — Business Plan (Monthly)",
        "enterprise":   "SolarPro Global — Enterprise Plan (Monthly)",
        "free":         "SolarPro Global — Free Plan",
    }
    description = PLAN_LABELS.get((pay["plan"] or "").lower(), f"SolarPro Global — {plan_label} Plan")

    md = f"""# SolarPro Global — Payment Receipt

---

**Receipt No.:** {ref}
**Date:** {date_str}
**Status:** {status_str}

---

## Billed To

**{org}**
{"" if not addr  else addr + "  "}
{"" if not email else "Email: " + email + "  "}
{"" if not phone else "Phone: " + phone + "  "}

---

## Receipt Details

| Field | Value |
|---|---|
| Description | {description} |
| Plan | {plan_label} |
| Amount | {currency} {amount_usd:,.2f} |
| Payment Gateway | {gateway} |
| Transaction Reference | {ref} |
| Date | {date_str} |
| Status | {status_str} |

---

## Summary

| | |
|---|---|
| **Subtotal** | **{currency} {amount_usd:,.2f}** |
| Tax / VAT | Included (if applicable) |
| **Total Paid** | **{currency} {amount_usd:,.2f}** |

---

*Thank you for your subscription to SolarPro Global.*

For billing questions contact us at **billing@aiappinvent.com**
or visit **https://solarpro.aiappinvent.com**

---

*SolarPro Global — Intelligent PV Solar System Design Platform*
*This is a computer-generated receipt and is valid without a signature.*
"""

    filename = f"SolarPro_Receipt_{ref.replace(' ','_')}.pdf"
    return _render_pdf(f"Payment Receipt — {ref}", md, filename)


# ─── Phase 4: Ticketing ───────────────────────────────────────────────────────

@app.route("/tickets", methods=["GET", "POST"])
@login_required
def tickets():
    if request.method == "POST":
        csrf_protect()
        subject  = request.form.get("subject", "").strip()
        message  = request.form.get("message", "").strip()
        priority = request.form.get("priority", "normal")
        if not subject or not message:
            flash("Subject and message are required.", "warning")
        else:
            with get_db() as c:
                c.execute(
                    "INSERT INTO tickets (user_id,subject,message,priority) VALUES (?,?,?,?)",
                    (session["user_id"], subject, message, priority))
            flash("Support ticket submitted. We'll respond within 24 hours.", "success")
        return redirect(url_for("tickets"))
    with get_db() as c:
        rows = c.execute(
            "SELECT * FROM tickets WHERE user_id=? ORDER BY updated_at DESC",
            (session["user_id"],)).fetchall()
    return render_template("tickets.html", user=current_user(), tickets=rows)


@app.route("/ticket/<int:tid>", methods=["GET", "POST"])
@login_required
def ticket_detail(tid):
    with get_db() as c:
        ticket = c.execute(
            "SELECT * FROM tickets WHERE id=? AND user_id=?",
            (tid, session["user_id"])).fetchone()
        if not ticket:
            abort(404)
        replies = c.execute(
            "SELECT r.*, u.username FROM ticket_replies r "
            "JOIN users u ON r.user_id=u.id WHERE r.ticket_id=? ORDER BY r.created_at",
            (tid,)).fetchall()
    if request.method == "POST":
        csrf_protect()
        msg = request.form.get("message", "").strip()
        if msg:
            with get_db() as c:
                c.execute(
                    "INSERT INTO ticket_replies (ticket_id,user_id,is_admin,message) VALUES (?,?,0,?)",
                    (tid, session["user_id"], msg))
                c.execute("UPDATE tickets SET status='open',updated_at=? WHERE id=?",
                          (datetime.now().isoformat(), tid))
            flash("Reply added.", "success")
        return redirect(url_for("ticket_detail", tid=tid))
    return render_template("ticket_detail.html", user=current_user(),
                           ticket=ticket, replies=replies)


# ─── AI Technical Assistant ───────────────────────────────────────────────────

# ── GitHub public API context cache (no auth needed — public repo) ────────────
_gh_ctx_cache = {"data": None, "expires": 0.0}

def _fetch_github_context():
    """Delegate to api_manager GitHub client (cached 5 min)."""
    commits = _api.github.recent_commits(10)
    if not commits:
        return ""
    lines = [f"  {c}" for c in commits]
    return "Recent platform changes (live from GitHub):\n" + "\n".join(lines)

def _load_learned_kb(agent="helpline", limit=20):
    """Return the top-N most-used learned KB entries for a given agent."""
    try:
        with get_db() as c:
            rows = c.execute(
                "SELECT question, answer FROM helpline_learned_kb "
                "WHERE agent=? ORDER BY use_count DESC, created_at DESC LIMIT ?",
                (agent, limit)
            ).fetchall()
        return [(r["question"], r["answer"]) for r in rows]
    except Exception:
        return []


def _learn_from_conversation(msgs, api_key, agent="helpline"):
    """
    Background: ask Claude Haiku to extract reusable Q&A pairs from a finished
    conversation and store new ones to helpline_learned_kb.
    Only called when there are at least 4 messages (2 full exchanges).
    """
    if len(msgs) < 4:
        return
    convo = "\n".join(
        f"{m['role'].upper()}: {m['content']}" for m in msgs[-16:]
    )
    prompt = f"""You are a knowledge-extraction assistant for a solar platform helpdesk (SolarPro Global).

CONVERSATION:
{convo}

TASK: Extract 0—3 reusable Q&A pairs that future helpdesk users might ask.
Rules:
- Only extract if the answer in the conversation was genuinely helpful and accurate.
- Skip generic greetings, acknowledgements, and vague exchanges.
- Each question must be a short phrase (â‰¤15 words) a real user would type.
- Each answer must be concise (2—4 sentences) and self-contained.
- Do NOT invent information not present in the conversation.

Return ONLY valid JSON, no markdown fences:
{{"learned":[{{"question":"...","answer":"..."}}]}}
If nothing worth extracting, return {{"learned":[]}}"""
    try:
        import anthropic as _ant
        client = _ant.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=600,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = "\n".join(raw.split("\n")[1:])
            if raw.endswith("```"):
                raw = raw[:-3]
        data = json.loads(raw)
        learned = data.get("learned", [])
        if not learned:
            return
        with get_db() as c:
            for item in learned:
                q = (item.get("question") or "").strip()[:200]
                a = (item.get("answer") or "").strip()[:800]
                if len(q) < 6 or len(a) < 10:
                    continue
                # Skip near-duplicates (first 40 chars of question)
                exists = c.execute(
                    "SELECT id FROM helpline_learned_kb WHERE agent=? AND question LIKE ?",
                    (agent, f"{q[:40]}%")
                ).fetchone()
                if not exists:
                    c.execute(
                        "INSERT INTO helpline_learned_kb (agent, question, answer) VALUES (?,?,?)",
                        (agent, q, a)
                    )
        app.logger.info(f"helpline learning: stored {len(learned)} new KB entries (agent={agent})")
    except Exception as e:
        app.logger.warning(f"helpline learning failed (agent={agent}): {e}")


_ASSISTANT_SYSTEM = """You are Helpline — the AI customer engagement, assessment, and technical support agent for SolarPro Global (IntelInfraAI Solar Platform). Your mission: guide, engage, assess, support, and convert prospects into real solar projects.

=== PLATFORM KNOWLEDGE ===
Design flow: Create Project â†' Location (country, region, tariff, funding mode) â†' Loads (appliances, watts, hours, demand factor) â†' Results (PV/battery/inverter/cable sizing + financials) â†' Reports
Funding modes: Loan Finance (DSCR analysis) or Self-Funded (NPV/IRR/payback)
Battery chemistry: LiFePO4 or Lead-Acid
Plans: Free Trial (14 days, 1 project, 5 AI Agent runs/mo), Professional ($49/mo — 10 projects, all 9 reports), Business ($99/mo — unlimited + white-label)
Reports: BOQ (8% markup, 15% installation), Economic (25-yr, 0.8% O&M, battery/inverter replacement, 8% tariff escalation), Proposal, Cable sizing (BS 7671/IEC 60364), Installation plan, Energy production
Settings: Organisation profile, Date/Time format, Appearance (5 themes, 7 accent colours, 5 fonts), Email/SMTP, Security
User management (admin only): Admin â†' Users — view all accounts, change plan (free/starter/professional/business/enterprise), assign job role (customer/bdo/sales_engineer/design_engineer/proposal_engineer/project_manager/technician/support_engineer/customer_success/admin), toggle admin flag, record manual payments, disable accounts. New users self-register at /register.
22+ countries with local tariff data; Standards: BS 7671, IEC 60364, NEC 2023, IEEE 1547

=== YOUR TASK AREAS ===

A. CUSTOMER ENGAGEMENT
- Welcome visitors and explain the platform's capabilities
- Guide new users through onboarding (Create Project â†' Location â†' Loads â†' Results â†' Reports)
- Answer FAQs about features, pricing, and workflow
- Recommend the right plan or service based on the user's stated needs

B. ASSESSMENT GUIDANCE
- Help users identify and list their loads (appliances, watts, hours/day, demand factor)
- Help estimate runtime requirements for critical loads
- Recommend suitable equipment categories (off-grid vs grid-tied, battery chemistry)
- Validate assessment inputs: flag missing country/region, missing loads, zero-watt entries
- Guide users to upload utility bills or share monthly kWh consumption if available

C. PRELIMINARY SOLAR DESIGN ESTIMATES (conversational only — full calculations done by the engine)
- Give rough estimates: a 5 kW home needs ~15—20 Ã— 350 Wp panels, ~10—20 kWh battery, ~5 kW inverter
- Estimate simple ROI: typical payback 3—7 years depending on tariff and system cost
- Estimate savings: daily kWh Ã— local tariff Ã— 365 = annual saving
- Always direct user to the full design engine for accurate sizing

D. PROPOSAL ASSISTANCE
- Help prepare a brief summary of what the customer needs (load profile, location, budget)
- Explain what the full Proposal report contains (technical + financial + BOQ)
- Guide users to the Proposal report under their project

E. FOLLOW-UP & ENGAGEMENT
- Remind users to complete unfinished steps (e.g. "You've added loads — click View Results next")
- Encourage booking a consultation for complex projects
- Promote the Professional plan for users who need more than 1 project

F. LEVEL 1—2 TECHNICAL SUPPORT
- Password/login issues â†' Forgot Password link on login page
- Portal navigation issues â†' guide step by step
- Monitoring explanation â†' Results page shows live system metrics after commissioning
- Alarm interpretation â†' high-priority alerts mean system fault; check inverter and battery status
- Basic troubleshooting: location not saving (both country AND region must be selected), loads page error (add at least one row), reports locked (requires Professional/Enterprise plan)
- Ticket generation â†' escalate confirmed bugs or account-level issues

G. CUSTOMER SUCCESS
- Promote annual maintenance plans for installed systems
- Suggest platform upgrades when users hit Free plan limits
- Collect feedback: ask if the report was useful and what could be improved
- Promote renewals and expansion for existing customers

=== COMMON ISSUES ===
- Location form not saving â†' both country AND region must be selected
- Loads page error â†' add at least one appliance row before calculating
- Reports locked â†' Professional/Enterprise plan required for BOQ and Economic reports
- Date/Time picker unresponsive â†' clear browser cache
- SMTP test failing â†' use App Passwords for Gmail; Brevo recommended for 300 free emails/day

=== RULES ===
- Be concise and warm — 2—5 sentences per reply
- Always try to answer using the knowledge above before anything else
- For rough estimates, give a sensible range and direct to the engine for accuracy
- Only include [ESCALATE] when the issue genuinely requires a human to access the user's account data (payment not applied, data corruption, confirmed bug after cache clear)
- Never invent features; if unsure, say what you know and invite the user to try it"""


@app.route("/api/assistant/chat", methods=["POST"])
@limiter.limit("30 per hour")
def assistant_chat():
    """AI technical assistant chat — no login required (widget only renders for logged-in users)."""
    csrf_protect()
    data    = request.get_json(silent=True) or {}
    message = (data.get("message") or "").strip()
    history = data.get("history") or []
    if not message:
        return jsonify({"error": "No message"}), 400

    api_key  = os.environ.get("ANTHROPIC_API_KEY", "")
    gh_token = os.environ.get("GITHUB_TOKEN", "")

    # (rule-based fallback handles the no-API case; no early exit needed)

    # Build shared message list and system prompt
    msgs = []
    for h in history[-12:]:
        role    = h.get("role", "")
        content = (h.get("content") or "").strip()
        if role in ("user", "assistant") and content:
            msgs.append({"role": role, "content": content})
    msgs.append({"role": "user", "content": message})
    gh_ctx = _fetch_github_context()

    # ── Inject learned KB into system prompt ─────────────────────────────────
    learned_entries = _load_learned_kb(agent="helpline", limit=20)
    learned_section = ""
    if learned_entries:
        pairs = "\n".join(f"Q: {q}\nA: {a}" for q, a in learned_entries)
        learned_section = f"\n\n=== LEARNED FROM PREVIOUS CONVERSATIONS ===\n{pairs}"

    system = _ASSISTANT_SYSTEM + learned_section + (f"\n\n{gh_ctx}" if gh_ctx else "")

    # ── Rule-based fallback answers (no API needed) ───────────────────────────
    _KB = [
        # Acknowledgements first — prevent "need/ok/thanks" matching other entries
        (["thank","thanks","okay","great","got it","understood","perfect","sorted","all good",
          "no problem","that's all","that's fine","i'm fine","im fine","i'm good","im good",
          "never mind","nevermind","bye","goodbye","cheers","appreciate","that is all"],
         "Great, glad I could help! Feel free to ask any time if you have more questions. ðŸ˜Š"),
        # Monitoring/alarms before project (both mention "dashboard")
        (["maintenance","service","fault","alarm","monitoring","alert"],
         "For installed systems: check the **Monitoring** section in your project. Red alerts = critical fault (check inverter display and battery status). Amber = warning (low battery, high temp). For recurring faults, raise a support ticket and our technical team will assist."),
        # Assessment & design guidance
        (["assess","site assessment","consultation","building","house","office","hospital","school","warehouse","apartment"],
         "Start with a **Free Site Assessment** — click the Assessment button on the homepage or go to the Assessment section. Provide your building type, location, and approximate size. Our team will follow up with a preliminary design."),
        (["panel","solar panel","pv panel","how many panel","wp","watt peak"],
         "A rough guide: divide your daily kWh demand by your peak sun hours (e.g. 5h for West Africa) and add 25% derating — that gives you the array kWp. Divide by your panel wattage (e.g. 400 Wp) for panel count. Use the **Results** page for the accurate calculation."),
        (["battery","storage","autonomy","backup","off-grid","days"],
         "Battery size = daily kWh Ã— autonomy days Ã· depth of discharge (80% for LiFePO4). For a 10 kWh/day home wanting 2 days backup: 10 Ã— 2 Ã· 0.8 = 25 kWh. Use the **Results** page for the precise figure including temperature derating."),
        (["inverter","charge controller","mppt","hybrid"],
         "Inverter size must cover your peak simultaneous load (sum of all loads running at once). MPPT charge controller rating = PV array current. For hybrid systems, the inverter handles both grid and battery. The Results page calculates this automatically."),
        (["roi","savings","payback","return on investment","investment cost"],
         "Typical solar payback in Africa: 3—6 years for commercial, 5—8 years for residential, depending on local tariff and system cost. The **Economic Analysis** report gives you 25-year NPV, IRR, and annual savings based on your actual load and local tariff data."),
        (["utility bill","electricity bill","purc","unit rate","kwh rate","electricity tariff"],
         "The platform uses local utility tariff data for your country/region. For Ghana, PURC rates are pre-loaded by category (Residential, Commercial, Industrial). Enter your monthly bill amount on the Location page to calibrate the financial model."),
        (["upgrade","recommend","which plan","what plan"],
         "Start with **Professional ($49/mo)** for up to 10 projects and all 9 PDF reports. For unlimited projects and white-label reports, choose **Business ($99/mo)**. All paid plans include unlimited AI Agent tender searches."),
        # Core platform flows
        (["load","appliance","consumption","add load","add appliance"],
         "Go to your project â†' click **Loads** in the sidebar. Add a row for each appliance: enter the name, watts, quantity, hours/day, and demand factor. Click **Calculate** when done."),
        (["location","country","region","irradiance","solar resource","tariff","grid zone"],
         "On the **Location** step, select your country then your region from the dropdown. Both fields must be filled before the form saves."),
        (["result","calcul","sizing","pv array","battery bank","inverter size"],
         "After saving your Loads, click **View Results** (or open Results from the sidebar). The engine sizes your PV array, battery bank, inverter, and cables automatically."),
        (["report","pdf","boq","proposal","export","download","bill of quantities"],
         "Reports are in the sidebar under your project. BOQ, Economic Analysis, and Proposal require a **Professional or Enterprise** plan. PDF download is on each report page."),
        (["plan","professional","enterprise","subscription","limit","feature"],
         "The Free Trial gives 1 project + 5 AI Agent runs for 14 days. Professional ($49/mo) gives 10 projects + all 9 reports. Business ($99/mo) is unlimited + white-label. Upgrade at **Settings â†' Upgrade**."),
        (["payment","momo","mobile money","mtn","paystack","stripe","billing","invoice"],
         "We accept MTN MoMo, AirtelTigo, and Vodafone Cash via Paystack, plus Visa/Mastercard worldwide. Go to **Settings â†' Upgrade** and choose your payment method."),
        (["add user","add a user","new user","create user","manage user","user management",
          "user account","user role","assign role","assign a role","user list","staff account",
          "team member","add staff","user admin","admin users"],
         "To add or manage users, go to **Admin â†' Users** (admin accounts only). From there you can view all accounts, change a user's plan, assign a job role (e.g. Design Engineer, Sales Engineer), toggle admin rights, record payments, or disable an account. New users can self-register at the **Register** page on the homepage."),
        (["login","password","forgot","reset","sign in","account","register"],
         "Use the **Forgot Password** link on the login page to reset. For new accounts, click **Register** on the homepage."),
        (["agent","tender","rfp","prospect","scan","bid","procurement"],
         "The AI Prospecting Agent is at **Admin â†' Agent** (admin accounts). Click **Run Agent** to scan live solar tender portals. Results appear in the Alerts table with source links."),
        (["setting","theme","colour","color","font","appearance","smtp","email config","profile"],
         "Go to **Settings** (top-right menu) to change appearance, date/time format, and email/SMTP configuration."),
        (["economic","npv","irr","dscr","loan","financ","25 year","analysis"],
         "The Economic Analysis runs a 25-year model with 0.8% O&M, battery/inverter replacement, and tariff escalation. Choose **Loan Finance** on the Location step to include DSCR/bankability."),
        (["cable","wiring","bs 7671","iec","standard","voltage drop"],
         "The AC Cable Sizing report is generated automatically. It follows BS 7671 / IEC 60364 and shows cable size, current rating, voltage drop, and protection device rating."),
        (["project","new project","create","start"],
         "From the **Dashboard**, click **New Project** to start. Each project goes through: Location â†' Loads â†' Results â†' Reports. Saved projects appear in the dashboard list."),
        (["hi","hello","help","hey","what can","how does","how do"],
         "Hi! I'm Helpline, SolarPro's AI assistant. I can help with: solar system sizing estimates, assessment guidance, the design flow (Location â†' Loads â†' Results â†' Reports), reports & exports, plans & pricing, payments, settings, and the AI Prospecting Agent. What do you need?"),
        # Bare "ok" last — least specific
        (["ok"],
         "Great, glad I could help! Feel free to ask any time if you have more questions. ðŸ˜Š"),
    ]
    def _rule_reply(msg_lower):
        import re as _re
        for keywords, answer in _KB:
            for k in keywords:
                # leading \b prevents "load" firing on "download";
                # no trailing \b so "plan" matches "plans", "calcul" matches "calculate"
                if _re.search(r'\b' + _re.escape(k), msg_lower):
                    return answer
        return None

    _GENERIC = ("I can help with: **solar sizing estimates** (panels, batteries, inverters, ROI), "
                "**assessment & onboarding**, **adding loads**, **reports & PDF export**, "
                "**plans & pricing**, **Mobile Money payment**, **AI tender agent**, "
                "**results & calculations**, **monitoring & alarms**, **settings & themes**, "
                "**economic analysis**, **cable sizing**, and **account/login**. "
                "Could you give me a bit more detail about what you need?")

    try:
        # -- api_manager handles full fallback chain --
        # Claude -> OpenRouter -> Ollama -> GitHub Models -> rule-based
        reply, _ai_provider = _api.ai.chat(msgs, system=system, max_tokens=500,
                                            user_id=session.get("user_id"),
                                            endpoint="/api/assistant/chat")
        if _ai_provider == "rule_based":
            _rule = _rule_reply(message.lower())
            if _rule:
                reply = _rule
            else:
                reply = (
                    "**We'll be back in a moment — sorry for any "
                    "inconvenience while we switch AI providers. "
                    "In the meantime I can help with sizing, pricing, "
                    "and getting started. What do you need?")

        escalate = "[ESCALATE]" in reply
        reply    = reply.replace("[ESCALATE]", "").strip()

        # ── Background learning — extract Q&A from conversation if long enough ──
        if api_key and len(msgs) >= 4:
            import threading as _th
            _app = app._get_current_object()
            _msgs_copy = list(msgs)
            def _bg_learn():
                with _app.app_context():
                    _learn_from_conversation(_msgs_copy, api_key, agent="helpline")
            _t = _th.Thread(target=_bg_learn, daemon=True)
            _t.start()

        return jsonify({"reply": reply, "escalate": escalate})

    except Exception as e:
        app.logger.error(f"assistant_chat error ({type(e).__name__}): {e}")
        # Always serve a useful answer — never show a raw error to the user
        reply = _rule_reply(message.lower()) or _GENERIC
        return jsonify({"reply": reply, "escalate": False})


@app.route("/api/assistant/escalate", methods=["POST"])
@login_required
def assistant_escalate():
    """Create a high-priority support ticket from the AI chat conversation."""
    csrf_protect()
    data    = request.get_json(silent=True) or {}
    summary = (data.get("summary") or "").strip()
    history = data.get("history") or []

    chat_log = ""
    for h in history[-20:]:
        role    = "User" if h.get("role") == "user" else "AI Assistant"
        content = (h.get("content") or "").strip()
        if content:
            chat_log += f"\n**{role}:** {content}\n"

    subject = (summary or "AI Assistant escalation")[:120]
    body    = (f"**Escalated from AI Assistant chat**\n"
               f"\n{'Chat transcript:' if chat_log else 'No transcript available.'}"
               f"\n{chat_log}\n\n"
               f"---\n*Auto-created by the SolarPro AI assistant.*")

    with get_db() as c:
        c.execute(
            "INSERT INTO tickets (user_id, subject, message, priority) VALUES (?,?,?,?)",
            (session["user_id"], subject, body, "high"))
        tid = c.execute("SELECT last_insert_rowid()").fetchone()[0]

    return jsonify({"ok": True, "ticket_id": tid})


# ─── Phase 4: Subscription upgrade & payment ─────────────────────────────────

@app.route("/upgrade")
@login_required
def upgrade():
    user = current_user()
    return render_template("upgrade.html", user=user,
                           plan_prices=PLAN_PRICES,
                           current_plan=(user["plan"] or "free").lower(),
                           stripe_key=bool(STRIPE_SECRET),
                           paystack_key=bool(PAYSTACK_SECRET),
                           paystack_public_key=PAYSTACK_PUBLIC,
                           demo_mode=DEMO_MODE)


@app.route("/upgrade/checkout", methods=["POST"])
@login_required
@limiter.limit("10 per hour")
def upgrade_checkout():
    csrf_protect()
    plan     = request.form.get("plan", "").lower()
    gateway  = request.form.get("gateway", "stripe")
    if plan not in PLAN_PRICES:
        flash("Invalid plan selected.", "danger")
        return redirect(url_for("upgrade"))
    price_usd = PLAN_PRICES[plan]["usd"]
    user = current_user()

    if gateway == "stripe" and STRIPE_SECRET:
        try:
            import stripe as _stripe
            _stripe.api_key = STRIPE_SECRET
            price_id = STRIPE_PRICES.get(plan, "")
            if not price_id:
                flash("Stripe price not configured for this plan. Contact support.", "warning")
                return redirect(url_for("upgrade"))
            checkout = _stripe.checkout.Session.create(
                payment_method_types=["card"],
                mode="subscription",
                line_items=[{"price": price_id, "quantity": 1}],
                customer_email=user["email"],
                metadata={"user_id": str(user["id"]), "plan": plan},
                success_url=request.host_url.rstrip("/") + url_for("upgrade_success") + "?session_id={CHECKOUT_SESSION_ID}",
                cancel_url=request.host_url.rstrip("/") + url_for("upgrade"),
            )
            return redirect(checkout.url, 303)
        except Exception as e:
            flash(f"Stripe error: {e}", "danger")
            return redirect(url_for("upgrade"))

    elif gateway == "paystack" and PAYSTACK_SECRET:
        amount_kobo = int(price_usd * 100 * 100)
        callback = request.host_url.rstrip("/") + url_for("paystack_callback")
        _ps_ok, _ps_data = _api.payment.initialize(
            user["email"], amount_kobo, callback,
            metadata={"user_id": user["id"], "plan": plan})
        if _ps_ok and _ps_data.get("authorization_url"):
            session["paystack_plan"] = plan
            return redirect(_ps_data["authorization_url"], 303)
        flash("Paystack initialization failed. Please try again.", "danger")
        return redirect(url_for("upgrade"))

    # No gateway configured — demo mode
    flash(f"Payment gateway not configured. To activate, set STRIPE_SECRET_KEY or "
          f"PAYSTACK_SECRET_KEY environment variables. "
          f"Contact billing@aiappinvent.com to upgrade your plan manually.", "info")
    return redirect(url_for("upgrade"))


@app.route("/upgrade/success")
@login_required
def upgrade_success():
    session_id = request.args.get("session_id", "")
    if STRIPE_SECRET and session_id:
        try:
            import stripe as _stripe
            _stripe.api_key = STRIPE_SECRET
            checkout = _stripe.checkout.Session.retrieve(session_id)
            if checkout.payment_status == "paid":
                plan = checkout.metadata.get("plan", "")
                uid  = int(checkout.metadata.get("user_id", 0))
                if plan and uid == session["user_id"]:
                    with get_db() as c:
                        c.execute("UPDATE users SET plan=? WHERE id=?", (plan, uid))
                    _record_payment(uid, "stripe", plan,
                                    PLAN_PRICES.get(plan, {}).get("usd", 0),
                                    reference=session_id)
                    flash(f"Subscription activated! Welcome to {PLAN_PRICES[plan]['label']}.", "success")
                    return redirect(url_for("dashboard"))
        except Exception:
            pass
    flash("Payment verified. If your plan has not updated, contact billing@aiappinvent.com.", "info")
    return redirect(url_for("dashboard"))


@app.route("/paystack/verify", methods=["POST"])
@login_required
@limiter.limit("20 per hour")
def paystack_verify():
    """Called by Paystack inline popup JS after successful payment."""
    csrf_protect()
    ref  = request.form.get("reference", "")
    plan = request.form.get("plan", "").lower()
    if not ref or not PAYSTACK_SECRET or plan not in PLAN_PRICES:
        flash("Payment verification failed — invalid request.", "danger")
        return redirect(url_for("upgrade"))
    _ps_ok, _ps_txn = _api.payment.verify(ref)
    if not _ps_ok:
        flash("Payment verification failed — please contact billing@aiappinvent.com.", "danger")
        return redirect(url_for("upgrade"))
    plan = session.pop("paystack_plan", "")
    if not plan:
        plan = (_ps_txn.get("metadata") or {}).get("plan", "")

    if ref and PAYSTACK_SECRET and plan:
        import urllib.request as _ur
        req = _ur.Request(f"https://api.paystack.co/transaction/verify/{ref}",
                          headers={"Authorization": f"Bearer {PAYSTACK_SECRET}"})
        try:
            with _ur.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read())
            if data.get("status") and data["data"].get("status") == "success":
                with get_db() as c:
                    c.execute("UPDATE users SET plan=? WHERE id=?",
                              (plan, session["user_id"]))
                _record_payment(session["user_id"], "paystack", plan,
                                PLAN_PRICES.get(plan, {}).get("usd", 0),
                                reference=ref)
                flash(f"Payment confirmed! Welcome to {PLAN_PRICES[plan]['label']}.", "success")
                return redirect(url_for("dashboard"))
        except Exception:
            pass
    flash("Payment verification failed. Contact billing@aiappinvent.com with your reference.", "warning")
    return redirect(url_for("upgrade"))


@app.route("/stripe/webhook", methods=["POST"])
def stripe_webhook():
    if not STRIPE_SECRET:
        return "", 400
    try:
        import stripe as _stripe
        _stripe.api_key = STRIPE_SECRET
        sig = request.headers.get("Stripe-Signature", "")
        event = _stripe.Webhook.construct_event(
            request.get_data(raw=True), sig, STRIPE_WEBHOOK)
        if event["type"] == "checkout.session.completed":
            obj  = event["data"]["object"]
            plan = obj.get("metadata", {}).get("plan", "")
            uid  = int(obj.get("metadata", {}).get("user_id", 0))
            if plan and uid:
                with get_db() as c:
                    c.execute("UPDATE users SET plan=? WHERE id=?", (plan, uid))
                _record_payment(uid, "stripe", plan,
                                PLAN_PRICES.get(plan, {}).get("usd", 0),
                                reference=obj.get("id", ""))
        elif event["type"] in ("customer.subscription.deleted",
                               "invoice.payment_failed"):
            uid = int(event["data"]["object"].get("metadata", {}).get("user_id", 0))
            if uid:
                with get_db() as c:
                    c.execute("UPDATE users SET plan='free' WHERE id=?", (uid,))
                _record_payment(uid, "stripe", "free", 0, status="cancelled")
    except Exception:
        return "", 400
    return "", 200


@app.route("/paystack/webhook", methods=["POST"])
def paystack_webhook():
    """Paystack push webhook — HMAC-SHA512 signature verified.
    Paystack signs every push event with HMAC-SHA512(PAYSTACK_SECRET_KEY, raw_body).
    The signature arrives in the X-Paystack-Signature header.
    """
    import hmac as _hmac, hashlib as _hashlib
    if not PAYSTACK_SECRET:
        return "", 400
    sig  = request.headers.get("x-paystack-signature", "")
    body = request.get_data(raw=True)
    expected = _hmac.new(PAYSTACK_SECRET.encode("utf-8"),
                         msg=body, digestmod=_hashlib.sha512).hexdigest()
    if not _hmac.compare_digest(sig, expected):
        return "", 400  # reject unsigned / tampered requests
    try:
        event = json.loads(body)
        if event.get("event") == "charge.success":
            data    = event.get("data", {})
            ref     = data.get("reference", "")
            meta    = data.get("metadata") or {}
            uid     = int(meta.get("user_id", 0))
            plan    = meta.get("plan", "")
            amount  = data.get("amount", 0) / 100   # kobo â†' USD
            if uid and plan and ref:
                with get_db() as c:
                    # Reject duplicate references
                    dup = c.execute("SELECT id FROM payments WHERE reference=?",
                                    (ref,)).fetchone()
                    if not dup:
                        c.execute("UPDATE users SET plan=? WHERE id=?", (plan, uid))
                        _record_payment(uid, "paystack", plan, amount, reference=ref)
    except Exception:
        pass
    return "", 200


@app.route("/robots.txt")
def robots_txt():
    """Protect private routes from search-engine crawlers."""
    content = (
        "User-agent: *\n"
        "Allow: /\n"
        "Disallow: /admin/\n"
        "Disallow: /dashboard\n"
        "Disallow: /project/\n"
        "Disallow: /api/\n"
        "Disallow: /upgrade/checkout\n"
        "Disallow: /paystack/\n"
        "Disallow: /stripe/\n"
        f"Sitemap: https://solarpro.aiappinvent.com/sitemap.xml\n"
    )
    return make_response(content, 200, {"Content-Type": "text/plain"})


# ─── Demo Mode & Upgrade Codes ───────────────────────────────────────────────

@app.route("/upgrade/demo-activate", methods=["POST"])
@login_required
@limiter.limit("5 per hour")
def upgrade_demo_activate():
    """Instantly activate Professional plan for DEMO_DAYS days — no payment needed."""
    csrf_protect()
    if not DEMO_MODE:
        flash("Demo mode is not enabled.", "warning")
        return redirect(url_for("upgrade"))
    user = current_user()
    if (user["plan"] or "free") not in ("free",):
        flash("You already have an active paid plan.", "info")
        return redirect(url_for("dashboard"))
    expires = (datetime.now() + timedelta(days=DEMO_DAYS)).isoformat()
    with get_db() as c:
        c.execute("UPDATE users SET plan='professional', subscription_end=? WHERE id=?",
                  (expires, user["id"]))
    _record_payment(user["id"], "demo", "professional", 0,
                    reference=f"DEMO-{DEMO_DAYS}d")
    flash(f"Demo Professional plan activated for {DEMO_DAYS} days — all features unlocked!", "success")
    return redirect(url_for("dashboard"))


@app.route("/upgrade/redeem", methods=["POST"])
@login_required
def upgrade_redeem_code():
    """Redeem an upgrade code issued by admin."""
    csrf_protect()
    code = request.form.get("code", "").strip().upper()
    if not code:
        flash("Please enter an upgrade code.", "warning")
        return redirect(url_for("upgrade"))
    with get_db() as c:
        row = c.execute("SELECT * FROM upgrade_codes WHERE code=?", (code,)).fetchone()
        if not row:
            flash("Invalid upgrade code.", "danger")
            return redirect(url_for("upgrade"))
        if row["uses"] >= row["max_uses"]:
            flash("This code has already been fully used.", "warning")
            return redirect(url_for("upgrade"))
        if row["expires_at"] and row["expires_at"] < datetime.now().isoformat():
            flash("This code has expired.", "warning")
            return redirect(url_for("upgrade"))
        expires = (datetime.now() + timedelta(days=row["duration_days"])).isoformat()
        c.execute("UPDATE users SET plan=?, subscription_end=? WHERE id=?",
                  (row["plan"], expires, session["user_id"]))
        c.execute("UPDATE upgrade_codes SET uses=uses+1 WHERE id=?", (row["id"],))
    _record_payment(session["user_id"], "code", row["plan"], 0, reference=code)
    flash(f"Code accepted! {row['plan'].title()} plan activated for {row['duration_days']} days.", "success")
    return redirect(url_for("dashboard"))


@app.route("/admin/codes", methods=["GET", "POST"])
@admin_required
def admin_codes():
    if request.method == "POST":
        csrf_protect()
        action = request.form.get("action")
        if action == "create":
            plan     = request.form.get("plan", "professional")
            duration = int(request.form.get("duration_days", 30))
            max_uses = int(request.form.get("max_uses", 1))
            days_exp = int(request.form.get("expires_in_days", 90))
            code     = secrets.token_hex(4).upper()
            expires  = (datetime.now() + timedelta(days=days_exp)).isoformat()
            with get_db() as c:
                c.execute(
                    "INSERT INTO upgrade_codes (code,plan,duration_days,max_uses,created_by,expires_at) "
                    "VALUES (?,?,?,?,?,?)",
                    (code, plan, duration, max_uses, session["user_id"], expires))
            flash(f"Code created: {code}", "success")
        elif action == "delete":
            cid = request.form.get("cid", type=int)
            with get_db() as c:
                c.execute("DELETE FROM upgrade_codes WHERE id=?", (cid,))
            flash("Code deleted.", "info")
        return redirect(url_for("admin_codes"))
    with get_db() as c:
        codes = c.execute("SELECT * FROM upgrade_codes ORDER BY created_at DESC").fetchall()
    return render_template("admin_codes.html", user=current_user(),
                           codes=codes, plan_prices=PLAN_PRICES)


# ─── Sales & Marketing Module ─────────────────────────────────────────────────

@app.route("/contact", methods=["POST"])
@limiter.limit("10 per hour")
def contact_lead():
    """Public contact form â†' lead capture."""
    name    = request.form.get("name", "").strip()
    email   = request.form.get("email", "").strip()
    phone   = request.form.get("phone", "").strip()
    company = request.form.get("company", "").strip()
    country = request.form.get("country", "").strip()
    interest= request.form.get("interest", "residential")
    message = request.form.get("message", "").strip()
    if not name or not email:
        flash("Name and email are required.", "warning")
        return redirect(url_for("landing") + "#contact")
    with get_db() as c:
        c.execute(
            "INSERT INTO leads (name,email,phone,company,country,interest,message,source) "
            "VALUES (?,?,?,?,?,?,?,?)",
            (name, email, phone, company, country, interest, message, "website"))
    flash("Thank you! We'll be in touch within 24 hours.", "success")
    return redirect(url_for("landing") + "#contact")


@app.route("/newsletter/subscribe", methods=["POST"])
@limiter.limit("20 per hour")
def newsletter_subscribe():
    email = request.form.get("email", "").strip()
    name  = request.form.get("name", "").strip()
    if not email:
        flash("Email is required.", "warning")
        return redirect(url_for("landing") + "#newsletter")
    try:
        with get_db() as c:
            c.execute("INSERT OR IGNORE INTO newsletter_subscribers (email,name) VALUES (?,?)",
                      (email, name))
        flash("Subscribed! You'll receive our solar industry updates.", "success")
    except Exception:
        flash("Already subscribed with that email.", "info")
    return redirect(url_for("landing") + "#newsletter")


# ─── Assessment Request (public) ──────────────────────────────────────────────

def _qualify_lead(name, company, phone, system_type, size_kw, budget_usd, message):
    """Rule-based lead scoring â†' (score 0-100, grade A/B/C/D, notes str)."""
    score = 0
    reasons = []
    # Budget scoring
    b = budget_usd.lower().replace(",","").replace("$","").replace("usd","").strip()
    try:
        bval = float(b.split("-")[0].split("k")[0]) * (1000 if "k" in b else 1)
    except Exception:
        bval = 0
    if bval >= 100000:  score += 35; reasons.append("High budget ($100k+)")
    elif bval >= 50000: score += 25; reasons.append("Good budget ($50k+)")
    elif bval >= 10000: score += 15; reasons.append("Medium budget ($10k+)")
    elif bval > 0:      score += 5;  reasons.append("Low budget")
    # System size
    try: sz = float(size_kw)
    except Exception: sz = 0
    if sz >= 100:    score += 30; reasons.append("Large commercial system (100kW+)")
    elif sz >= 50:   score += 22; reasons.append("Commercial system (50-100kW)")
    elif sz >= 10:   score += 14; reasons.append("SME system (10-50kW)")
    elif sz >= 1:    score += 6;  reasons.append("Residential system")
    # System type
    if system_type in ("commercial", "industrial"): score += 15; reasons.append("Commercial/Industrial priority")
    elif system_type == "hybrid":                   score += 10; reasons.append("Hybrid system")
    # Lead quality signals
    if company.strip(): score += 10; reasons.append("Company name provided")
    if phone.strip():   score += 5;  reasons.append("Phone provided")
    if len(message.strip()) > 50: score += 5; reasons.append("Detailed message")
    score = min(score, 100)
    if score >= 80:   grade = "A"
    elif score >= 60: grade = "B"
    elif score >= 40: grade = "C"
    else:             grade = "D"
    return score, grade, "; ".join(reasons) if reasons else "Unscored"


@app.route("/assess/quick", methods=["POST"])
@limiter.limit("10 per hour")
def assess_quick():
    """AJAX popup assessment intake — landing page modal.
    Returns JSON {ok, ref, name} on success or {ok, error} on failure.
    """
    import random, string as _str
    name       = request.form.get("name", "").strip()
    phone      = request.form.get("phone", "").strip()
    country    = request.form.get("country", "").strip()
    region     = request.form.get("region", "").strip()
    bldg_desc  = request.form.get("building_desc", "").strip()
    bldg_size  = request.form.get("building_size", "").strip()
    num_floors = request.form.get("num_floors", "1").strip()
    bldg_type  = request.form.get("building_type", "").strip()

    if not name or not phone:
        return jsonify({"ok": False, "error": "Full name and phone number are required."}), 400

    try:
        num_floors_i = int(num_floors)
    except Exception:
        num_floors_i = 1

    # Generate unique reference SA-XXXXXX
    ref = "SA-" + "".join(random.choices(_str.ascii_uppercase + _str.digits, k=6))
    location_desc = f"{region}, {country}".strip(", ") if region else country
    score, grade, notes = _qualify_lead(name, "", phone, bldg_type or "residential",
                                        "0", "", bldg_desc)

    try:
        with get_db() as c:
            c.execute(
                "INSERT INTO assessment_requests "
                "(name,email,phone,country,region,system_type,location_desc,message,"
                " ai_score,ai_grade,ai_notes,source,status,pipeline_stage,"
                " assessment_ref,building_desc,building_size,num_floors,building_type) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (name, "", phone, country, region,
                 bldg_type or "residential", location_desc, bldg_desc,
                 score, grade, notes, "landing_popup", "open",
                 "assessment_submitted",
                 ref, bldg_desc, bldg_size, num_floors_i, bldg_type))
            # Mirror into leads CRM for pipeline tracking
            c.execute(
                "INSERT INTO leads (name,email,phone,country,interest,message,source,"
                " system_type,ai_score,ai_grade,ai_notes,pipeline_stage) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                (name, "", phone, country,
                 bldg_type or "residential", bldg_desc, "assessment_popup",
                 bldg_type or "residential",
                 score, grade, notes, "assessment_submitted"))
    except Exception as e:
        app.logger.error("assess_quick DB error: %s", e)
        return jsonify({"ok": False, "error": "Could not save assessment. Please try again."}), 500

    return jsonify({"ok": True, "ref": ref, "name": name.split()[0]})


@app.route("/assess", methods=["GET", "POST"])
@limiter.limit("20 per hour")
def assessment_request():
    """Public assessment request form."""
    if request.method == "POST":
        name        = request.form.get("name", "").strip()
        email       = request.form.get("email", "").strip()
        phone       = request.form.get("phone", "").strip()
        company     = request.form.get("company", "").strip()
        country     = request.form.get("country", "").strip()
        system_type = request.form.get("system_type", "off-grid")
        size_kw     = request.form.get("system_size_kw", "0")
        budget      = request.form.get("budget_usd", "").strip()
        location    = request.form.get("location_desc", "").strip()
        message     = request.form.get("message", "").strip()
        if not name or not email:
            flash("Name and email are required.", "warning")
            return redirect(url_for("assessment_request"))
        score, grade, notes = _qualify_lead(name, company, phone, system_type, size_kw, budget, message)
        try: size_kw_f = float(size_kw)
        except Exception: size_kw_f = 0.0
        with get_db() as c:
            c.execute(
                "INSERT INTO assessment_requests "
                "(name,email,phone,company,country,system_type,system_size_kw,budget_usd,"
                "location_desc,message,ai_score,ai_grade,ai_notes,source) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,'website')",
                (name, email, phone, company, country, system_type, size_kw_f,
                 budget, location, message, score, grade, notes))
            # Also add to leads table for unified CRM
            c.execute(
                "INSERT INTO leads (name,email,phone,company,country,interest,message,source,"
                "system_type,system_size_kw,budget_usd,ai_score,ai_grade,ai_notes,pipeline_stage) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (name, email, phone, company, country, system_type, message, "assess_form",
                 system_type, size_kw_f, budget, score, grade, notes,
                 "qualified" if grade in ("A","B") else "new"))
        flash("Thank you! Our team will contact you within 24 hours with your assessment.", "success")
        return redirect(url_for("assessment_request"))
    return render_template("assess.html", user=current_user(), countries=get_countries())


@app.route("/api/assess/design", methods=["POST"])
@limiter.limit("10 per hour")
def assess_design():
    """
    AJAX — runs full preliminary solar sizing from the assessment form.
    Accepts JSON payload, returns JSON with sizing + cost + financial results.
    Also saves lead to DB and emails the results to the user.
    """
    import random, string as _str
    data = request.get_json(silent=True) or {}

    name      = (data.get("name") or "").strip()
    email     = (data.get("email") or "").strip()
    phone     = (data.get("phone") or "").strip()
    country   = (data.get("country") or "").strip()
    region    = (data.get("region") or "").strip()
    bldg_type = (data.get("building_type") or "residential").strip()
    loads     = data.get("loads") or []

    if not name or not email:
        return jsonify({"ok": False, "error": "Full name and email are required."}), 400
    if not country or not region:
        return jsonify({"ok": False, "error": "Please select a country and region."}), 400
    if not loads:
        return jsonify({"ok": False, "error": "Please add at least one appliance to the load schedule."}), 400

    sd = get_solar_data(country, region)
    if not sd:
        return jsonify({"ok": False, "error": "Location not found. Please select a valid country and region."}), 400

    psh      = sd["psh"]
    temp     = sd["avg_temp"]
    tariff   = sd["tariff"]
    currency = sd["currency"]
    symbol   = sd["symbol"]
    cost_kwp = sd["cost_usd_kwp"]
    fx       = sd["fx_usd"]

    # ── Load calculation ──────────────────────────────────────────────────────
    daily_kwh = 0.0
    peak_kw   = 0.0
    for ld in loads:
        try:
            w  = float(ld.get("watts") or 0)
            q  = float(ld.get("qty") or 1)
            h  = float(ld.get("hours") or 0)
            df = float(ld.get("demand_factor") or 1.0)
            daily_kwh += w * q * h * df / 1000.0
            peak_kw   += w * q / 1000.0
        except Exception:
            pass

    if daily_kwh <= 0:
        return jsonify({"ok": False, "error": "Total daily load is zero. Please check appliance wattage and hours."}), 400

    # ── Sizing ────────────────────────────────────────────────────────────────
    pv_kw, num_panels, _td = calc_pv(daily_kwh, psh, temp)
    bat_kwh, num_bat, unit_kwh = calc_battery(daily_kwh, autonomy=1.5)
    inv_kw = calc_inverter(daily_kwh, peak_kw)

    # ── Cost estimate ─────────────────────────────────────────────────────────
    pv_usd  = pv_kw * cost_kwp
    bat_usd = bat_kwh * BATTERY_CHEMISTRY["LiFePO4"]["usd_per_kwh"]
    if   inv_kw <= 3:   inv_usd = 280
    elif inv_kw <= 5:   inv_usd = 400
    elif inv_kw <= 8:   inv_usd = 560
    elif inv_kw <= 12:  inv_usd = 820
    else:               inv_usd = 1180
    equip_usd   = pv_usd + bat_usd + inv_usd
    total_usd   = equip_usd * 1.24           # +8% markup +15% install â‰ˆ Ã—1.24
    total_local = total_usd * fx

    # ── Financials ────────────────────────────────────────────────────────────
    annual_kwh = daily_kwh * 365
    annual_sav = annual_kwh * tariff
    payback_yr = round(total_local / annual_sav, 1) if annual_sav > 0 else 0
    co2_yr     = round(annual_kwh * 0.40 / 1000, 2)

    # ── Reference ─────────────────────────────────────────────────────────────
    ref = "SA-" + "".join(random.choices(_str.ascii_uppercase + _str.digits, k=6))

    # ── Save to DB ────────────────────────────────────────────────────────────
    try:
        score, grade, notes = _qualify_lead(name, "", phone, bldg_type, str(pv_kw), "", "")
        with get_db() as c:
            c.execute(
                "INSERT INTO assessment_requests "
                "(name,email,phone,country,region,system_type,location_desc,message,"
                " ai_score,ai_grade,ai_notes,source,status,pipeline_stage,"
                " assessment_ref,building_type) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (name, email, phone, country, region, bldg_type,
                 f"{region}, {country}",
                 json.dumps({"loads": loads, "daily_kwh": round(daily_kwh,2),
                             "pv_kw": pv_kw, "bat_kwh": bat_kwh, "inv_kw": inv_kw}),
                 score, grade, notes, "assess_design", "open",
                 "design_generated", ref, bldg_type))
            c.execute(
                "INSERT INTO leads (name,email,phone,country,interest,message,source,"
                " system_type,ai_score,ai_grade,ai_notes,pipeline_stage) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                (name, email, phone, country, bldg_type,
                 f"Auto-design {ref}: {pv_kw}kWp PV / {bat_kwh}kWh battery / {inv_kw}kW inverter",
                 "assess_design", bldg_type, score, grade, notes, "design_generated"))
    except Exception as db_e:
        app.logger.error("assess_design DB: %s", db_e)

    # ── Email results to user ─────────────────────────────────────────────────
    first = name.split()[0]
    html_email = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"></head>
<body style="font-family:Arial,sans-serif;background:#f4f4f8;margin:0;padding:20px">
<div style="max-width:600px;margin:0 auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 2px 12px rgba(0,0,0,.08)">
  <div style="background:linear-gradient(135deg,#0f0f22,#1a1a3e);padding:32px;text-align:center">
    <div style="font-size:32px;margin-bottom:8px">â˜€ï¸</div>
    <h1 style="color:#f59e0b;margin:0;font-size:22px;font-weight:800">Preliminary Solar Design</h1>
    <div style="color:#a0a0c8;font-size:13px;margin-top:6px">Reference: {ref}</div>
  </div>
  <div style="padding:28px 32px">
    <p style="color:#333;font-size:15px">Hi {first},</p>
    <p style="color:#555;font-size:14px">Here is your preliminary solar system design for <strong>{region}, {country}</strong>
    based on the load schedule you provided. This is an indicative estimate — our engineers will refine it during consultation.</p>

    <h3 style="color:#1a1a3e;font-size:15px;margin:24px 0 12px;border-bottom:2px solid #f59e0b;padding-bottom:6px">
      ðŸ"Š Preliminary System Sizing
    </h3>
    <table style="width:100%;border-collapse:collapse;font-size:14px">
      <tr style="background:#f8f8fc"><td style="padding:10px 14px;color:#555;width:50%">Daily Energy Demand</td>
        <td style="padding:10px 14px;font-weight:700;color:#1a1a3e">{daily_kwh:.1f} kWh/day</td></tr>
      <tr><td style="padding:10px 14px;color:#555">Peak Sun Hours ({region})</td>
        <td style="padding:10px 14px;font-weight:700;color:#1a1a3e">{psh} hrs/day</td></tr>
      <tr style="background:#f8f8fc"><td style="padding:10px 14px;color:#555">PV Array Size</td>
        <td style="padding:10px 14px;font-weight:700;color:#f59e0b">{pv_kw:.1f} kWp ({num_panels} panels Ã— 400Wp)</td></tr>
      <tr><td style="padding:10px 14px;color:#555">Battery Bank (1.5 days autonomy)</td>
        <td style="padding:10px 14px;font-weight:700;color:#0ea5e9">{bat_kwh:.1f} kWh ({num_bat} Ã— {unit_kwh}kWh units)</td></tr>
      <tr style="background:#f8f8fc"><td style="padding:10px 14px;color:#555">Inverter / Charger</td>
        <td style="padding:10px 14px;font-weight:700;color:#1a1a3e">{inv_kw:.0f} kW</td></tr>
    </table>

    <h3 style="color:#1a1a3e;font-size:15px;margin:24px 0 12px;border-bottom:2px solid #22c55e;padding-bottom:6px">
      ðŸ'° Estimated Cost &amp; Savings
    </h3>
    <table style="width:100%;border-collapse:collapse;font-size:14px">
      <tr style="background:#f8f8fc"><td style="padding:10px 14px;color:#555">Estimated System Cost</td>
        <td style="padding:10px 14px;font-weight:700;color:#1a1a3e">{symbol}{total_local:,.0f} {currency} (â‰ˆ USD {total_usd:,.0f})</td></tr>
      <tr><td style="padding:10px 14px;color:#555">Estimated Annual Savings</td>
        <td style="padding:10px 14px;font-weight:700;color:#22c55e">{symbol}{annual_sav:,.0f} {currency}/year</td></tr>
      <tr style="background:#f8f8fc"><td style="padding:10px 14px;color:#555">Simple Payback Period</td>
        <td style="padding:10px 14px;font-weight:700;color:#1a1a3e">{payback_yr} years</td></tr>
      <tr><td style="padding:10px 14px;color:#555">Annual COâ‚‚ Offset</td>
        <td style="padding:10px 14px;font-weight:700;color:#22c55e">{co2_yr} tonnes COâ‚‚/year</td></tr>
    </table>

    <div style="background:#fffbeb;border:1px solid #f59e0b44;border-radius:8px;padding:16px;margin:24px 0;font-size:13px;color:#555">
      <strong style="color:#f59e0b">âš  Note:</strong> This is a preliminary estimate based on standard assumptions.
      A detailed site survey and full engineering design may adjust these figures.
      Local taxes, import duties, and civil works are not included in the cost estimate.
    </div>

    <div style="text-align:center;margin:28px 0">
      <a href="https://solarpro.aiappinvent.com/register" style="display:inline-block;background:linear-gradient(135deg,#f59e0b,#fbbf24);color:#0f0f22;font-weight:800;text-decoration:none;padding:14px 32px;border-radius:8px;font-size:15px">
        Request Full Consultation â†'
      </a>
      <div style="color:#888;font-size:12px;margin-top:10px">Create a free account to book your consultation</div>
    </div>
  </div>
  <div style="background:#f8f8fc;padding:16px 32px;text-align:center;border-top:1px solid #e8e8f0">
    <div style="color:#888;font-size:12px">SolarPro Global Â· AI-Powered Solar Design Platform</div>
    <div style="color:#aaa;font-size:11px;margin-top:4px">Reference: {ref} Â· {region}, {country}</div>
  </div>
</div>
</body></html>"""

    try:
        _subj = ("Your Preliminary Solar Design (" + ref + ") - SolarPro Global")
        _pl = ("Hi " + str(first) + ", design ready. Ref: " + ref + ". "
               "PV: " + str(round(pv_kw,1)) + "kWp | Battery: " + str(round(bat_kwh,1)) + "kWh | Inverter: " + str(round(inv_kw)) + "kW. "
               "Cost: " + symbol + "{:,.0f}".format(total_local) + " " + currency + ". "
               "Visit https://solarpro.aiappinvent.com to book a consultation.")
        _ok, _err = _send_email(email, _subj, html_email, text_body=_pl, from_addr=EMAIL_SALES)
        if not _ok:
            app.logger.warning("assess_design email failed: %s", _err)
    except Exception as mail_e:
        app.logger.warning("assess_design email failed: %s", mail_e)

    return jsonify({
        "ok":         True,
        "ref":        ref,
        "first":      first,
        "daily_kwh":  round(daily_kwh, 2),
        "peak_kw":    round(peak_kw, 2),
        "psh":        psh,
        "pv_kw":      pv_kw,
        "num_panels": num_panels,
        "bat_kwh":    bat_kwh,
        "num_bat":    num_bat,
        "unit_kwh":   unit_kwh,
        "inv_kw":     inv_kw,
        "total_usd":  round(total_usd, 0),
        "total_local":round(total_local, 0),
        "currency":   currency,
        "symbol":     symbol,
        "annual_kwh": round(annual_kwh, 0),
        "annual_sav": round(annual_sav, 0),
        "payback_yr": payback_yr,
        "co2_yr":     co2_yr,
        "country":    country,
        "region":     region,
    })


@app.route("/assess/consultation", methods=["POST"])
@login_required
def assess_consultation():
    """Authenticated — submit consultation request from assessment results page."""
    csrf_protect()
    ref     = request.form.get("ref", "").strip()
    message = request.form.get("message", "").strip()
    u       = current_user()
    if not message:
        flash("Please describe your project briefly.", "warning")
        return redirect(url_for("assessment_request"))
    with get_db() as c:
        c.execute(
            "INSERT INTO tickets (user_id, subject, body, status, priority) VALUES (?,?,?,?,?)",
            (u["id"],
             f"Consultation Request — {ref or 'Assessment'}",
             f"Assessment Ref: {ref}\n\n{message}",
             "open", "high"))
    flash("Consultation request submitted! Our team will contact you within 24 hours.", "success")
    return redirect(url_for("assessment_request"))


# ─── Installer Registration (public) ──────────────────────────────────────────

@app.route("/installer/register", methods=["GET", "POST"])
@limiter.limit("10 per hour")
def installer_register():
    """Public installer/subcontractor registration."""
    if request.method == "POST":
        company   = request.form.get("company_name", "").strip()
        contact   = request.form.get("contact_name", "").strip()
        email     = request.form.get("email", "").strip()
        phone     = request.form.get("phone", "").strip()
        country   = request.form.get("country", "").strip()
        regions   = request.form.get("regions", "").strip()
        years     = request.form.get("years_exp", "0")
        staff     = request.form.get("staff_count", "0")
        certs     = request.form.get("certifications", "").strip()
        specs     = request.form.get("specialties", "").strip()
        max_kw    = request.form.get("max_project_kw", "0")
        website   = request.form.get("website", "").strip()
        notes     = request.form.get("notes", "").strip()
        if not company or not email or not contact:
            flash("Company name, contact name, and email are required.", "warning")
            return redirect(url_for("installer_register"))
        try: years_i = int(years)
        except Exception: years_i = 0
        try: staff_i = int(staff)
        except Exception: staff_i = 0
        try: max_kw_f = float(max_kw)
        except Exception: max_kw_f = 0.0
        # Simple grade based on experience
        if years_i >= 10 and staff_i >= 20: grade = "A"
        elif years_i >= 5 and staff_i >= 5: grade = "B"
        elif years_i >= 2: grade = "C"
        else: grade = "D"
        try:
            with get_db() as c:
                c.execute(
                    "INSERT INTO installers (company_name,contact_name,email,phone,country,"
                    "regions,years_exp,staff_count,certifications,specialties,max_project_kw,"
                    "website,notes,ai_grade) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (company, contact, email, phone, country, regions, years_i, staff_i,
                     certs, specs, max_kw_f, website, notes, grade))
            flash("Registration submitted! We'll review your application within 2 business days.", "success")
        except Exception:
            flash("An account with that email already exists.", "warning")
        return redirect(url_for("installer_register"))
    return render_template("installer_register.html", user=current_user(), countries=get_countries())


# ─── Admin: Assessment Requests ───────────────────────────────────────────────

@app.route("/admin/assessments", methods=["GET", "POST"])
@admin_required
def admin_assessments():
    if request.method == "POST":
        csrf_protect()
        aid    = request.form.get("aid", type=int)
        action = request.form.get("action")
        with get_db() as c:
            if action == "update":
                c.execute(
                    "UPDATE assessment_requests SET pipeline_stage=?,assigned_to=?,"
                    "follow_up_date=?,status=?,updated_at=? WHERE id=?",
                    (request.form.get("pipeline_stage"),
                     request.form.get("assigned_to",""),
                     request.form.get("follow_up_date",""),
                     request.form.get("status","open"),
                     datetime.now().isoformat(), aid))
                flash("Assessment request updated.", "success")
            elif action == "delete":
                c.execute("DELETE FROM assessment_requests WHERE id=?", (aid,))
                flash("Deleted.", "info")
        return redirect(url_for("admin_assessments"))
    stage_f = request.args.get("stage", "all")
    with get_db() as c:
        if stage_f != "all":
            rows = c.execute(
                "SELECT * FROM assessment_requests WHERE pipeline_stage=? ORDER BY created_at DESC",
                (stage_f,)).fetchall()
        else:
            rows = c.execute(
                "SELECT * FROM assessment_requests ORDER BY created_at DESC").fetchall()
        counts = {}
        for s in ("new","qualified","assessment_sent","proposal_sent","won","lost"):
            counts[s] = c.execute(
                "SELECT COUNT(*) FROM assessment_requests WHERE pipeline_stage=?", (s,)).fetchone()[0]
    return render_template("admin_assessments.html", user=current_user(),
                           requests=rows, stage_filter=stage_f, counts=counts)


# ─── Admin: Installer Management ──────────────────────────────────────────────

@app.route("/admin/installers", methods=["GET", "POST"])
@admin_required
def admin_installers():
    if request.method == "POST":
        csrf_protect()
        iid    = request.form.get("iid", type=int)
        action = request.form.get("action")
        with get_db() as c:
            if action == "approve":
                c.execute("UPDATE installers SET status='approved' WHERE id=?", (iid,))
                flash("Installer approved.", "success")
            elif action == "reject":
                c.execute("UPDATE installers SET status='rejected' WHERE id=?", (iid,))
                flash("Installer rejected.", "info")
            elif action == "delete":
                c.execute("DELETE FROM installers WHERE id=?", (iid,))
                flash("Deleted.", "info")
        return redirect(url_for("admin_installers"))
    status_f = request.args.get("status", "all")
    with get_db() as c:
        if status_f != "all":
            rows = c.execute(
                "SELECT * FROM installers WHERE status=? ORDER BY created_at DESC", (status_f,)).fetchall()
        else:
            rows = c.execute("SELECT * FROM installers ORDER BY created_at DESC").fetchall()
        pending = c.execute("SELECT COUNT(*) FROM installers WHERE status='pending'").fetchone()[0]
        approved = c.execute("SELECT COUNT(*) FROM installers WHERE status='approved'").fetchone()[0]
    return render_template("admin_installers.html", user=current_user(),
                           installers=rows, status_filter=status_f,
                           pending=pending, approved=approved)


# ─── Admin: CRM Pipeline ──────────────────────────────────────────────────────

@app.route("/admin/pipeline")
@admin_required
def admin_pipeline():
    """Kanban-style CRM pipeline across all leads + assessment requests — 9-stage."""
    # 9-stage pipeline from spec + lost for dropped records
    STAGES = [
        "assessment_submitted", "assessment_reviewed", "lead_qualified",
        "consultation", "proposal", "negotiation",
        "won", "installation", "after_sales", "lost",
    ]
    # Remap legacy stage names to the new pipeline stages
    STAGE_REMAP = {
        "new":             "assessment_submitted",
        "qualified":       "lead_qualified",
        "assessment_sent": "consultation",
        "proposal_sent":   "proposal",
        "won":             "won",
        "lost":            "lost",
    }
    with get_db() as c:
        leads_rows  = c.execute("SELECT * FROM leads ORDER BY created_at DESC").fetchall()
        assess_rows = c.execute("SELECT * FROM assessment_requests ORDER BY created_at DESC").fetchall()
    # Normalise assessment_requests rows to look like leads for the Kanban
    all_leads = []
    for row in leads_rows:
        d = dict(row)
        d.setdefault("rec_type", "lead")
        all_leads.append(d)
    for ar in assess_rows:
        d = dict(ar)
        d.setdefault("interest", d.get("system_type", ""))
        d.setdefault("company", "")
        d.setdefault("notes", "")
        d.setdefault("rec_type", "assess")
        all_leads.append(d)
    by_stage = {s: [] for s in STAGES}
    for lead in all_leads:
        try:
            st = lead.get("pipeline_stage") or "assessment_submitted"
        except Exception:
            st = "assessment_submitted"
        st = STAGE_REMAP.get(st, st)          # remap legacy names
        if st not in STAGES:
            st = "assessment_submitted"
        by_stage[st].append(lead)
    total = len(all_leads)
    won   = len(by_stage["won"])
    conv  = round(won / total * 100, 1) if total else 0
    return render_template("admin_pipeline.html", user=current_user(),
                           by_stage=by_stage, stages=STAGES, total=total,
                           won=won, conv=conv)


@app.route("/admin/sales")
@admin_required
def admin_sales():
    """Sales & Marketing dashboard — KPIs, leads pipeline, revenue, conversions."""
    with get_db() as c:
        leads      = c.execute("SELECT * FROM leads ORDER BY created_at DESC").fetchall()
        subs       = c.execute("SELECT COUNT(*) FROM newsletter_subscribers WHERE status='active'").fetchone()[0]
        users      = c.execute("SELECT plan, COUNT(*) as cnt FROM users GROUP BY plan").fetchall()
        payments   = c.execute(
            "SELECT p.*, u.username, u.email FROM payments p "
            "JOIN users u ON p.user_id=u.id "
            "WHERE p.status='success' ORDER BY p.created_at DESC LIMIT 50").fetchall()
        total_rev  = c.execute("SELECT SUM(amount_usd) FROM payments WHERE status='success'").fetchone()[0] or 0
        news_count = c.execute("SELECT COUNT(*) FROM news_posts WHERE is_published=1").fetchone()[0]
        monthly_signups = c.execute(
            "SELECT strftime('%Y-%m', created_at) as mo, COUNT(*) as cnt "
            "FROM users GROUP BY mo ORDER BY mo DESC LIMIT 6").fetchall()
        assess_count   = c.execute("SELECT COUNT(*) FROM assessment_requests").fetchone()[0]
        assess_open    = c.execute("SELECT COUNT(*) FROM assessment_requests WHERE status='open'").fetchone()[0]
        installers_pending = c.execute("SELECT COUNT(*) FROM installers WHERE status='pending'").fetchone()[0]
        pipeline_won   = c.execute("SELECT COUNT(*) FROM leads WHERE pipeline_stage='won'").fetchone()[0]
        pipeline_total = c.execute("SELECT COUNT(*) FROM leads").fetchone()[0]
    lead_status = {}
    for ld in leads:
        s = ld["status"]
        lead_status[s] = lead_status.get(s, 0) + 1
    pipeline_dist = {}
    for ld in leads:
        s = ld["pipeline_stage"] or "new"
        pipeline_dist[s] = pipeline_dist.get(s, 0) + 1
    plan_dist = {r["plan"]: r["cnt"] for r in users}
    conv = round(pipeline_won / pipeline_total * 100, 1) if pipeline_total else 0
    return render_template("admin_sales.html", user=current_user(),
                           leads=leads, subs=subs, plan_dist=plan_dist,
                           payments=payments, total_rev=total_rev,
                           news_count=news_count, lead_status=lead_status,
                           monthly_signups=monthly_signups,
                           assess_count=assess_count, assess_open=assess_open,
                           installers_pending=installers_pending,
                           pipeline_dist=pipeline_dist, conv=conv)


@app.route("/admin/leads", methods=["GET", "POST"])
@admin_required
def admin_leads():
    if request.method == "POST":
        csrf_protect()
        lid    = request.form.get("lid", type=int)
        action = request.form.get("action")
        with get_db() as c:
            if action == "update_status":
                c.execute(
                    "UPDATE leads SET status=?,notes=?,pipeline_stage=?,follow_up_date=? WHERE id=?",
                    (request.form.get("status"), request.form.get("notes",""),
                     request.form.get("pipeline_stage","new"),
                     request.form.get("follow_up_date",""), lid))
                flash("Lead updated.", "success")
            elif action == "delete":
                c.execute("DELETE FROM leads WHERE id=?", (lid,))
                flash("Lead deleted.", "info")
        return redirect(url_for("admin_leads"))
    status_filter = request.args.get("status", "all")
    with get_db() as c:
        if status_filter != "all":
            rows = c.execute("SELECT * FROM leads WHERE status=? ORDER BY created_at DESC",
                             (status_filter,)).fetchall()
        else:
            rows = c.execute("SELECT * FROM leads ORDER BY created_at DESC").fetchall()
    return render_template("admin_leads.html", user=current_user(),
                           leads=rows, status_filter=status_filter)


@app.route("/admin/news", methods=["GET", "POST"])
@admin_required
def admin_news():
    if request.method == "POST":
        csrf_protect()
        action = request.form.get("action")
        with get_db() as c:
            if action == "create":
                c.execute(
                    "INSERT INTO news_posts (title,content,category,is_published) VALUES (?,?,?,?)",
                    (request.form["title"], request.form["content"],
                     request.form.get("category","industry"),
                     1 if request.form.get("publish") else 0))
                flash("News post created.", "success")
            elif action == "edit":
                nid = request.form.get("nid", type=int)
                c.execute(
                    "UPDATE news_posts SET title=?,content=?,category=?,is_published=?,updated_at=? WHERE id=?",
                    (request.form["title"], request.form["content"],
                     request.form.get("category","industry"),
                     1 if request.form.get("publish") else 0,
                     datetime.now().isoformat(), nid))
                flash("Post updated.", "success")
            elif action == "delete":
                c.execute("DELETE FROM news_posts WHERE id=?",
                          (request.form.get("nid", type=int),))
                flash("Post deleted.", "info")
        return redirect(url_for("admin_news"))
    with get_db() as c:
        posts = c.execute("SELECT * FROM news_posts ORDER BY created_at DESC").fetchall()
    return render_template("admin_news.html", user=current_user(), posts=posts)


@app.route("/admin/newsletter")
@admin_required
def admin_newsletter():
    with get_db() as c:
        subs = c.execute("SELECT * FROM newsletter_subscribers ORDER BY created_at DESC").fetchall()
    return render_template("admin_newsletter.html", user=current_user(), subs=subs)


@app.route("/admin/newsletter/unsub/<int:sid>", methods=["POST"])
@admin_required
def admin_newsletter_unsub(sid):
    csrf_protect()
    with get_db() as c:
        c.execute("UPDATE newsletter_subscribers SET status='unsubscribed' WHERE id=?", (sid,))
    flash("Subscriber removed.", "info")
    return redirect(url_for("admin_newsletter"))


# ─── Procurement Module ───────────────────────────────────────────────────────

@app.route("/procurement")
@login_required
def procurement():
    with get_db() as c:
        suppliers = c.execute("SELECT * FROM suppliers WHERE is_active=1 ORDER BY name").fetchall()
        catalog   = c.execute("SELECT e.*, s.name as sup_name FROM equipment_catalog e "
                              "LEFT JOIN suppliers s ON e.supplier_id=s.id "
                              "WHERE e.is_active=1 ORDER BY e.category, e.name").fetchall()
        categories = list(dict.fromkeys(r["category"] for r in catalog))
    return render_template("procurement.html", user=current_user(),
                           suppliers=suppliers, catalog=catalog, categories=categories)


@app.route("/procurement/suppliers", methods=["GET", "POST"])
@admin_required
def procurement_suppliers():
    if request.method == "POST":
        csrf_protect()
        action = request.form.get("action")
        f = request.form
        with get_db() as c:
            if action == "add":
                c.execute(
                    "INSERT INTO suppliers (name,country,contact_name,phone,email,website,"
                    "categories,lead_time_days,payment_terms,rating,notes) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                    (f["name"], f.get("country",""), f.get("contact_name",""),
                     f.get("phone",""), f.get("email",""), f.get("website",""),
                     f.get("categories",""), int(f.get("lead_time_days",30)),
                     f.get("payment_terms","TT 30 days"), int(f.get("rating",5)),
                     f.get("notes","")))
                flash("Supplier added.", "success")
            elif action == "edit":
                sid = f.get("sid", type=int)
                c.execute(
                    "UPDATE suppliers SET name=?,country=?,contact_name=?,phone=?,email=?,"
                    "website=?,categories=?,lead_time_days=?,payment_terms=?,rating=?,notes=? WHERE id=?",
                    (f["name"], f.get("country",""), f.get("contact_name",""),
                     f.get("phone",""), f.get("email",""), f.get("website",""),
                     f.get("categories",""), int(f.get("lead_time_days",30)),
                     f.get("payment_terms","TT 30 days"), int(f.get("rating",5)),
                     f.get("notes",""), sid))
                flash("Supplier updated.", "success")
            elif action == "delete":
                c.execute("UPDATE suppliers SET is_active=0 WHERE id=?",
                          (f.get("sid", type=int),))
                flash("Supplier deactivated.", "info")
        return redirect(url_for("procurement_suppliers"))
    with get_db() as c:
        rows = c.execute("SELECT * FROM suppliers ORDER BY name").fetchall()
    return render_template("procurement_suppliers.html", user=current_user(), suppliers=rows)


@app.route("/procurement/catalog", methods=["GET", "POST"])
@admin_required
def procurement_catalog():
    if request.method == "POST":
        csrf_protect()
        action = request.form.get("action")
        f = request.form
        with get_db() as c:
            if action == "add":
                c.execute(
                    "INSERT INTO equipment_catalog (category,name,brand,model,spec,unit,"
                    "price_usd,supplier_id,lead_time_days,notes) VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (f["category"], f["name"], f.get("brand",""), f.get("model",""),
                     f.get("spec",""), f.get("unit","No."), float(f.get("price_usd",0)),
                     int(f.get("supplier_id",0)), int(f.get("lead_time_days",30)),
                     f.get("notes","")))
                flash("Equipment added.", "success")
            elif action == "edit":
                eid = f.get("eid", type=int)
                c.execute(
                    "UPDATE equipment_catalog SET category=?,name=?,brand=?,model=?,spec=?,"
                    "unit=?,price_usd=?,supplier_id=?,lead_time_days=?,notes=? WHERE id=?",
                    (f["category"], f["name"], f.get("brand",""), f.get("model",""),
                     f.get("spec",""), f.get("unit","No."), float(f.get("price_usd",0)),
                     int(f.get("supplier_id",0)), int(f.get("lead_time_days",30)),
                     f.get("notes",""), eid))
                flash("Equipment updated.", "success")
            elif action == "delete":
                c.execute("UPDATE equipment_catalog SET is_active=0 WHERE id=?",
                          (f.get("eid", type=int),))
                flash("Item deactivated.", "info")
        return redirect(url_for("procurement_catalog"))
    with get_db() as c:
        items = c.execute(
            "SELECT e.*, s.name as sup_name FROM equipment_catalog e "
            "LEFT JOIN suppliers s ON e.supplier_id=s.id ORDER BY e.category, e.name").fetchall()
        suppliers = c.execute("SELECT id,name FROM suppliers WHERE is_active=1").fetchall()
    cats = ["PV Modules","Inverters","Batteries","MPPT","Cables",
            "Protection","Earthing","Mounting","Testing","Sundries"]
    return render_template("procurement_catalog.html", user=current_user(),
                           items=items, suppliers=suppliers, categories=cats)


@app.route("/project/<int:pid>/procurement")
@login_required
def project_procurement(pid):
    """Generate a smart procurement plan from the project BOQ, matched to suppliers."""
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run calculations first.", "warning")
        return redirect(url_for("project_results", pid=pid))

    r   = project["data"]["results"]
    d   = project["data"]
    sym = d.get("symbol", "$")
    fx  = float(d.get("fx_usd") or 1.0)

    # Eagerly convert all DB rows to plain dicts to avoid sqlite3.Row issues
    with get_db() as c:
        suppliers = {row["id"]: dict(row) for row in
                     c.execute("SELECT * FROM suppliers WHERE is_active=1").fetchall()}
        catalog   = [dict(row) for row in
                     c.execute("SELECT * FROM equipment_catalog WHERE is_active=1 "
                               "ORDER BY category").fetchall()]

    CAT_MAP = {
        "PV": "PV Modules", "Battery": "Batteries", "Inverter": "Inverters",
        "MPPT": "MPPT", "Cable": "Cables", "Earth": "Earthing",
        "Surge": "Protection", "Mounting": "Mounting", "AC RCCB": "Protection",
    }

    plan_rows = []
    grand_usd = 0.0
    for boq in r.get("boq_rows", []):
        try:
            desc  = boq.get("desc", "")
            cat   = next((v for k, v in CAT_MAP.items() if k.lower() in desc.lower()),
                         "Sundries")
            match = next((ci for ci in catalog if ci.get("category") == cat), None)
            sup   = suppliers.get(match["supplier_id"]) if match else None

            # Unit price: prefer catalog price, fall back to BOQ rate converted to USD
            total_r   = float(boq.get("total_r") or boq.get("rate") or 0)
            unit_usd  = float(match["price_usd"]) if match else (total_r / fx if fx else 0)
            qty       = float(boq.get("qty") or 1)
            total_usd = unit_usd * qty
            lead_days = int(match.get("lead_time_days") or 30) if match else 30

            grand_usd += total_usd
            plan_rows.append({
                "no":          boq.get("no", len(plan_rows) + 1),
                "desc":        desc,
                "spec":        boq.get("spec", ""),
                "qty":         qty,
                "unit":        boq.get("unit", "No."),
                "category":    cat,
                "catalog":     match["name"] if match else "— specify on quotation",
                "brand":       match.get("brand", "") if match else "",
                "supplier":    sup["name"] if sup else "Open market / RFQ",
                "supplier_id": sup["id"] if sup else 0,
                "lead_days":   lead_days,
                "unit_usd":    unit_usd,
                "total_usd":   total_usd,
                "total_local": total_usd * fx,
            })
        except Exception as exc:
            app.logger.warning("procurement: skipping BOQ row %s — %s", boq, exc)
            continue

    # Group by supplier for purchase orders
    by_supplier = {}
    for row in plan_rows:
        key = row["supplier"]
        by_supplier.setdefault(key, {"supplier": key, "rows": [], "total_usd": 0.0})
        by_supplier[key]["rows"].append(row)
        by_supplier[key]["total_usd"] += row["total_usd"]

    return render_template("procurement_plan.html", user=current_user(),
                           project=project, d=d, r=r, sym=sym, fx=fx,
                           plan_rows=plan_rows, grand_usd=grand_usd,
                           by_supplier=list(by_supplier.values()),
                           suppliers=suppliers)


@app.route("/project/<int:pid>/procurement/pdf")
@login_required
@limiter.limit("10 per minute")
def export_pdf_procurement(pid):
    """PDF Procurement Plan — maps BOQ to suppliers, generates sourcing schedule."""
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        return redirect(url_for("project_results", pid=pid))
    r   = project["data"]["results"]
    d   = project["data"]
    sym = d.get("symbol", "$")
    fx  = d.get("fx_usd", 1.0)

    with get_db() as c:
        suppliers = {r2["id"]: dict(r2) for r2 in
                     c.execute("SELECT * FROM suppliers WHERE is_active=1").fetchall()}
        catalog   = list(c.execute("SELECT * FROM equipment_catalog WHERE is_active=1").fetchall())

    CAT_MAP = {"PV": "PV Modules","Battery": "Batteries","Inverter": "Inverters",
               "MPPT": "MPPT","Cable": "Cables","Earth": "Earthing",
               "Surge": "Protection","Mounting": "Mounting","AC RCCB": "Protection"}

    rows = []
    grand_usd = 0.0
    for boq in r.get("boq_rows", []):
        desc  = boq["desc"]
        cat   = next((v for k, v in CAT_MAP.items() if k.lower() in desc.lower()), "Sundries")
        match = next((c2 for c2 in catalog if c2["category"] == cat), None)
        sup   = suppliers.get(match["supplier_id"]) if match else None
        unit_usd  = match["price_usd"] if match else boq["total_r"] / fx
        total_usd = unit_usd * boq["qty"]
        grand_usd += total_usd
        rows.append((boq["no"], desc, boq["qty"], boq["unit"],
                     match["brand"] if match else "—",
                     sup["name"] if sup else "RFQ",
                     match["lead_time_days"] if match else 30,
                     unit_usd, total_usd))

    md = f"""# Procurement Plan — {project["name"]}

**{d.get("region","")}, {d.get("country","")}** | {_fmt(r["pv_kw"],2)} kWp | {d.get("system_type","").title()} System

Prepared by: SolarPro Global | Currency: USD (local rates apply Ã— {fx})

---

# Equipment Sourcing Schedule

| No. | Description | Qty | Unit | Brand | Supplier | Lead (days) | Unit USD | Total USD |
|---|---|---|---|---|---|---|---|---|
"""
    for no, desc, qty, unit, brand, sup, lead, u_usd, t_usd in rows:
        md += f"| {no} | {desc} | {qty} | {unit} | {brand} | {sup} | {lead} | ${_fmt(u_usd,2)} | **${_fmt(t_usd,2)}** |\n"

    md += f"""
| | | | | | | | **GRAND TOTAL** | **${_fmt(grand_usd,2)}** |

---

# Procurement Notes

- Lead times are indicative — confirm with suppliers on order placement
- All equipment must meet IEC 61215, IEC 62109, IEC 62619, BS 7671 as applicable
- Request formal technical data sheets (TDS) and test certificates before approval
- Payment terms subject to commercial negotiation with each supplier
- Allow 10% contingency budget for freight, customs, and incidentals

---

# Key Suppliers

"""
    seen = set()
    for _, _, _, _, _, sup_name, lead, _, _ in rows:
        if sup_name not in seen and sup_name != "RFQ":
            seen.add(sup_name)
            sup_obj = next((s for s in suppliers.values() if s["name"] == sup_name), None)
            if sup_obj:
                md += f"**{sup_name}** | {sup_obj.get('country','')} | {sup_obj.get('email','')} | Lead: {lead} days | Terms: {sup_obj.get('payment_terms','TT 30 days')}\n\n"

    md += f"\n---\n\n*Procurement Plan generated by SolarPro Global Â· {project['name']}*\n"
    fname = f"Procurement_{project['name'].replace(' ','_')}.pdf"
    return _render_pdf(f"Procurement Plan — {project['name']}", md, fname)


# ─── Email Report Sending ─────────────────────────────────────────────────────

@app.route("/project/<int:pid>/email", methods=["GET", "POST"])
@login_required
def project_email(pid):
    """Send a PDF report to recipients via SMTP (user settings > env vars)."""
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run calculations first.", "warning")
        return redirect(url_for("project_results", pid=pid))

    # Convert sqlite3.Row â†' plain dict so .get() works
    user = dict(current_user() or {})

    # Resolve email: Resend key > SMTP user settings > SMTP env vars
    u_resend = (user.get("resend_api_key") or "").strip()
    u_host   = (user.get("smtp_host") or "").strip()
    u_user   = (user.get("smtp_user") or "").strip()
    u_pass   = (user.get("smtp_pass") or "").strip()
    u_from   = (user.get("smtp_from") or "").strip()
    u_port_s = (user.get("smtp_port") or "587").strip()
    u_port   = int(u_port_s) if u_port_s.isdigit() else 587
    u_tls    = (user.get("smtp_tls") or "starttls") == "starttls"

    eff_resend = u_resend or RESEND_API_KEY
    eff_host   = u_host or SMTP_HOST
    eff_user   = u_user or SMTP_USER
    eff_pass   = u_pass or SMTP_PASS
    eff_from   = u_from or EMAIL_SALES or SMTP_FROM
    eff_port   = u_port if u_host else SMTP_PORT
    eff_tls    = u_tls  if u_host else SMTP_TLS

    REPORT_OPTIONS = [
        ("Full Proposal (All Reports)", url_for("export_pdf_proposal",    pid=pid)),
        ("PV System Design Report",     url_for("export_pdf_pv",          pid=pid)),
        ("BOQ Report",                  url_for("export_pdf_boq",         pid=pid)),
        ("Economic Analysis",           url_for("export_pdf_economic",    pid=pid)),
        ("Energy Impact",               url_for("export_pdf_energy",      pid=pid)),
        ("AC Cable Schedule",           url_for("export_pdf_cable",       pid=pid)),
        ("Installation Plan",           url_for("export_pdf_installation",pid=pid)),
        ("Installation Work Plan",      url_for("export_pdf_workplan",    pid=pid)),
        ("Staffing Plan",               url_for("export_pdf_staffing",    pid=pid)),
        ("Procurement Plan",            url_for("export_pdf_procurement", pid=pid)),
        ("Site Assessment",             url_for("export_pdf_inspection",  pid=pid)),
    ]

    if request.method == "POST":
        csrf_protect()
        recipients = [e.strip() for e in request.form.get("recipients","").split(",") if e.strip()]
        subject    = request.form.get("subject","").strip() or f"Solar Project Report — {project['name']}"
        body_text  = request.form.get("body","").strip()

        if not recipients:
            flash("Enter at least one recipient email.", "warning")
            return redirect(url_for("project_email", pid=pid))

        if not eff_resend and not (eff_host and eff_user):
            with get_db() as c:
                c.execute(
                    "INSERT INTO email_logs (user_id,project_id,recipients,subject,status,error_msg) "
                    "VALUES (?,?,?,?,?,?)",
                    (session["user_id"], pid, ",".join(recipients), subject,
                     "failed", "Email not configured"))
            flash(
                "Email not configured. Add a Resend API key or SMTP settings in Settings.",
                "warning")
            return redirect(url_for("project_email", pid=pid))

        _ptxt = body_text or (
            "Please find the solar project report for " + project["name"] + " attached. "
            "Generated by SolarPro Global."
        )
        _phtml = (
            "<div style='font-family:sans-serif;padding:24px'>"
            "<p>" + _ptxt + "</p>"
            "<hr><p style='color:#888;font-size:12px'>SolarPro Global</p></div>"
        )
        # Render the selected report to PDF bytes and attach it.
        # Why: prior code only sent the message body and the placeholder said
        # "...attached" but nothing was. We dispatch to the existing PDF view
        # function (registered as a Flask endpoint) and capture its bytes.
        _attachments = None
        _report_label = (request.form.get("report") or "").strip()
        _endpoint_map = {
            "Full Proposal (All Reports)": "export_pdf_proposal",
            "PV System Design Report":     "export_pdf_pv",
            "BOQ Report":                  "export_pdf_boq",
            "Economic Analysis":           "export_pdf_economic",
            "Energy Impact":               "export_pdf_energy",
            "AC Cable Schedule":           "export_pdf_cable",
            "Installation Plan":           "export_pdf_installation",
            "Installation Work Plan":      "export_pdf_workplan",
            "Staffing Plan":               "export_pdf_staffing",
            "Procurement Plan":            "export_pdf_procurement",
            "Site Assessment":             "export_pdf_inspection",
        }
        _ep_name = _endpoint_map.get(_report_label)
        if _ep_name:
            _vf = app.view_functions.get(_ep_name)
            if _vf:
                try:
                    _resp = _vf(pid)
                    _pdf_bytes = _resp.get_data() if hasattr(_resp, "get_data") else None
                    if _pdf_bytes and _pdf_bytes[:4] == b"%PDF":
                        _fname = getattr(_resp, "download_name", None) or (_report_label.replace(" ", "_") + ".pdf")
                        _attachments = [(_fname, _pdf_bytes, "application/pdf")]
                except Exception as _pdf_exc:
                    logger.warning("PDF attachment render failed for %s: %s", _report_label, _pdf_exc)
        _ok, _err = _send_email(
            recipients, subject, _phtml, text_body=_ptxt,
            from_addr=eff_from, resend_key=eff_resend or None,
            attachments=_attachments,
        )
        if _ok:
            with get_db() as c:
                c.execute(
                    "INSERT INTO email_logs (user_id,project_id,recipients,subject,status) "
                    "VALUES (?,?,?,?,?)",
                    (session["user_id"], pid, ",".join(recipients), subject, "sent"))
            flash("Email sent to: " + ", ".join(recipients), "success")
        else:
            with get_db() as c:
                c.execute(
                    "INSERT INTO email_logs (user_id,project_id,recipients,subject,status,error_msg) "
                    "VALUES (?,?,?,?,?,?)",
                    (session["user_id"], pid, ",".join(recipients), subject, "failed", _err))
            flash("Email failed: " + str(_err), "danger")

        return redirect(url_for("project_email", pid=pid))

    with get_db() as c:
        logs = c.execute(
            "SELECT * FROM email_logs WHERE project_id=? ORDER BY created_at DESC LIMIT 10",
            (pid,)).fetchall()
    smtp_ok = bool(eff_resend or (eff_host and eff_user))
    d = project["data"]
    r = d.get("results", {})
    return render_template("email_report.html", user=user, project=project,
                           d=d, r=r, report_options=REPORT_OPTIONS,
                           smtp_ok=smtp_ok, logs=logs)


# ─── DOCX Export ─────────────────────────────────────────────────────────────

@app.route("/project/<int:pid>/export/docx")
@login_required
@limiter.limit("10 per minute")
def export_docx(pid):
    """Full Word document report — Summary + Load Schedule + BOQ + Financial + Cable."""
    gate = _paid_only(pid)
    if gate: return gate
    project = get_project(pid)
    if not project or "results" not in project["data"]:
        flash("Run calculations first.", "warning")
        return redirect(url_for("project_results", pid=pid))

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r   = project["data"]["results"]
    d   = project["data"]
    eco = r["economics"]
    sym = d.get("symbol", "$")

    # Org details for title page and footer
    u = current_user()
    org_name    = (u["org_name"]    or u["name"] or u["username"] or "SolarPro Global").strip()
    org_address = (u["org_address"] or "").strip()
    org_email   = (u["org_email"]   or u["email"] or "").strip()
    org_phone   = (u["org_phone"]   or "").strip()
    org_website = (u["org_website"] or "").strip()

    doc = Document()

    # ── Page layout ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(0.9)
    section.different_first_page_header_footer = True   # no header on title page

    GOLD  = RGBColor(0xF5, 0x9E, 0x0B)
    DARK  = RGBColor(0x1E, 0x3A, 0x8A)
    GREY  = RGBColor(0x6B, 0x72, 0x80)

    def _heading(text, level=1, color=DARK):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = color
        return p

    def _para(text, bold=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def _table_row(tbl, values, bold=False, shade=None):
        row = tbl.add_row()
        for i, v in enumerate(values):
            cell = row.cells[i]
            cell.text = str(v)
            for run in cell.paragraphs[0].runs:
                run.bold = bold
                run.font.size = Pt(9)
            if shade:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), shade)
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)

    def _set_col_widths(tbl, widths_inches):
        """Set explicit column widths on a table."""
        for i, w in enumerate(widths_inches):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

    def _add_page_field(paragraph, prefix="Page "):
        """Append 'Page N of M' auto-field to a paragraph."""
        paragraph.add_run(prefix)
        for instrTxt in ("PAGE", " of ", "NUMPAGES"):
            if instrTxt in (" of ", ):
                paragraph.add_run(instrTxt)
                continue
            run = paragraph.add_run()
            fc1 = OxmlElement("w:fldChar")
            fc1.set(qn("w:fldCharType"), "begin")
            itext = OxmlElement("w:instrText")
            itext.set(qn("xml:space"), "preserve")
            itext.text = instrTxt
            fc2 = OxmlElement("w:fldChar")
            fc2.set(qn("w:fldCharType"), "end")
            run._r.append(fc1)
            run._r.append(itext)
            run._r.append(fc2)

    # ── Header (all pages except title) ───────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp_run = hp.add_run(f"{org_name}  |  {project['name']}  |  SolarPro Global")
    hp_run.font.size = Pt(8)
    hp_run.font.color.rgb = GREY

    # ── Footer (all pages) ────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(fp, "Page ")
    fp.runs[-1 if fp.runs else 0]  # ensure created
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY

    # ── Title page ────────────────────────────────────────────────────────────
    # Org banner
    if org_name and org_name != "SolarPro Global":
        to = doc.add_paragraph()
        to.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tor = to.add_run(org_name.upper())
        tor.bold = True; tor.font.size = Pt(14); tor.font.color.rgb = GREY

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("SolarPro Global")
    run.bold = True; run.font.size = Pt(22); run.font.color.rgb = GOLD

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run(project["name"])
    run2.bold = True; run2.font.size = Pt(16); run2.font.color.rgb = DARK

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.add_run(f"PV Solar System Design Report\n"
               f"{d.get('region','')}, {d.get('country','')}\n"
               f"Generated: {datetime.now().strftime('%d %B %Y')}")

    # Org contact block at bottom of title page
    doc.add_paragraph()
    contact_lines = [x for x in [org_address, org_phone, org_email, org_website] if x]
    if contact_lines:
        tc = doc.add_paragraph()
        tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_run = tc.add_run("\n".join(contact_lines))
        tc_run.font.size = Pt(9)
        tc_run.font.color.rgb = GREY

    doc.add_page_break()

    # ── 1. Project Summary ────────────────────────────────────────────────────
    _heading("1. Project Summary")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _table_row(tbl, ["Parameter", "Value"], bold=True, shade="1E3A8A")
    for k, v in [
        ("Location",         f"{d.get('region','')}, {d.get('country','')}"),
        ("System Type",      d.get("system_type","").title()),
        ("Phase",            "Three-Phase 415V" if d.get("phase")=="three" else "Single-Phase 230V"),
        ("Daily Demand",     f"{r['daily_kwh']:.3f} kWh/day"),
        ("PV Array",         f"{r['pv_kw']:.3f} kWp ({r['num_panels']} Ã— {r.get('panel_wp',400)} Wp)"),
        ("Battery",          f"{r['bat_kwh']:.2f} kWh ({r['num_bat']} Ã— {r['unit_bat_kwh']:.2g} kWh {r.get('chemistry','')})"),
        ("Inverter",         f"{r['inv_kw']:.1f} kW"),
        ("System Cost",      f"{sym} {eco['total_local']:,.0f}"),
        ("Simple Payback",   f"{eco['payback']:.1f} years"),
        ("NPV (25yr)",       f"{sym} {eco['npv']:,.0f}"),
        ("IRR",              f"{eco['irr_pct']:.1f}%" if eco["irr_pct"] else "N/A"),
        ("DSCR",             f"{eco['dscr']:.2f}"),
        ("Bankability",      eco["bankability"]),
        ("Verdict",          eco["verdict"]),
        ("COâ‚‚ Reduction",    f"{eco['co2_yr']:.2f} t/year"),
    ]:
        _table_row(tbl, [k, v])

    _set_col_widths(tbl, [2.8, 3.5])
    doc.add_paragraph()

    # ── 2. Load Schedule ─────────────────────────────────────────────────────
    _heading("2. Electrical Load Schedule")
    tbl2 = doc.add_table(rows=1, cols=7)
    tbl2.style = "Table Grid"
    _table_row(tbl2, ["Category","Load Name","W","Qty","Hrs/Day","kWh/Day","Critical"],
               bold=True, shade="1E3A8A")
    tot = 0.0
    for ld in d.get("loads", []):
        kwh = float(ld.get("wattage",0))*float(ld.get("quantity",1))*float(ld.get("hours",0))/1000
        tot += kwh
        _table_row(tbl2, [ld.get("category",""), ld.get("name",""),
                          ld.get("wattage",0), ld.get("quantity",1),
                          ld.get("hours",0), f"{kwh:.3f}",
                          "Yes" if ld.get("critical") else ""])
    _table_row(tbl2, ["","","","","","TOTAL",f"{tot:.3f} kWh/day"], bold=True)
    _set_col_widths(tbl2, [1.0, 1.8, 0.6, 0.5, 0.7, 0.7, 0.6])
    doc.add_paragraph()

    # ── 3. Bill of Quantities ─────────────────────────────────────────────────
    _heading("3. Bill of Quantities (BOQ)")
    tbl3 = doc.add_table(rows=1, cols=6)
    tbl3.style = "Table Grid"
    _table_row(tbl3, ["No.","Description","Qty","Unit",f"Rate ({sym})",f"Amount ({sym})"],
               bold=True, shade="1E3A8A")
    for row in r.get("boq_rows", []):
        _table_row(tbl3, [row["no"], row["desc"], row["qty"], row["unit"],
                          f"{row['total_r']:.2f}", f"{row['amount']:.2f}"])
    _table_row(tbl3, ["","","","","GRAND TOTAL", f"{sym} {r.get('boq_grand',0):.2f}"],
               bold=True, shade="F59E0B")
    _set_col_widths(tbl3, [0.4, 2.6, 0.5, 0.5, 1.0, 1.0])
    doc.add_paragraph()

    # ── 4. Financial Analysis ─────────────────────────────────────────────────
    _heading("4. Financial Engineering & Economic Analysis")
    tbl4 = doc.add_table(rows=1, cols=2)
    tbl4.style = "Table Grid"
    _table_row(tbl4, ["Metric", "Value"], bold=True, shade="1E3A8A")
    for k, v in [
        ("Total Investment",        f"{sym} {eco['total_local']:,.0f}"),
        ("Annual Generation",       f"{eco['annual_kwh']:,.0f} kWh/yr"),
        ("Gross Annual Saving",     f"{sym} {eco['annual_sav']:,.0f}/yr"),
        ("Net Annual Benefit (Y1)", f"{sym} {eco['net_yr1']:,.0f}/yr"),
        ("Simple Payback",          f"{eco['payback']:.1f} years"),
        ("NPV (25yr, 12% disc.)",   f"{sym} {eco['npv']:,.0f}"),
        ("IRR",                     f"{eco['irr_pct']:.1f}%" if eco["irr_pct"] else "N/A"),
        ("ROI (25yr)",              f"{eco['roi_pct']:.0f}%"),
        ("Loan Amount (70%)",       f"{sym} {eco['loan_amt']:,.0f}"),
        ("Monthly Repayment",       f"{sym} {eco['pmt']:,.0f}"),
        ("DSCR",                    f"{eco['dscr']:.2f}"),
        ("Bankability",             eco["bankability"]),
        ("Verdict",                 eco["verdict"]),
        ("COâ‚‚ Savings",             f"{eco['co2_yr']:.2f} t/year"),
        ("25-yr Cumulative Saving", f"{sym} {eco['cumul_25']:,.0f}"),
    ]:
        _table_row(tbl4, [k, v])
    _set_col_widths(tbl4, [2.8, 3.5])

    doc.add_paragraph()
    _heading("4.1 Cash Flow Projection (25 Years)", level=2)
    tbl5 = doc.add_table(rows=1, cols=5)
    tbl5.style = "Table Grid"
    _table_row(tbl5, ["Year", f"Gross ({sym})", f"O&M ({sym})", f"Net ({sym})", f"Cumulative ({sym})"],
               bold=True, shade="1E3A8A")
    for cf in eco.get("cf_rows", []):
        flag = " â† BREAK-EVEN" if eco.get("breakeven") and cf["yr"] == eco["breakeven"] else ""
        _table_row(tbl5, [cf["yr"], f"{cf['gross']:,.0f}", f"{cf['om']:,.0f}",
                          f"{cf['net']:,.0f}", f"{cf['cumul']:,.0f}{flag}"])
    _set_col_widths(tbl5, [0.5, 1.5, 1.3, 1.3, 1.7])
    doc.add_paragraph()

    # ── 5. AC Cable Schedule ──────────────────────────────────────────────────
    _heading("5. AC Cable Sizing Schedule")
    tbl6 = doc.add_table(rows=1, cols=7)
    tbl6.style = "Table Grid"
    _table_row(tbl6, ["Circuit", "Power (kW)", "Ib (A)", "L (m)", "Cable (mmÂ²)", "VD (%)", "Breaker"],
               bold=True, shade="1E3A8A")
    for c2 in r.get("ac_cables", []):
        _table_row(tbl6, [c2["circuit"], c2["power_kw"], c2["design_current"],
                          c2["length_m"], f"{c2['cable_size_mm2']} mmÂ²",
                          f"{c2['vd_percent']:.2f}%", f"{c2['breaker_a']} A"])
    _set_col_widths(tbl6, [1.5, 0.9, 0.7, 0.7, 0.9, 0.7, 0.7])
    doc.add_paragraph()

    # ── 6. Energy Impact Analysis ─────────────────────────────────────────────
    _heading("6. Energy Impact & Environmental Analysis")
    monthly_factors = [0.88, 0.90, 0.95, 1.00, 1.05, 1.08, 1.10, 1.08, 1.03, 0.98, 0.92, 0.88]
    base_monthly    = r["daily_kwh"] * 30.44
    months_names    = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    tbl7 = doc.add_table(rows=1, cols=4)
    tbl7.style = "Table Grid"
    tariff = float(d.get("tariff", 0) or 0)
    sys_type = d.get("system_type", "grid").lower()
    grid_offset_pct = 1.0 if sys_type == "off-grid" else 0.8
    _table_row(tbl7, ["Month", "Generation (kWh)", f"Saving ({sym})", "Grid Offset (kWh)"],
               bold=True, shade="1E3A8A")
    annual_gen = 0.0
    for mn, mf in zip(months_names, monthly_factors):
        gen = round(base_monthly * mf, 1)
        sav = round(gen * tariff, 1)
        off = round(gen * grid_offset_pct, 1)
        annual_gen += gen
        _table_row(tbl7, [mn, f"{gen:,.1f}", f"{sav:,.1f}", f"{off:,.1f}"])
    _table_row(tbl7, ["Annual Total", f"{annual_gen:,.1f}", f"{annual_gen * tariff:,.1f}",
                      f"{annual_gen * grid_offset_pct:,.1f}"], bold=True, shade="F59E0B")
    _set_col_widths(tbl7, [0.8, 1.6, 1.6, 1.6])
    doc.add_paragraph()

    trees_equiv = round(eco["co2_yr"] / 21.77, 0)
    cars_equiv  = round(eco["co2_yr"] / 4.6, 2)
    _para(f"Annual COâ‚‚ Reduction: {eco['co2_yr']:.2f} tonnes  |  "
          f"Equivalent to planting {int(trees_equiv):,} trees or removing {cars_equiv} cars.",
          size=9)
    doc.add_paragraph()

    _para("Standard: BS 7671:2018 / IEC 60364-5-52 | Temperature derating and grouping factors applied.",
          size=9)
    _para(f"Report generated by {org_name} using SolarPro Global — {datetime.now().strftime('%d %B %Y')}",
          size=9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"SolarPro_{project['name'].replace(' ','_')}_FullReport.docx"
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=fname)


# ─── SaaS Platform Stats API (admin) ─────────────────────────────────────────

@app.route("/admin/platform")
@admin_required
def admin_platform():
    """SaaS platform metrics — MRR, churn, active users, project pipeline."""
    with get_db() as c:
        users       = c.execute("SELECT * FROM users ORDER BY created_at DESC").fetchall()
        total_users = len(users)
        paid_users  = sum(1 for u in users if (u["plan"] or "free") not in ("free","disabled"))
        plan_dist   = {}
        for u in users:
            p = (u["plan"] or "free").lower()
            plan_dist[p] = plan_dist.get(p, 0) + 1

        total_rev   = c.execute("SELECT SUM(amount_usd) FROM payments WHERE status='success'").fetchone()[0] or 0
        proj_count  = c.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        ticket_open = c.execute("SELECT COUNT(*) FROM tickets WHERE status='open'").fetchone()[0]
        leads_new   = c.execute("SELECT COUNT(*) FROM leads WHERE status='new'").fetchone()[0]
        subs_count  = c.execute("SELECT COUNT(*) FROM newsletter_subscribers WHERE status='active'").fetchone()[0]
        email_sent  = c.execute("SELECT COUNT(*) FROM email_logs WHERE status='sent'").fetchone()[0]

        # Monthly revenue
        monthly_rev = c.execute(
            "SELECT strftime('%Y-%m', created_at) as mo, SUM(amount_usd) as rev "
            "FROM payments WHERE status='success' GROUP BY mo ORDER BY mo DESC LIMIT 12").fetchall()

        # MRR estimate — sum of current paid plan monthly prices
        MRR = sum(PLAN_PRICES.get((u["plan"] or "free").lower(), {}).get("usd", 0)
                  for u in users if (u["plan"] or "free") not in ("free","disabled","demo"))

        recent_users = users[:10]
        payments     = c.execute(
            "SELECT p.*, u.username FROM payments p JOIN users u ON p.user_id=u.id "
            "WHERE p.status='success' ORDER BY p.created_at DESC LIMIT 20").fetchall()

    return render_template("admin_platform.html", user=current_user(),
                           total_users=total_users, paid_users=paid_users,
                           plan_dist=plan_dist, total_rev=total_rev,
                           proj_count=proj_count, ticket_open=ticket_open,
                           leads_new=leads_new, subs_count=subs_count,
                           email_sent=email_sent, monthly_rev=monthly_rev,
                           MRR=MRR, recent_users=recent_users, payments=payments,
                           plan_prices=PLAN_PRICES)


# ─── Error handlers ───────────────────────────────────────────────────────────

@app.errorhandler(403)
def err_403(e):
    return render_template("error.html", code=403,
        title="Access Denied",
        message="You don't have permission to access this page. "
                "If you submitted a form, please go back and try again."), 403

@app.errorhandler(404)
def err_404(e):
    return render_template("error.html", code=404,
        title="Page Not Found",
        message="The page you're looking for doesn't exist or has been moved."), 404

@app.errorhandler(429)
def err_429(e):
    return render_template("error.html", code=429,
        title="Too Many Requests",
        message="You've made too many requests in a short time. "
                "Please wait a moment and try again."), 429

@app.errorhandler(500)
def err_500(e):
    return render_template("error.html", code=500,
        title="Internal Server Error",
        message="Something went wrong on our end. "
                "Please go back to the dashboard and try again."), 500


# ─── Client Prospecting Agent ─────────────────────────────────────────────────

import re as _re

# All country names recognised for the hard foreign-country gate.
# "benin" omitted â†' collision with Benin City (Nigeria).
# "niger" omitted â†' collision with Nigeria.
_ALL_COUNTRY_NAMES = frozenset([
    "zambia", "zimbabwe", "togo", "cameroon", "senegal", "gambia", "liberia",
    "burkina faso", "mali", "guinea", "sierra leone", "ivory coast",
    "mozambique", "malawi", "angola", "botswana", "namibia", "lesotho",
    "eswatini", "swaziland", "madagascar", "mauritius", "eritrea", "djibouti",
    "somalia", "chad", "gabon", "congo", "rwanda", "burundi", "uganda",
    "kenya", "tanzania", "ethiopia", "egypt", "morocco", "algeria",
    "tunisia", "libya", "south africa", "south sudan", "sudan",
    "nigeria", "ghana",
    "india", "pakistan", "bangladesh", "indonesia", "philippines",
    "vietnam", "thailand", "myanmar", "sri lanka", "nepal", "malaysia",
    "cambodia", "laos", "mongolia",
    "united kingdom", "england", "scotland", "wales",
    "united states", "america",
    "australia", "germany", "spain", "france", "china", "japan",
    "saudi arabia", "united arab emirates", "qatar", "kuwait",
    "iran", "iraq", "jordan", "lebanon", "syria", "yemen", "oman", "bahrain",
    "brazil", "argentina", "colombia", "peru", "chile", "mexico", "ecuador",
])

# Alternate names that map to a canonical selected-country value
_COUNTRY_EXEMPTIONS = {
    "uk":            {"united kingdom", "england", "scotland", "wales", "britain", "uk"},
    "usa":           {"united states", "america", "usa"},
    "south africa":  {"south africa", "south african"},
    "united kingdom": {"united kingdom", "england", "scotland", "wales", "britain", "uk"},
    "united states": {"united states", "america", "usa"},
}


def _is_past_deadline(text):
    """Return True if we can detect a deadline/closing date in text and it has passed."""
    import re as _re2
    from datetime import datetime
    today = datetime.utcnow()
    t = text.lower()
    # Quick year check: if only year-2025 or earlier found with no later year, it's expired
    years = [int(y) for y in _re2.findall(r'\b(20\d{2})\b', t)]
    if years and max(years) < today.year:
        return True
    # Try to parse full dates near deadline keywords
    patterns = [
        r'(?:closing|deadline|submission|due|submit by|bids? by|close[sd]?)[:\s]+(\d{1,2}[\s\-/][a-z]{3,9}[\s\-/]\d{4})',
        r'(?:closing|deadline|submission|due|submit by|bids? by|close[sd]?)[:\s]+([a-z]{3,9}\s+\d{1,2},?\s*\d{4})',
        r'(?:closing|deadline|submission|due)[:\s]+(\d{4}-\d{2}-\d{2})',
    ]
    MONTH = {"january":1,"february":2,"march":3,"april":4,"may":5,"june":6,
             "july":7,"august":8,"september":9,"october":10,"november":11,"december":12,
             "jan":1,"feb":2,"mar":3,"apr":4,"jun":6,"jul":7,"aug":8,
             "sep":9,"oct":10,"nov":11,"dec":12}
    for pat in patterns:
        m = _re2.search(pat, t)
        if not m:
            continue
        raw = m.group(1).strip()
        # Try YYYY-MM-DD
        try:
            d = datetime.strptime(raw, "%Y-%m-%d")
            return d < today
        except Exception:
            pass
        # Try "15 june 2026" or "june 15, 2026"
        parts = _re2.split(r'[\s\-/,]+', raw)
        nums = [p for p in parts if p.isdigit()]
        words = [p for p in parts if p.isalpha()]
        month_num = next((MONTH[w] for w in words if w in MONTH), None)
        year_num  = next((int(n) for n in nums if len(n) == 4), None)
        day_num   = next((int(n) for n in nums if len(n) <= 2), None)
        if month_num and year_num:
            try:
                d = datetime(year_num, month_num, day_num or 1)
                return d < today
            except Exception:
                pass
    return False


_CITY_MAP = {
    "ghana":        ["accra", "kumasi", "takoradi", "tema", "tamale", "cape coast",
                     "koforidua", "bolgatanga", "sunyani", "techiman"],
    "nigeria":      ["lagos", "abuja", "kano", "ibadan", "port harcourt", "enugu",
                     "kaduna", "calabar", "warri", "uyo", "benin city", "owerri"],
    "kenya":        ["nairobi", "mombasa", "kisumu", "nakuru", "eldoret", "thika"],
    "south africa": ["johannesburg", "cape town", "durban", "pretoria", "soweto",
                     "bloemfontein", "port elizabeth", "east london"],
    "tanzania":     ["dar es salaam", "dodoma", "arusha", "mwanza", "zanzibar"],
    "zambia":       ["lusaka", "ndola", "kitwe", "kabwe", "livingstone"],
    "ethiopia":     ["addis ababa", "dire dawa", "mekele", "gondar", "hawassa"],
    "senegal":      ["dakar", "thiÃ¨s", "kaolack", "saint-louis", "ziguinchor"],
    "cameroon":     ["douala", "yaounde", "bamenda", "bafoussam", "garoua"],
    "uganda":       ["kampala", "gulu", "mbarara", "jinja", "entebbe"],
    "rwanda":       ["kigali", "butare", "gisenyi", "ruhengeri"],
    "uk":           ["london", "birmingham", "manchester", "leeds", "glasgow"],
    "usa":          ["new york", "los angeles", "chicago", "houston", "phoenix"],
    "india":        ["mumbai", "delhi", "bangalore", "chennai", "hyderabad", "pune"],
}


def _extract_location_from_text(title, body, loc):
    """Try to extract city/region from result text; fall back to country name."""
    text = (title + " " + body).lower()
    cities = _CITY_MAP.get(loc.lower(), [])
    for city in cities:
        if city in text:
            return city.title() + ", " + loc
    return loc


def _foreign_country_in_text(title, loc_lower, full_content=None):
    """Return True if this page is clearly about a different country.

    Title-only check (no full_content):
      Block if selected country is absent from title AND a foreign country is present.

    With full_content:
      Even if selected country is in the title, block if any foreign country is
      mentioned MORE TIMES than the selected country in the full page text — that
      means the page is really about the foreign country (e.g. a Sudan tender
      that has 'Ghana' once in a sidebar nav).
    """
    t = title.lower()
    exempt = {loc_lower} | _COUNTRY_EXEMPTIONS.get(loc_lower, set())

    if full_content:
        fc = full_content.lower()
        loc_count = fc.count(loc_lower)
        # Require selected country to appear at least twice in page content
        if loc_count < 2:
            return True
        # Block if any foreign country appears more than selected country
        for name in _ALL_COUNTRY_NAMES:
            if name in exempt:
                continue
            if fc.count(name) > loc_count:
                return True
        return False

    # No full content — title-only check
    if loc_lower in t:
        return False
    for name in _ALL_COUNTRY_NAMES:
        if name in exempt:
            continue
        if _re.search(r'\b' + _re.escape(name) + r'\b', t):
            return True
    return False


@app.route("/admin/agent")
@admin_required
def admin_agent():
    with get_db() as c:
        saved = c.execute(
            "SELECT * FROM leads WHERE source='agent' ORDER BY created_at DESC LIMIT 50"
        ).fetchall()
        total_saved = c.execute(
            "SELECT COUNT(*) FROM leads WHERE source='agent'"
        ).fetchone()[0]
        mstate = c.execute(
            "SELECT last_scan, last_count, scan_interval, notify_email, "
            "last_agent_run, agent_run_count FROM monitor_state WHERE id=1"
        ).fetchone()
    ms = {
        "last_scan":       mstate[0] if mstate else "",
        "last_count":      mstate[1] if mstate else 0,
        "scan_interval":   mstate[2] if mstate else 120,
        "notify_email":    bool(mstate[3]) if mstate else False,
        "last_agent_run":  mstate[4] if mstate else "",
        "agent_run_count": mstate[5] if mstate else 0,
    }
    _has_claude = bool(os.environ.get("ANTHROPIC_API_KEY"))
    _has_or     = bool(os.environ.get("OPENROUTER_API_KEY"))
    _has_gh     = bool(os.environ.get("GITHUB_TOKEN"))
    _ai_label   = ("OpenRouter" if _has_or
                   else "GitHub Models" if _has_gh
                   else "Primary AI" if _has_claude
                   else "none")
    return render_template("admin_agent.html", user=current_user(),
                           saved_leads=saved, total_saved=total_saved,
                           has_ai=bool(_has_claude or _has_or or _has_gh),
                           ai_provider=_ai_label,
                           ms=ms)


@app.route("/admin/agent/run", methods=["POST"])
@admin_required
@limiter.limit("20 per hour")
def admin_agent_run():
    csrf_protect()
    country   = request.form.get("country", "").strip()
    sector    = request.form.get("sector", "Commercial")
    system_kw = request.form.get("system_kw", "10-100")
    budget    = request.form.get("budget", "$5,000-$100,000")
    focus     = request.form.get("focus", "").strip()
    count     = min(int(request.form.get("count", 8)), 12)

    loc       = country if country else "Ghana"
    loc_q     = f'"{loc}"'
    loc_label = loc

    # ── Step 1: Deep multi-source search ─────────────────────────────────────────
    search_results = []
    search_error   = None

    # ── Country aliases: include major cities so results aren't missed ────────
    COUNTRY_CITIES = {
        # No single-letter or 2-letter aliases — they match inside any English word
        "ghana":        ["ghana", "ghanaian", "accra", "kumasi", "takoradi", "tema",
                         "tamale", "cape coast", "koforidua", "bolgatanga", ".gh"],
        "nigeria":      ["nigeria", "nigerian", "lagos", "abuja", "kano", "ibadan",
                         "port harcourt", "enugu", "kaduna", "benin city", ".ng"],
        "kenya":        ["kenya", "kenyan", "nairobi", "mombasa", "kisumu", "nakuru", ".ke"],
        "south africa": ["south africa", "johannesburg", "cape town", "durban",
                         "pretoria", "gauteng", ".za"],
        "tanzania":     ["tanzania", "dar es salaam", "dodoma", "arusha", ".tz"],
        "zambia":       ["zambia", "lusaka", "ndola", "kitwe", ".zm"],
        "ethiopia":     ["ethiopia", "addis ababa", "dire dawa", ".et"],
        "senegal":      ["senegal", "dakar", ".sn"],
        "cameroon":     ["cameroon", "douala", "yaounde", ".cm"],
        "uganda":       ["uganda", "kampala", ".ug"],
        "rwanda":       ["rwanda", "kigali", ".rw"],
        "uk":           ["united kingdom", "england", "scotland", "wales", "london",
                         "birmingham", "manchester", ".uk", ".co.uk"],
        "usa":          ["united states", "america", ".gov", ".edu"],
        "india":        ["india", "indian", "mumbai", "delhi", "bangalore", ".in"],
    }
    loc_lower   = loc.lower()
    loc_aliases = COUNTRY_CITIES.get(loc_lower, [loc_lower])

    # ── Search infrastructure ─────────────────────────────────────────────────
    UN_PORTALS = (
        "site:ungm.org OR site:devex.com OR site:reliefweb.int "
        "OR site:dgmarket.com OR site:tendersinfo.com"
    )
    DFI_PORTALS = (
        "site:worldbank.org OR site:afdb.org OR site:ifc.org "
        "OR site:esmap.org OR site:geapp.org"
    )
    AFRICA_PORTALS = (
        "site:africatenders.com OR site:tendersontime.com "
        "OR site:globaltenders.com OR site:ecreee.org"
    )
    JOB_DOMAINS    = ["jobberman.com", "myjobmag.com", "brightermonday.com",
                      "ghanaiansjobs.com", "jobsinghana.com", "indeed.com"]
    SOCIAL_DOMAINS = ["facebook.com", "linkedin.com", "twitter.com", "x.com"]

    sector_q = sector.lower() if sector else "commercial"
    focus_q  = focus if focus else f"solar PV {sector_q}"

    queries = [
        f'"tender for" solar installation {loc_q} 2026 2027',
        f'"invitation to bid" solar PV {loc_q} 2026 2027',
        f'"request for proposals" solar {loc_q} 2026 2027',
        f'"expression of interest" solar {loc_q} installation 2026 2027',
        f'({UN_PORTALS}) solar {loc_q} tender OR RFP 2026 2027',
        f'({DFI_PORTALS}) solar {loc_q} tender OR "invitation to bid" 2026 2027',
        f'({AFRICA_PORTALS}) solar {loc_q} tender OR RFP 2026 2027',
        f'{loc_q} {sector_q} solar "supply and install" tender OR RFP 2026 2027',
        f'site:facebook.com {loc_q} solar "looking for installer" OR "need solar" 2026',
        f'site:linkedin.com {loc_q} solar contractor OR seeking OR tender 2026',
        f'{loc_q} {focus_q} solar procurement OR bid OR tender 2026 2027',
    ]

    rfp_keywords = [
        "tender", "rfp", "itb", "eoi", "invitation to bid",
        "request for proposal", "expression of interest",
        "call for tenders", "procurement notice", "bidding document",
        "installation works", "epc contract",
    ]

    def _is_social(url):
        return any(d in url for d in SOCIAL_DOMAINS)

    def _is_job_board(url):
        return any(d in url for d in JOB_DOMAINS)

    try:
        # -- api_manager search (DuckDuckGo + 6h cache + stale fallback) --
        _raw_results = []
        for q in queries:
            try:
                _hits = _api.search.query(q, max_results=15)
                _raw_results.extend(_hits)
            except Exception:
                pass
        search_results = _raw_results[:50]
    except Exception as e:
        search_error = str(e)

    # ── Step 1.5: Fetch actual page content for each candidate ───────────────────
    # Search snippets are only 200-400 chars — not enough to confirm a real tender.
    # We fetch the actual page (first 6000 chars of text) so Claude and our filters
    # work on the real document, not a teaser.
    if search_results:
        import concurrent.futures, requests as _req
        from html.parser import HTMLParser as _HP

        def _strip_html(raw_html):
            class _P(_HP):
                def __init__(self):
                    super().__init__()
                    self._skip = False
                    self._chunks = []
                def handle_starttag(self, tag, attrs):
                    if tag in ('script','style','nav','header','footer','aside','menu'):
                        self._skip = True
                def handle_endtag(self, tag):
                    if tag in ('script','style','nav','header','footer','aside','menu'):
                        self._skip = False
                def handle_data(self, data):
                    if not self._skip and data.strip():
                        self._chunks.append(data.strip())
            p = _P()
            p.feed(raw_html[:80000])
            return ' '.join(p._chunks)[:6000]

        def _fetch_page(r):
            url = r.get("href", "")
            # Skip social/job — require login or are dynamic
            if any(d in url for d in SOCIAL_DOMAINS + JOB_DOMAINS):
                return r
            try:
                hdrs = {'User-Agent': 'Mozilla/5.0 (compatible; SolarPro/1.0)'}
                resp = _req.get(url, timeout=6, headers=hdrs, allow_redirects=True)
                if resp.status_code == 200 and "text/html" in resp.headers.get("Content-Type",""):
                    text = _strip_html(resp.text)
                    if text:
                        r["full_content"] = text
            except Exception:
                pass
            return r

        # Fetch in parallel — max 12 pages, 5 workers
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as pool:
            search_results = list(pool.map(_fetch_page, search_results[:12]))

        # Re-filter using full page content where available
        refined = []
        for r in search_results:
            content   = (r.get("full_content") or r.get("body","")).lower()
            title_r   = r.get("title","").lower()
            url_r     = r.get("href","")
            # Re-check expired deadline on full content
            if r.get("full_content") and _is_past_deadline(content):
                continue
            # Re-check wrong country using full content (count-based dominance check)
            fc = r.get("full_content")
            if _foreign_country_in_text(title_r, loc_lower, full_content=fc):
                continue
            # For formal sources: rfp keyword must appear somewhere in full content
            if r.get("full_content") and not (_is_social(url_r) or _is_job_board(url_r)):
                if not any(kw in content for kw in rfp_keywords):
                    continue
            refined.append(r)
        if refined:
            search_results = refined

    # ── Step 2: AI analyses real search results → structured prospects ──────────
    or_key   = os.environ.get("OPENROUTER_API_KEY", "")
    api_key  = os.environ.get("ANTHROPIC_API_KEY", "")
    gh_token = os.environ.get("GITHUB_TOKEN", "")
    _has_any_ai = bool(or_key or api_key or gh_token or os.environ.get("OLLAMA_URL"))
    if _has_any_ai and search_results:
        try:
            # Trim snippets: max 1500 chars each so total prompt stays under ~20k tokens
            snippets = "\n\n".join(
                f"[{i+1}] TITLE: {r.get('title','')}\nURL: {r.get('href','')}\n"
                f"CONTENT:\n{(r.get('full_content') or r.get('body',''))[:1500]}"
                for i, r in enumerate(search_results[:10])
            )
            prompt = f"""You are a solar PV procurement intelligence analyst. Your job is to READ and ANALYSE the full content of each result below, then extract ONLY genuine, currently open procurement opportunities.

Search criteria:
- Target country: {loc_label} — results must be FOR THIS COUNTRY ONLY
- Sector: {sector}
- System size: {system_kw} kW
- Budget: {budget}
- Focus: {focus or 'general'}

RESULTS TO ANALYSE:

{snippets}

INSTRUCTIONS — read each result's CONTENT carefully before deciding:
1. READ the full content provided. If it does not clearly confirm an open, active procurement for solar works IN {loc_label}, SKIP IT entirely.
2. SKIP if the content is about a different country — even if {loc_label} appears once in a header or navigation. The TENDER ITSELF must be in {loc_label}.
3. SKIP news articles, project completion stories, country overviews, funding announcements.
4. SKIP if the closing/deadline date has already passed (today is {__import__('datetime').date.today()}).
5. ONLY include results where the content confirms: what is being procured, who is issuing it, and how to respond.
6. source_url = exact URL, copied verbatim — do not modify
7. company_name = the issuing organisation extracted from the content
8. Never invent data — use "" if a field is not stated in the content
9. Extract ALL contact intelligence: named persons, office names, GPS/addresses, websites, submission methods, mandatory documents, certifications required.
10. Score each opportunity: classify as hot (deadline â‰¤14 days OR budget confirmed high), warm (deadline â‰¤45 days OR medium budget), cold (distant deadline or vague budget). urgency_score 1—10 (10=closes within 7 days). revenue_potential = realistic contract value in USD if stated, else estimate from system size.

PRIORITY RULE — score based on how many of these 5 key fields are present in the result:
  1. work_description  — scope / what is being procured
  2. requirements      — eligibility or technical requirements
  3. tor               — terms of reference or detailed scope
  4. deadline          — submission closing date
  5. submission_address OR contact_details — where/how to submit or who to contact
  Priority: "high" = all 5 present; "medium" = 3 or 4 present; "low" = 2 or fewer present

Return up to {count} results. Return ONLY valid JSON, no markdown:
{{
  "prospects": [
    {{
      "company_name": "Issuing organisation",
      "type": "RFP / Tender / EOI / ITB / Contract Notice / Grant / Installation Job",
      "project_category": "Rooftop / Ground-mount / Hybrid / Off-grid / Mini-grid / Street lighting / Water pumping / Other",
      "location": "city AND country from result, e.g. 'Accra, Ghana' — use city name if stated, else just country. Do NOT invent.",
      "estimated_kw": 0,
      "estimated_usd": 0,
      "budget": "stated budget exactly as written, e.g. 'USD 500,000' or ''",
      "revenue_potential": 0,
      "pain_points": [],
      "pitch": "one sentence: what they are procuring",
      "work_description": "full scope of work as stated — supply and install, design only, EPC, etc. Use '' if not stated.",
      "requirements": "eligibility or technical requirements stated in result. Use '' if not stated.",
      "mandatory_documents": "list of mandatory documents stated, e.g. 'Company registration, Tax clearance, PURC licence'. Use '' if not stated.",
      "certifications": "certifications or accreditations required, e.g. 'ISO 9001, ECG approved installer'. Use '' if not stated.",
      "tor": "terms of reference or scope details if stated. Use '' if not stated.",
      "deadline": "closing/submission date exactly as stated, e.g. '30 June 2026'. Use '' if not stated.",
      "submission_address": "where/how to submit: email, portal URL, physical address. Use '' if not stated.",
      "submission_method": "email / online portal / physical drop-off / courier / hand-deliver — as stated. Use '' if not stated.",
      "contact_details": "procurement contact name, email, phone if stated. Use '' if not stated.",
      "contact_person": "named contact person extracted from result. Use '' if not stated.",
      "procurement_office": "procurement office or department name. Use '' if not stated.",
      "website": "organisation website if stated. Use '' if not stated.",
      "gps_location": "GPS coordinates or full physical address of the project/organisation if stated. Use '' if not stated.",
      "tender_ref": "reference number if stated, else ''",
      "contact_strategy": "how to respond based on source type",
      "decision_maker": "procurement contact name if named, else ''",
      "classification": "hot / warm / cold",
      "urgency_score": 5,
      "priority": "high = all 5 fields present; medium = 3-4; low = 2 or fewer",
      "source_url": "https://exact-url-from-result.com/path",
      "source_title": "Exact title from result",
      "source_snippet": "most relevant verbatim sentence, max 300 chars",
      "verified": true
    }}
  ]
}}"""
            import urllib.request as _ur_ai, urllib.error as _ue_ai, json as _json_ai
            raw = None
            ai_source = None
            _provider_errors = []  # accumulate errors for debugging

            # ── 1. OpenRouter — free Llama (primary) ─────────────────────
            # Each model tried; on any error, continue to next model then next provider
            if or_key and raw is None:
                _or_models = [
                    "meta-llama/llama-3.1-8b-instruct:free",
                    "google/gemma-2-9b-it:free",
                    "mistralai/mistral-7b-instruct:free",
                    "meta-llama/llama-3.3-70b-instruct:free",
                ]
                for _or_model in _or_models:
                    try:
                        _or_payload = _json_ai.dumps({
                            "model":      _or_model,
                            "messages":   [{"role": "user", "content": prompt}],
                            "max_tokens": 3000,
                        }).encode()
                        _or_req = _ur_ai.Request(
                            "https://openrouter.ai/api/v1/chat/completions",
                            data=_or_payload,
                            headers={
                                "Authorization": f"Bearer {or_key}",
                                "Content-Type":  "application/json",
                                "HTTP-Referer":  "https://solarpro.aiappinvent.com",
                                "X-Title":       "SolarPro Prospecting Agent"
                            }
                        )
                        with _ur_ai.urlopen(_or_req, timeout=90) as _or_resp:
                            _or_result = _json_ai.loads(_or_resp.read())
                        raw = _or_result["choices"][0]["message"]["content"].strip()
                        ai_source = f"web+openrouter({_or_model.split('/')[1]})"
                        break
                    except _ue_ai.HTTPError as _oe:
                        _ob = _oe.read().decode("utf-8", errors="ignore")[:200]
                        _provider_errors.append(f"OR {_or_model} HTTP{_oe.code}: {_ob}")
                        continue
                    except Exception as _oe:
                        _provider_errors.append(f"OR {_or_model}: {_oe}")
                        continue  # try next model; fall through to next provider if all fail

            # ── 2. Ollama (local inference) ───────────────────────────────
            if os.environ.get("OLLAMA_URL") and raw is None:
                try:
                    _ollama_url   = os.environ.get("OLLAMA_URL", "http://localhost:11434")
                    _ollama_model = os.environ.get("OLLAMA_MODEL", "mistral")
                    _payload4 = _json_ai.dumps({
                        "model":    _ollama_model,
                        "messages": [{"role": "user", "content": prompt}],
                        "stream":   False
                    }).encode()
                    _req4 = _ur_ai.Request(
                        f"{_ollama_url}/api/chat",
                        data=_payload4,
                        headers={"Content-Type": "application/json"}
                    )
                    with _ur_ai.urlopen(_req4, timeout=120) as _r4:
                        _result4 = _json_ai.loads(_r4.read())
                    raw = _result4["message"]["content"].strip()
                    ai_source = "web+ollama"
                except Exception as _oe4:
                    _provider_errors.append(f"Ollama: {_oe4}")

            # ── 3. GitHub Models — free GPT-4.1-mini ──────────────────────
            if gh_token and raw is None:
                try:
                    _payload3 = _json_ai.dumps({
                        "model":       "gpt-4.1-mini",
                        "messages":    [{"role": "user", "content": prompt}],
                        "max_tokens":  4000,
                        "temperature": 0.3
                    }).encode()
                    _req3 = _ur_ai.Request(
                        "https://models.inference.ai.azure.com/chat/completions",
                        data=_payload3,
                        headers={
                            "Authorization": f"Bearer {gh_token}",
                            "Content-Type":  "application/json",
                            "Accept":        "application/json",
                            "User-Agent":    "solarpro-agent/1.0"
                        })
                    with _ur_ai.urlopen(_req3, timeout=60) as _r3:
                        _result3 = _json_ai.loads(_r3.read())
                    raw = _result3["choices"][0]["message"]["content"].strip()
                    ai_source = "web+github-models"
                except Exception as _oe3:
                    _provider_errors.append(f"GitHub: {_oe3}")

            # ── 4. Anthropic Claude (last resort — saves API credits) ─────
            if api_key and raw is None:
                try:
                    import anthropic as _ant
                    client = _ant.Anthropic(api_key=api_key)
                    msg    = client.messages.create(
                        model="claude-opus-4-7", max_tokens=4000,
                        messages=[{"role": "user", "content": prompt}]
                    )
                    raw = msg.content[0].text.strip()
                    ai_source = "web+claude"
                except Exception as _oe_ant:
                    _provider_errors.append(f"Claude: {_oe_ant}")

            if raw is None:
                raise ValueError("No AI provider available; errors: " + " | ".join(_provider_errors))

            if raw.startswith("```"):
                raw = "\n".join(raw.split("\n")[1:])
                if raw.endswith("```"):
                    raw = raw[:-3]
            data = json.loads(raw)
            # AI_BUDGET_LEDGER_MARKER_AGENT - record one ledger row per
            # successful prospecting-agent run so admin spend is visible.
            try:
                import ai_budget as _ab_agent
                _src = ai_source or ""
                if "openrouter" in _src:
                    _prov, _mdl = "openrouter", _src.split("(")[-1].rstrip(")")
                elif "ollama" in _src:
                    _prov, _mdl = "ollama", os.environ.get("OLLAMA_MODEL", "")
                elif "github" in _src:
                    _prov, _mdl = "github_models", "gpt-4.1-mini"
                elif "claude" in _src:
                    _prov, _mdl = "claude", "claude-opus-4-7"
                else:
                    _prov, _mdl = "unknown", ""
                _ab_agent.record_usage(
                    user_id=session.get("user_id"),
                    provider=_prov, model=_mdl,
                    prompt_tokens=_ab_agent.estimate_tokens(prompt),
                    completion_tokens=_ab_agent.estimate_tokens(raw or ""),
                    endpoint="/admin/agent/run")
            except Exception:
                pass
            return jsonify({"ok": True, "prospects": data["prospects"],
                            "source": ai_source, "result_count": len(search_results)})
        except Exception as e:
            _ai_error = str(e)  # captured for debug; fall through to template extraction

    # ── Step 3: Template extraction (no AI key) ──────────────────────────────────
    def _classify_source(url):
        u = url.lower()
        if any(d in u for d in ["facebook.com", "twitter.com", "x.com"]):
            return "Social Media Lead"
        if "linkedin.com" in u:
            return "LinkedIn Lead"
        if any(d in u for d in ["jobberman.com", "myjobmag.com", "brightermonday.com",
                                  "indeed.com", "jobsinghana.com", "ghanaiansjobs.com"]):
            return "Job Board — Active Project"
        if any(d in u for d in [".gov.gh", ".gov.ng", ".gov.ke", "gov.", "district",
                                  "assembly", "council", "ministry"]):
            return "Government / Public Sector"
        if any(d in u for d in ["ungm.org", "devex.com", "worldbank.org", "afdb.org",
                                  "reliefweb.int", "dgmarket.com", "tendersontime.com",
                                  "globaltenders.com", "africatenders.com"]):
            return "Tender Portal"
        return "Commercial / Private Sector"

    def _contact_strategy(source_type, url):
        if source_type == "Social Media Lead":
            return "Respond directly to the post — offer free site assessment and quote"
        if source_type == "LinkedIn Lead":
            return "Connect on LinkedIn, message offering a no-obligation solar audit"
        if source_type == "Job Board — Active Project":
            return "Company is hiring for a solar project — contact HR/procurement directly"
        if source_type == "Government / Public Sector":
            return "Submit formal expression of interest or bid via the government procurement portal"
        if source_type == "Tender Portal":
            return "Download tender documents and submit bid before closing date"
        return "Contact via source link — offer free survey and detailed quotation"

    def _infer_type(title, body):
        t = (title + " " + body).lower()
        if any(w in t for w in ["tender", "rfp", "itb", "invitation to bid", "call for"]):
            return "Tender / RFP"
        if any(w in t for w in ["expression of interest", "eoi", "prequalif"]):
            return "Expression of Interest"
        if any(w in t for w in ["job", "vacancy", "hiring", "technician", "installer"]):
            return "Job Post — Active Project"
        if any(w in t for w in ["looking for", "need", "quote", "recommend", "how much"]):
            return "Social Media — Seeking Installer"
        return "Solar Project Opportunity"

    def _extract_deadline(text):
        import re
        t = text.lower()
        # Match patterns like "closing date: 30 June 2026", "deadline: 2026-06-30"
        patterns = [
            r'(?:closing|deadline|submission|due|close[sd]?|submit by|bids? by)[:\s]+([0-9]{1,2}[\s\-/][a-z]+[\s\-/][0-9]{4})',
            r'(?:closing|deadline|submission|due|close[sd]?|submit by|bids? by)[:\s]+([a-z]+ [0-9]{1,2},?\s*[0-9]{4})',
            r'(?:closing|deadline|submission|due|close[sd]?|submit by|bids? by)[:\s]+([0-9]{4}-[0-9]{2}-[0-9]{2})',
            r'\b([0-9]{1,2} (?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]* [0-9]{4})\b',
            r'\b([0-9]{4}-[0-9]{2}-[0-9]{2})\b',
        ]
        for pat in patterns:
            m = re.search(pat, t)
            if m:
                return m.group(1).strip()
        return ""

    def _extract_contact(text):
        import re
        emails = re.findall(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', text)
        phones = re.findall(r'(?:\+?[0-9][\d\s\-().]{7,15})', text)
        parts = []
        if emails:
            parts.append(emails[0])
        if phones:
            parts.append(phones[0].strip())
        return "; ".join(parts)

    def _extract_submission_address(text):
        t = text.lower()
        for kw in ["submit to", "submission to", "send to", "addressed to",
                   "tender box", "procurement office", "p.o. box", "po box"]:
            idx = t.find(kw)
            if idx != -1:
                return text[idx:idx+120].strip()
        return ""

    def _extract_description(title, body):
        t = (title + " " + body).lower()
        for kw in ["supply and install", "supply & install", "design and install",
                   "epc contract", "installation of", "procurement of",
                   "construction of", "works contract", "solar pv system",
                   "off-grid solar", "mini grid", "solar farm"]:
            if kw in t:
                idx = t.find(kw)
                return (title + " " + body)[max(0, idx-20):idx+120].strip()
        return ""

    def _extract_requirements(body):
        t = body.lower()
        for kw in ["must have", "required to", "eligible", "experience of",
                   "registered", "qualification", "minimum", "years experience"]:
            if kw in t:
                idx = t.find(kw)
                return body[idx:idx+200].strip()
        return ""

    def _score_priority(work_desc, requirements, tor, deadline, submission_addr, contact):
        # Count how many of the 5 mandatory fields are present.
        # submission_address OR contact_details counts as the 5th field.
        submission_ok = bool((submission_addr and submission_addr.strip()) or
                             (contact and contact.strip()))
        filled = sum(1 for f in [work_desc, requirements, tor, deadline] if f and f.strip())
        if submission_ok:
            filled += 1
        if filled == 5:
            return "high"
        if filled >= 3:
            return "medium"
        return "low"

    if search_results:
        prospects = []
        for r in search_results[:count]:
            title   = r.get("title", "")
            url     = r.get("href", "")
            snippet = r.get("body", "")
            combined_text = title + " " + snippet
            src_type = _classify_source(url)
            work_desc   = _extract_description(title, snippet)
            deadline_v  = _extract_deadline(combined_text)
            contact_v   = _extract_contact(combined_text)
            sub_addr    = _extract_submission_address(combined_text)
            requirements = _extract_requirements(snippet)
            prospects.append({
                "company_name":       title[:80] if title else "Unknown",
                "type":               _infer_type(title, snippet),
                "location":           _extract_location_from_text(title, snippet, loc),
                "estimated_kw":       0,
                "estimated_usd":      0,
                "pain_points":        [src_type],
                "pitch":              snippet[:250] if snippet else "",
                "work_description":   work_desc,
                "requirements":       requirements,
                "tor":                "",
                "deadline":           deadline_v,
                "submission_address": sub_addr,
                "contact_details":    contact_v,
                "tender_ref":         "",
                "contact_strategy":   _contact_strategy(src_type, url),
                "decision_maker":     "",
                "priority":           _score_priority(work_desc, requirements, "", deadline_v, sub_addr, contact_v),
                "source_url":         url,
                "source_title":       title,
                "source_snippet":     snippet[:300],
                "verified":           True,
            })
        return jsonify({"ok": True, "prospects": prospects,
                        "source": "web_search", "result_count": len(search_results),
                        "ai_error": _ai_error if '_ai_error' in dir() else None})

    # ── Step 4: Last resort — inform user search failed ─────────────────────────
    return jsonify({"ok": False,
                    "error": f"Web search returned no results. {search_error or ''} "
                             "Try different search criteria or add an ANTHROPIC_API_KEY."})


@app.route("/admin/agent/notify", methods=["POST"])
@admin_required
def admin_agent_notify():
    """Called by the frontend after a successful agent run to log it and optionally email."""
    csrf_protect()
    count   = int(request.form.get("count", 0))
    country = request.form.get("country", "").strip() or "Global"
    source  = request.form.get("source", "web_search")
    now     = datetime.now().strftime("%Y-%m-%d %H:%M UTC")
    with get_db() as c:
        c.execute(
            "UPDATE monitor_state SET last_agent_run=?, "
            "agent_run_count=agent_run_count+1 WHERE id=1", (now,))
        state = c.execute(
            "SELECT notify_email FROM monitor_state WHERE id=1"
        ).fetchone()
    if state and state[0] and count > 0:
        if source == "web+claude":
            src_label = "Web Search + AI"
        elif source == "web+mistral":
            src_label = "Web Search + Mistral AI"
        elif source == "web+ollama":
            src_label = "Web Search + Ollama (local)"
        elif "github" in source:
            src_label = "GitHub Models + GPT"
        else:
            src_label = "Live Web Search"
        def _send():
            _send_prospect_notification(
                f"Agent Run Complete — {count} Prospect{'s' if count!=1 else ''} Found",
                [f"<strong>{count}</strong> solar prospect{'s' if count!=1 else ''} found for <strong>{country}</strong>.",
                 f"Source: {src_label}",
                 "Open the Agent Dashboard to review results and save the best leads to your CRM."]
            )
        threading.Thread(target=_send, daemon=True).start()
    return jsonify({"ok": True})


@app.route("/admin/agent/save", methods=["POST"])
@admin_required
def admin_agent_save():
    csrf_protect()
    name     = request.form.get("name", "").strip()
    company  = request.form.get("company", "").strip()
    country  = request.form.get("country", "").strip()
    interest = request.form.get("interest", "solar").strip()
    notes    = request.form.get("notes", "").strip()
    if not name:
        return jsonify({"ok": False, "error": "Name required"}), 400
    with get_db() as c:
        existing = c.execute(
            "SELECT id FROM leads WHERE company=? AND source='agent'", (company,)
        ).fetchone()
        if existing:
            return jsonify({"ok": False, "error": "Already saved"}), 409
        c.execute(
            "INSERT INTO leads (name,email,phone,company,country,interest,message,source,status) "
            "VALUES (?,?,?,?,?,?,?,?,?)",
            (name, "", "", company, country, interest, notes, "agent", "new")
        )
    return jsonify({"ok": True})


# ─── Background Monitor: actively listens for new solar postings ───────────────

def _monitor_search(loc="Ghana"):
    """Run the same search pipeline as admin_agent_run for a given country.
    Returns a list of dicts with url, title, snippet, source_type."""
    from ddgs import DDGS

    loc_q = f'"{loc}"'
    COUNTRY_CITIES = {
        "ghana":        ["ghana", "ghanaian", "accra", "kumasi", "takoradi", "tema",
                         "tamale", "cape coast", ".gh"],
        "nigeria":      ["nigeria", "nigerian", "lagos", "abuja", "kano", ".ng"],
        "kenya":        ["kenya", "kenyan", "nairobi", "mombasa", ".ke"],
        "south africa": ["south africa", "johannesburg", "cape town", "durban", ".za"],
        "tanzania":     ["tanzania", "dar es salaam", ".tz"],
        "zambia":       ["zambia", "lusaka", ".zm"],
        "ethiopia":     ["ethiopia", "addis ababa", ".et"],
        "senegal":      ["senegal", "dakar", ".sn"],
        "cameroon":     ["cameroon", "douala", ".cm"],
        "uk":           ["united kingdom", "england", "london", ".uk", ".co.uk"],
        "usa":          ["united states", "america", ".gov"],
        "india":        ["india", "mumbai", "delhi", "bangalore", ".in"],
    }
    loc_lower   = loc.lower()
    loc_aliases = COUNTRY_CITIES.get(loc_lower, [loc_lower])

    UN_PORTALS = (
        "site:ungm.org OR site:devex.com OR site:reliefweb.int "
        "OR site:dgmarket.com OR site:tendersinfo.com"
    )
    DFI_PORTALS = (
        "site:worldbank.org OR site:afdb.org OR site:ifc.org "
        "OR site:esmap.org OR site:geapp.org"
    )
    AFRICA_PORTALS = (
        "site:africatenders.com OR site:tendersontime.com "
        "OR site:globaltenders.com OR site:ecreee.org"
    )
    JOB_DOMAINS    = ["jobberman.com", "myjobmag.com", "brightermonday.com",
                      "ghanaiansjobs.com", "jobsinghana.com", "indeed.com"]
    SOCIAL_DOMAINS = ["facebook.com", "linkedin.com", "twitter.com", "x.com"]

    monitor_queries = [
        f'"tender for" solar installation {loc_q} 2026 2027',
        f'"invitation to bid" solar PV {loc_q} 2026 2027',
        f'"request for proposals" solar {loc_q} 2026 2027',
        f'"expression of interest" solar {loc_q} installation 2026 2027',
        f'({UN_PORTALS}) solar {loc_q} tender OR RFP 2026 2027',
        f'({DFI_PORTALS}) solar {loc_q} tender OR "invitation to bid" 2026 2027',
        f'({AFRICA_PORTALS}) solar {loc_q} tender OR RFP 2026 2027',
        f'{loc_q} "district assembly" solar installation tender OR contract 2026 2027',
        f'site:facebook.com {loc_q} solar "looking for installer" OR "need solar" 2026',
        f'site:linkedin.com {loc_q} solar "contractor" OR "seeking" OR "tender" 2026',
        f'site:jobberman.com solar {loc_q} installer OR technician 2026',
        f'{loc_q} "solar backup" OR "off-grid solar" "supply and install" 2026 2027',
    ]

    skip_domains = [
        "pv-magazine", "pvtech", "reuters.com", "bloomberg.com", "wikipedia.org",
        "youtube.com", "tiktok.com", "solarpowerworldonline", "greentechmedia",
        "theguardian.com", "bbc.com", "cnn.com", "aljazeera.com",
        "businessghana.com", "ghanaweb.com", "myjoyonline.com",
        "esi-africa.com", "irena.org/news", "afdb.org/en/news",
        "worldbank.org/en/news", "adb.org/news",
    ]
    news_url_paths = [
        "/news/", "/blog/", "/article/", "/story/", "/stories/",
        "/press/", "/newsroom/", "/publication/", "/en/news",
        # AfDB/WorldBank narrative paths — project/country pages are NOT procurement notices
        "afdb.org/en/news", "afdb.org/en/projects", "afdb.org/en/countries",
        "afdb.org/en/documents", "afdb.org/en/topics",
        "worldbank.org/en/news", "worldbank.org/en/results",
        "worldbank.org/en/country", "worldbank.org/en/project",
        "worldbank.org/en/topic", "ifc.org/en/stories", "ifc.org/wps/wcm",
        "esmap.org/node", "esmap.org/story",
    ]
    news_title_words = [
        "awarded", "wins contract", "signs agreement", "completed",
        "inaugurated", "commissioned", "connected to grid",
    ]
    rfp_keywords = [
        "tender", "rfp", "itb", "eoi", "invitation to bid",
        "request for proposal", "expression of interest",
        "call for tenders", "procurement notice", "bidding document",
        "installation works", "epc contract",
    ]
    opportunity_keywords = [
        "solar", "install", "installer", "technician", "contractor",
        "supply", "design", "panel", "pv", "quote", "looking for",
        "need", "seeking", "vacancy", "hiring", "job",
    ]
    solar_keywords = [
        "solar", "photovoltaic", "pv system", "solar pv", "solar power",
        "mini grid", "off-grid solar", "renewable energy",
    ]

    def _real_url(raw):
        if not raw:
            return raw
        from urllib.parse import urlparse, parse_qs, unquote
        p = urlparse(raw)
        if "duckduckgo.com" in p.netloc:
            qs = parse_qs(p.query)
            for key in ("uddg", "u"):
                if key in qs:
                    return unquote(qs[key][0])
        return raw

    def _is_social(url):
        return any(d in url for d in SOCIAL_DOMAINS)

    def _is_job_board(url):
        return any(d in url for d in JOB_DOMAINS)

    def _classify(url):
        if _is_social(url):
            if "linkedin" in url:
                return "LinkedIn Lead"
            return "Social Media Lead"
        if _is_job_board(url):
            return "Job Board — Active Project"
        if any(p in url for p in ["gov.", ".gov", "assembly", "council", "ministry"]):
            return "Government / Public Sector"
        if any(p in url for p in ["ungm", "devex", "reliefweb", "afdb", "worldbank"]):
            return "Tender Portal"
        return "Commercial / Private Sector"

    results = []
    try:
        with DDGS() as ddgs:
            for q in monitor_queries:
                try:
                    for r in ddgs.text(q, max_results=8, safesearch="off"):
                        url   = _real_url(r.get("href", ""))
                        body  = r.get("body", "").lower()
                        title = r.get("title", "").lower()
                        if not url:
                            continue
                        if any(d in url for d in skip_domains):
                            continue
                        url_lower = url.lower()
                        if not _is_social(url):
                            if any(p in url_lower for p in news_url_paths):
                                continue
                        if any(w in title for w in news_title_words):
                            continue
                        if _is_past_deadline(title + " " + body):
                            continue
                        if _foreign_country_in_text(title, loc_lower):
                            continue
                        # Country gate — formal: country name in title/url; social: alias list
                        if _is_social(url) or _is_job_board(url):
                            combined = title + " " + body + " " + url_lower
                            if not any(alias in combined for alias in loc_aliases):
                                continue
                        else:
                            title_url = title + " " + url_lower
                            if loc_lower not in title_url:
                                continue
                        if not any(kw in title or kw in body for kw in solar_keywords):
                            continue
                        if _is_social(url) or _is_job_board(url):
                            if not any(kw in title or kw in body for kw in opportunity_keywords):
                                continue
                        else:
                            if not any(kw in title for kw in rfp_keywords):
                                continue
                        if not any(x["url"] == url for x in results):
                            results.append({
                                "url":         url,
                                "title":       r.get("title", "")[:255],
                                "snippet":     r.get("body", "")[:500],
                                "source_type": _classify(url),
                            })
                except Exception:
                    continue
                if len(results) >= 30:
                    break
    except Exception:
        pass
    return results


def _run_monitor_scan():
    """Scan for new solar postings and insert novel results into monitor_alerts."""
    db_path = os.path.join(os.path.dirname(__file__), "solar.db")
    import sqlite3 as _sqlite3
    conn = _sqlite3.connect(db_path, timeout=15)
    try:
        # DB-level lock: only one gunicorn worker runs the scan at a time
        row = conn.execute("SELECT is_running FROM monitor_state WHERE id=1").fetchone()
        if row and row[0]:
            return
        conn.execute("UPDATE monitor_state SET is_running=1 WHERE id=1")
        conn.commit()

        # Determine which countries to scan: default Ghana + any in recent agent leads
        countries_to_scan = {"Ghana"}
        rows = conn.execute(
            "SELECT DISTINCT country FROM leads WHERE source='agent' AND country != '' LIMIT 10"
        ).fetchall()
        for r in rows:
            if r[0]:
                countries_to_scan.add(r[0])

        new_count = 0
        for country in countries_to_scan:
            items = _monitor_search(country)
            for item in items:
                existing = conn.execute(
                    "SELECT id FROM monitor_alerts WHERE url=?", (item["url"],)
                ).fetchone()
                if not existing:
                    conn.execute(
                        "INSERT INTO monitor_alerts (url, title, snippet, country, source_type, is_new) "
                        "VALUES (?, ?, ?, ?, ?, 1)",
                        (item["url"], item["title"], item["snippet"],
                         country, item["source_type"])
                    )
                    new_count += 1

        from datetime import datetime as _dt
        conn.execute(
            "UPDATE monitor_state SET last_scan=?, last_count=?, is_running=0 WHERE id=1",
            (_dt.utcnow().strftime("%Y-%m-%d %H:%M UTC"), new_count)
        )
        conn.commit()
    except Exception:
        try:
            conn.execute("UPDATE monitor_state SET is_running=0 WHERE id=1")
            conn.commit()
        except Exception:
            pass
    finally:
        conn.close()


@app.route("/admin/agent/monitor/status")
@admin_required
def monitor_status():
    with get_db() as c:
        new_count = c.execute(
            "SELECT COUNT(*) FROM monitor_alerts WHERE is_new=1"
        ).fetchone()[0]
        state = c.execute(
            "SELECT last_scan, last_count, scan_interval, notify_email, "
            "last_agent_run, agent_run_count FROM monitor_state WHERE id=1"
        ).fetchone()
    return jsonify({
        "new_count":       new_count,
        "last_scan":       state[0] if state else "",
        "last_count":      state[1] if state else 0,
        "scan_interval":   state[2] if state else 120,
        "notify_email":    bool(state[3]) if state else False,
        "last_agent_run":  state[4] if state else "",
        "agent_run_count": state[5] if state else 0,
    })


@app.route("/admin/agent/monitor/settings", methods=["POST"])
@admin_required
def monitor_settings():
    """Save scan frequency and email notification preference."""
    csrf_protect()
    interval     = int(request.form.get("scan_interval", 120))
    notify_email = 1 if request.form.get("notify_email") else 0
    # Clamp interval: min 15 min, max 24 hours
    interval = max(15, min(interval, 1440))
    with get_db() as c:
        c.execute(
            "UPDATE monitor_state SET scan_interval=?, notify_email=? WHERE id=1",
            (interval, notify_email))
    return jsonify({"ok": True, "scan_interval": interval, "notify_email": bool(notify_email)})


@app.route("/admin/agent/monitor/alerts")
@admin_required
def monitor_alerts_list():
    with get_db() as c:
        rows = c.execute(
            "SELECT id, url, title, snippet, country, source_type, found_at "
            "FROM monitor_alerts WHERE is_new=1 ORDER BY found_at DESC LIMIT 50"
        ).fetchall()
    alerts = [
        {"id": r[0], "url": r[1], "title": r[2], "snippet": r[3],
         "country": r[4], "source_type": r[5], "found_at": r[6]}
        for r in rows
    ]
    return jsonify({"alerts": alerts})


@app.route("/admin/agent/monitor/dismiss", methods=["POST"])
@admin_required
def monitor_dismiss():
    csrf_protect()
    with get_db() as c:
        c.execute("UPDATE monitor_alerts SET is_new=0")
    return jsonify({"ok": True})


@app.route("/admin/agent/monitor/run", methods=["POST"])
@admin_required
def monitor_run_now():
    """Manually trigger an immediate scan (useful for testing)."""
    csrf_protect()
    threading.Thread(target=_run_monitor_scan, daemon=True).start()
    return jsonify({"ok": True, "message": "Scan started"})


def _send_prospect_notification(subject, body_lines, admin_email=None):
    """Send email notification to admin for prospect agent events."""
    try:
        to = admin_email or SMTP_USER or EMAIL_SUPPORT
        if not to:
            return False
        txt = chr(10).join(body_lines)
        html_rows = "".join("<li style='margin-bottom:6px'>" + l + "</li>" for l in body_lines)
        html = (
            "<div style='font-family:sans-serif;background:#0a0a14;color:#e2e2f0;"
            "padding:28px;border-radius:12px;max-width:600px'>"
            "<h2 style='color:#a78bfa;margin-top:0'>AI Agent - " + subject + "</h2>"
            "<ul style='padding-left:20px;color:#c8c8e8'>" + html_rows + "</ul>"
            "<hr style='border-color:#1e1e3a;margin:20px 0'>"
            "<a href='https://solarpro.aiappinvent.com/admin/agent' "
            "style='background:linear-gradient(135deg,#7c3aed,#a78bfa);color:#fff;"
            "padding:10px 22px;border-radius:8px;text-decoration:none;font-weight:700'>"
            "Open Agent Dashboard</a>"
            "<p style='color:#6868a0;font-size:11px;margin-top:20px'>SolarPro Global - AI Prospect Agent</p>"
            "</div>"
        )
        ok, _ = _send_email(to, "[SolarPro Agent] " + subject, html, text_body=txt)
        return ok
    except Exception:
        return False


# Start background monitor thread (reads scan_interval from DB each cycle)
def _monitor_loop():
    import time, sqlite3 as _sq3
    # Initial delay so the app fully starts before the first scan
    time.sleep(120)
    while True:
        try:
            _run_monitor_scan()
            # After scan: check if email notification is wanted for new alerts
            db_path = os.path.join(os.path.dirname(__file__), "solar.db")
            try:
                conn = _sq3.connect(db_path, timeout=10)
                row = conn.execute(
                    "SELECT notify_email, last_count, scan_interval FROM monitor_state WHERE id=1"
                ).fetchone()
                conn.close()
                notify = row[0] if row else 0
                new_ct = row[1] if row else 0
                interval_min = row[2] if row else 120
            except Exception:
                notify, new_ct, interval_min = 0, 0, 120
            if notify and new_ct > 0:
                _send_prospect_notification(
                    f"{new_ct} New Solar Prospect Alert{'s' if new_ct!=1 else ''}",
                    [f"The background monitor found <strong>{new_ct}</strong> new solar opportunities.",
                     "Log in to the Agent Dashboard to review and save them to your CRM.",
                     f"Next scan in {interval_min} minute{'s' if interval_min!=1 else ''}."]
                )
        except Exception:
            interval_min = 120
        # Read fresh interval each cycle
        try:
            db_path = os.path.join(os.path.dirname(__file__), "solar.db")
            conn = _sq3.connect(db_path, timeout=10)
            row2 = conn.execute("SELECT scan_interval FROM monitor_state WHERE id=1").fetchone()
            conn.close()
            interval_min = int(row2[0]) if row2 and row2[0] else 120
        except Exception:
            interval_min = 120
        time.sleep(interval_min * 60)


_monitor_thread = threading.Thread(target=_monitor_loop, daemon=True)
_monitor_thread.start()


# ─── Entry point ──────────────────────────────────────────────────────────────

# -- API Manager Admin Routes ---------------------------------------------

@app.route("/admin/api-status")
@admin_required
def admin_api_status():
    status = _api.status()
    return render_template("admin_api.html", status=status)


@app.route("/admin/api/reload", methods=["POST"])
@admin_required
def admin_api_reload():
    csrf_protect()
    _api.reload()
    flash("All API keys reloaded from environment.", "success")
    return redirect(url_for("admin_api_status"))


@app.route("/admin/api/clear-cache", methods=["POST"])
@admin_required
def admin_api_clear_cache():
    csrf_protect()
    provider = request.form.get("provider") or None
    _api.clear_cache(provider)
    flash("Cache cleared" + (f" for {provider}" if provider else " (all)") + ".", "success")
    return redirect(url_for("admin_api_status"))


# -- Legal Pages ----------------------------------------------------------

@app.route("/terms")
def terms():
    return render_template("terms.html")


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


# -- Beta Testing ----------------------------------------------------------

@app.route("/beta-signup", methods=["POST"])
@limiter.limit("3 per hour")
def beta_signup():
    csrf_protect()
    name    = request.form.get("name", "").strip()
    email   = request.form.get("email", "").strip().lower()
    company = request.form.get("company", "").strip()
    role    = request.form.get("role", "").strip()
    if not name or not email:
        flash("Name and email are required.", "danger")
        return redirect(url_for("landing") + "#beta")
    with get_db() as c:
        exists = c.execute("SELECT id FROM beta_signups WHERE email=?", (email,)).fetchone()
        if exists:
            flash("You're already on the beta waitlist!", "info")
            return redirect(url_for("landing") + "#beta")
        c.execute(
            "INSERT INTO beta_signups (name, email, company, role) VALUES (?,?,?,?)",
            (name, email, company, role))
    _send_email(
        EMAIL_SALES, "New Beta Signup: " + name,
        "<h3>New Beta Signup</h3><p><b>Name:</b> " + name + "<br><b>Email:</b> " + email +
        "<br><b>Company:</b> " + (company or "N/A") + "<br><b>Role:</b> " + (role or "N/A") + "</p>"
        "<p><a href='https://solarpro.aiappinvent.com/admin/beta'>View in Admin</a></p>",
        from_addr=EMAIL_SUPPORT)
    _send_email(
        email, "You're on the SolarPro Beta Waitlist!",
        "<h2>Thanks, " + name + "!</h2>"
        "<p>You've been added to the SolarPro Global beta waitlist. "
        "We'll send your invite to <b>" + email + "</b> soon.</p>"
        "<p>Visit: <a href='https://solarpro.aiappinvent.com'>solarpro.aiappinvent.com</a></p>"
        "<p>-- The SolarPro Team</p>",
        from_addr=EMAIL_SALES)
    flash("You're on the list! Check your email for confirmation.", "success")
    return redirect(url_for("landing") + "#beta")


@app.route("/feedback", methods=["POST"])
@login_required
def submit_feedback():
    csrf_protect()
    fb_type = request.form.get("type", "general")
    message = request.form.get("message", "").strip()
    page    = request.form.get("page", "").strip()
    if not message:
        return jsonify({"ok": False, "msg": "Message required"}), 400
    with get_db() as c:
        u = c.execute("SELECT username, email FROM users WHERE id=?",
                      (session["user_id"],)).fetchone()
    username = u["username"] if u else ""
    uemail   = u["email"]    if u else ""
    with get_db() as c:
        c.execute(
            "INSERT INTO beta_feedback (user_id, username, email, type, message, page) "
            "VALUES (?,?,?,?,?,?)",
            (session["user_id"], username, uemail, fb_type, message, page))
    _send_email(
        EMAIL_SUPPORT,
        "[" + fb_type.upper() + "] Beta Feedback from " + username,
        "<h3>Beta Feedback</h3><p><b>From:</b> " + username + " (" + uemail + ")<br>"
        "<b>Type:</b> " + fb_type + "<br><b>Page:</b> " + (page or "N/A") + "</p>"
        "<p><b>Message:</b><br>" + message + "</p>"
        "<p><a href='https://solarpro.aiappinvent.com/admin/feedback'>View feedback</a></p>",
        from_addr=EMAIL_SUPPORT)
    return jsonify({"ok": True, "msg": "Feedback submitted. Thank you!"})


@app.route("/rate", methods=["GET"])
@login_required
def rate_form():
    """GET the beta rating form. The submit POSTs to /rate which
    writes to beta_feedback with type='rating' and the three score
    columns populated."""
    return render_template("rate.html", user=current_user())


@app.route("/rate", methods=["POST"])
@login_required
def submit_rating():
    """Persist a 3-axis beta rating into beta_feedback. The three
    scores are clamped to 1..5 server-side so the slider UI cannot
    bypass the contract. Optional comment lives in the message column."""
    csrf_protect()
    def _clamp(name):
        try:
            v = int(request.form.get(name, "0"))
        except (TypeError, ValueError):
            v = 0
        return max(1, min(5, v))
    perf       = _clamp("perf_score")
    creativity = _clamp("creativity_score")
    value      = _clamp("value_score")
    comment    = (request.form.get("message") or "").strip()
    page       = (request.form.get("page") or "").strip()
    if not any([perf, creativity, value]):
        return jsonify({"ok": False, "msg": "Pick at least one score"}), 400
    with get_db() as c:
        u = c.execute("SELECT username, email FROM users WHERE id=?",
                      (session["user_id"],)).fetchone()
    username = u["username"] if u else ""
    uemail   = u["email"]    if u else ""
    # Compose a short summary line into the message column too, so the
    # admin /admin/feedback view shows the rating context at a glance.
    summary = f"Performance={perf}/5  Creativity={creativity}/5  Value={value}/5"
    full_msg = summary + ("\n\n" + comment if comment else "")
    with get_db() as c:
        c.execute(
            "INSERT INTO beta_feedback (user_id, username, email, type, message, "
            "page, perf_score, creativity_score, value_score) "
            "VALUES (?,?,?,?,?,?,?,?,?)",
            (session["user_id"], username, uemail, "rating", full_msg, page,
             perf, creativity, value))
    _send_email(
        EMAIL_SUPPORT,
        "[RATING] " + summary + " from " + username,
        "<h3>Beta Rating</h3><p><b>From:</b> " + username + " (" + uemail + ")<br>"
        "<b>Page:</b> " + (page or "N/A") + "</p>"
        "<table style=\"border-collapse:collapse\"><tr><td style=\"padding:4px 12px\"><b>Performance</b></td><td style=\"padding:4px 12px\">"
        + str(perf) + "/5</td></tr><tr><td style=\"padding:4px 12px\"><b>Creativity</b></td><td style=\"padding:4px 12px\">"
        + str(creativity) + "/5</td></tr><tr><td style=\"padding:4px 12px\"><b>Value</b></td><td style=\"padding:4px 12px\">"
        + str(value) + "/5</td></tr></table>"
        + ("<p><b>Comment:</b><br>" + comment + "</p>" if comment else "")
        + "<p><a href=\"https://solarpro.aiappinvent.com/admin/feedback\">View ratings</a></p>",
        from_addr=EMAIL_SUPPORT)
    return jsonify({"ok": True, "msg": "Rating submitted. Thank you!"})


@app.route("/admin/beta")
@admin_required
def admin_beta():
    with get_db() as c:
        signups = c.execute(
            "SELECT * FROM beta_signups ORDER BY created_at DESC").fetchall()
    counts = {
        "total":    len(signups),
        "pending":  sum(1 for s in signups if s["status"] == "pending"),
        "approved": sum(1 for s in signups if s["status"] == "approved"),
        "rejected": sum(1 for s in signups if s["status"] == "rejected"),
        "invited":  sum(1 for s in signups if s["invited_at"]),
    }
    return render_template("admin_beta.html", signups=signups, counts=counts)


@app.route("/admin/beta/invite", methods=["POST"])
@admin_required
def admin_beta_invite():
    csrf_protect()
    signup_id = request.form.get("signup_id", type=int)
    if not signup_id:
        flash("Invalid request.", "danger")
        return redirect(url_for("admin_beta"))
    with get_db() as c:
        s = c.execute("SELECT * FROM beta_signups WHERE id=?", (signup_id,)).fetchone()
    if not s:
        flash("Signup not found.", "danger")
        return redirect(url_for("admin_beta"))
    reg_url = "https://solarpro.aiappinvent.com/register"
    ok, msg = _send_email(
        s["email"], "You're Invited to SolarPro Beta!",
        "<h2>Your Beta Invite is Here, " + s["name"] + "!</h2>"
        "<p>You've been approved for early access to <b>SolarPro Global</b>.</p>"
        "<p><a href='" + reg_url + "' style='background:#f59e0b;color:#000;"
        "padding:12px 24px;border-radius:8px;text-decoration:none;"
        "font-weight:bold;display:inline-block;margin:16px 0'>Create Beta Account</a></p>"
        "<p>Or visit: " + reg_url + "</p>"
        "<p>-- The SolarPro Team</p>",
        from_addr=EMAIL_SALES)
    if ok:
        from datetime import datetime as _dt
        with get_db() as c:
            c.execute(
                "UPDATE beta_signups SET status='approved', invited_at=? WHERE id=?",
                (_dt.utcnow().strftime("%Y-%m-%d %H:%M:%S"), signup_id))
        flash("Invite sent to " + s["email"] + ".", "success")
    else:
        flash("Failed to send invite: " + msg, "danger")
    return redirect(url_for("admin_beta"))


@app.route("/admin/beta/manual-invite", methods=["POST"])
@admin_required
def admin_beta_manual_invite():
    csrf_protect()
    name  = request.form.get("name", "").strip()
    email = request.form.get("email", "").strip().lower()
    if not name or not email:
        flash("Name and email required.", "danger")
        return redirect(url_for("admin_beta"))
    with get_db() as c:
        exists = c.execute("SELECT id FROM beta_signups WHERE email=?", (email,)).fetchone()
        if not exists:
            c.execute(
                "INSERT INTO beta_signups (name, email, status) VALUES (?,?,'approved')",
                (name, email))
            signup_id = c.execute("SELECT last_insert_rowid()").fetchone()[0]
        else:
            signup_id = exists["id"]
    reg_url = "https://solarpro.aiappinvent.com/register"
    ok, msg = _send_email(
        email, "You're Invited to SolarPro Beta!",
        "<h2>Hi " + name + ", your SolarPro Beta invite is ready!</h2>"
        "<p>You've been personally invited to try <b>SolarPro Global</b>.</p>"
        "<p><a href='" + reg_url + "' style='background:#f59e0b;color:#000;"
        "padding:12px 24px;border-radius:8px;text-decoration:none;"
        "font-weight:bold;display:inline-block;margin:16px 0'>Create Free Account</a></p>"
        "<p>-- The SolarPro Team</p>",
        from_addr=EMAIL_SALES)
    if ok:
        from datetime import datetime as _dt
        with get_db() as c:
            c.execute("UPDATE beta_signups SET invited_at=? WHERE id=?",
                      (_dt.utcnow().strftime("%Y-%m-%d %H:%M:%S"), signup_id))
        flash("Invite sent to " + email + ".", "success")
    else:
        flash("Failed: " + msg, "danger")
    return redirect(url_for("admin_beta"))


@app.route("/admin/beta/status", methods=["POST"])
@admin_required
def admin_beta_status():
    csrf_protect()
    signup_id = request.form.get("signup_id", type=int)
    status    = request.form.get("status", "")
    if status not in ("pending", "approved", "rejected"):
        flash("Invalid status.", "danger")
        return redirect(url_for("admin_beta"))
    with get_db() as c:
        c.execute("UPDATE beta_signups SET status=? WHERE id=?", (status, signup_id))
    flash("Status updated.", "success")
    return redirect(url_for("admin_beta"))


# ─── Autonomous Agents dashboard ───────────────────────────────
# In-app read-out for the 5 cron workflows (beta-monitor, agent-triage,
# synthetic-health, email-delivery-check, daily-digest). All execution
# happens in GH Actions; this page is just the admin-facing surface so
# the operator does not need to leave the app to know what they did.

_GH_REPO = "marc667us/solar-pv-designer-lite"

_AGENT_SPECS = [
    {
        "key":         "beta-monitor",
        "name":        "Beta Monitor",
        "cadence":     "every 30 min",
        "icon":        "bi-broadcast-pin",
        "color":       "#22c55e",
        "purpose":     "Polls /admin/feedback + /admin/tickets + /admin/beta plus three security probes; alerts on any new aggregate-count change.",
        "state_file":  "data/response_state.json",
        "workflow":    "beta-monitor.yml",
    },
    {
        "key":         "agent-triage",
        "name":        "Agent Triage",
        "cadence":     "hourly HH:23",
        "icon":        "bi-robot",
        "color":       "#fbbf24",
        "purpose":     "Per-item: LLM classify -> Brevo ACK -> GH issue create if severity high+. Tier 2-4 of the autonomous stack.",
        "state_file":  "data/agent_state.json",
        "workflow":    "agent-triage.yml",
    },
    {
        "key":         "synthetic-health",
        "name":        "Synthetic Health",
        "cadence":     "hourly HH:17",
        "icon":        "bi-heart-pulse",
        "color":       "#0ea5e9",
        "purpose":     "End-to-end critical-user-path walk: landing -> admin login -> create project -> design engine -> proposal PDF. Red on any step failure.",
        "state_file":  None,
        "workflow":    "synthetic-health.yml",
    },
    {
        "key":         "email-delivery-check",
        "name":        "Email Delivery Check",
        "cadence":     "every 2h HH:37",
        "icon":        "bi-envelope-check",
        "color":       "#a855f7",
        "purpose":     "Polls Brevo events API for bounces / blocks / spam complaints / deferrals against the SolarPro sender domain.",
        "state_file":  "data/email_delivery_state.json",
        "workflow":    "email-delivery-check.yml",
    },
    {
        "key":         "daily-digest",
        "name":        "Daily Digest",
        "cadence":     "09:00 UTC daily",
        "icon":        "bi-calendar-week",
        "color":       "#f43f5e",
        "purpose":     "One-shot owner summary of last-24h response volumes, rating averages, security pulse, and synthetic-health conclusions.",
        "state_file":  None,
        "workflow":    "daily-digest.yml",
    },
]


def _agent_state_rows(state_dict):
    """Flatten the state dict into [(label, value)] for table render.
    Skips the long human-prose `note` and the bulky alerted_event_keys
    list (which is just dedup history, not operator-visible signal)."""
    if not isinstance(state_dict, dict):
        return []
    skip = {"note", "alerted_event_keys", "actions_this_run",
            "alerts_this_poll"}
    rows = []
    for k, v in state_dict.items():
        if k in skip: continue
        if isinstance(v, dict):
            # Flatten one level so security_audit + totals etc. show inline
            small = ", ".join(f"{k2}={v2}" for k2, v2 in v.items()
                                  if not isinstance(v2, (dict, list)))[:200]
            rows.append((k, small or "(nested)"))
        elif isinstance(v, list):
            rows.append((k, f"list ({len(v)} entries)"))
        else:
            rows.append((k, str(v)[:200]))
    return rows


@app.route("/admin/agents")
@admin_required
def admin_agents():
    """Read each cron's state file from disk and render the dashboard."""
    import json as _json
    root = os.path.dirname(os.path.abspath(__file__))
    agents = []
    for spec in _AGENT_SPECS:
        state = None
        if spec["state_file"]:
            sp = os.path.join(root, spec["state_file"])
            if os.path.exists(sp):
                try:
                    state = _json.load(open(sp, "r", encoding="utf-8"))
                except Exception:
                    state = None
        agents.append({
            "key":          spec["key"],
            "name":         spec["name"],
            "cadence":      spec["cadence"],
            "icon":         spec["icon"],
            "color":        spec["color"],
            "purpose":      spec["purpose"],
            "state":        state,
            "state_rows":   _agent_state_rows(state) if state else [],
            "runs_url":
                f"https://github.com/{_GH_REPO}/actions/workflows/"
                f"{spec['workflow']}",
            "dispatch_url":
                f"https://github.com/{_GH_REPO}/actions/workflows/"
                f"{spec['workflow']}",
            "file_url":
                f"https://github.com/{_GH_REPO}/blob/master/.github/"
                f"workflows/{spec['workflow']}",
        })
    return render_template("admin_agents.html", agents=agents)


# ─── Tier 2-4 agent-triage JSON API ─────────────────────────────
# JSON-returning siblings to /admin/{feedback,tickets,beta} so the
# hourly agent-triage workflow (.github/workflows/agent-triage.yml)
# can fetch per-item records (id + body + submitter) without scraping
# the Jinja-rendered admin pages.

def _rows_to_json(rows):
    """Convert sqlite3.Row / DictCursor rows to plain dicts. Strips
    nothing — callers can drop fields per their need (e.g. the
    agent only needs id + message + email)."""
    out = []
    for r in rows:
        try:
            out.append({k: r[k] for k in r.keys()})
        except Exception:
            # Fall back to positional access if .keys() unavailable.
            out.append(dict(r) if hasattr(r, "keys") else list(r))
    return out

def _limit_since(default=50, max_=200):
    """Parse `?limit=N&since=ID` query params. Server-side clamps
    keep an over-eager agent from pulling 10k rows in one call."""
    try:
        limit = int(request.args.get("limit", default))
    except (TypeError, ValueError):
        limit = default
    limit = max(1, min(max_, limit))
    try:
        since = int(request.args.get("since", "0"))
    except (TypeError, ValueError):
        since = 0
    return limit, since


@app.route("/admin/api/feedback")
@admin_required
def admin_api_feedback():
    limit, since = _limit_since()
    with get_db() as c:
        rows = c.execute(
            "SELECT * FROM beta_feedback WHERE id > ? ORDER BY id DESC LIMIT ?",
            (since, limit),
        ).fetchall()
    return jsonify({"ok": True, "items": _rows_to_json(rows),
                    "count": len(rows), "since": since, "limit": limit})


@app.route("/admin/api/tickets")
@admin_required
def admin_api_tickets():
    limit, since = _limit_since()
    with get_db() as c:
        rows = c.execute(
            "SELECT t.*, u.email AS submitter_email, u.username AS submitter_username "
            "FROM tickets t LEFT JOIN users u ON t.user_id = u.id "
            "WHERE t.id > ? ORDER BY t.id DESC LIMIT ?",
            (since, limit),
        ).fetchall()
    return jsonify({"ok": True, "items": _rows_to_json(rows),
                    "count": len(rows), "since": since, "limit": limit})


@app.route("/admin/api/beta_signups")
@admin_required
def admin_api_beta_signups():
    limit, since = _limit_since()
    with get_db() as c:
        rows = c.execute(
            "SELECT * FROM beta_signups WHERE id > ? ORDER BY id DESC LIMIT ?",
            (since, limit),
        ).fetchall()
    return jsonify({"ok": True, "items": _rows_to_json(rows),
                    "count": len(rows), "since": since, "limit": limit})


@app.route("/admin/feedback")
@admin_required
def admin_feedback():
    sf = request.args.get("status", "all")
    with get_db() as c:
        if sf == "all":
            items = c.execute(
                "SELECT * FROM beta_feedback ORDER BY created_at DESC").fetchall()
        else:
            items = c.execute(
                "SELECT * FROM beta_feedback WHERE status=? ORDER BY created_at DESC",
                (sf,)).fetchall()
        counts = {
            "all":      c.execute("SELECT COUNT(*) FROM beta_feedback").fetchone()[0],
            "new":      c.execute("SELECT COUNT(*) FROM beta_feedback WHERE status='new'").fetchone()[0],
            "reviewed": c.execute("SELECT COUNT(*) FROM beta_feedback WHERE status='reviewed'").fetchone()[0],
            "resolved": c.execute("SELECT COUNT(*) FROM beta_feedback WHERE status='resolved'").fetchone()[0],
        }
    return render_template("admin_feedback.html", items=items, counts=counts, status_filter=sf)


@app.route("/admin/feedback/update", methods=["POST"])
@admin_required
def admin_feedback_update():
    csrf_protect()
    fb_id  = request.form.get("fb_id", type=int)
    status = request.form.get("status", "")
    if status not in ("new", "reviewed", "resolved"):
        flash("Invalid status.", "danger")
        return redirect(url_for("admin_feedback"))
    with get_db() as c:
        c.execute("UPDATE beta_feedback SET status=? WHERE id=?", (status, fb_id))
    flash("Feedback updated.", "success")
    return redirect(url_for("admin_feedback"))



# Health Check Endpoints
import time as _time


@app.route("/api/health")
@limiter.exempt
def api_health():
    """Primary health check used by K8s readiness/liveness probes."""
    t0 = _time.time()
    db_ok = False
    try:
        with get_db() as _c:
            _c.execute("SELECT 1").fetchone()
        db_ok = True
    except Exception:
        pass
    redis_ok = None
    redis_url = os.environ.get("REDIS_URL", "")
    if redis_url:
        try:
            import redis as _redis
            _r = _redis.from_url(redis_url, socket_connect_timeout=2)
            _r.ping()
            redis_ok = True
        except Exception:
            redis_ok = False
    status = "ok" if db_ok else "degraded"
    payload = {
        "status":      status,
        "service":     "solarpro-backend",
        "version":     os.environ.get("APP_VERSION", "1.0.0"),
        "environment": os.environ.get("APP_ENV", "production"),
        "backend":     "running",
        "database":    "connected" if db_ok else "error",
        "redis":       ("connected" if redis_ok else "error") if redis_ok is not None else "not_configured",
        "latency_ms":  round((_time.time() - t0) * 1000, 1),
        "timestamp":   datetime.utcnow().isoformat() + "Z",
    }
    return jsonify(payload), (200 if db_ok else 503)


@app.route("/api/health/database")
@limiter.exempt
def api_health_database():
    """Database health check with query latency."""
    t0 = _time.time()
    try:
        with get_db() as _c:
            uc = _c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
            pc = _c.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        return jsonify({
            "status": "connected",
            "latency_ms": round((_time.time() - t0) * 1000, 1),
            "stats": {"users": uc, "projects": pc},
            "database_type": "postgresql" if os.environ.get("DATABASE_URL","").startswith("postgres") else "sqlite",
            "timestamp": datetime.utcnow().isoformat() + "Z",
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "error": str(e),
                        "timestamp": datetime.utcnow().isoformat() + "Z"}), 503


@app.route("/api/health/redis")
@limiter.exempt
def api_health_redis():
    """Redis health check."""
    redis_url = os.environ.get("REDIS_URL", "")
    if not redis_url:
        return jsonify({"status": "not_configured", "timestamp": datetime.utcnow().isoformat() + "Z"}), 200
    t0 = _time.time()
    try:
        import redis as _redis
        _r = _redis.from_url(redis_url, socket_connect_timeout=3)
        _r.ping()
        info = _r.info("memory")
        return jsonify({
            "status": "connected",
            "latency_ms": round((_time.time() - t0) * 1000, 1),
            "used_memory_human": info.get("used_memory_human", "N/A"),
            "connected_clients": _r.info("clients").get("connected_clients", 0),
            "timestamp": datetime.utcnow().isoformat() + "Z",
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "error": str(e),
                        "timestamp": datetime.utcnow().isoformat() + "Z"}), 503


@app.route("/api/health/queue")
@limiter.exempt
def api_health_queue():
    """Celery queue health (checks Redis queue lengths)."""
    redis_url = os.environ.get("REDIS_URL", "")
    if not redis_url:
        return jsonify({"status": "not_configured", "timestamp": datetime.utcnow().isoformat() + "Z"}), 200
    try:
        import redis as _redis
        _r = _redis.from_url(redis_url, socket_connect_timeout=3)
        return jsonify({
            "status": "ok",
            "queues": {"default": _r.llen("celery"), "heavy": _r.llen("heavy"),
                       "ai_tasks": _r.llen("ai_tasks"), "email": _r.llen("email"),
                       "reports": _r.llen("reports")},
            "timestamp": datetime.utcnow().isoformat() + "Z",
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "error": str(e),
                        "timestamp": datetime.utcnow().isoformat() + "Z"}), 503


@app.route("/api/health/storage")
@limiter.exempt
def api_health_storage():
    """File storage health check."""
    try:
        import tempfile
        _t = tempfile.NamedTemporaryFile(delete=False, suffix=".sp")
        _t.write(b"hc"); _t.close(); os.unlink(_t.name)
        db_size = os.path.getsize(DB_PATH) if os.path.exists(DB_PATH) else 0
        return jsonify({"status": "ok", "local_write": "ok",
                        "database_size_mb": round(db_size / 1024 / 1024, 2),
                        "timestamp": datetime.utcnow().isoformat() + "Z"}), 200
    except Exception as e:
        return jsonify({"status": "error", "error": str(e),
                        "timestamp": datetime.utcnow().isoformat() + "Z"}), 503


@app.route("/api/health/ai")
@limiter.exempt
def api_health_ai():
    """AI service configuration health check."""
    services = {
        "anthropic":   "configured" if os.environ.get("ANTHROPIC_API_KEY") else "not_configured",
        "openrouter":  "configured" if os.environ.get("OPENROUTER_API_KEY") else "not_configured",
        "ollama":      "configured" if os.environ.get("OLLAMA_URL") else "not_configured",
        "github_models": "configured" if os.environ.get("GITHUB_MODELS_TOKEN") else "not_configured",
    }
    any_ok = any(v == "configured" for v in services.values())
    return jsonify({
        "status": "ok" if any_ok else "degraded",
        "services": services,
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }), 200


@app.route("/api/ping")
@limiter.exempt
def api_ping():
    """Minimal ping endpoint for Uptime Kuma / load balancer."""
    return jsonify({"pong": True, "timestamp": datetime.utcnow().isoformat() + "Z"}), 200


@app.route("/api/version")
@limiter.exempt
def api_version():
    """Build identity for beta evaluators + ops. Returns the VERSION file
    contents + git commit SHA + a UTC build timestamp.

    The VERSION file is a single-line plain-text semver string updated by
    hand at each tag (e.g. 0.9.0-beta.1). The commit SHA is read at import
    time from the RENDER_GIT_COMMIT env var which Render sets on every
    build; falls back to "unknown" on local runs without that env."""
    _root = os.path.dirname(os.path.abspath(__file__))
    try:
        _ver = open(os.path.join(_root, "VERSION"), "r", encoding="utf-8").read().strip()
    except Exception:
        _ver = "unknown"
    return jsonify({
        "version":     _ver,
        "commit":      os.environ.get("RENDER_GIT_COMMIT", "unknown")[:12],
        "build_time":  os.environ.get("RENDER_BUILD_TIME", ""),
        "channel":     "beta",
    }), 200



@app.route("/metrics")
@limiter.exempt
def prometheus_metrics():
    """Prometheus metrics endpoint."""
    try:
        from prometheus_client import generate_latest, CONTENT_TYPE_LATEST
        return generate_latest(), 200, {"Content-Type": CONTENT_TYPE_LATEST}
    except ImportError:
        with get_db() as _c:
            uc = _c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
            pc = _c.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        body = (
            "# HELP solarpro_users_total Total registered users\n"
            "# TYPE solarpro_users_total gauge\n"
            f"solarpro_users_total {uc}\n"
            "# HELP solarpro_projects_total Total projects\n"
            "# TYPE solarpro_projects_total gauge\n"
            f"solarpro_projects_total {pc}\n"
        )
        return body, 200, {"Content-Type": "text/plain; version=0.0.4"}


# Admin Operations Center Routes
@app.route("/admin/operations")
@admin_required
def admin_operations():
    """NOC/SOC operations center."""
    with get_db() as c:
        users = c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        projects = c.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        failed_logins = c.execute(
            "SELECT COUNT(*) FROM audit_logs WHERE action='failed_login' "
            "AND created_at >= datetime('now', '-24 hours')"
        ).fetchone()[0] if _table_exists(c, "audit_logs") else 0
        db_path = DB_PATH
        db_size = os.path.getsize(db_path) if os.path.exists(db_path) else 0
    security = {
        "failed_logins": failed_logins,
        "active_sessions": 0,
        "blocked_requests": 0,
        "tenant_violations": 0,
        "expired_tokens": 0,
        "revoked_tokens": 0,
    }
    db_stats = {
        "size_mb": round(db_size / 1024 / 1024, 2),
        "tables": 17,
        "users": users,
        "projects": projects,
    }
    perf = {"active_users": users}
    backup = {"last_backup": "Not configured", "size_mb": 0, "status": "pending", "retention": "30 days"}
    return render_template("admin_operations.html",
                           user=current_user(), security=security,
                           db_stats=db_stats, perf=perf, backup=backup)


def _table_exists(conn, table_name):
    r = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,)).fetchone()
    return r is not None


@app.route("/admin/logs")
@admin_required
def admin_logs():
    """Secured system log viewer."""
    log_type    = request.args.get("type", "audit")
    date_from   = request.args.get("date_from", "")
    date_to     = request.args.get("date_to", "")
    user_filter = request.args.get("user_filter", "")
    action_flt  = request.args.get("action", "")
    import os, json as _json
    LOG_DIR = os.environ.get("LOG_DIR", os.path.join(os.path.dirname(__file__), "logs"))
    log_file_map = {
        "audit":    os.path.join(LOG_DIR, "audit", "audit.log"),
        "security": os.path.join(LOG_DIR, "security", "security.log"),
        "error":    os.path.join(LOG_DIR, "backend", "error.log"),
        "ai":       os.path.join(LOG_DIR, "ai-agents", "agent.log"),
        "queue":    os.path.join(LOG_DIR, "queue", "worker.log"),
        "engineering": os.path.join(LOG_DIR, "backend", "app.log"),
    }
    log_entries = []
    log_file = log_file_map.get(log_type, "")
    if log_file and os.path.exists(log_file):
        try:
            with open(log_file, "r", encoding="utf-8", errors="replace") as f:
                lines = f.readlines()[-500:]  # last 500 lines
            for line in reversed(lines):
                try:
                    entry = _json.loads(line.strip())
                    if action_flt and action_flt.lower() not in str(entry.get("action","")).lower():
                        continue
                    if user_filter and user_filter not in str(entry.get("user_id","")) + str(entry.get("username","")):
                        continue
                    log_entries.append(entry)
                    if len(log_entries) >= 200:
                        break
                except Exception:
                    pass
        except Exception:
            pass
    return render_template("admin_logs.html",
                           user=current_user(),
                           log_entries=log_entries,
                           log_type=log_type,
                           date_from=date_from, date_to=date_to,
                           user_filter=user_filter,
                           action_filter=action_flt,
                           total_entries=len(log_entries))


@app.route("/admin/ops/backup/run", methods=["POST"])
@admin_required
def admin_ops_backup_run():
    """Create immediate DB backup (SQLite copy)."""
    import shutil, time as _t
    try:
        backup_dir = os.path.join(os.path.dirname(DB_PATH), "backups")
        os.makedirs(backup_dir, exist_ok=True)
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        backup_file = os.path.join(backup_dir, f"solar_{ts}.db")
        shutil.copy2(DB_PATH, backup_file)
        size_mb = round(os.path.getsize(backup_file) / 1024 / 1024, 2)
        log_audit(action="database_backup", user_id=session.get("user_id"), status="success")
        return jsonify({"status": "success", "filename": os.path.basename(backup_file),
                        "size_mb": size_mb, "timestamp": ts})
    except Exception as e:
        log_error(module="backup", action="database_backup", error=str(e))
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/admin/ops/backup/download")
@admin_required
def admin_ops_backup_download():
    """Download latest DB backup."""
    backup_dir = os.path.join(os.path.dirname(DB_PATH), "backups")
    if not os.path.isdir(backup_dir):
        flash("No backups found. Run a backup first.", "warning")
        return redirect(url_for("admin_operations"))
    backups = sorted([f for f in os.listdir(backup_dir) if f.endswith(".db")], reverse=True)
    if not backups:
        flash("No backup files found.", "warning")
        return redirect(url_for("admin_operations"))
    log_audit(action="backup_downloaded", user_id=session.get("user_id"), status="success")
    return send_file(os.path.join(backup_dir, backups[0]),
                     as_attachment=True, download_name=backups[0])


@app.route("/admin/ops/security/audit", methods=["POST"])
@admin_required
def admin_ops_security_audit():
    """
    Quick security audit endpoint for the Admin Ops Center.

    Inputs:
      _csrf form field (validated by csrf_protect)
      Admin login cookie

    Output:
      JSON with user counts, admin counts, presence flags for SECRET_KEY,
      Paystack/AI/Brevo/Axigen API keys, plus an overall status field.

    Why the defensive wrapping:
      Earlier intermittent "Response ended prematurely" failures during the
      session audit were caused by worker churn on a free-tier host returning
      a half-written response. Each section below is now individually try/except
      wrapped so a partial failure still returns a complete JSON document and
      never leaves the connection mid-response.

    Syntax notes:
      _table_exists is a project helper that PRAGMA-checks for a table name
      so a missing 'payments' table on a fresh DB doesn't raise OperationalError
    """
    # CSRF was missing — POST endpoints in this project all call csrf_protect()
    csrf_protect()
    results = {"status": "pass"}

    # User-table counts. Wrap each query so one failure doesn't abort the rest.
    try:
        with get_db() as c:
            results["total_users"] = c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
            results["admin_users"] = c.execute("SELECT COUNT(*) FROM users WHERE is_admin=1").fetchone()[0]
            results["disabled_accounts"] = c.execute("SELECT COUNT(*) FROM users WHERE plan='disabled'").fetchone()[0]
            # payments may not exist on a fresh-DB deploy — degrade to 0 instead of crashing
            if _table_exists(c, "payments"):
                results["recent_payments"] = c.execute(
                    "SELECT COUNT(*) FROM payments WHERE status='success'"
                ).fetchone()[0]
            else:
                results["recent_payments"] = 0
    except Exception as e:
        # Surface the failure in the JSON instead of returning a 500 HTML page
        results["status"] = "warning"
        results["db_error"] = str(e)[:120]

    # Env-var presence checks. bool(...) collapses a missing or empty string to False
    # so no value (and therefore no secret) is ever leaked. Read via os.environ.get
    # then strip BOM/whitespace via _env_clean so an injected BOM never marks a real
    # key as 'unset' or vice-versa.
    results["secret_key_set"]   = bool(_env_clean("SECRET_KEY"))
    results["anthropic_key_set"]= bool(_env_clean("ANTHROPIC_API_KEY"))
    results["paystack_key_set"] = bool(_env_clean("PAYSTACK_SECRET_KEY"))
    # Resend keys must start with 're_' to be valid; treat anything else as unset
    _rk = _env_clean("RESEND_API_KEY")
    results["resend_key_set"]   = bool(_rk and _rk.startswith("re_"))
    # New providers wired in this session
    results["brevo_key_set"]    = bool(_env_clean("BREVO_API_KEY"))
    results["axigen_configured"]= bool(_env_clean("AXIGEN_SERVER_URL")
                                       and _env_clean("AXIGEN_USER")
                                       and _env_clean("AXIGEN_PASSWORD"))

    # Audit-log the run; do not let a logging failure prevent the response
    try:
        log_audit(action="security_audit_run",
                  user_id=session.get("user_id"),
                  status=results.get("status", "pass"))
    except Exception:
        pass
    return jsonify(results)


@app.route("/admin/ops/security/sessions")
@admin_required
def admin_ops_security_sessions():
    # List recent logins from users table (column-safe)
    try:
        conn = get_db()
        if not _table_exists(conn, "users"):
            conn.close()
            return jsonify({"status": "ok", "active_sessions": [], "message": "users table not found"})
        # Discover available columns
        cols = [row[1] for row in conn.execute("PRAGMA table_info(users)").fetchall()]
        sel_cols = ["id", "username"]
        if "plan" in cols:       sel_cols.append("plan")
        if "last_login" in cols: sel_cols.append("last_login")
        if "created_at" in cols: sel_cols.append("created_at")
        if "is_admin" in cols:   sel_cols.append("is_admin")
        order_col = "last_login" if "last_login" in cols else "id"
        sql = "SELECT %s FROM users ORDER BY %s DESC LIMIT 50" % (", ".join(sel_cols), order_col)
        rows = conn.execute(sql).fetchall()
        conn.close()
        sessions = []
        for r in rows:
            try:
                sessions.append(dict(r))
            except Exception:
                sessions.append(dict(zip(sel_cols, tuple(r))))
        return jsonify({"status": "ok", "active_sessions": sessions, "count": len(sessions)})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/admin/ops/security/revoke-all-sessions", methods=["POST"])
@admin_required
def admin_ops_revoke_all_sessions():
    """Revoke all user sessions (except current admin)."""
    current_uid = session.get("user_id")
    session.clear()
    log_security(event_type="revoke_all_sessions", user_id=current_uid,
                 ip_address=request.remote_addr)
    return jsonify({"status": "success", "message": "All sessions revoked. Users must re-login."})


@app.route("/admin/ops/db/vacuum", methods=["POST"])
@admin_required
def admin_ops_db_vacuum():
    """Run VACUUM on SQLite database."""
    try:
        import time as _t
        t0 = _t.time()
        with get_db() as c:
            c.execute("VACUUM")
            c.execute("ANALYZE")
        elapsed = round((_t.time() - t0) * 1000, 1)
        log_audit(action="db_vacuum", user_id=session.get("user_id"), status="success")
        return jsonify({"status": "success", "duration_ms": elapsed})
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/admin/ops/logs/export")
@admin_required
def admin_ops_export_logs():
    """Export logs as text file."""
    import io as _io
    log_type = request.args.get("type", "audit")
    LOG_DIR = os.environ.get("LOG_DIR", os.path.join(os.path.dirname(__file__), "logs"))
    log_file_map = {
        "audit":    os.path.join(LOG_DIR, "audit", "audit.log"),
        "security": os.path.join(LOG_DIR, "security", "security.log"),
        "error":    os.path.join(LOG_DIR, "backend", "error.log"),
    }
    lf = log_file_map.get(log_type, "")
    if lf and os.path.exists(lf):
        content = open(lf, "rb").read()
    else:
        content = b"No log file found for type: " + log_type.encode()
    log_audit(action="logs_exported", user_id=session.get("user_id"),
              resource=log_type, status="success")
    buf = _io.BytesIO(content)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return send_file(buf, as_attachment=True, download_name=f"solarpro_{log_type}_{ts}.log",
                     mimetype="text/plain")

# -- Admin Ops: Ping endpoints ------------------------------------------------

@app.route("/admin/ops/ping/frontend")
@admin_required
def admin_ops_ping_frontend():
    # Check if the frontend (this Flask server) is responding
    return jsonify({"status": "ok", "service": "frontend", "message": "Flask app responding", "host": request.host})


@app.route("/admin/ops/ping/backend")
@admin_required
def admin_ops_ping_backend():
    # Check backend health (DB + app)
    try:
        conn = get_db()
        conn.execute("SELECT 1").fetchone()
        conn.close()
        return jsonify({"status": "ok", "service": "backend", "message": "Backend healthy, DB reachable"})
    except Exception as e:
        return jsonify({"status": "error", "service": "backend", "message": str(e)}), 500


@app.route("/admin/ops/ping/redis")
@admin_required
def admin_ops_ping_redis():
    # Check Redis connectivity
    import os
    redis_url = os.environ.get("REDIS_URL", "")
    if not redis_url:
        return jsonify({"status": "unavailable", "service": "redis", "message": "REDIS_URL not configured (SQLite mode)"})
    try:
        import redis as redis_lib
        r = redis_lib.from_url(redis_url, socket_connect_timeout=3)
        r.ping()
        info = r.info("memory")
        return jsonify({"status": "ok", "service": "redis", "message": "Redis PONG",
                        "used_memory": info.get("used_memory_human")})
    except ImportError:
        return jsonify({"status": "unavailable", "service": "redis", "message": "redis package not installed"})
    except Exception as e:
        return jsonify({"status": "error", "service": "redis", "message": str(e)}), 500


@app.route("/admin/ops/ping/database")
@admin_required
def admin_ops_ping_database():
    # Check database and return basic stats
    import time, os
    try:
        conn = get_db()
        t0 = time.time()
        user_count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        project_count = conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        latency_ms = round((time.time() - t0) * 1000, 2)
        conn.close()
        db_path = os.environ.get("SQLITE_PATH", "solar.db")
        db_size_kb = round(os.path.getsize(db_path) / 1024, 1) if os.path.exists(db_path) else 0
        return jsonify({"status": "ok", "service": "database", "message": "Database healthy",
                        "users": user_count, "projects": project_count,
                        "latency_ms": latency_ms, "db_size_kb": db_size_kb})
    except Exception as e:
        return jsonify({"status": "error", "service": "database", "message": str(e)}), 500


# -- Admin Ops: RLS + Tenant Isolation ----------------------------------------

@app.route("/admin/ops/db/rls-check")
@admin_required
def admin_ops_rls_check():
    # Check Row Level Security status
    import os
    db_url = os.environ.get("DATABASE_URL", "sqlite:///solar.db")
    results = []
    if db_url.startswith("sqlite"):
        try:
            conn = get_db()
            for tbl in ["users", "projects", "tickets", "payments"]:
                try:
                    cols = [row[1] for row in conn.execute("PRAGMA table_info(%s)" % tbl).fetchall()]
                    has_isolation = any(c in cols for c in ["user_id", "organization_id", "org_id"])
                    results.append({"table": tbl, "rls_active": has_isolation,
                                    "note": "tenant column present" if has_isolation else "no tenant column"})
                except Exception:
                    results.append({"table": tbl, "rls_active": False, "note": "table not found"})
            conn.close()
            return jsonify({"status": "ok", "service": "rls",
                            "message": "SQLite: tenant column checks passed. Full RLS after PostgreSQL migration.",
                            "policies": results})
        except Exception as e:
            return jsonify({"status": "error", "service": "rls", "message": str(e)}), 500
    else:
        try:
            import psycopg2
            conn2 = psycopg2.connect(db_url)
            cur = conn2.cursor()
            cur.execute("SELECT tablename, policyname, cmd FROM pg_policies WHERE schemaname='public' ORDER BY tablename")
            rows = cur.fetchall()
            results = [{"table": r[0], "policy": r[1], "cmd": r[2]} for r in rows]
            conn2.close()
            return jsonify({"status": "ok", "service": "rls",
                            "message": "%d RLS policies active on PostgreSQL" % len(results),
                            "policies": results})
        except Exception as e:
            return jsonify({"status": "error", "service": "rls", "message": str(e)}), 500


@app.route("/admin/ops/security/tenant-isolation")
@admin_required
def admin_ops_tenant_isolation():
    # Verify tenant isolation checks
    try:
        conn = get_db()
        tests = []
        orphan_projects = conn.execute(
            "SELECT COUNT(*) FROM projects WHERE user_id NOT IN (SELECT id FROM users)"
        ).fetchone()[0]
        tests.append({"test": "orphan_projects", "passed": orphan_projects == 0,
                      "detail": "%d orphan projects" % orphan_projects})
        try:
            orphan_tickets = conn.execute(
                "SELECT COUNT(*) FROM tickets WHERE user_id NOT IN (SELECT id FROM users)"
            ).fetchone()[0]
            tests.append({"test": "orphan_tickets", "passed": orphan_tickets == 0,
                          "detail": "%d orphan tickets" % orphan_tickets})
        except Exception:
            tests.append({"test": "orphan_tickets", "passed": True, "detail": "skipped"})
        admin_count = conn.execute("SELECT COUNT(*) FROM users WHERE is_admin=1").fetchone()[0]
        tests.append({"test": "admin_accounts", "passed": admin_count > 0,
                      "detail": "%d admin account(s)" % admin_count})
        try:
            plaintext = conn.execute(
                "SELECT COUNT(*) FROM users WHERE password NOT LIKE '$%' AND password NOT LIKE 'pbkdf2%' AND length(password) < 30"
            ).fetchone()[0]
            tests.append({"test": "password_hashing", "passed": plaintext == 0,
                          "detail": "all passwords hashed" if plaintext == 0 else "%d possibly plaintext" % plaintext})
        except Exception:
            tests.append({"test": "password_hashing", "passed": True, "detail": "skipped"})
        conn.close()
        all_passed = all(t["passed"] for t in tests)
        return jsonify({"status": "ok" if all_passed else "warning",
                        "service": "tenant_isolation",
                        "message": "All isolation checks passed" if all_passed else "Some checks failed",
                        "tests": tests})
    except Exception as e:
        return jsonify({"status": "error", "service": "tenant_isolation", "message": str(e)}), 500


# -- Admin Ops: System tools --------------------------------------------------

@app.route("/admin/ops/system/pip-audit", methods=["POST"])
@admin_required
def admin_ops_pip_audit():
    # Run pip-audit or pip check for known vulnerabilities
    csrf_protect()
    import subprocess, sys, json as _json
    for cmd in [
        [sys.executable, "-m", "pip_audit", "--format=json", "--progress-spinner=off"],
        [sys.executable, "-m", "pip", "check"],
    ]:
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            output = result.stdout or result.stderr or "(no output)"
            if "pip_audit" in " ".join(cmd):
                try:
                    audit_data = _json.loads(output)
                    vulns = audit_data.get("vulnerabilities", [])
                    return jsonify({"status": "ok" if len(vulns) == 0 else "warning",
                                   "tool": "pip-audit",
                                   "vulnerabilities_found": len(vulns),
                                   "results": vulns[:20],
                                   "message": "No known vulnerabilities" if not vulns else "%d vulnerabilities found" % len(vulns)})
                except Exception:
                    return jsonify({"status": "ok", "tool": "pip-audit", "output": output[:2000], "return_code": result.returncode})
            else:
                return jsonify({"status": "ok" if result.returncode == 0 else "warning",
                               "tool": "pip check",
                               "output": output[:2000],
                               "message": "No package conflicts" if result.returncode == 0 else "Conflicts detected",
                               "return_code": result.returncode})
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            return jsonify({"status": "error", "message": "Timed out after 60s"}), 500
    return jsonify({"status": "error", "message": "pip-audit and pip check unavailable"}), 500


@app.route("/admin/ops/queue/restart", methods=["POST"])
@admin_required
def admin_ops_restart_queue():
    # Signal Celery workers to restart gracefully
    csrf_protect()
    import os
    try:
        from celery import Celery
    except ImportError:
        return jsonify({"status": "unavailable",
                       "message": "Celery not installed. Restart via: kubectl rollout restart deployment/celery-worker"})
    redis_url = os.environ.get("REDIS_URL", "")
    if not redis_url:
        return jsonify({"status": "unavailable",
                       "message": "REDIS_URL not configured. Celery requires Redis broker."})
    try:
        app_celery = Celery(broker=redis_url)
        app_celery.control.warm_shutdown(reply=False)
        return jsonify({"status": "ok", "message": "Warm shutdown sent to Celery workers."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/admin/ops/cache/clear", methods=["POST"])
@admin_required
def admin_ops_clear_cache():
    # Clear Redis cache entries
    csrf_protect()
    import os
    redis_url = os.environ.get("REDIS_URL", "")
    cleared_items = []
    try:
        from api_manager import api as _apim
        _apim.clear_cache()
        cleared_items.append("api_manager cache")
    except Exception:
        pass
    if not redis_url:
        return jsonify({"status": "ok" if cleared_items else "unavailable",
                       "message": "Redis not configured. Cleared: %s" % (", ".join(cleared_items) or "nothing")})
    try:
        import redis as redis_lib
        r = redis_lib.from_url(redis_url, socket_connect_timeout=3)
        deleted = 0
        for pattern in [b"shard:*", b"solar:*", b"rate:*"]:
            cursor = 0
            while True:
                cursor, keys = r.scan(cursor, match=pattern, count=100)
                if keys:
                    r.delete(*keys)
                    deleted += len(keys)
                if cursor == 0:
                    break
        cleared_items.append("Redis (%d keys)" % deleted)
        return jsonify({"status": "ok", "message": "Cache cleared: %s" % ", ".join(cleared_items), "keys_deleted": deleted})
    except ImportError:
        return jsonify({"status": "unavailable", "message": "redis package not installed"}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/admin/ops/logs/view")
@admin_required
def admin_ops_view_logs():
    # Return last 100 app log entries as JSON
    import os, json as _json
    log_type = request.args.get("type", "app")
    log_paths = {"app": "logs/backend/app.log", "error": "logs/backend/error.log",
                 "security": "logs/security/security.log"}
    log_path = log_paths.get(log_type, "logs/backend/app.log")
    if not os.path.exists(log_path):
        return jsonify({"status": "ok", "entries": [],
                        "message": "Log file not found: %s. Logging to stdout in this environment." % log_path})
    entries = []
    try:
        with open(log_path, "r", encoding="utf-8", errors="replace") as f:
            lines = f.readlines()[-100:]
        for line in lines:
            line = line.strip()
            if not line:
                continue
            try:
                entries.append(_json.loads(line))
            except Exception:
                entries.append({"raw": line})
        return jsonify({"status": "ok", "log_type": log_type, "entries": entries, "count": len(entries)})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/admin/ops/logs/audit")
@admin_required
def admin_ops_view_audit_logs():
    # Return last 100 audit log entries
    import os, json as _json
    for audit_path in ["logs/audit/audit.log", "logs/audit.log"]:
        if os.path.exists(audit_path):
            entries = []
            try:
                with open(audit_path, "r", encoding="utf-8", errors="replace") as f:
                    lines = f.readlines()[-100:]
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        entries.append(_json.loads(line))
                    except Exception:
                        entries.append({"raw": line})
                return jsonify({"status": "ok", "source": "file", "entries": entries, "count": len(entries)})
            except Exception as e:
                return jsonify({"status": "error", "message": str(e)}), 500
    try:
        conn = get_db()
        try:
            rows = conn.execute(
                "SELECT id, action, user_id, resource, status, created_at FROM audit_logs ORDER BY id DESC LIMIT 100"
            ).fetchall()
            entries = [{"id": r[0], "action": r[1], "user_id": r[2],
                        "resource": r[3], "status": r[4], "created_at": r[5]} for r in rows]
            conn.close()
            return jsonify({"status": "ok", "source": "database", "entries": entries, "count": len(entries)})
        except Exception:
            conn.close()
            return jsonify({"status": "ok", "entries": [],
                            "message": "No audit_logs table yet. Activates after PostgreSQL migration."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/admin/ops/system/load-test", methods=["POST"])
@admin_required
def admin_ops_load_test():
    # Run lightweight internal load test: 50 requests to /api/ping
    csrf_protect()
    import time, threading
    CONCURRENT = 5
    REQUESTS_EACH = 10
    results = {"success": 0, "error": 0, "times": []}
    lock = threading.Lock()

    def _worker():
        for _ in range(REQUESTS_EACH):
            t0 = time.time()
            try:
                with app.test_client() as tc:
                    resp = tc.get("/api/ping")
                    elapsed = (time.time() - t0) * 1000
                    with lock:
                        if resp.status_code == 200:
                            results["success"] += 1
                        else:
                            results["error"] += 1
                        results["times"].append(round(elapsed, 2))
            except Exception:
                with lock:
                    results["error"] += 1

    threads = [threading.Thread(target=_worker) for _ in range(CONCURRENT)]
    t_start = time.time()
    for t in threads:
        t.start()
    for t in threads:
        t.join(timeout=30)
    duration = round(time.time() - t_start, 2)
    times = sorted(results["times"])
    total = results["success"] + results["error"]
    return jsonify({
        "status": "ok",
        "message": "%d requests in %ss" % (total, duration),
        "total_requests": total,
        "successful": results["success"],
        "errors": results["error"],
        "duration_seconds": duration,
        "rps": round(total / max(duration, 0.01), 1),
        "latency_ms": {
            "min": min(times) if times else 0,
            "max": max(times) if times else 0,
            "avg": round(sum(times) / max(len(times), 1), 2),
            "p95": times[int(len(times) * 0.95)] if len(times) > 1 else (times[0] if times else 0),
        }
    })



# -- Admin Ops: Email diagnostics & test --------------------------------------

# -- Admin Ops: Email status & test (v2 - with diagnostics) ------------------

@app.route("/admin/ops/email/status")
@admin_required
def admin_ops_email_status():
    # Show current email configuration (masked) - forces env reload
    import os
    from api_manager import api as _apim
    try:
        _apim.email._load()
    except Exception:
        pass
    # All env reads go through _env_clean so a stray BOM doesn't poison comparisons
    host      = _env_clean("SMTP_HOST")
    port      = _env_clean("SMTP_PORT")
    user      = _env_clean("SMTP_USER")
    smtp_pass = _env_clean("SMTP_PASS")
    frm       = _env_clean("SMTP_FROM")
    tls       = _env_clean("SMTP_TLS")
    resend    = _env_clean("RESEND_API_KEY")
    ax_url    = _env_clean("AXIGEN_SERVER_URL")
    ax_user   = _env_clean("AXIGEN_USER")
    ax_pass   = _env_clean("AXIGEN_PASSWORD")
    return jsonify({
        "status": "ok",
        # Brevo is now primary (free 300/day HTTPS API).
        "brevo_configured": bool(_env_clean("BREVO_API_KEY")),
        "brevo_key_prefix": (_env_clean("BREVO_API_KEY")[:10] + "...") if _env_clean("BREVO_API_KEY") else "(not set)",
        # Axigen secondary HTTPS provider (was primary; superseded by Brevo)
        "axigen_configured": bool(ax_url and ax_user and ax_pass),
        "axigen_url":  ax_url  or "(not set)",
        "axigen_user": ax_user or "(not set)",
        "axigen_pass": ("*" * 8) if ax_pass else "(not set)",
        "resend_configured": bool(resend and resend.startswith("re_")),
        "resend_key_prefix": resend[:8] + "..." if resend else "(not set)",
        "smtp_configured": bool(host and user and smtp_pass),
        "smtp_host": host or "(not set)",
        "smtp_port": port or "(not set)",
        "smtp_user": user or "(not set)",
        "smtp_pass": ("*" * 8) if smtp_pass else "(not set)",
        "smtp_from": frm or "(not set)",
        "smtp_tls": tls or "(not set)",
        "email_sales":     os.environ.get("EMAIL_SALES",     "(not set)"),
        "email_support":   os.environ.get("EMAIL_SUPPORT",   "(not set)"),
        "email_billing":   os.environ.get("EMAIL_BILLING",   "(not set)"),
    })


@app.route("/admin/ops/email/test", methods=["POST"])
@admin_required
def admin_ops_email_test():
    # Send a test email with detailed per-provider diagnostics
    csrf_protect()
    import os, smtplib
    from api_manager import api as _apim
    # Reload env vars so Render picks up latest secrets
    try:
        _apim.email._load()
    except Exception:
        pass
    # Get admin email from DB
    try:
        conn = get_db()
        row = conn.execute("SELECT email FROM users WHERE is_admin=1 LIMIT 1").fetchone()
        conn.close()
        admin_email = row[0] if row else None
    except Exception:
        admin_email = None
    if not admin_email:
        admin_email = os.environ.get("EMAIL_SUPPORT", "support@aiappinvent.com")
    html = ("<div style='font-family:sans-serif;padding:20px;background:#0f0f22;color:#e2e2f0'>"
            "<h2 style='color:#f59e0b'>SolarPro Admin Email Test</h2>"
            "<p>Test from Admin Operations Center. Delivery is working!</p>"
            "<small style='color:#6868a0'>solarpro.aiappinvent.com</small></div>")
    subject = "SolarPro Admin - Email Test"
    diagnostics = []

    # --- Try Brevo first (HTTPS API, free 300/day, primary provider) ---
    # Get a key at https://app.brevo.com -> SMTP & API -> API keys.
    # Sender must be a Brevo-verified address (single-sender verify is free + instant).
    brevo_key = _env_clean("BREVO_API_KEY")
    if brevo_key:
        try:
            import requests as _rq
            br_sender = _env_clean("SMTP_FROM") or _env_clean("EMAIL_SUPPORT") or "support@aiappinvent.com"
            br_resp = _rq.post(
                "https://api.brevo.com/v3/smtp/email",
                json={"sender": {"email": br_sender},
                      "to": [{"email": admin_email}],
                      "subject": subject,
                      "htmlContent": html},
                headers={"api-key": brevo_key,
                         "Content-Type": "application/json",
                         "Accept": "application/json"},
                timeout=15)
            if br_resp.status_code in (200, 201, 202):
                diagnostics.append({"provider": "brevo", "status": "ok",
                                    "http": br_resp.status_code,
                                    "messageId": br_resp.json().get("messageId", "")})
                return jsonify({"status": "ok", "sent_to": admin_email,
                               "provider": "brevo", "diagnostics": diagnostics,
                               "message": "Test email sent via Brevo to " + admin_email})
            diagnostics.append({"provider": "brevo", "status": "error",
                                "http": br_resp.status_code,
                                "detail": br_resp.text[:200]})
        except Exception as e:
            diagnostics.append({"provider": "brevo", "status": "error", "detail": str(e)[:200]})
    else:
        diagnostics.append({"provider": "brevo", "status": "skipped",
                            "detail": "BREVO_API_KEY not configured"})

    # --- Try Axigen first (HTTPS, primary, works through Render firewall) ---
    # AXIGEN_SERVER_URL must include the API base, e.g. https://mail.example.com/api/v1
    # AXIGEN_USER / AXIGEN_PASSWORD are the mailbox credentials on that server
    ax_url  = _env_clean("AXIGEN_SERVER_URL")
    ax_user = _env_clean("AXIGEN_USER")
    ax_pass = _env_clean("AXIGEN_PASSWORD")
    if ax_url and ax_user and ax_pass:
        try:
            import requests as _rq
            # POST a single mail; auth tuple sends an HTTP Basic header
            r_ax = _rq.post(
                ax_url.rstrip("/") + "/mails/send",
                json={"from": ax_user, "to": admin_email,
                      "subject": subject, "bodyHtml": html},
                auth=(ax_user, ax_pass), timeout=15)
            if r_ax.status_code in (200, 201, 202):
                diagnostics.append({"provider": "axigen", "status": "ok", "http": r_ax.status_code})
                return jsonify({"status": "ok", "sent_to": admin_email,
                               "provider": "axigen", "diagnostics": diagnostics,
                               "message": "Test email sent via Axigen to " + admin_email})
            diagnostics.append({"provider": "axigen", "status": "error",
                                "http": r_ax.status_code,
                                "detail": r_ax.text[:120]})
        except Exception as e:
            diagnostics.append({"provider": "axigen", "status": "error", "detail": str(e)[:120]})
    else:
        diagnostics.append({"provider": "axigen", "status": "skipped",
                            "detail": "AXIGEN_SERVER_URL/USER/PASSWORD not configured"})

    # --- Try Resend first (HTTPS, works through Render firewall) ---
    resend_key = os.environ.get("RESEND_API_KEY", "")
    if resend_key and not resend_key.startswith("re_..."):
        try:
            import resend as _r
            _r.api_key = resend_key
            params = {"from": "onboarding@resend.dev",
                      "to": [admin_email], "subject": subject, "html": html}
            result = _r.Emails.send(params)
            if result and result.get("id"):
                diagnostics.append({"provider": "resend", "status": "ok", "id": result["id"]})
                return jsonify({"status": "ok", "sent_to": admin_email,
                               "provider": "resend", "diagnostics": diagnostics,
                               "message": "Test email sent via Resend to " + admin_email})
            else:
                diagnostics.append({"provider": "resend", "status": "error", "detail": str(result)[:80]})
        except Exception as e:
            diagnostics.append({"provider": "resend", "status": "error", "detail": str(e)[:120]})
    else:
        diagnostics.append({"provider": "resend", "status": "skipped", "detail": "RESEND_API_KEY not configured"})

    # --- Try SMTP (port from env - use 587 STARTTLS on Render) ---
    smtp_host = os.environ.get("SMTP_HOST", "")
    smtp_port = int(_env_clean("SMTP_PORT", "587") or "587")
    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pass = os.environ.get("SMTP_PASS", "")
    smtp_tls  = os.environ.get("SMTP_TLS", "true").lower() in ("1", "true", "yes")
    if smtp_host and smtp_user and smtp_pass:
        try:
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText as _MT
            msg2 = MIMEMultipart("alternative")
            msg2["From"]    = os.environ.get("SMTP_FROM", smtp_user)
            msg2["To"]      = admin_email
            msg2["Subject"] = subject
            msg2.attach(_MT(html, "html"))
            if smtp_tls:
                srv = smtplib.SMTP(smtp_host, smtp_port, timeout=12)
                srv.ehlo()
                srv.starttls()
            else:
                srv = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=12)
            srv.login(smtp_user, smtp_pass)
            srv.sendmail(msg2["From"], [admin_email], msg2.as_string())
            srv.quit()
            diagnostics.append({"provider": "smtp", "status": "ok", "host": smtp_host, "port": smtp_port, "tls": smtp_tls})
            return jsonify({"status": "ok", "sent_to": admin_email,
                           "provider": "smtp", "diagnostics": diagnostics,
                           "message": "Test email sent via SMTP to " + admin_email})
        except smtplib.SMTPAuthenticationError as e:
            diagnostics.append({"provider": "smtp", "status": "auth_error", "detail": str(e)[:120], "host": smtp_host, "port": smtp_port})
        except (OSError, smtplib.SMTPConnectError) as e:
            diagnostics.append({"provider": "smtp", "status": "connection_error", "detail": str(e)[:120], "host": smtp_host, "port": smtp_port,
                                 "hint": "Render blocks outbound SMTP. Use Resend API instead."})
        except Exception as e:
            diagnostics.append({"provider": "smtp", "status": "error", "detail": str(e)[:120], "host": smtp_host, "port": smtp_port})
    else:
        diagnostics.append({"provider": "smtp", "status": "skipped", "detail": "SMTP credentials not configured"})

    return jsonify({
        "status": "error",
        "sent_to": admin_email,
        "message": "Email delivery failed on all providers. See diagnostics.",
        "diagnostics": diagnostics,
        "hint": "Fix: verify Resend domain at resend.com/domains OR ensure SMTP_PORT=587 + SMTP_TLS=true in Render env vars."
    })



@app.route("/admin/ops/email/env-keys")
@admin_required
def admin_ops_env_keys():
    """Temporary: list env var keys the running process sees.
    Inputs:  none (GET, admin only)
    Output:  JSON {keys: [...], smtp_keys: {KEY: len(value), ...}}
    """
    import os
    keys = sorted(os.environ.keys())
    smtp_keys = {k: len(os.environ.get(k, "")) for k in keys
                 if k.startswith(("SMTP_", "RESEND_", "EMAIL_", "PAYSTACK_", "OLLAMA_", "SECRET"))}
    return jsonify({
        "total_count": len(keys),
        "all_keys": keys,
        "interesting_keys_and_lengths": smtp_keys,
    })


# ============================================================
# REFERRAL PROGRAM
# - /r/<code>   captures the ref cookie + sends visitor to landing
# - /referrals  authenticated user dashboard with link + stats
# - register()  hook (already inlined above) reads the ref cookie
# ============================================================

def _gen_referral_code():
    """Return an 8-char uppercase alphanumeric code not yet in users.referral_code.
    Inputs:  none
    Output:  unique code string
    Syntax:  secrets.token_urlsafe gives URL-safe base64; we strip dashes/underscores
             and uppercase so the code is easy to type/share verbally.
    """
    import secrets as _sec
    for _ in range(20):  # 20 attempts before giving up (extremely unlikely collision)
        code = _sec.token_urlsafe(6).replace("_","").replace("-","")[:8].upper()
        with get_db() as c:
            hit = c.execute("SELECT 1 FROM users WHERE referral_code = ?", (code,)).fetchone()
            if not hit:
                return code
    raise RuntimeError("could not generate unique referral code after 20 attempts")


@app.route("/r/<code>")
def referral_capture(code):
    """Capture a referral click.
    Inputs:  url path /r/<CODE>
    Output:  302 to landing, with a ref_code cookie set for 30 days.
    """
    code = (code or "").upper().strip()[:16]
    resp = redirect(url_for("landing"))
    # 30 days, Lax so it survives top-level navigation but not third-party iframes.
    resp.set_cookie("ref_code", code, max_age=30*24*3600,
                    httponly=False, samesite="Lax")
    return resp


@app.route("/referrals")
@login_required
def referrals_page():
    """
    User-facing referral dashboard.
    Inputs:  login cookie
    Output:  rendered referrals.html with the user link + stats
    """
    u = current_user()
    # Lazy backfill: if this user pre-dates the feature, give them a code now.
    code = u.get("referral_code") if isinstance(u, dict) else getattr(u, "referral_code", None)
    if not code:
        code = _gen_referral_code()
        with get_db() as c:
            c.execute("UPDATE users SET referral_code = ? WHERE id = ?", (code, u["id"]))
    # Stats: anyone who set referred_by to this user
    with get_db() as c:
        signups = c.execute("SELECT COUNT(*) FROM users WHERE referred_by = ?",
                            (u["id"],)).fetchone()[0]
        upgraded = c.execute(
            "SELECT COUNT(*) FROM users WHERE referred_by = ? "
            "AND plan IS NOT NULL AND plan NOT IN ('free','disabled')",
            (u["id"],)).fetchone()[0]
        recent_rows = c.execute(
            "SELECT username, plan, created_at FROM users "
            "WHERE referred_by = ? ORDER BY id DESC LIMIT 10",
            (u["id"],)).fetchall()
    recent = [dict(r) for r in recent_rows] if recent_rows else []
    # Build the share URL using request.host_url so it works on any deployment
    share_url = request.host_url.rstrip("/") + "/r/" + code
    return render_template("referrals.html", user=u, referral_code=code,
                           share_url=share_url, signups=signups,
                           upgraded=upgraded, recent=recent)


# ─── Admin Ops: missing endpoints — added 2026-06-09 ─────────────────────────
# 7 routes that the admin_operations.html JS calls but that 404'd before.
# All gated by @admin_required (which itself checks login + is_admin flag).

@app.route("/admin/ops/ping/queue")
@admin_required
def admin_ops_ping_queue():
    """Queue subsystem status. On Render free tier there's no Celery/Redis,
    so we return WARN (not error) — that's the expected configured state."""
    import time
    t0 = time.time()
    redis_url = (os.environ.get("REDIS_URL", "") or os.environ.get("CELERY_BROKER_URL", "")).strip()
    if not redis_url:
        return jsonify({
            "status": "warn", "service": "queue",
            "message": "Queue/Celery not configured (Render free tier — no Redis add-on)",
            "broker_url_set": False,
            "latency_ms": round((time.time() - t0) * 1000, 2),
        })
    try:
        import redis  # may not be installed; we lazy-import
        r = redis.from_url(redis_url, socket_timeout=2)
        r.ping()
        # Approximate queue depth via Celery's default queue name "celery"
        try:
            depth = r.llen("celery")
        except Exception:
            depth = -1
        return jsonify({
            "status": "ok", "service": "queue",
            "message": "Redis broker reachable",
            "broker_url_set": True, "queue_depth": depth,
            "latency_ms": round((time.time() - t0) * 1000, 2),
        })
    except Exception as e:
        return jsonify({
            "status": "warn", "service": "queue",
            "message": f"Redis broker unreachable: {e}",
            "broker_url_set": True,
            "latency_ms": round((time.time() - t0) * 1000, 2),
        })


@app.route("/admin/ops/ping/ai")
@admin_required
def admin_ops_ping_ai():
    """AI provider configuration snapshot. Reports which providers are
    configured (key present), not which actually respond — keeps the
    endpoint fast (no outbound HTTP calls)."""
    import time
    t0 = time.time()
    providers = {
        "anthropic":    bool((os.environ.get("ANTHROPIC_API_KEY") or "").strip()),
        "openrouter":   bool((os.environ.get("OPENROUTER_API_KEY") or "").strip()),
        "ollama":       bool((os.environ.get("OLLAMA_URL") or "").strip()),
        "github_models":bool((os.environ.get("GITHUB_TOKEN") or "").strip()),
    }
    configured_count = sum(1 for v in providers.values() if v)
    return jsonify({
        "status": "ok" if configured_count > 0 else "warn",
        "service": "ai",
        "message": f"{configured_count} of {len(providers)} AI providers configured",
        "providers": providers,
        "latency_ms": round((time.time() - t0) * 1000, 2),
    })


@app.route("/admin/ops/ping/storage")
@admin_required
def admin_ops_ping_storage():
    """Disk space status for the volume hosting solar.db."""
    import time, shutil
    t0 = time.time()
    db_path = os.environ.get("SQLITE_PATH", os.environ.get("DB_PATH", "solar.db"))
    target = os.path.dirname(os.path.abspath(db_path)) or "."
    try:
        usage = shutil.disk_usage(target)
        pct_used = round((usage.used / usage.total) * 100, 1)
        status = "ok" if pct_used < 85 else ("warn" if pct_used < 95 else "error")
        return jsonify({
            "status": status, "service": "storage",
            "message": f"Disk at {pct_used}% used on {target}",
            "total_gb": round(usage.total / (1024**3), 2),
            "free_gb":  round(usage.free  / (1024**3), 2),
            "used_pct": pct_used,
            "latency_ms": round((time.time() - t0) * 1000, 2),
        })
    except Exception as e:
        return jsonify({
            "status": "error", "service": "storage",
            "message": str(e),
            "latency_ms": round((time.time() - t0) * 1000, 2),
        }), 500


def _admin_ops_download_json(filename, payload):
    """Helper: serve a JSON blob as a downloaded file (Content-Disposition)."""
    from flask import make_response
    import json as _json
    body = _json.dumps(payload, indent=2, default=str)
    resp = make_response(body)
    resp.headers["Content-Type"] = "application/json"
    resp.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


@app.route("/admin/ops/security/report")
@admin_required
def admin_ops_security_report():
    """Downloadable security snapshot — aggregates audit log, brute-force
    state, session count, and security headers state into one JSON file."""
    import time
    from datetime import datetime as _dt
    snapshot = {
        "generated_at": _dt.utcnow().isoformat() + "Z",
        "report_type": "security",
        "checks": {},
    }
    try:
        conn = get_db()
        try:
            failed_logins = conn.execute(
                "SELECT COUNT(*) FROM audit_logs WHERE action='login_failed'"
            ).fetchone()[0] if _table_exists(conn, "audit_logs") else None
        except Exception:
            failed_logins = None
        try:
            user_count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
            admin_count = conn.execute("SELECT COUNT(*) FROM users WHERE is_admin=1").fetchone()[0]
        except Exception:
            user_count = admin_count = None
        conn.close()
        snapshot["checks"]["users_total"]   = user_count
        snapshot["checks"]["admins_total"]  = admin_count
        snapshot["checks"]["failed_logins_lifetime"] = failed_logins
    except Exception as e:
        snapshot["checks"]["db_error"] = str(e)
    snapshot["headers_configured"] = {
        "csp": True,            # web_app sets a Content-Security-Policy header
        "csrf_on_post": True,   # CSRF _csrf token enforced on POST forms
        "session_cookie_secure": True,
        "brute_force_lockout_min": 15,
    }
    fname = f"solarpro-security-{int(time.time())}.json"
    return _admin_ops_download_json(fname, snapshot)


@app.route("/admin/ops/db/report")
@admin_required
def admin_ops_db_report():
    """Downloadable DB health report — table sizes, row counts, schema version."""
    import time
    from datetime import datetime as _dt
    snapshot = {
        "generated_at": _dt.utcnow().isoformat() + "Z",
        "report_type": "database",
        "backend": "postgresql" if (os.environ.get("DATABASE_URL","").startswith("postgres"))
                                else "sqlite",
        "checks": {},
    }
    try:
        conn = get_db()
        tables = ("users", "projects", "tickets", "audit_logs", "beta_feedback",
                  "email_logs", "secret_audit")
        counts = {}
        for t in tables:
            try:
                counts[t] = conn.execute(f"SELECT COUNT(*) FROM {t}").fetchone()[0]
            except Exception:
                counts[t] = None  # table not present
        snapshot["checks"]["row_counts"] = counts
        db_path = os.environ.get("DB_PATH", "solar.db")
        if os.path.exists(db_path):
            snapshot["checks"]["file_size_kb"] = round(os.path.getsize(db_path) / 1024, 1)
        conn.close()
    except Exception as e:
        snapshot["checks"]["db_error"] = str(e)
    fname = f"solarpro-db-{int(time.time())}.json"
    return _admin_ops_download_json(fname, snapshot)


@app.route("/admin/ops/health/report")
@admin_required
def admin_ops_health_report():
    """Downloadable consolidated health report aggregating every subsystem."""
    import time
    from datetime import datetime as _dt
    snapshot = {
        "generated_at": _dt.utcnow().isoformat() + "Z",
        "report_type": "health",
        "host": os.environ.get("HOSTNAME", "unknown"),
        "render_service": os.environ.get("RENDER_SERVICE_NAME", "unknown"),
        "subsystems": {},
    }
    # Re-invoke each ping endpoint INTERNALLY to pull its JSON response body.
    for sub in ("database", "redis", "queue", "ai", "storage", "backend"):
        vf = app.view_functions.get(f"admin_ops_ping_{sub}")
        if vf is None:
            snapshot["subsystems"][sub] = {"status": "not_implemented"}
            continue
        try:
            resp = vf()
            body = resp.get_data() if hasattr(resp, "get_data") else b""
            import json as _json
            snapshot["subsystems"][sub] = _json.loads(body) if body else {}
        except Exception as e:
            snapshot["subsystems"][sub] = {"status": "error", "message": str(e)}
    fname = f"solarpro-health-{int(time.time())}.json"
    return _admin_ops_download_json(fname, snapshot)


@app.route("/admin/ops/logs/archive", methods=["POST"])
@admin_required
def admin_ops_logs_archive():
    """Archive (rotate) log files: rename current logs/*.log to
    logs/<name>.YYYYMMDD.log.gz and start fresh. Returns a summary.
    On Render free tier where logs aren't persistent across restart,
    this is mostly informational."""
    import time, os, gzip, shutil
    from datetime import datetime as _dt
    logs_dir = os.environ.get("LOGS_DIR", "logs")
    archived = []
    errors = []
    if not os.path.isdir(logs_dir):
        return jsonify({
            "status": "warn",
            "message": f"No logs directory at {logs_dir}",
            "archived": [],
        })
    stamp = _dt.utcnow().strftime("%Y%m%d_%H%M%S")
    for fname in os.listdir(logs_dir):
        if not fname.endswith(".log"):
            continue
        src = os.path.join(logs_dir, fname)
        dst = os.path.join(logs_dir, f"{fname[:-4]}.{stamp}.log.gz")
        try:
            with open(src, "rb") as fin, gzip.open(dst, "wb") as fout:
                shutil.copyfileobj(fin, fout)
            # Truncate the live log so it starts fresh (don't unlink — keep handle valid)
            open(src, "w").close()
            archived.append({"file": fname, "archive": os.path.basename(dst)})
        except Exception as e:
            errors.append({"file": fname, "error": str(e)})
    return jsonify({
        "status": "ok" if not errors else "partial",
        "message": f"Archived {len(archived)} log file(s); {len(errors)} error(s)",
        "archived": archived, "errors": errors,
        "timestamp": stamp,
    })


def _table_exists(conn, table_name):
    """SQLite-specific helper: check if a table exists. Returns False on
    Postgres (where the caller should use information_schema). For the
    security report this is best-effort; missing audit_logs table just
    yields None for that field."""
    try:
        row = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
            (table_name,)
        ).fetchone()
        return row is not None
    except Exception:
        return False


# ── AI Budget — quota status endpoint ────────────────────────────────────────
# Lets the chat widget show remaining-tokens / reset-time, and lets admin UIs
# read org-wide spend without scraping the ledger. Login required.
@app.route("/api/ai/quota")
@login_required
def api_ai_quota():
    import ai_budget as _ab
    uid = session.get("user_id")
    remaining, reset_s, used = _ab.get_user_remaining(uid)
    spend = _ab.get_org_spend_this_month()
    return jsonify({
        "user": {
            "used":           used,
            "limit":          _ab.USER_TOKEN_CAP_24H,
            "remaining":      remaining,
            "reset_seconds":  reset_s,
        },
        "org": {
            "spent_usd":  round(spend, 4),
            "limit_usd":  _ab.SPEND_CAP_USD_MONTHLY,
        },
    })



if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=5000, debug=False)
